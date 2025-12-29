"""
File management routes for the Flask application.
"""
import json
from datetime import datetime  # Import datetime
from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify
from flask_login import login_required, current_user
from models import (File, PREntry, InquiryDetails, TraceDetails, DisciplinaryAction, Institution, 
                    ReportSoughtDetails, ReportAskedDetails, PreliminaryStatement, Rule15Statement,
                    RTIApplication, CourtCase, WomenHarassmentCase, ComplaintDetails, CMOPortalDetails,
                    RVUDetails, SocialSecurityPension, KESCPCRDetails, KHRCDetails, SCSTDetails,
                    KWCDetails, VigilanceACDetails, RajyaLokNiyamasabhaDetails, PoliceCaseDetails,
                    AttackOnDoctorsCase, AttackOnStaffsCase)
from extensions import db, csrf
from sqlalchemy import or_

files_bp = Blueprint('files', __name__)

# Constants
FILE_TYPES = [
    "Women Harassment", "Police Case", "Medical Negligence",
    "Attack on Doctors", "Attack on Staffs", "Unauthorised Absence",
    "RTI", "Duty Lapse", "Private Practice", "Denial of Treatment",
    "Social Security Pension", "Others"
]

CATEGORIES = [
    "CMO Portal", "RVU", "KeSCPCR", "KHRC", "NKS",
    "SC/ST", "KWC", "Court Case", "Rajya/Lok/Niyamasabha",
    "Vig & Anti Corruption", "Complaint", "Others"
]

FILE_STATUSES = [
    "With Clerk", "Submitted", "Despatched", "Closed", "Mailed & Despatched",
    "Parked", "Stock File", "Attached", "Speak", "Handed Over"
]

INSTITUTION_TYPES = [
    "CHC", "GH", "FHC", "PHC", "DH", "THQH", "TH", "DMO", "Other Units", "Others"
]


@files_bp.route('/management')
@login_required
def file_management():
    """File Management page - similar to desktop app's first tab."""
    # Get active files (not closed and not handed over), ordered by last modified
    active_files = File.query.filter(
        or_(File.is_closed == 0, File.is_closed == None),
        or_(File.status != 'Handed Over', File.status == None)
    ).order_by(File.last_modified.desc()).all()
    
    # Get closed files
    closed_files = File.query.filter(File.is_closed == 1).order_by(File.last_modified.desc()).all()
    
    # Get handed over files
    handed_over_files = File.query.filter(
        or_(File.is_closed == 0, File.is_closed == None),
        File.status == 'Handed Over'
    ).order_by(File.last_modified.desc()).all()
    
    # Get current file if specified
    current_file = None
    inquiry_details = None
    report_sought_list = []
    report_asked_list = []
    file_number = request.args.get('file')
    trace_details = None
    if file_number:
        current_file = File.query.filter_by(file_number=file_number).first()
        if current_file:
            # Get inquiry details for this file
            inquiry_details = InquiryDetails.query.filter_by(file_number=file_number).first()
            # Get report sought list for this file
            report_sought_list = ReportSoughtDetails.query.filter_by(file_number=file_number).all()
            # Get report asked list for this file
            report_asked_list = ReportAskedDetails.query.filter_by(file_number=file_number).all()
            # Get trace details for this file
            trace_details = TraceDetails.query.filter_by(file_number=file_number).first()
    
    # Get institutions for autocomplete
    institutions = Institution.query.order_by(Institution.name.asc()).all()
    
    # Get current year for default values
    current_year = datetime.now().year
    
    return render_template('files/management.html',
                          active_files=active_files,
                          closed_files=closed_files,
                          handed_over_files=handed_over_files,
                          current_file=current_file,
                          inquiry_details=inquiry_details,
                          report_sought_list=report_sought_list,
                          report_asked_list=report_asked_list,
                          trace_details=trace_details,
                          file_types=FILE_TYPES,
                          categories=CATEGORIES,
                          statuses=FILE_STATUSES,
                          institution_types=INSTITUTION_TYPES,
                          institutions=institutions,
                          current_year=current_year)




@files_bp.route('/management/action', methods=['POST'])
@login_required
def file_management_action():
    """Handle file management actions (create, update, reopen)."""
    action = request.form.get('action', 'create')
    file_number = request.form.get('file_number', '').strip()
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    if action == 'create':
        # Check if file already exists
        existing = File.query.filter_by(file_number=file_number).first()
        if existing:
            flash(f'File "{file_number}" already exists.', 'danger')
            return redirect(url_for('files.file_management'))
        
        # Create new file
        selected_types = request.form.getlist('type_of_file')
        selected_categories = request.form.getlist('category')
        
        new_file = File(
            file_number=file_number,
            file_type=request.form.get('file_type', 'E-Office'),
            subject=request.form.get('subject', '').strip(),
            details_of_file=request.form.get('details_of_file', '').strip(),
            status=request.form.get('status', ''),
            category=json.dumps(selected_categories) if selected_categories else None,
            type_of_file=json.dumps(selected_types) if selected_types else None,
            disciplinary_action=request.form.get('disciplinary_action', 'No'),
            institution_type=request.form.get('institution_type', ''),
            institution_name=request.form.get('institution_name', ''),
            follow_up_date=request.form.get('follow_up_date', ''),
            remarks=request.form.get('remarks', ''),
            almirah=request.form.get('almirah', ''),
            rack=request.form.get('rack', ''),
            row=request.form.get('row', ''),
            file_year=request.form.get('file_year', ''),
            last_modified=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        )
        
        db.session.add(new_file)
        db.session.commit()
        
        flash(f'File "{file_number}" created successfully.', 'success')
        return redirect(url_for('files.file_management', file=file_number))
    
    elif action == 'update':
        file = File.query.filter_by(file_number=file_number).first()
        if not file:
            flash(f'File "{file_number}" not found.', 'danger')
            return redirect(url_for('files.file_management'))
        
        # Update file
        selected_types = request.form.getlist('type_of_file')
        selected_categories = request.form.getlist('category')
        
        file.file_type = request.form.get('file_type', file.file_type)
        file.subject = request.form.get('subject', '').strip()
        file.details_of_file = request.form.get('details_of_file', '').strip()
        file.status = request.form.get('status', '')
        file.category = json.dumps(selected_categories) if selected_categories else None
        file.type_of_file = json.dumps(selected_types) if selected_types else None
        file.disciplinary_action = request.form.get('disciplinary_action', 'No')
        file.institution_type = request.form.get('institution_type', '')
        file.institution_name = request.form.get('institution_name', '')
        file.follow_up_date = request.form.get('follow_up_date', '')
        file.remarks = request.form.get('remarks', '')
        file.almirah = request.form.get('almirah', '')
        file.rack = request.form.get('rack', '')
        file.row = request.form.get('row', '')
        file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        db.session.commit()
        
        flash(f'File "{file_number}" updated successfully.', 'success')
        return redirect(url_for('files.file_management', file=file_number))
    
    elif action == 'reopen':
        file = File.query.filter_by(file_number=file_number).first()
        if not file:
            flash(f'File "{file_number}" not found.', 'danger')
            return redirect(url_for('files.file_management'))
        
        file.is_closed = False
        file.closed_date = None
        file.closing_remarks = None
        file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        db.session.commit()
        
        flash(f'File "{file_number}" reopened successfully.', 'success')
        return redirect(url_for('files.file_management', file=file_number))
    
    return redirect(url_for('files.file_management'))


@files_bp.route('/close', methods=['POST'])
@login_required
def close_file():
    """Close a file."""
    file_number = request.form.get('file_number', '').strip()
    closing_remarks = request.form.get('closing_remarks', '').strip()
    
    file = File.query.filter_by(file_number=file_number).first()
    if not file:
        flash(f'File "{file_number}" not found.', 'danger')
        return redirect(url_for('files.file_management'))
    
    file.is_closed = True
    file.closed_date = datetime.now().strftime('%Y-%m-%d')
    file.closing_remarks = closing_remarks
    file.status = 'Closed'
    file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    db.session.commit()
    
    flash(f'File "{file_number}" closed successfully.', 'success')
    return redirect(url_for('files.file_management'))


@files_bp.route('/delete', methods=['POST'])
@login_required
def delete_file():
    """Delete a file."""
    if not current_user.is_admin():
        flash('Only administrators can delete files.', 'danger')
        return redirect(url_for('files.file_management'))
    
    file_number = request.form.get('file_number', '').strip()
    
    file = File.query.filter_by(file_number=file_number).first()
    if not file:
        flash(f'File "{file_number}" not found.', 'danger')
        return redirect(url_for('files.file_management'))
    
    db.session.delete(file)
    db.session.commit()
    
    flash(f'File "{file_number}" deleted successfully.', 'success')
    return redirect(url_for('files.file_management'))


@files_bp.route('/change-mode', methods=['POST'])
@login_required
def change_file_mode():
    """Change file mode between E-Office and Physical, updating file number format."""
    old_file_number = request.form.get('old_file_number', '').strip()
    new_mode = request.form.get('new_mode', '').strip()
    
    if not old_file_number or not new_mode:
        return jsonify({'success': False, 'error': 'Missing required parameters'}), 400
    
    # Get the file using raw SQL to avoid session issues
    result = db.session.execute(
        db.text('SELECT file_number, file_type FROM files WHERE file_number = :fn'),
        {'fn': old_file_number}
    ).fetchone()
    
    if not result:
        return jsonify({'success': False, 'error': f'File "{old_file_number}" not found'}), 404
    
    old_file_type = result[1]
    
    # Parse the old file number and convert to new format
    try:
        new_file_number = convert_file_number_format(old_file_number, new_mode)
        if not new_file_number:
            return jsonify({'success': False, 'error': 'Could not convert file number format. Please ensure the file number follows the correct format.'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error converting file number: {str(e)}'}), 400
    
    # Check if new file number already exists
    existing = db.session.execute(
        db.text('SELECT file_number FROM files WHERE file_number = :fn'),
        {'fn': new_file_number}
    ).fetchone()
    
    if existing:
        return jsonify({'success': False, 'error': f'File number "{new_file_number}" already exists'}), 400
    
    # Update file number in all related tables using raw SQL for foreign key constraints
    try:
        # Get all related records first
        related_tables = [
            ('pr_entries', 'file_number'),
            ('inquiry_details', 'file_number'),
            ('trace_details', 'file_number'),
            ('disciplinary_action_details', 'file_number'),
            ('report_sought_details', 'file_number'),
            ('report_asked_details', 'file_number'),
            ('rti_applications', 'file_number'),
            ('court_case_details', 'file_number'),
            ('women_harassment_cases', 'file_number'),
            ('complaint_details', 'file_number'),
            ('cmo_portal_details', 'file_number'),
            ('rvu_details', 'file_number'),
            ('social_security_pension_details', 'file_number'),
            ('kescpcr_details', 'file_number'),
            ('khrc_details', 'file_number'),
            ('preliminary_statements', 'file_number'),
            ('rule15_statements', 'file_number'),
        ]
        
        # Temporarily disable foreign key checks for SQLite
        db.session.execute(db.text('PRAGMA foreign_keys=OFF'))
        
        # Update file number in main files table
        db.session.execute(
            db.text('UPDATE files SET file_number = :new, file_type = :mode, last_modified = :modified WHERE file_number = :old'),
            {'new': new_file_number, 'mode': new_mode, 'old': old_file_number, 'modified': datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        )
        
        # Update all related tables
        for table_name, column_name in related_tables:
            try:
                db.session.execute(
                    db.text(f'UPDATE {table_name} SET {column_name} = :new WHERE {column_name} = :old'),
                    {'new': new_file_number, 'old': old_file_number}
                )
            except Exception as table_error:
                # Table might not exist, continue
                print(f"Note: Could not update table {table_name}: {table_error}")
                pass
        
        # Re-enable foreign key checks
        db.session.execute(db.text('PRAGMA foreign_keys=ON'))
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': f'File mode changed successfully from {old_file_type} to {new_mode}',
            'old_file_number': old_file_number,
            'new_file_number': new_file_number
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': f'Database error: {str(e)}'}), 500


def convert_file_number_format(file_number, new_mode):
    """
    Convert file number between E-Office and Physical formats.
    
    E-Office format: DMOH-TVM/[Number]/[Year]-[Section]
    Physical format: [Section]-[Number]/[Year]/DMO(H)T
    """
    import re
    
    # Try to parse as E-Office format: DMOH-TVM/Number/Year-Section
    eoffice_pattern = r'^DMOH-TVM/(\d+)/(\d{4})-([A-Za-z0-9]+)$'
    eoffice_match = re.match(eoffice_pattern, file_number)
    
    # Try to parse as Physical format: Section-Number/Year/DMO(H)T
    physical_pattern = r'^([A-Za-z0-9]+)-(\d+)/(\d{4})/DMO\(H\)T$'
    physical_match = re.match(physical_pattern, file_number)
    
    if eoffice_match and new_mode == 'Physical':
        # Convert from E-Office to Physical
        number = eoffice_match.group(1)
        year = eoffice_match.group(2)
        section = eoffice_match.group(3)
        return f'{section}-{number}/{year}/DMO(H)T'
    
    elif physical_match and new_mode == 'E-Office':
        # Convert from Physical to E-Office
        section = physical_match.group(1)
        number = physical_match.group(2)
        year = physical_match.group(3)
        return f'DMOH-TVM/{number}/{year}-{section}'
    
    # Try alternative patterns for variations
    # E-Office without section: DMOH-TVM/Number/Year
    eoffice_no_section = r'^DMOH-TVM/(\d+)/(\d{4})$'
    eoffice_ns_match = re.match(eoffice_no_section, file_number)
    if eoffice_ns_match and new_mode == 'Physical':
        number = eoffice_ns_match.group(1)
        year = eoffice_ns_match.group(2)
        return f'A6-{number}/{year}/DMO(H)T'  # Default section A6
    
    # Physical without section: Number/Year/DMO(H)T
    physical_no_section = r'^(\d+)/(\d{4})/DMO\(H\)T$'
    physical_ns_match = re.match(physical_no_section, file_number)
    if physical_ns_match and new_mode == 'E-Office':
        number = physical_ns_match.group(1)
        year = physical_ns_match.group(2)
        return f'DMOH-TVM/{number}/{year}-A6'  # Default section A6
    
    return None


@files_bp.route('/')
@login_required
def list_files():
    """List all files with search and filtering."""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Get filter parameters
    search_query = request.args.get('q', '').strip()
    file_type_filter = request.args.get('file_type', '')
    status_filter = request.args.get('status', '')
    category_filter = request.args.get('category', '')
    is_closed_filter = request.args.get('is_closed', '')
    file_format = request.args.get('file_format', '')  # Physical or E-Office
    sort_by = request.args.get('sort', '')
    sort_order = request.args.get('order', 'desc')
    
    # Build query
    query = File.query
    
    if search_query:
        query = query.filter(
            or_(
                File.file_number.ilike(f'%{search_query}%'),
                File.subject.ilike(f'%{search_query}%'),
                File.details_of_file.ilike(f'%{search_query}%'),
                File.institution_name.ilike(f'%{search_query}%')
            )
        )
    
    if file_type_filter:
        query = query.filter(File.type_of_file.ilike(f'%{file_type_filter}%'))
    
    if status_filter:
        query = query.filter(File.status == status_filter)
    
    if category_filter:
        query = query.filter(File.category.ilike(f'%{category_filter}%'))
    
    if is_closed_filter == 'active':
        query = query.filter(or_(File.is_closed == 0, File.is_closed == None))
    elif is_closed_filter == 'closed':
        query = query.filter(File.is_closed == 1)
    
    if file_format:
        query = query.filter(File.file_type == file_format)
    
    # Apply sorting
    if sort_by:
        sort_column = None
        if sort_by == 'file_number':
            sort_column = File.file_number
        elif sort_by == 'subject':
            sort_column = File.subject
        elif sort_by == 'status':
            sort_column = File.status
        elif sort_by == 'institution_name':
            sort_column = File.institution_name
        elif sort_by == 'follow_up_date':
            sort_column = File.follow_up_date
        
        if sort_column is not None:
            if sort_order == 'desc':
                query = query.order_by(sort_column.desc())
            else:
                query = query.order_by(sort_column.asc())
        else:
            query = query.order_by(File.file_number.desc())
    else:
        query = query.order_by(File.file_number.desc())
    
    # Paginate
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    files = pagination.items
    
    return render_template('files/list.html', 
                          files=files, 
                          pagination=pagination,
                          search_query=search_query,
                          file_type_filter=file_type_filter,
                          status_filter=status_filter,
                          category_filter=category_filter,
                          is_closed_filter=is_closed_filter,
                          file_format=file_format,
                          file_types=FILE_TYPES,
                          categories=CATEGORIES,
                          statuses=FILE_STATUSES,
                          sort_by=sort_by,
                          sort_order=sort_order)


@files_bp.route('/create', methods=['GET', 'POST'])
@login_required
def create_file():
    """Create a new file."""
    if request.method == 'POST':
        file_number = request.form.get('file_number', '').strip()
        
        if not file_number:
            flash('File number is required.', 'danger')
            return render_template('files/create.html', 
                                  file_types=FILE_TYPES,
                                  categories=CATEGORIES,
                                  statuses=FILE_STATUSES)
        
        # Check if file already exists
        existing = File.query.filter_by(file_number=file_number).first()
        if existing:
            flash(f'File number "{file_number}" already exists.', 'danger')
            return render_template('files/create.html', 
                                  file_types=FILE_TYPES,
                                  categories=CATEGORIES,
                                  statuses=FILE_STATUSES)
        
        # Get form data
        selected_types = request.form.getlist('type_of_file')
        selected_categories = request.form.getlist('category')
        
        new_file = File(
            file_number=file_number,
            file_type=request.form.get('file_type', 'Physical'),
            subject=request.form.get('subject', '').strip(),
            details_of_file=request.form.get('details_of_file', '').strip(),
            status=request.form.get('status', ''),
            category=json.dumps(selected_categories) if selected_categories else None,
            type_of_file=json.dumps(selected_types) if selected_types else None,
            disciplinary_action=request.form.get('disciplinary_action', 'No'),
            institution_type=request.form.get('institution_type', ''),
            institution_name=request.form.get('institution_name', ''),
            follow_up_date=request.form.get('follow_up_date', ''),
            remarks=request.form.get('remarks', ''),
            almirah=request.form.get('almirah', ''),
            rack=request.form.get('rack', ''),
            row=request.form.get('row', ''),
            migrated_to_eoffice=request.form.get('migrated_to_eoffice', 'No'),
            eoffice_file_number=request.form.get('eoffice_file_number', '')
        )
        
        db.session.add(new_file)
        db.session.commit()
        
        flash(f'File "{file_number}" created successfully.', 'success')
        return redirect(url_for('files.view_file', file_number=file_number))
    
    return render_template('files/create.html', 
                          file_types=FILE_TYPES,
                          categories=CATEGORIES,
                          statuses=FILE_STATUSES)


@files_bp.route('/<path:file_number>')
@login_required
def view_file(file_number):
    """View file details."""
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    # Parse JSON fields
    categories = []
    types = []
    try:
        if file.category:
            categories = json.loads(file.category)
    except:
        categories = [file.category] if file.category else []
    
    try:
        if file.type_of_file:
            types = json.loads(file.type_of_file)
    except:
        types = [file.type_of_file] if file.type_of_file else []
    
    # Get related data
    pr_entries = PREntry.query.filter_by(file_number=file_number).all()
    inquiry = InquiryDetails.query.filter_by(file_number=file_number).first()
    disciplinary_actions = DisciplinaryAction.query.filter_by(file_number=file_number).all()
    
    return render_template('files/view.html', 
                          file=file,
                          categories=categories,
                          types=types,
                          pr_entries=pr_entries,
                          inquiry=inquiry,
                          disciplinary_actions=disciplinary_actions)


@files_bp.route('/<path:file_number>/edit', methods=['GET', 'POST'])
@login_required
def edit_file(file_number):
    """Edit file details."""
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    if request.method == 'POST':
        # Get form data
        file.subject = request.form.get('subject', '').strip()
        file.details_of_file = request.form.get('details_of_file', '').strip()
        file.status = request.form.get('status', '')
        file.category = request.form.get('category', '')
        file.type_of_file = request.form.get('type_of_file', '')
        file.file_type = request.form.get('file_type', '')
        file.file_year = request.form.get('file_year', '')
        file.disciplinary_action = request.form.get('disciplinary_action', '')
        file.institution_type = request.form.get('institution_type', '')
        file.institution_name = request.form.get('institution_name', '')
        file.follow_up_date = request.form.get('follow_up_date', '')
        file.remarks = request.form.get('remarks', '')
        file.almirah = request.form.get('almirah', '')
        file.rack = request.form.get('rack', '')
        file.row = request.form.get('row', '')
        file.migrated_to_eoffice = request.form.get('migrated_to_eoffice', '')
        file.eoffice_file_number = request.form.get('eoffice_file_number', '')
        file.original_file_number = request.form.get('original_file_number', '')
        file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        db.session.commit()
        
        flash(f'File "{file_number}" updated successfully.', 'success')
        return redirect(url_for('files.view_file', file_number=file_number))
    
    return render_template('files/edit.html', 
                          file=file,
                          file_types=FILE_TYPES,
                          categories=CATEGORIES,
                          statuses=FILE_STATUSES)


@files_bp.route('/<path:file_number>/close', methods=['POST'])
@login_required
def close_file_by_number(file_number):
    """Close a file by file number."""
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    file.is_closed = True
    file.closed_date = request.form.get('closed_date', '')
    file.closing_remarks = request.form.get('closing_remarks', '')
    file.status = 'Closed'
    file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    db.session.commit()
    
    flash(f'File "{file_number}" has been closed.', 'success')
    return redirect(url_for('files.view_file', file_number=file_number))


@files_bp.route('/<path:file_number>/reopen', methods=['POST'])
@login_required
def reopen_file(file_number):
    """Reopen a closed file."""
    if not current_user.is_admin():
        flash('Only administrators can reopen files.', 'danger')
        return redirect(url_for('files.view_file', file_number=file_number))
    
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    file.is_closed = False
    file.status = 'With Clerk'
    file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    db.session.commit()
    
    flash(f'File "{file_number}" has been reopened.', 'success')
    return redirect(url_for('files.view_file', file_number=file_number))


@files_bp.route('/<path:file_number>/delete', methods=['POST'])
@login_required
def delete_file_by_number(file_number):
    """Delete a file by file number."""
    if not current_user.is_admin():
        flash('Only administrators can delete files.', 'danger')
        return redirect(url_for('files.view_file', file_number=file_number))
    
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    db.session.delete(file)
    db.session.commit()
    
    flash(f'File "{file_number}" has been deleted.', 'success')
    return redirect(url_for('files.list_files'))


@files_bp.route('/<path:file_number>/pr-entries', methods=['GET', 'POST'])
@login_required
def manage_pr_entries(file_number):
    """Manage PR entries for a file."""
    file = File.query.filter_by(file_number=file_number).first_or_404()
    
    if request.method == 'POST':
        entry_id = request.form.get('entry_id', '').strip()
        
        # Get form values
        serial_number = request.form.get('serial_number', '').strip()
        current_number = request.form.get('current_number', '').strip()
        date_receipt_clerk = request.form.get('date_receipt_clerk', '').strip()
        title = request.form.get('title', '').strip()
        from_whom_outside_name = request.form.get('from_whom_outside_name', '').strip()
        
        # Validation - require at least one field to be filled
        if not any([serial_number, current_number, date_receipt_clerk, title, from_whom_outside_name]):
            flash('Please fill at least one field before adding an entry.', 'warning')
            return redirect(url_for('files.manage_pr_entries', file_number=file_number))
        
        try:
            if entry_id:
                # Update existing entry
                pr_entry = PREntry.query.get_or_404(int(entry_id))
                pr_entry.serial_number = serial_number
                pr_entry.current_number = current_number
                pr_entry.date_receipt_clerk = date_receipt_clerk
                pr_entry.title = title
                pr_entry.from_whom_outside_name = from_whom_outside_name
                pr_entry.from_whom_outside_number = request.form.get('from_whom_outside_number', '')
                pr_entry.from_whom_outside_date = request.form.get('from_whom_outside_date', '')
                pr_entry.submitted_by_clerk_date = request.form.get('submitted_by_clerk_date', '')
                pr_entry.return_to_clerk_date = request.form.get('return_to_clerk_date', '')
                pr_entry.reference_issued_to_whom = request.form.get('reference_issued_to_whom', '')
                pr_entry.reference_issued_date = request.form.get('reference_issued_date', '')
                pr_entry.reply_fresh_current_from_whom = request.form.get('reply_fresh_current_from_whom', '')
                pr_entry.reply_fresh_current_number = request.form.get('reply_fresh_current_number', '')
                pr_entry.reply_fresh_current_date = request.form.get('reply_fresh_current_date', '')
                pr_entry.date_receipt_clerk_fresh = request.form.get('date_receipt_clerk_fresh', '')
                pr_entry.disposal_nature = request.form.get('disposal_nature', '')
                pr_entry.disposal_date = request.form.get('disposal_date', '')
                
                # Update file's last_modified
                file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                
                db.session.commit()
                flash('PR Entry updated successfully.', 'success')
            else:
                # Create new entry
                pr_entry = PREntry(
                    file_number=file_number,
                    serial_number=serial_number,
                    current_number=current_number,
                    date_receipt_clerk=date_receipt_clerk,
                    title=title,
                    from_whom_outside_name=from_whom_outside_name,
                    from_whom_outside_number=request.form.get('from_whom_outside_number', ''),
                    from_whom_outside_date=request.form.get('from_whom_outside_date', ''),
                    submitted_by_clerk_date=request.form.get('submitted_by_clerk_date', ''),
                    return_to_clerk_date=request.form.get('return_to_clerk_date', ''),
                    reference_issued_to_whom=request.form.get('reference_issued_to_whom', ''),
                    reference_issued_date=request.form.get('reference_issued_date', ''),
                    reply_fresh_current_from_whom=request.form.get('reply_fresh_current_from_whom', ''),
                    reply_fresh_current_number=request.form.get('reply_fresh_current_number', ''),
                    reply_fresh_current_date=request.form.get('reply_fresh_current_date', ''),
                    date_receipt_clerk_fresh=request.form.get('date_receipt_clerk_fresh', ''),
                    disposal_nature=request.form.get('disposal_nature', ''),
                    disposal_date=request.form.get('disposal_date', '')
                )
                
                db.session.add(pr_entry)
                
                # Update file's last_modified
                file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                
                db.session.commit()
                flash('PR Entry added successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Database error: {str(e)}. Please contact administrator to fix sequence.', 'danger')
        
        return redirect(url_for('files.manage_pr_entries', file_number=file_number))
    
    pr_entries = PREntry.query.filter_by(file_number=file_number).order_by(PREntry.serial_number.asc()).all()
    
    return render_template('files/pr_entries.html', file=file, pr_entries=pr_entries)


@files_bp.route('/pr-entry/<int:entry_id>/json')
@login_required
def get_pr_entry_json(entry_id):
    """Get PR entry data as JSON for AJAX editing."""
    entry = PREntry.query.get_or_404(entry_id)
    return jsonify({
        'success': True,
        'entry': {
            'id': entry.id,
            'file_number': entry.file_number,
            'serial_number': entry.serial_number or '',
            'current_number': entry.current_number or '',
            'date_receipt_clerk': entry.date_receipt_clerk or '',
            'title': entry.title or '',
            'from_whom_outside_name': entry.from_whom_outside_name or '',
            'from_whom_outside_number': entry.from_whom_outside_number or '',
            'from_whom_outside_date': entry.from_whom_outside_date or '',
            'submitted_by_clerk_date': entry.submitted_by_clerk_date or '',
            'return_to_clerk_date': entry.return_to_clerk_date or '',
            'reference_issued_to_whom': entry.reference_issued_to_whom or '',
            'reference_issued_date': entry.reference_issued_date or '',
            'reply_fresh_current_from_whom': entry.reply_fresh_current_from_whom or '',
            'reply_fresh_current_number': entry.reply_fresh_current_number or '',
            'reply_fresh_current_date': entry.reply_fresh_current_date or '',
            'date_receipt_clerk_fresh': entry.date_receipt_clerk_fresh or '',
            'disposal_nature': entry.disposal_nature or '',
            'disposal_date': entry.disposal_date or ''
        }
    })


@files_bp.route('/pr-entry/<int:entry_id>/delete', methods=['POST'])
@login_required
def delete_pr_entry(entry_id):
    """Delete a PR entry."""
    entry = PREntry.query.get_or_404(entry_id)
    file_number = entry.file_number
    
    db.session.delete(entry)
    db.session.commit()
    
    flash('PR Entry deleted successfully.', 'success')
    return redirect(url_for('files.manage_pr_entries', file_number=file_number))


@files_bp.route('/api/search')
@login_required
def api_search_files():
    """API endpoint for file search (autocomplete)."""
    query = request.args.get('q', '').strip()
    
    if len(query) < 2:
        return jsonify([])
    
    files = File.query.filter(
        or_(
            File.file_number.ilike(f'%{query}%'),
            File.subject.ilike(f'%{query}%')
        )
    ).limit(10).all()
    
    results = [{
        'file_number': f.file_number,
        'subject': f.subject or '',
        'status': f.status or ''
    } for f in files]
    
    return jsonify(results)


@files_bp.route('/export')
@login_required
def export_files():
    """Export files to CSV."""
    import csv
    from io import StringIO
    from flask import Response
    
    # Get filter parameters (same as list_files)
    search_query = request.args.get('q', '').strip()
    file_type_filter = request.args.get('file_type', '')
    status_filter = request.args.get('status', '')
    
    # Build query
    query = File.query
    
    if search_query:
        query = query.filter(
            or_(
                File.file_number.ilike(f'%{search_query}%'),
                File.subject.ilike(f'%{search_query}%'),
                File.details_of_file.ilike(f'%{search_query}%')
            )
        )
    
    if file_type_filter:
        query = query.filter(File.type_of_file.ilike(f'%{file_type_filter}%'))
    
    if status_filter:
        query = query.filter(File.status == status_filter)
    
    files = query.all()
    
    # Create CSV
    output = StringIO()
    # Add UTF-8 BOM for Excel to recognize Unicode characters (Malayalam etc.)
    output.write('\ufeff')
    writer = csv.writer(output)
    
    # Header row
    writer.writerow([
        'File Number', 'File Type', 'Subject', 'Details', 'Status',
        'Category', 'Type of File', 'Institution', 'Is Closed'
    ])
    
    # Data rows
    for f in files:
        writer.writerow([
            f.file_number,
            f.file_type or '',
            f.subject or '',
            f.details_of_file or '',
            f.status or '',
            f.category or '',
            f.type_of_file or '',
            f.institution_name or '',
            'Yes' if f.is_closed else 'No'
        ])
    
    output.seek(0)
    return Response(
        output.getvalue().encode('utf-8-sig'),
        mimetype='text/csv; charset=utf-8-sig',
        headers={'Content-Disposition': 'attachment; filename=files_export.csv'}
    )


# =====================================================
# INQUIRY DETAILS ROUTES
# =====================================================

@files_bp.route('/inquiry-details/save', methods=['POST'])
@login_required
def save_inquiry_details():
    """Save or update inquiry details for a file."""
    file_number = request.form.get('file_number', '').strip()
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    try:
        # Check if inquiry details already exist for this file
        inquiry = InquiryDetails.query.filter_by(file_number=file_number).first()
        
        if not inquiry:
            inquiry = InquiryDetails(file_number=file_number)
            db.session.add(inquiry)
        
        # Update preliminary inquiry fields (use 1/0 for PostgreSQL integer columns)
        inquiry.prelim_conducted = 1 if 'prelim_conducted' in request.form else 0
        inquiry.prelim_io_name = request.form.get('prelim_io_name', '').strip()
        inquiry.prelim_inquiry_date = request.form.get('prelim_inquiry_date', '')
        inquiry.prelim_venue = request.form.get('prelim_venue', '').strip()
        inquiry.prelim_report_submitted = 1 if 'prelim_report_submitted' in request.form else 0
        inquiry.prelim_report_to = request.form.get('prelim_report_to', '').strip()
        inquiry.prelim_report_date = request.form.get('prelim_report_date', '')
        
        # Update Rule 15(ii) inquiry fields (use 1/0 for PostgreSQL integer columns)
        inquiry.rule15_conducted = 1 if 'rule15_conducted' in request.form else 0
        inquiry.rule15_io_name = request.form.get('rule15_io_name', '').strip()
        inquiry.rule15_inquiry_date = request.form.get('rule15_inquiry_date', '')
        inquiry.rule15_venue = request.form.get('rule15_venue', '').strip()
        inquiry.rule15_report_submitted = 1 if 'rule15_report_submitted' in request.form else 0
        inquiry.rule15_report_to = request.form.get('rule15_report_to', '').strip()
        inquiry.rule15_report_date = request.form.get('rule15_report_date', '')
        
        # Update file's last_modified
        file = File.query.filter_by(file_number=file_number).first()
        if file:
            file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        db.session.commit()
        flash('Inquiry details saved successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error saving inquiry details: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


# =====================================================
# REPORT SOUGHT ROUTES
# =====================================================

@files_bp.route('/report-sought/save', methods=['POST'])
@login_required
def save_report_sought():
    """Save report sought record(s) for a file - handles multiple bodies."""
    import json
    file_number = request.form.get('file_number', '').strip()
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    # Get the file to get subject
    file = File.query.filter_by(file_number=file_number).first()
    subject = file.subject if file else ''
    institution = request.form.get('institution', '').strip()
    
    # Check if multiple bodies data is provided (from desktop-style form)
    details_json = request.form.get('details', '').strip()
    
    try:
        if details_json and details_json.startswith('['):
            # Parse JSON data for multiple bodies
            bodies_data = json.loads(details_json)
            records_added = 0
            
            for body_data in bodies_data:
                report_sought = ReportSoughtDetails(
                    file_number=file_number,
                    subject=subject,
                    body=body_data.get('body', ''),
                    report_sought_date=body_data.get('sought_date', ''),
                    status='Pending' if body_data.get('submitted') != 'Yes' else 'Submitted',
                    institution=institution,
                    details='',
                    submitted=body_data.get('submitted', 'No'),
                    submitted_date=body_data.get('submitted_date', '')
                )
                db.session.add(report_sought)
                records_added += 1
            
            if records_added > 0:
                # Update file's last_modified
                if file:
                    file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                db.session.commit()
                flash(f'{records_added} Report Sought record(s) added successfully.', 'success')
            else:
                flash('No bodies selected. Please select at least one body.', 'warning')
        else:
            # Fallback to single record (from edit form or legacy)
            report_sought = ReportSoughtDetails(
                file_number=file_number,
                subject=subject,
                body=request.form.get('body', '').strip(),
                report_sought_date=request.form.get('report_sought_date', ''),
                status=request.form.get('status', '').strip(),
                institution=institution,
                details=details_json,
                submitted=request.form.get('submitted', 'No'),
                submitted_date=request.form.get('submitted_date', '')
            )
            db.session.add(report_sought)
            # Update file's last_modified
            if file:
                file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            db.session.commit()
            flash('Report Sought record added successfully.', 'success')
            
    except Exception as e:
        db.session.rollback()
        flash(f'Error saving Report Sought record: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


@files_bp.route('/report-sought/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def edit_report_sought(id):
    """Edit a report sought record."""
    record = ReportSoughtDetails.query.get_or_404(id)
    
    if request.method == 'POST':
        record.body = request.form.get('body', '').strip()
        record.report_sought_date = request.form.get('report_sought_date', '')
        record.status = request.form.get('status', '').strip()
        record.institution = request.form.get('institution', '').strip()
        record.details = request.form.get('details', '').strip()
        record.submitted = request.form.get('submitted', 'No')
        record.submitted_date = request.form.get('submitted_date', '')
        
        try:
            # Update file's last_modified
            file = File.query.filter_by(file_number=record.file_number).first()
            if file:
                file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            db.session.commit()
            flash('Report Sought record updated successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating Report Sought record: {str(e)}', 'danger')
        
        return redirect(url_for('files.file_management', file=record.file_number))
    
    # GET request - render edit form
    return render_template('files/edit_report_sought.html', record=record)


@files_bp.route('/report-sought/<int:id>/delete')
@login_required
def delete_report_sought(id):
    """Delete a report sought record."""
    record = ReportSoughtDetails.query.get_or_404(id)
    file_number = record.file_number
    
    db.session.delete(record)
    
    try:
        db.session.commit()
        flash('Report Sought record deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting Report Sought record: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


# =====================================================
# REPORT ASKED ROUTES
# =====================================================

@files_bp.route('/report-asked/save', methods=['POST'])
@login_required
def save_report_asked():
    """Save a new report asked record for a file."""
    file_number = request.form.get('file_number', '').strip()
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    # Create new report asked record
    report_asked = ReportAskedDetails(
        file_number=file_number,
        whether_report_asked=request.form.get('whether_report_asked', 'No'),
        asked_date=request.form.get('asked_date', ''),
        institution_name=request.form.get('institution_name', '').strip(),
        report_submitted=request.form.get('report_submitted', 'No'),
        received_date=request.form.get('received_date', '')
    )
    
    db.session.add(report_asked)
    
    try:
        # Update file's last_modified
        file = File.query.filter_by(file_number=file_number).first()
        if file:
            file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        db.session.commit()
        flash('Report Asked record added successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error saving Report Asked record: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


@files_bp.route('/report-asked/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def edit_report_asked(id):
    """Edit a report asked record."""
    record = ReportAskedDetails.query.get_or_404(id)
    
    if request.method == 'POST':
        record.whether_report_asked = request.form.get('whether_report_asked', 'No')
        record.asked_date = request.form.get('asked_date', '')
        record.institution_name = request.form.get('institution_name', '').strip()
        record.report_submitted = request.form.get('report_submitted', 'No')
        record.received_date = request.form.get('received_date', '')
        
        try:
            # Update file's last_modified
            file = File.query.filter_by(file_number=record.file_number).first()
            if file:
                file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            db.session.commit()
            flash('Report Asked record updated successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating Report Asked record: {str(e)}', 'danger')
        
        return redirect(url_for('files.file_management', file=record.file_number))
    
    # GET request - render edit form
    return render_template('files/edit_report_asked.html', record=record)


@files_bp.route('/report-asked/<int:id>/delete')
@login_required
def delete_report_asked(id):
    """Delete a report asked record."""
    record = ReportAskedDetails.query.get_or_404(id)
    file_number = record.file_number
    
    db.session.delete(record)
    
    try:
        db.session.commit()
        flash('Report Asked record deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting Report Asked record: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


# =====================================================
# TRACE DETAILS ROUTES
# =====================================================

@files_bp.route('/trace-details/save', methods=['POST'])
@login_required
def save_trace_details():
    """Save or update trace details for a file."""
    file_number = request.form.get('file_number', '').strip()
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    # Get trace data from form
    almirah = request.form.get('almirah', '').strip()
    rack = request.form.get('rack', '').strip()
    row = request.form.get('row', '').strip()
    migrated_to_eoffice = request.form.get('migrated_to_eoffice', 'No')
    eoffice_file_number = request.form.get('eoffice_file_number', '').strip()
    
    # Check if this is a migration request
    migration_requested = (migrated_to_eoffice == 'Yes' and eoffice_file_number)
    
    if migration_requested:
        # Validate E-Office file number format
        if not eoffice_file_number.startswith('DMOH-TVM/'):
            flash('E-Office file number must be in format: DMOH-TVM/Number/Year-Section', 'danger')
            return redirect(url_for('files.file_management', file=file_number))
        
        # Check if E-Office file already exists
        existing_eoffice = File.query.filter_by(file_number=eoffice_file_number).first()
        
        if existing_eoffice:
            flash(f'E-Office file {eoffice_file_number} already exists. Migration cannot proceed.', 'danger')
            return redirect(url_for('files.file_management', file=file_number))
        
        # Get the physical file
        physical_file = File.query.filter_by(file_number=file_number).first()
        if not physical_file:
            flash('Physical file not found.', 'danger')
            return redirect(url_for('files.file_management'))
        
        try:
            # Create the E-Office file with the same data
            eoffice_file = File(
                file_number=eoffice_file_number,
                subject=physical_file.subject,
                status=physical_file.status,
                follow_up_date=physical_file.follow_up_date,
                details_of_file=physical_file.details_of_file,
                file_type='E-Office',
                category=physical_file.category,
                type_of_file=physical_file.type_of_file,
                disciplinary_action=physical_file.disciplinary_action,
                institution_type=physical_file.institution_type,
                institution_name=physical_file.institution_name,
                remarks=f"Migrated from physical file: {file_number}. " + (physical_file.remarks or ''),
                is_closed=physical_file.is_closed,
                closing_remarks=physical_file.closing_remarks
            )
            db.session.add(eoffice_file)
            
            # Update the physical file to mark as migrated
            physical_file.remarks = f"Migrated to E-Office: {eoffice_file_number}. " + (physical_file.remarks or '')
            
            # Update or create trace details for physical file
            trace_details = TraceDetails.query.filter_by(file_number=file_number).first()
            if not trace_details:
                trace_details = TraceDetails(file_number=file_number)
                db.session.add(trace_details)
            
            trace_details.almirah = almirah
            trace_details.rack = rack
            trace_details.row = row
            trace_details.migrated_to_eoffice = 'Yes'
            trace_details.eoffice_file_number = eoffice_file_number
            
            db.session.commit()
            flash(f'File successfully migrated from {file_number} to {eoffice_file_number}.', 'success')
            return redirect(url_for('files.file_management', file=eoffice_file_number))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error during migration: {str(e)}', 'danger')
            return redirect(url_for('files.file_management', file=file_number))
    
    # Regular trace details save (no migration)
    trace_details = TraceDetails.query.filter_by(file_number=file_number).first()
    
    if not trace_details:
        trace_details = TraceDetails(file_number=file_number)
        db.session.add(trace_details)
    
    trace_details.almirah = almirah
    trace_details.rack = rack
    trace_details.row = row
    trace_details.migrated_to_eoffice = migrated_to_eoffice
    trace_details.eoffice_file_number = eoffice_file_number if migrated_to_eoffice == 'Yes' else ''
    
    try:
        # Update file's last_modified
        file = File.query.filter_by(file_number=file_number).first()
        if file:
            file.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        db.session.commit()
        flash('Trace details saved successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error saving trace details: {str(e)}', 'danger')
    
    return redirect(url_for('files.file_management', file=file_number))


@files_bp.route('/trace-details/<file_number>')
@login_required
def get_trace_details(file_number):
    """Get trace details for a file (API endpoint for AJAX)."""
    trace_details = TraceDetails.query.filter_by(file_number=file_number).first()
    
    if trace_details:
        return {
            'success': True,
            'data': {
                'almirah': trace_details.almirah or '',
                'rack': trace_details.rack or '',
                'row': trace_details.row or '',
                'migrated_to_eoffice': trace_details.migrated_to_eoffice or 'No',
                'eoffice_file_number': trace_details.eoffice_file_number or ''
            }
        }
    else:
        return {
            'success': True,
            'data': {
                'almirah': '',
                'rack': '',
                'row': '',
                'migrated_to_eoffice': 'No',
                'eoffice_file_number': ''
            }
        }


# =====================================================
# STATEMENT GENERATOR ROUTES
# =====================================================

@files_bp.route('/statement-generator')
@login_required
def statement_generator():
    """Statement Generator page for preliminary or Rule 15(ii) inquiries."""
    file_number = request.args.get('file', '')
    inquiry_type = request.args.get('type', 'preliminary')  # 'preliminary' or 'rule15'
    
    if not file_number:
        flash('File number is required.', 'danger')
        return redirect(url_for('files.file_management'))
    
    # Get the file
    file = File.query.filter_by(file_number=file_number).first()
    if not file:
        flash('File not found.', 'danger')
        return redirect(url_for('files.file_management'))
    
    # Get inquiry details
    inquiry_details = InquiryDetails.query.filter_by(file_number=file_number).first()
    
    # Get existing statement data
    statement = None
    if inquiry_type == 'preliminary':
        statement = PreliminaryStatement.query.filter_by(file_number=file_number).first()
    else:
        statement = Rule15Statement.query.filter_by(file_number=file_number).first()
    
    # Parse JSON data if exists
    basic_info = {}
    respondents = []
    qa_pairs = []
    
    if statement:
        basic_info = json.loads(statement.basic_info or '{}')
        respondents = json.loads(statement.respondents or '[]')
        qa_pairs = json.loads(statement.qa_pairs or '[]')
    
    return render_template('files/statement_generator.html',
                          file=file,
                          file_number=file_number,
                          inquiry_type=inquiry_type,
                          inquiry_details=inquiry_details,
                          basic_info=basic_info,
                          respondents=respondents,
                          qa_pairs=qa_pairs,
                          current_year=datetime.now().year)


@files_bp.route('/statement-generator/save', methods=['POST'])
@login_required
def save_statement():
    """Save statement generator data."""
    file_number = request.form.get('file_number', '').strip()
    inquiry_type = request.form.get('inquiry_type', 'preliminary')
    
    if not file_number:
        return jsonify({'success': False, 'message': 'File number is required.'})
    
    # Get basic_info from JSON
    basic_info_json = request.form.get('basic_info', '{}')
    try:
        basic_info = json.loads(basic_info_json)
    except json.JSONDecodeError:
        basic_info = {
            'subject': request.form.get('subject', ''),
            'date': request.form.get('statement_date', ''),
            'officer': request.form.get('officer', ''),
            'time': request.form.get('statement_time', '')
        }
    
    # Get respondents and QA pairs from JSON
    respondents_json = request.form.get('respondents', '[]')
    qa_pairs_json = request.form.get('qa_pairs', '[]')
    
    try:
        respondents = json.loads(respondents_json)
        qa_pairs = json.loads(qa_pairs_json)
    except json.JSONDecodeError:
        respondents = []
        qa_pairs = []
    
    # Save to database
    try:
        if inquiry_type == 'preliminary':
            statement = PreliminaryStatement.query.filter_by(file_number=file_number).first()
            if not statement:
                statement = PreliminaryStatement(file_number=file_number)
                db.session.add(statement)
        else:
            statement = Rule15Statement.query.filter_by(file_number=file_number).first()
            if not statement:
                statement = Rule15Statement(file_number=file_number)
                db.session.add(statement)
        
        statement.basic_info = json.dumps(basic_info)
        statement.respondents = json.dumps(respondents)
        statement.qa_pairs = json.dumps(qa_pairs)
        statement.last_modified = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        if not statement.created_date:
            statement.created_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Statement saved successfully.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error saving statement: {str(e)}'})


@files_bp.route('/statement-generator/load/<file_number>/<inquiry_type>')
@login_required
def load_statement(file_number, inquiry_type):
    """Load statement data for AJAX."""
    if inquiry_type == 'preliminary':
        statement = PreliminaryStatement.query.filter_by(file_number=file_number).first()
    else:
        statement = Rule15Statement.query.filter_by(file_number=file_number).first()
    
    if statement:
        return jsonify({
            'success': True,
            'data': {
                'basic_info': json.loads(statement.basic_info or '{}'),
                'respondents': json.loads(statement.respondents or '[]'),
                'qa_pairs': json.loads(statement.qa_pairs or '[]')
            }
        })
    else:
        return jsonify({
            'success': True,
            'data': {
                'basic_info': {},
                'respondents': [],
                'qa_pairs': []
            }
        })


@files_bp.route('/statement-generator/delete', methods=['POST'])
@login_required
def delete_statement():
    """Delete statement data from database."""
    data = request.get_json()
    file_number = data.get('file_number', '').strip()
    inquiry_type = data.get('inquiry_type', 'preliminary')
    
    if not file_number:
        return jsonify({'success': False, 'message': 'File number is required.'})
    
    try:
        if inquiry_type == 'preliminary':
            statement = PreliminaryStatement.query.filter_by(file_number=file_number).first()
        else:
            statement = Rule15Statement.query.filter_by(file_number=file_number).first()
        
        if statement:
            db.session.delete(statement)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Statement data deleted successfully.'})
        else:
            return jsonify({'success': False, 'message': 'No statement data found for this file.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error deleting statement: {str(e)}'})


@files_bp.route('/statement-generator/generate-document', methods=['POST'])
@login_required
def generate_statement_document():
    """Generate Word document for statement."""
    from io import BytesIO
    from flask import send_file
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({'success': False, 'message': 'python-docx library is not installed. Please install it with: pip install python-docx'})
    
    data = request.get_json()
    file_number = data.get('file_number', '')
    inquiry_type = data.get('inquiry_type', 'preliminary')
    basic_info = data.get('basic_info', {})
    respondent = data.get('respondent', {})
    qa_pairs = data.get('qa_pairs', [])
    
    # Create Word document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('മൊഴി' if inquiry_type == 'preliminary' else 'Rule 15(ii) മൊഴി')
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # File Information
    info_para = doc.add_paragraph()
    info_para.add_run('ഫയൽ നമ്പർ: ').bold = True
    info_para.add_run(file_number)
    
    subject_para = doc.add_paragraph()
    subject_para.add_run('വിഷയം: ').bold = True
    subject_para.add_run(basic_info.get('subject', ''))
    
    date_para = doc.add_paragraph()
    date_para.add_run('തീയതി: ').bold = True
    date_para.add_run(basic_info.get('date', ''))
    
    time_para = doc.add_paragraph()
    time_para.add_run('സമയം: ').bold = True
    time_para.add_run(basic_info.get('time', ''))
    
    officer_para = doc.add_paragraph()
    officer_para.add_run('അന്വേഷണ ഉദ്യോഗസ്ഥൻ: ').bold = True
    officer_para.add_run(basic_info.get('officer', ''))
    
    doc.add_paragraph()
    
    # Respondent Information
    resp_title = doc.add_paragraph()
    resp_title_run = resp_title.add_run('മൊഴി നൽകുന്നയാൾ:')
    resp_title_run.bold = True
    resp_title_run.underline = True
    
    if respondent.get('type') == 'employee':
        if respondent.get('employee_type') == 'permanent':
            doc.add_paragraph(f"പേര്: {respondent.get('perm_emp_name', '')}")
            doc.add_paragraph(f"PEN: {respondent.get('perm_emp_pen', '')}")
            doc.add_paragraph(f"ഉദ്യോഗപ്പേര്: {respondent.get('perm_emp_desig', '')}")
            doc.add_paragraph(f"സ്ഥാപനം: {respondent.get('perm_emp_inst', '')}")
        else:
            doc.add_paragraph(f"പേര്: {respondent.get('temp_emp_name', '')}")
            doc.add_paragraph(f"ഉദ്യോഗപ്പേര്: {respondent.get('temp_emp_desig', '')}")
            doc.add_paragraph(f"{respondent.get('temp_emp_id_type', 'ആധാർ')}: {respondent.get('temp_emp_id_no', '')}")
            doc.add_paragraph(f"മേൽവിലാസം: {respondent.get('temp_emp_address', '')}")
    else:
        doc.add_paragraph(f"പേര്: {respondent.get('public_name', '')}")
        doc.add_paragraph(f"{respondent.get('public_id_type', 'ആധാർ')}: {respondent.get('public_id_no', '')}")
        doc.add_paragraph(f"മേൽവിലാസം: {respondent.get('public_address', '')}")
    
    doc.add_paragraph()
    
    # Questions and Answers
    qa_title = doc.add_paragraph()
    qa_title_run = qa_title.add_run('ചോദ്യോത്തരങ്ങൾ:')
    qa_title_run.bold = True
    qa_title_run.underline = True
    
    for i, qa in enumerate(qa_pairs, 1):
        q_para = doc.add_paragraph()
        q_para.add_run(f'ചോദ്യം {i}: ').bold = True
        question = qa.get('question', '') if isinstance(qa, dict) else (qa[0] if isinstance(qa, list) else str(qa))
        q_para.add_run(question)
        
        a_para = doc.add_paragraph()
        a_para.add_run(f'ഉത്തരം {i}: ').bold = True
        answer = qa.get('answer', '') if isinstance(qa, dict) else (qa[1] if isinstance(qa, list) and len(qa) > 1 else '')
        a_para.add_run(answer)
        
        doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.add_run('മൊഴി നൽകിയ ആളിന്റെ ഒപ്പ്')
    
    doc.add_paragraph()
    
    officer_sig = doc.add_paragraph()
    officer_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    officer_sig.add_run('അന്വേഷണ ഉദ്യോഗസ്ഥന്റെ ഒപ്പ്')
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Generate filename
    safe_filename = file_number.replace('/', '_').replace('\\', '_').replace(':', '_')
    filename = f'Statement_{safe_filename}.docx'
    
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@files_bp.route('/cmo-portal/save', methods=['POST'])
@csrf.exempt
@login_required
def save_cmo_portal():
    """Save or update CMO Portal details."""
    try:
        data = request.get_json(force=True, silent=True)
        if not data:
            return jsonify({'success': False, 'message': 'No data received.'})
        
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number is required.'})
        
        docket_number = data.get('docket_number', '').strip()
        date_of_receipt = data.get('date_of_receipt', '').strip()
        finalised = 'Yes' if data.get('finalised') else 'No'
        finalised_date = data.get('finalised_date', '').strip()
        
        # Check if record exists
        cmo_details = CMOPortalDetails.query.filter_by(file_number=file_number).first()
        
        if cmo_details:
            # Update existing
            cmo_details.docket_number = docket_number
            cmo_details.date_of_receipt = date_of_receipt
            cmo_details.finalised = finalised
            cmo_details.finalised_date = finalised_date
        else:
            # Create new
            cmo_details = CMOPortalDetails(
                file_number=file_number,
                docket_number=docket_number,
                date_of_receipt=date_of_receipt,
                finalised=finalised,
                finalised_date=finalised_date
            )
            db.session.add(cmo_details)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'CMO Portal details saved successfully.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@files_bp.route('/cmo-portal/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_cmo_portal():
    """Delete CMO Portal details."""
    try:
        data = request.get_json(force=True, silent=True)
        if not data:
            return jsonify({'success': False, 'message': 'No data received.'})
        
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number is required.'})
        
        cmo_details = CMOPortalDetails.query.filter_by(file_number=file_number).first()
        if cmo_details:
            db.session.delete(cmo_details)
            db.session.commit()
            return jsonify({'success': True, 'message': 'CMO Portal details deleted successfully.'})
        else:
            return jsonify({'success': False, 'message': 'No CMO Portal details found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@files_bp.route('/cmo-portal/get/<path:file_number>')
@login_required
def get_cmo_portal(file_number):
    """Get CMO Portal details for a file."""
    cmo_details = CMOPortalDetails.query.filter_by(file_number=file_number).first()
    if cmo_details:
        return jsonify({
            'success': True,
            'data': {
                'docket_number': cmo_details.docket_number or '',
                'date_of_receipt': cmo_details.date_of_receipt or '',
                'finalised': cmo_details.finalised == 'Yes',
                'finalised_date': cmo_details.finalised_date or ''
            }
        })
    return jsonify({'success': True, 'data': None})


# ============================================================================
# RVU Details Routes
# ============================================================================
@files_bp.route('/rvu/save', methods=['POST'])
@csrf.exempt
@login_required
def save_rvu():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        
        details = RVUDetails.query.filter_by(file_number=file_number).first()
        if details:
            details.dhs_file_number = data.get('dhs_file_number', '')
            details.receipt_date = data.get('receipt_date', '')
            details.receipt_date_inquiry_status = data.get('inquiry_status', '')
            details.receipt_date_inquiry_date = data.get('inquiry_date', '')
            details.report_status = data.get('report_status', '')
            details.interim_report_date = data.get('interim_report_date', '')
            details.sent_date = data.get('sent_date', '')
        else:
            details = RVUDetails(
                file_number=file_number,
                dhs_file_number=data.get('dhs_file_number', ''),
                receipt_date=data.get('receipt_date', ''),
                receipt_date_inquiry_status=data.get('inquiry_status', ''),
                receipt_date_inquiry_date=data.get('inquiry_date', ''),
                report_status=data.get('report_status', ''),
                interim_report_date=data.get('interim_report_date', ''),
                sent_date=data.get('sent_date', '')
            )
            db.session.add(details)
        db.session.commit()
        return jsonify({'success': True, 'message': 'RVU details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/rvu/get/<path:file_number>')
@login_required
def get_rvu(file_number):
    details = RVUDetails.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'dhs_file_number': details.dhs_file_number or '',
            'receipt_date': details.receipt_date or '',
            'inquiry_status': details.receipt_date_inquiry_status or '',
            'inquiry_date': details.receipt_date_inquiry_date or '',
            'report_status': details.report_status or '',
            'interim_report_date': details.interim_report_date or '',
            'sent_date': details.sent_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/rvu/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_rvu():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = RVUDetails.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Generic Category Details Routes (KESCPCR, KHRC, SCST, KWC)
# ============================================================================
@files_bp.route('/category-details/save', methods=['POST'])
@csrf.exempt
@login_required
def save_category_details():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        category = data.get('category', '').strip()
        if not file_number or not category:
            return jsonify({'success': False, 'message': 'File number and category required.'})
        
        model_map = {
            'kescpcr': (KESCPCRDetails, 'kescpcr_case_no'),
            'khrc': (KHRCDetails, 'khrc_case_no'),
            'scst': (SCSTDetails, 'scst_case_no'),
            'kwc': (KWCDetails, 'kwc_case_no')
        }
        if category not in model_map:
            return jsonify({'success': False, 'message': 'Invalid category.'})
        
        Model, case_no_field = model_map[category]
        details = Model.query.filter_by(file_number=file_number).first()
        
        if details:
            setattr(details, case_no_field, data.get('case_no', ''))
            details.receipt_date = data.get('receipt_date', '')
            details.report_status = data.get('report_status', '')
            details.interim_report_date = data.get('interim_report_date', '')
            details.sent_date = data.get('sent_date', '')
            details.finalised = 'Yes' if data.get('finalised') else 'No'
            details.finalised_date = data.get('finalised_date', '')
        else:
            details = Model(file_number=file_number)
            setattr(details, case_no_field, data.get('case_no', ''))
            details.receipt_date = data.get('receipt_date', '')
            details.report_status = data.get('report_status', '')
            details.interim_report_date = data.get('interim_report_date', '')
            details.sent_date = data.get('sent_date', '')
            details.finalised = 'Yes' if data.get('finalised') else 'No'
            details.finalised_date = data.get('finalised_date', '')
            db.session.add(details)
        
        db.session.commit()
        return jsonify({'success': True, 'message': f'{category.upper()} details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/category-details/get/<category>/<path:file_number>')
@login_required
def get_category_details(category, file_number):
    model_map = {
        'kescpcr': (KESCPCRDetails, 'kescpcr_case_no'),
        'khrc': (KHRCDetails, 'khrc_case_no'),
        'scst': (SCSTDetails, 'scst_case_no'),
        'kwc': (KWCDetails, 'kwc_case_no')
    }
    if category not in model_map:
        return jsonify({'success': False, 'message': 'Invalid category.'})
    
    Model, case_no_field = model_map[category]
    details = Model.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'case_no': getattr(details, case_no_field, '') or '',
            'receipt_date': details.receipt_date or '',
            'report_status': details.report_status or '',
            'interim_report_date': details.interim_report_date or '',
            'sent_date': details.sent_date or '',
            'finalised': details.finalised == 'Yes',
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/category-details/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_category_details():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        category = data.get('category', '').strip()
        model_map = {'kescpcr': KESCPCRDetails, 'khrc': KHRCDetails, 'scst': SCSTDetails, 'kwc': KWCDetails}
        if category not in model_map:
            return jsonify({'success': False, 'message': 'Invalid category.'})
        Model = model_map[category]
        details = Model.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Vigilance AC Details Routes
# ============================================================================
@files_bp.route('/vigilance/save', methods=['POST'])
@csrf.exempt
@login_required
def save_vigilance():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = VigilanceACDetails.query.filter_by(file_number=file_number).first()
        if details:
            details.vigilance_ac_case_no = data.get('case_no', '')
            details.receipt_date = data.get('receipt_date', '')
            details.finalised = 'Yes' if data.get('finalised') else 'No'
            details.finalised_date = data.get('finalised_date', '')
        else:
            details = VigilanceACDetails(
                file_number=file_number,
                vigilance_ac_case_no=data.get('case_no', ''),
                receipt_date=data.get('receipt_date', ''),
                finalised='Yes' if data.get('finalised') else 'No',
                finalised_date=data.get('finalised_date', '')
            )
            db.session.add(details)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Vigilance details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/vigilance/get/<path:file_number>')
@login_required
def get_vigilance(file_number):
    details = VigilanceACDetails.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'case_no': details.vigilance_ac_case_no or '',
            'receipt_date': details.receipt_date or '',
            'finalised': details.finalised == 'Yes',
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/vigilance/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_vigilance():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = VigilanceACDetails.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Rajya Lok Niyamasabha Details Routes
# ============================================================================
@files_bp.route('/sabha/save', methods=['POST'])
@csrf.exempt
@login_required
def save_sabha():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = RajyaLokNiyamasabhaDetails.query.filter_by(file_number=file_number).first()
        if details:
            details.sabha_type = data.get('sabha_type', '')
            details.receipt_date = data.get('receipt_date', '')
            details.report_status = data.get('report_status', '')
            details.sent_date = data.get('sent_date', '')
            details.finalised = 'Yes' if data.get('finalised') else 'No'
            details.finalised_date = data.get('finalised_date', '')
        else:
            details = RajyaLokNiyamasabhaDetails(
                file_number=file_number,
                sabha_type=data.get('sabha_type', ''),
                receipt_date=data.get('receipt_date', ''),
                report_status=data.get('report_status', ''),
                sent_date=data.get('sent_date', ''),
                finalised='Yes' if data.get('finalised') else 'No',
                finalised_date=data.get('finalised_date', '')
            )
            db.session.add(details)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Sabha details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/sabha/get/<path:file_number>')
@login_required
def get_sabha(file_number):
    details = RajyaLokNiyamasabhaDetails.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'sabha_type': details.sabha_type or '',
            'receipt_date': details.receipt_date or '',
            'report_status': details.report_status or '',
            'sent_date': details.sent_date or '',
            'finalised': details.finalised == 'Yes',
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})


# ============================================================================
# Court Case Routes
# ============================================================================
@files_bp.route('/court-case/save', methods=['POST'])
@csrf.exempt
@login_required
def save_court_case():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = CourtCase.query.filter_by(file_number=file_number).first()
        if not details:
            details = CourtCase(file_number=file_number)
            db.session.add(details)
        details.name_of_forum = data.get('name_of_forum', '')
        details.forum_specify = data.get('forum_specify', '')
        details.case_no = data.get('case_no', '')
        details.month = data.get('month', '')
        details.year = data.get('year', '')
        details.receipt_date = data.get('receipt_date', '')
        details.related_to_mo = data.get('related_to_mo', '')
        details.mo_pen = data.get('mo_pen', '')
        details.sf_forwarded = data.get('sf_forwarded', '')
        details.sf_forwarding_date = data.get('sf_forwarding_date', '')
        details.affidavit_filed = data.get('affidavit_filed', '')
        details.affidavit_filed_date = data.get('affidavit_filed_date', '')
        details.present_status = data.get('present_status', '')
        details.disposed_against_dhs = data.get('disposed_against_dhs', '')
        details.disposal_date = data.get('disposal_date', '')
        details.remarks = data.get('remarks', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Court case saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/court-case/get/<path:file_number>')
@login_required
def get_court_case(file_number):
    details = CourtCase.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'name_of_forum': details.name_of_forum or '',
            'forum_specify': details.forum_specify or '',
            'case_no': details.case_no or '',
            'month': details.month or '',
            'year': details.year or '',
            'receipt_date': details.receipt_date or '',
            'related_to_mo': details.related_to_mo or '',
            'mo_pen': details.mo_pen or '',
            'sf_forwarded': details.sf_forwarded or '',
            'sf_forwarding_date': details.sf_forwarding_date or '',
            'affidavit_filed': details.affidavit_filed or '',
            'affidavit_filed_date': details.affidavit_filed_date or '',
            'present_status': details.present_status or '',
            'disposed_against_dhs': details.disposed_against_dhs or '',
            'disposal_date': details.disposal_date or '',
            'remarks': details.remarks or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/court-case/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_court_case():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = CourtCase.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Complaint Routes
# ============================================================================
@files_bp.route('/complaint/save', methods=['POST'])
@csrf.exempt
@login_required
def save_complaint():
    try:
        # Try to get JSON data - handle both JSON and form data
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()
        
        if not data:
            return jsonify({'success': False, 'message': 'No data received.'})
        
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number is required.'})
        
        # Verify file exists
        file_obj = File.query.filter_by(file_number=file_number).first()
        if not file_obj:
            return jsonify({'success': False, 'message': f'File {file_number} does not exist.'})
        
        details = ComplaintDetails.query.filter_by(file_number=file_number).first()
        if not details:
            details = ComplaintDetails(file_number=file_number)
            db.session.add(details)
        
        # Update all fields
        details.plaintiff_type = data.get('plaintiff_type', '')
        details.plaintiff_pen = data.get('plaintiff_pen', '')
        details.plaintiff_name = data.get('plaintiff_name', '')
        details.plaintiff_designation = data.get('plaintiff_designation', '')
        details.plaintiff_institution = data.get('plaintiff_institution', '')
        details.plaintiff_address = data.get('plaintiff_address', '')
        details.plaintiff_contact_number = data.get('plaintiff_contact_number', '')
        details.plaintiff_email = data.get('plaintiff_email', '')
        details.respondent_type = data.get('respondent_type', '')
        details.respondent_pen = data.get('respondent_pen', '')
        details.respondent_name = data.get('respondent_name', '')
        details.respondent_designation = data.get('respondent_designation', '')
        details.respondent_institution = data.get('respondent_institution', '')
        details.respondent_address = data.get('respondent_address', '')
        details.respondent_contact_number = data.get('respondent_contact_number', '')
        details.respondent_email = data.get('respondent_email', '')
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Complaint details saved successfully.'})
    except Exception as e:
        db.session.rollback()
        import traceback
        error_msg = str(e)
        tb = traceback.format_exc()
        print(f"Error saving complaint: {error_msg}\n{tb}")  # Log for debugging
        return jsonify({'success': False, 'message': f'Error: {error_msg}'})

@files_bp.route('/complaint/get/<path:file_number>')
@login_required
def get_complaint(file_number):
    details = ComplaintDetails.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'plaintiff_type': details.plaintiff_type or '',
            'plaintiff_pen': details.plaintiff_pen or '',
            'plaintiff_name': details.plaintiff_name or '',
            'plaintiff_designation': details.plaintiff_designation or '',
            'plaintiff_institution': details.plaintiff_institution or '',
            'plaintiff_address': details.plaintiff_address or '',
            'plaintiff_contact_number': details.plaintiff_contact_number or '',
            'plaintiff_email': details.plaintiff_email or '',
            'respondent_type': details.respondent_type or '',
            'respondent_pen': details.respondent_pen or '',
            'respondent_name': details.respondent_name or '',
            'respondent_designation': details.respondent_designation or '',
            'respondent_institution': details.respondent_institution or '',
            'respondent_address': details.respondent_address or '',
            'respondent_contact_number': details.respondent_contact_number or '',
            'respondent_email': details.respondent_email or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/complaint/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_complaint():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = ComplaintDetails.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Women Harassment Routes
# ============================================================================
@files_bp.route('/women-harassment/save', methods=['POST'])
@csrf.exempt
@login_required
def save_women_harassment():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = WomenHarassmentCase.query.filter_by(file_number=file_number).first()
        if not details:
            details = WomenHarassmentCase(file_number=file_number)
            db.session.add(details)
        details.icc_report_attached = 1 if data.get('icc_report_attached', False) else 0
        details.icc_report_date = data.get('icc_report_date', '')
        details.finalised = 1 if data.get('finalised', False) else 0
        details.finalised_date = data.get('finalised_date', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Women harassment details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/women-harassment/get/<path:file_number>')
@login_required
def get_women_harassment(file_number):
    details = WomenHarassmentCase.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'icc_report_attached': details.icc_report_attached == 1,
            'icc_report_date': details.icc_report_date or '',
            'finalised': details.finalised == 1,
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/women-harassment/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_women_harassment():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = WomenHarassmentCase.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Police Case Routes
# ============================================================================
@files_bp.route('/police-case/save', methods=['POST'])
@csrf.exempt
@login_required
def save_police_case():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = PoliceCaseDetails.query.filter_by(file_number=file_number).first()
        if not details:
            details = PoliceCaseDetails(file_number=file_number)
            db.session.add(details)
        details.case_no = data.get('case_no', '')
        details.fir_no = data.get('fir_no', '')
        details.case_date = data.get('case_date', '')
        details.police_station = data.get('police_station', '')
        details.detained_over_48_hours = data.get('detained_over_48_hours', '')
        details.suspended_from_service = data.get('suspended_from_service', '')
        details.present_status = data.get('present_status', '')
        details.finalised = data.get('finalised', '')
        details.finalised_date = data.get('finalised_date', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Police case saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/police-case/get/<path:file_number>')
@login_required
def get_police_case(file_number):
    details = PoliceCaseDetails.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'case_no': details.case_no or '',
            'fir_no': details.fir_no or '',
            'case_date': details.case_date or '',
            'police_station': details.police_station or '',
            'detained_over_48_hours': details.detained_over_48_hours or '',
            'suspended_from_service': details.suspended_from_service or '',
            'present_status': details.present_status or '',
            'finalised': details.finalised or '',
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/police-case/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_police_case():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = PoliceCaseDetails.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Attack on Doctors Routes
# ============================================================================
@files_bp.route('/attack-doctors/save', methods=['POST'])
@csrf.exempt
@login_required
def save_attack_doctors():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = AttackOnDoctorsCase.query.filter_by(file_number=file_number).first()
        if not details:
            details = AttackOnDoctorsCase(file_number=file_number)
            db.session.add(details)
        details.police_informed = data.get('police_informed', '')
        details.police_station = data.get('police_station', '')
        details.reported_date = data.get('reported_date', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Attack on doctors details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/attack-doctors/get/<path:file_number>')
@login_required
def get_attack_doctors(file_number):
    details = AttackOnDoctorsCase.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'police_informed': details.police_informed or '',
            'police_station': details.police_station or '',
            'reported_date': details.reported_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/attack-doctors/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_attack_doctors():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = AttackOnDoctorsCase.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Attack on Staffs Routes
# ============================================================================
@files_bp.route('/attack-staffs/save', methods=['POST'])
@csrf.exempt
@login_required
def save_attack_staffs():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = AttackOnStaffsCase.query.filter_by(file_number=file_number).first()
        if not details:
            details = AttackOnStaffsCase(file_number=file_number)
            db.session.add(details)
        details.police_informed = data.get('police_informed', '')
        details.police_station = data.get('police_station', '')
        details.reported_date = data.get('reported_date', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Attack on staffs details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/attack-staffs/get/<path:file_number>')
@login_required
def get_attack_staffs(file_number):
    details = AttackOnStaffsCase.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'police_informed': details.police_informed or '',
            'police_station': details.police_station or '',
            'reported_date': details.reported_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/attack-staffs/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_attack_staffs():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = AttackOnStaffsCase.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ============================================================================
# Social Security Pension Routes
# ============================================================================
@files_bp.route('/social-security/save', methods=['POST'])
@csrf.exempt
@login_required
def save_social_security():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        if not file_number:
            return jsonify({'success': False, 'message': 'File number required.'})
        details = SocialSecurityPension.query.filter_by(file_number=file_number).first()
        if not details:
            details = SocialSecurityPension(file_number=file_number)
            db.session.add(details)
        details.main_list_sl_no = data.get('main_list_sl_no', '')
        details.pen = data.get('pen', '')
        details.name = data.get('name', '')
        details.sevana_pensioner_id = data.get('sevana_pensioner_id', '')
        details.aadhar_no = data.get('aadhar_no', '')
        details.amount = float(data.get('amount', 0) or 0)
        details.refunded_status = data.get('refunded_status', '')
        details.refunded_amount = float(data.get('refunded_amount', 0) or 0)
        details.letter_no = data.get('letter_no', '')
        details.letter_date = data.get('letter_date', '')
        details.receipt_no = data.get('receipt_no', '')
        details.finalised = data.get('finalised', '')
        details.finalised_date = data.get('finalised_date', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Social security pension details saved.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@files_bp.route('/social-security/get/<path:file_number>')
@login_required
def get_social_security(file_number):
    details = SocialSecurityPension.query.filter_by(file_number=file_number).first()
    if details:
        return jsonify({'success': True, 'data': {
            'main_list_sl_no': details.main_list_sl_no or '',
            'pen': details.pen or '',
            'name': details.name or '',
            'sevana_pensioner_id': details.sevana_pensioner_id or '',
            'aadhar_no': details.aadhar_no or '',
            'amount': details.amount or 0,
            'refunded_status': details.refunded_status or '',
            'refunded_amount': details.refunded_amount or 0,
            'letter_no': details.letter_no or '',
            'letter_date': details.letter_date or '',
            'receipt_no': details.receipt_no or '',
            'finalised': details.finalised or '',
            'finalised_date': details.finalised_date or ''
        }})
    return jsonify({'success': True, 'data': None})

@files_bp.route('/social-security/delete', methods=['POST'])
@csrf.exempt
@login_required
def delete_social_security():
    try:
        data = request.get_json(force=True, silent=True)
        file_number = data.get('file_number', '').strip()
        details = SocialSecurityPension.query.filter_by(file_number=file_number).first()
        if details:
            db.session.delete(details)
            db.session.commit()
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Not found.'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})