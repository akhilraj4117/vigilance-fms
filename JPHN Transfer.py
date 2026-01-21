import sys
import sqlite3
import csv
from datetime import datetime, date
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                               QHBoxLayout, QTableWidget, QTableWidgetItem, 
                               QPushButton, QLineEdit, QLabel, QComboBox, 
                               QDateEdit, QDialog, QFormLayout, QMessageBox,
                               QHeaderView, QCheckBox, QTextEdit, QFileDialog,
                               QFrame, QSizePolicy, QProgressDialog, QTabWidget,
                               QGroupBox, QSpinBox, QAbstractItemView, QMenu,
                               QTextBrowser, QScrollArea, QInputDialog, QGridLayout)
from PySide6.QtCore import Qt, QDate, Signal, QStringListModel
from PySide6.QtGui import QFont, QIcon, QColor, QShortcut, QAction, QTextDocument, QPixmap
from PySide6.QtWidgets import QCompleter
from PySide6.QtGui import QKeySequence
from PySide6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
from PySide6.QtGui import QPageLayout, QPageSize
import os


def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)


class NoScrollDateEdit(QDateEdit):
    """Custom QDateEdit that disables mouse wheel scrolling"""
    def wheelEvent(self, event):
        event.ignore()


class NoScrollComboBox(QComboBox):
    """Custom QComboBox that disables mouse wheel scrolling"""
    def wheelEvent(self, event):
        event.ignore()


class SearchableComboBox(QComboBox):
    """Custom QComboBox with type-ahead search functionality"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setEditable(True)
        self.setInsertPolicy(QComboBox.NoInsert)
        
        # Setup completer for type-ahead search
        self.completer_obj = QCompleter(self)
        self.completer_obj.setCaseSensitivity(Qt.CaseInsensitive)
        self.completer_obj.setFilterMode(Qt.MatchContains)
        self.completer_obj.setCompletionMode(QCompleter.PopupCompletion)
        self.setCompleter(self.completer_obj)
        
        # Connect completer activation to select the item
        self.completer_obj.activated.connect(self._on_completer_activated)
        
        # Initial update of completer
        self._update_completer()
        
        # Update completer model when items change
        self.model().rowsInserted.connect(self._update_completer)
        self.model().rowsRemoved.connect(self._update_completer)
        self.model().modelReset.connect(self._update_completer)
    
    def _on_completer_activated(self, text):
        """When user selects from completer dropdown, set that item"""
        index = self.findText(text, Qt.MatchFixedString)
        if index >= 0:
            self.setCurrentIndex(index)
        
    def _update_completer(self):
        """Update completer model with current items"""
        items = [self.itemText(i) for i in range(self.count())]
        model = QStringListModel(items, self.completer_obj)
        self.completer_obj.setModel(model)
    
    def refresh_completer(self):
        """Force refresh the completer model - call after bulk updates"""
        self._update_completer()
    
    def wheelEvent(self, event):
        event.ignore()
    
    def focusOutEvent(self, event):
        """Validate text on focus out - reset if invalid"""
        super().focusOutEvent(event)
        text = self.currentText().strip()
        if text:
            index = self.findText(text, Qt.MatchFixedString)
            if index >= 0:
                self.setCurrentIndex(index)
            else:
                # Try case-insensitive match
                for i in range(self.count()):
                    if self.itemText(i).lower() == text.lower():
                        self.setCurrentIndex(i)
                        return
                # If text doesn't match any item, reset to empty
                self.setCurrentIndex(0)


class LoginDialog(QDialog):
    """Login screen with Department of Health Services branding"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Login - Department of Health Services")
        self.setFont(QFont("Calibri", 11))
        self.authenticated = False
        
        # Set window icon
        icon_path = resource_path("favicon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # Set fixed size and center on screen
        self.setWindowFlags(Qt.Window)
        self.setFixedSize(650, 580)
        
        self.setup_ui()
        self.apply_styles()
        self.center_on_screen()
    
    def center_on_screen(self):
        """Center the dialog on the screen"""
        screen = QApplication.primaryScreen().geometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)
    
    def setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Center container
        center_widget = QWidget()
        center_layout = QVBoxLayout(center_widget)
        center_layout.setSpacing(8)
        center_layout.setContentsMargins(30, 20, 30, 20)
        center_layout.setAlignment(Qt.AlignCenter)
        
        # Add stretch to center vertically
        main_layout.addStretch(1)
        
        # Title: Department of Health Services
        title_label = QLabel("Department of Health Services")
        title_label.setFont(QFont("Calibri", 22, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #1a5276;")
        center_layout.addWidget(title_label)
        
        # Kerala
        kerala_label = QLabel("Kerala")
        kerala_label.setFont(QFont("Calibri", 18, QFont.Bold))
        kerala_label.setAlignment(Qt.AlignCenter)
        kerala_label.setStyleSheet("color: #27ae60;")
        center_layout.addWidget(kerala_label)
        
        # Spacer
        center_layout.addSpacing(10)
        
        # Logo
        logo_label = QLabel()
        logo_label.setAlignment(Qt.AlignCenter)
        # Try to load logo from file, otherwise show placeholder
        logo_path = resource_path("Health Logo.png")
        if os.path.exists(logo_path):
            pixmap = QPixmap(logo_path)
            scaled_pixmap = pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            logo_label.setPixmap(scaled_pixmap)
        else:
            # Placeholder if logo not found
            logo_label.setText("🏥")
            logo_label.setFont(QFont("Segoe UI Emoji", 48))
            logo_label.setStyleSheet("color: #1a5276;")
        center_layout.addWidget(logo_label)
        
        # Spacer
        center_layout.addSpacing(10)
        
        # Directorate of Health Services
        directorate_label = QLabel("Directorate of Health Services")
        directorate_label.setFont(QFont("Calibri", 16, QFont.Bold))
        directorate_label.setAlignment(Qt.AlignCenter)
        directorate_label.setStyleSheet("color: #1a5276;")
        center_layout.addWidget(directorate_label)
        
        # Transfer and Posting
        transfer_label = QLabel("Transfer and Posting")
        transfer_label.setFont(QFont("Calibri", 14))
        transfer_label.setAlignment(Qt.AlignCenter)
        transfer_label.setStyleSheet("color: #2c3e50;")
        center_layout.addWidget(transfer_label)
        
        # Spacer before login form
        center_layout.addSpacing(15)
        
        # Login form container
        form_widget = QWidget()
        form_layout = QFormLayout(form_widget)
        form_layout.setSpacing(10)
        form_layout.setContentsMargins(80, 10, 80, 10)
        form_layout.setLabelAlignment(Qt.AlignRight)
        
        # User ID
        self.user_id_input = QLineEdit()
        self.user_id_input.setPlaceholderText("Enter User ID")
        self.user_id_input.setMinimumWidth(250)
        self.user_id_input.setMinimumHeight(35)
        self.user_id_input.setFont(QFont("Calibri", 11))
        user_id_label = QLabel("User ID:")
        user_id_label.setFont(QFont("Calibri", 12, QFont.Bold))
        form_layout.addRow(user_id_label, self.user_id_input)
        
        # Password
        self.password_input = QLineEdit()
        self.password_input.setPlaceholderText("Enter Password")
        self.password_input.setEchoMode(QLineEdit.Password)
        self.password_input.setMinimumWidth(250)
        self.password_input.setMinimumHeight(35)
        self.password_input.setFont(QFont("Calibri", 11))
        password_label = QLabel("Password:")
        password_label.setFont(QFont("Calibri", 12, QFont.Bold))
        form_layout.addRow(password_label, self.password_input)
        
        center_layout.addWidget(form_widget, alignment=Qt.AlignCenter)
        
        # Login button
        self.login_btn = QPushButton("Login")
        self.login_btn.setMinimumWidth(200)
        self.login_btn.setMinimumHeight(45)
        self.login_btn.setFont(QFont("Calibri", 14, QFont.Bold))
        self.login_btn.clicked.connect(self.handle_login)
        self.login_btn.setDefault(True)
        center_layout.addWidget(self.login_btn, alignment=Qt.AlignCenter)
        
        # Connect Enter key to login
        self.user_id_input.returnPressed.connect(self.handle_login)
        self.password_input.returnPressed.connect(self.handle_login)
        
        main_layout.addWidget(center_widget, alignment=Qt.AlignCenter)
        main_layout.addStretch(1)
    
    def apply_styles(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #e8f6f3, stop:0.5 #ebf5fb, stop:1 #e8f8f5);
            }
            QLineEdit {
                border: 2px solid #85c1e9;
                border-radius: 8px;
                padding: 8px 15px;
                background-color: white;
            }
            QLineEdit:focus {
                border: 2px solid #27ae60;
            }
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 30px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0e2f44, stop:1 #196f3d);
            }
        """)
    
    def handle_login(self):
        user_id = self.user_id_input.text().strip()
        password = self.password_input.text().strip()
        
        if not user_id or not password:
            QMessageBox.warning(self, "Login Error", "Please enter both User ID and Password.")
            return
        
        # Simple authentication (can be enhanced with database)
        valid_users = {
            "revathy": "4117"
        }
        
        if user_id in valid_users and valid_users[user_id] == password:
            self.authenticated = True
            self.accept()
        else:
            QMessageBox.warning(self, "Login Failed", "Invalid User ID or Password.\n\nPlease try again.")
            self.password_input.clear()
            self.password_input.setFocus()
    
    def is_authenticated(self):
        return self.authenticated


class WelcomeDialog(QDialog):
    """Welcome screen to choose between General Transfer and Regular Transfer"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("JPHN Gr I Transfer Management System")
        self.setFont(QFont("Calibri", 11))
        self.selected_type = None
        
        # Set window icon
        icon_path = resource_path("favicon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # Set fixed size and center on screen
        self.setWindowFlags(Qt.Window)
        self.setFixedSize(650, 500)
        
        self.setup_ui()
        self.apply_styles()
        self.center_on_screen()
    
    def center_on_screen(self):
        """Center the dialog on the screen"""
        screen = QApplication.primaryScreen().geometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(40, 40, 40, 40)
        
        # Title
        title_label = QLabel("JPHN Gr I Transfer & Posting")
        title_label.setFont(QFont("Calibri", 22, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #1a5276;")
        layout.addWidget(title_label)
        
        # Subtitle
        subtitle_label = QLabel("Kerala Health Services Department")
        subtitle_label.setFont(QFont("Calibri", 12))
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_label.setStyleSheet("color: #27ae60;")
        layout.addWidget(subtitle_label)
        
        layout.addSpacing(20)
        
        # Instruction
        inst_label = QLabel("Select Transfer Type to Continue:")
        inst_label.setFont(QFont("Calibri", 12))
        inst_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(inst_label)
        
        layout.addSpacing(10)
        
        # Buttons container
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(30)
        
        # General Transfer Button
        general_frame = QFrame()
        general_frame.setObjectName("optionFrame")
        general_layout = QVBoxLayout(general_frame)
        general_layout.setSpacing(15)
        general_layout.setContentsMargins(20, 25, 20, 25)
        
        general_icon = QLabel("📋")
        general_icon.setFont(QFont("Segoe UI Emoji", 36))
        general_icon.setAlignment(Qt.AlignCenter)
        general_layout.addWidget(general_icon)
        
        general_title = QLabel("General Transfer")
        general_title.setFont(QFont("Calibri", 14, QFont.Bold))
        general_title.setAlignment(Qt.AlignCenter)
        general_layout.addWidget(general_title)
        
        general_desc = QLabel("For general transfer\noperations and postings")
        general_desc.setFont(QFont("Calibri", 10))
        general_desc.setAlignment(Qt.AlignCenter)
        general_desc.setStyleSheet("color: #666;")
        general_layout.addWidget(general_desc)
        
        general_btn = QPushButton("Open General Transfer")
        general_btn.setMinimumHeight(40)
        general_btn.setCursor(Qt.PointingHandCursor)
        general_btn.clicked.connect(lambda: self.select_type("general"))
        general_layout.addWidget(general_btn)
        
        buttons_layout.addWidget(general_frame)
        
        # Regular Transfer Button
        regular_frame = QFrame()
        regular_frame.setObjectName("optionFrame")
        regular_layout = QVBoxLayout(regular_frame)
        regular_layout.setSpacing(15)
        regular_layout.setContentsMargins(20, 25, 20, 25)
        
        regular_icon = QLabel("📑")
        regular_icon.setFont(QFont("Segoe UI Emoji", 36))
        regular_icon.setAlignment(Qt.AlignCenter)
        regular_layout.addWidget(regular_icon)
        
        regular_title = QLabel("Regular Transfer")
        regular_title.setFont(QFont("Calibri", 14, QFont.Bold))
        regular_title.setAlignment(Qt.AlignCenter)
        regular_layout.addWidget(regular_title)
        
        regular_desc = QLabel("For regular/routine\ntransfer management")
        regular_desc.setFont(QFont("Calibri", 10))
        regular_desc.setAlignment(Qt.AlignCenter)
        regular_desc.setStyleSheet("color: #666;")
        regular_layout.addWidget(regular_desc)
        
        regular_btn = QPushButton("Open Regular Transfer")
        regular_btn.setMinimumHeight(40)
        regular_btn.setCursor(Qt.PointingHandCursor)
        regular_btn.clicked.connect(lambda: self.select_type("regular"))
        regular_layout.addWidget(regular_btn)
        
        buttons_layout.addWidget(regular_frame)
        
        layout.addLayout(buttons_layout)
        
        layout.addStretch()
        
        # Footer
        footer_label = QLabel("Each transfer type uses a separate database")
        footer_label.setFont(QFont("Calibri", 9))
        footer_label.setAlignment(Qt.AlignCenter)
        footer_label.setStyleSheet("color: #999;")
        layout.addWidget(footer_label)
    
    def apply_styles(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #e8f6f3, stop:0.5 #ebf5fb, stop:1 #e8f8f5);
            }
            QFrame#optionFrame {
                background-color: white;
                border: 2px solid #85c1e9;
                border-radius: 12px;
            }
            QFrame#optionFrame:hover {
                border: 2px solid #27ae60;
                background-color: #f0fff4;
            }
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0e2f44, stop:1 #196f3d);
            }
        """)
    
    def select_type(self, transfer_type):
        self.selected_type = transfer_type
        self.accept()
    
    def get_selected_type(self):
        return self.selected_type


class YearSelectionDialog(QDialog):
    """Dialog to select year for General Transfer"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Year - General Transfer")
        self.setFont(QFont("Calibri", 11))
        self.selected_year = None
        
        # Set window icon
        icon_path = resource_path("favicon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # Set fixed size and center on screen
        self.setWindowFlags(Qt.Window)
        self.setFixedSize(800, 500)
        
        self.setup_ui()
        self.apply_styles()
        self.center_on_screen()
    
    def center_on_screen(self):
        """Center the dialog on the screen"""
        screen = QApplication.primaryScreen().geometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(50, 30, 50, 30)
        
        # Title
        title_label = QLabel("General Transfer - Year Selection")
        title_label.setFont(QFont("Calibri", 22, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #1a5276;")
        layout.addWidget(title_label)
        
        # Subtitle
        subtitle_label = QLabel("Select an existing year or create a new one")
        subtitle_label.setFont(QFont("Calibri", 11))
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_label.setStyleSheet("color: #5d6d7e;")
        layout.addWidget(subtitle_label)
        
        layout.addSpacing(20)
        
        # Get available years from existing database files
        self.available_years = self.get_available_years()
        
        # Previous Years section
        years_frame = QFrame()
        years_frame.setObjectName("yearsFrame")
        years_layout = QVBoxLayout(years_frame)
        years_layout.setSpacing(15)
        years_layout.setContentsMargins(30, 20, 30, 20)
        
        years_title = QLabel("📂 Open Previous Year:")
        years_title.setFont(QFont("Calibri", 13, QFont.Bold))
        years_title.setStyleSheet("color: #1a5276;")
        years_layout.addWidget(years_title)
        
        # Dropdown row
        dropdown_row = QHBoxLayout()
        dropdown_row.setSpacing(15)
        
        # Dropdown for previous years
        self.previous_year_combo = QComboBox()
        self.previous_year_combo.setMinimumSize(350, 42)
        self.previous_year_combo.setFont(QFont("Calibri", 11))
        
        if self.available_years:
            self.previous_year_combo.addItem("-- Select a Year --")
            # Sort years descending
            for year in sorted(self.available_years, reverse=True):
                self.previous_year_combo.addItem(f"📅 {year}", year)
        else:
            self.previous_year_combo.addItem("No previous years found")
            self.previous_year_combo.setEnabled(False)
        
        dropdown_row.addWidget(self.previous_year_combo, 1)
        
        # Open button
        self.open_year_btn = QPushButton("  Open Selected  ")
        self.open_year_btn.setMinimumSize(160, 42)
        self.open_year_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        self.open_year_btn.setCursor(Qt.PointingHandCursor)
        self.open_year_btn.setObjectName("openYearBtn")
        self.open_year_btn.clicked.connect(self.open_selected_year)
        self.open_year_btn.setEnabled(len(self.available_years) > 0)
        dropdown_row.addWidget(self.open_year_btn)
        
        years_layout.addLayout(dropdown_row)
        
        layout.addWidget(years_frame)
        
        layout.addSpacing(15)
        
        # Add new year section
        new_year_frame = QFrame()
        new_year_frame.setObjectName("newYearFrame")
        new_year_layout = QVBoxLayout(new_year_frame)
        new_year_layout.setSpacing(15)
        new_year_layout.setContentsMargins(30, 20, 30, 20)
        
        new_year_label = QLabel("➕ Create New Year:")
        new_year_label.setFont(QFont("Calibri", 13, QFont.Bold))
        new_year_label.setStyleSheet("color: #1a5276;")
        new_year_layout.addWidget(new_year_label)
        
        # Form row
        form_row = QHBoxLayout()
        form_row.setSpacing(15)
        
        year_label = QLabel("Year:")
        year_label.setFont(QFont("Calibri", 11))
        form_row.addWidget(year_label)
        
        self.new_year_spin = QSpinBox()
        self.new_year_spin.setRange(2020, 2100)
        self.new_year_spin.setValue(datetime.now().year)
        self.new_year_spin.setMinimumSize(120, 42)
        self.new_year_spin.setFont(QFont("Calibri", 11))
        form_row.addWidget(self.new_year_spin)
        
        form_row.addSpacing(20)
        
        add_year_btn = QPushButton("  Create & Open  ")
        add_year_btn.setMinimumSize(150, 42)
        add_year_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        add_year_btn.setCursor(Qt.PointingHandCursor)
        add_year_btn.setObjectName("addYearBtn")
        add_year_btn.clicked.connect(self.add_new_year)
        form_row.addWidget(add_year_btn)
        
        form_row.addStretch()
        new_year_layout.addLayout(form_row)
        
        layout.addWidget(new_year_frame)
        
        layout.addStretch()
        
        # Back button
        back_btn = QPushButton("← Back to Selection")
        back_btn.setMinimumHeight(45)
        back_btn.setMaximumWidth(220)
        back_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        back_btn.setCursor(Qt.PointingHandCursor)
        back_btn.setObjectName("backBtn")
        back_btn.clicked.connect(self.reject)
        layout.addWidget(back_btn, alignment=Qt.AlignCenter)
        
        layout.addSpacing(10)
        
        # Footer
        footer_label = QLabel("Each year has its own separate database with no connection to other years")
        footer_label.setFont(QFont("Calibri", 9))
        footer_label.setAlignment(Qt.AlignCenter)
        footer_label.setStyleSheet("color: #999;")
        layout.addWidget(footer_label)
    
    def get_available_years(self):
        """Scan for existing general transfer database files"""
        import os
        years = []
        # Look for files like jphn_general_2024.db, jphn_general_2025.db
        for filename in os.listdir('.'):
            if filename.startswith('jphn_general_') and filename.endswith('.db'):
                try:
                    year_str = filename.replace('jphn_general_', '').replace('.db', '')
                    year = int(year_str)
                    if 2000 <= year <= 2100:
                        years.append(year)
                except ValueError:
                    continue
        return years
    
    def apply_styles(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #e8f6f3, stop:0.5 #ebf5fb, stop:1 #e8f8f5);
            }
            QFrame#yearsFrame, QFrame#newYearFrame {
                background-color: white;
                border: 2px solid #85c1e9;
                border-radius: 12px;
            }
            QPushButton#openYearBtn {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton#openYearBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton#openYearBtn:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0e2f44, stop:1 #196f3d);
            }
            QPushButton#openYearBtn:disabled {
                background-color: #bdc3c7;
            }
            QPushButton#addYearBtn {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton#addYearBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton#backBtn {
                background-color: #1a5276;
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton#backBtn:hover {
                background-color: #154360;
            }
            QComboBox, QSpinBox {
                border: 2px solid #85c1e9;
                border-radius: 6px;
                padding: 8px 12px;
                background-color: white;
            }
            QComboBox:focus, QSpinBox:focus {
                border: 2px solid #27ae60;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 10px;
            }
        """)
    
    def open_selected_year(self):
        """Open the selected year from dropdown"""
        idx = self.previous_year_combo.currentIndex()
        if idx <= 0:  # First item is placeholder
            QMessageBox.warning(self, "Warning", "Please select a year from the dropdown!")
            return
        
        data = self.previous_year_combo.currentData()
        if data:
            self.selected_year = data
            self.accept()
    
    def select_year(self, year):
        self.selected_year = year
        self.accept()
    
    def add_new_year(self):
        year = self.new_year_spin.value()
        if year in self.available_years:
            # Year already exists, just open it
            reply = QMessageBox.question(self, "Year Exists", 
                f"Year {year} already exists. Do you want to open it?",
                QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.selected_year = year
                self.accept()
        else:
            # New year - confirm creation
            reply = QMessageBox.question(self, "Create New Year",
                f"This will create a new database for year {year}.\nContinue?",
                QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.selected_year = year
                self.accept()
    
    def get_selected_year(self):
        return self.selected_year


class RegularTransferSelectionDialog(QDialog):
    """Dialog to select month/year for Regular Transfer"""
    
    MONTHS = ["January", "February", "March", "April", "May", "June",
              "July", "August", "September", "October", "November", "December"]
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Transfer List - Regular Transfer")
        self.setFont(QFont("Calibri", 11))
        self.selected_month = None
        self.selected_year = None
        
        # Set window icon
        icon_path = resource_path("favicon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # Set fixed size and center on screen
        self.setWindowFlags(Qt.Window)
        self.setFixedSize(800, 550)
        
        self.setup_ui()
        self.apply_styles()
        self.center_on_screen()
    
    def center_on_screen(self):
        """Center the dialog on the screen"""
        screen = QApplication.primaryScreen().geometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(50, 30, 50, 30)
        
        # Title
        title_label = QLabel("Regular Transfer - Select Transfer List")
        title_label.setFont(QFont("Calibri", 22, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #1a5276;")
        layout.addWidget(title_label)
        
        # Subtitle
        subtitle_label = QLabel("Select an existing transfer list or create a new one")
        subtitle_label.setFont(QFont("Calibri", 11))
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_label.setStyleSheet("color: #5d6d7e;")
        layout.addWidget(subtitle_label)
        
        layout.addSpacing(20)
        
        # Get available transfer lists
        self.available_lists = self.get_available_lists()
        
        # Previous Transfer Lists section
        lists_frame = QFrame()
        lists_frame.setObjectName("listsFrame")
        lists_layout = QVBoxLayout(lists_frame)
        lists_layout.setSpacing(15)
        lists_layout.setContentsMargins(30, 20, 30, 20)
        
        lists_title = QLabel("📂 Open Previous Transfer List:")
        lists_title.setFont(QFont("Calibri", 13, QFont.Bold))
        lists_title.setStyleSheet("color: #1a5276;")
        lists_layout.addWidget(lists_title)
        
        # Dropdown row
        dropdown_row = QHBoxLayout()
        dropdown_row.setSpacing(15)
        
        # Dropdown for previous lists
        self.previous_list_combo = QComboBox()
        self.previous_list_combo.setMinimumSize(350, 42)
        self.previous_list_combo.setFont(QFont("Calibri", 11))
        
        if self.available_lists:
            self.previous_list_combo.addItem("-- Select a Transfer List --")
            # Sort by year (desc) then month (desc)
            sorted_lists = sorted(self.available_lists, 
                                  key=lambda x: (x[1], self.MONTHS.index(x[0])), 
                                  reverse=True)
            for month, year in sorted_lists:
                self.previous_list_combo.addItem(f"📅 {month} {year}", (month, year))
        else:
            self.previous_list_combo.addItem("No previous transfer lists found")
            self.previous_list_combo.setEnabled(False)
        
        dropdown_row.addWidget(self.previous_list_combo, 1)
        
        # Open button
        self.open_list_btn = QPushButton("  Open Selected  ")
        self.open_list_btn.setMinimumSize(160, 42)
        self.open_list_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        self.open_list_btn.setCursor(Qt.PointingHandCursor)
        self.open_list_btn.setObjectName("openListBtn")
        self.open_list_btn.clicked.connect(self.open_selected_list)
        self.open_list_btn.setEnabled(len(self.available_lists) > 0)
        dropdown_row.addWidget(self.open_list_btn)
        
        lists_layout.addLayout(dropdown_row)
        
        layout.addWidget(lists_frame)
        
        layout.addSpacing(15)
        
        # Add new transfer list section
        new_list_frame = QFrame()
        new_list_frame.setObjectName("newListFrame")
        new_list_layout = QVBoxLayout(new_list_frame)
        new_list_layout.setSpacing(15)
        new_list_layout.setContentsMargins(30, 20, 30, 20)
        
        new_list_label = QLabel("➕ Create New Transfer List:")
        new_list_label.setFont(QFont("Calibri", 13, QFont.Bold))
        new_list_label.setStyleSheet("color: #1a5276;")
        new_list_layout.addWidget(new_list_label)
        
        # Form row
        form_row = QHBoxLayout()
        form_row.setSpacing(15)
        
        # Month combo
        month_label = QLabel("Month:")
        month_label.setFont(QFont("Calibri", 11))
        form_row.addWidget(month_label)
        
        self.new_month_combo = QComboBox()
        self.new_month_combo.addItems(self.MONTHS)
        self.new_month_combo.setCurrentIndex(datetime.now().month - 1)
        self.new_month_combo.setMinimumSize(140, 42)
        self.new_month_combo.setFont(QFont("Calibri", 11))
        form_row.addWidget(self.new_month_combo)
        
        form_row.addSpacing(20)
        
        # Year spin
        year_label = QLabel("Year:")
        year_label.setFont(QFont("Calibri", 11))
        form_row.addWidget(year_label)
        
        self.new_year_spin = QSpinBox()
        self.new_year_spin.setRange(2020, 2100)
        self.new_year_spin.setValue(datetime.now().year)
        self.new_year_spin.setMinimumSize(100, 42)
        self.new_year_spin.setFont(QFont("Calibri", 11))
        form_row.addWidget(self.new_year_spin)
        
        form_row.addSpacing(20)
        
        add_list_btn = QPushButton("  Create & Open  ")
        add_list_btn.setMinimumSize(150, 42)
        add_list_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        add_list_btn.setCursor(Qt.PointingHandCursor)
        add_list_btn.setObjectName("addListBtn")
        add_list_btn.clicked.connect(self.add_new_list)
        form_row.addWidget(add_list_btn)
        
        form_row.addStretch()
        new_list_layout.addLayout(form_row)
        
        layout.addWidget(new_list_frame)
        
        layout.addStretch()
        
        # Back button
        back_btn = QPushButton("← Back to Selection")
        back_btn.setMinimumHeight(45)
        back_btn.setMaximumWidth(220)
        back_btn.setFont(QFont("Calibri", 11, QFont.Bold))
        back_btn.setCursor(Qt.PointingHandCursor)
        back_btn.setObjectName("backBtn")
        back_btn.clicked.connect(self.reject)
        layout.addWidget(back_btn, alignment=Qt.AlignCenter)
        
        layout.addSpacing(10)
        
        # Footer
        footer_label = QLabel("Each transfer list has its own separate database")
        footer_label.setFont(QFont("Calibri", 9))
        footer_label.setAlignment(Qt.AlignCenter)
        footer_label.setStyleSheet("color: #999;")
        layout.addWidget(footer_label)
    
    def get_available_lists(self):
        """Scan for existing regular transfer database files"""
        import os
        lists = []
        # Look for files like jphn_regular_January_2025.db
        for filename in os.listdir('.'):
            if filename.startswith('jphn_regular_') and filename.endswith('.db'):
                try:
                    # Extract month and year from filename
                    parts = filename.replace('jphn_regular_', '').replace('.db', '')
                    # Split by underscore to get month and year
                    underscore_idx = parts.rfind('_')
                    if underscore_idx > 0:
                        month = parts[:underscore_idx]
                        year = int(parts[underscore_idx + 1:])
                        if month in self.MONTHS and 2000 <= year <= 2100:
                            lists.append((month, year))
                except (ValueError, IndexError):
                    continue
        return lists
    
    def apply_styles(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #e8f6f3, stop:0.5 #ebf5fb, stop:1 #e8f8f5);
            }
            QFrame#listsFrame, QFrame#newListFrame {
                background-color: white;
                border: 2px solid #85c1e9;
                border-radius: 12px;
            }
            QPushButton#openListBtn {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton#openListBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton#openListBtn:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0e2f44, stop:1 #196f3d);
            }
            QPushButton#openListBtn:disabled {
                background-color: #bdc3c7;
            }
            QPushButton#addListBtn {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton#addListBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
            QPushButton#backBtn {
                background-color: #1a5276;
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton#backBtn:hover {
                background-color: #154360;
            }
            QComboBox, QSpinBox {
                border: 2px solid #85c1e9;
                border-radius: 6px;
                padding: 8px 12px;
                background-color: white;
            }
            QComboBox:focus, QSpinBox:focus {
                border: 2px solid #27ae60;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 10px;
            }
        """)
    
    def open_selected_list(self):
        """Open the selected transfer list from dropdown"""
        idx = self.previous_list_combo.currentIndex()
        if idx <= 0:  # First item is placeholder
            QMessageBox.warning(self, "Warning", "Please select a transfer list from the dropdown!")
            return
        
        data = self.previous_list_combo.currentData()
        if data:
            self.selected_month, self.selected_year = data
            self.accept()
    
    def select_list(self, month, year):
        self.selected_month = month
        self.selected_year = year
        self.accept()
    
    def add_new_list(self):
        month = self.new_month_combo.currentText()
        year = self.new_year_spin.value()
        
        # Check if already exists
        if (month, year) in self.available_lists:
            reply = QMessageBox.question(self, "Transfer List Exists", 
                f"Transfer list for {month} {year} already exists. Do you want to open it?",
                QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.selected_month = month
                self.selected_year = year
                self.accept()
        else:
            # New list - confirm creation
            reply = QMessageBox.question(self, "Create New Transfer List",
                f"This will create a new database for {month} {year}.\nContinue?",
                QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.selected_month = month
                self.selected_year = year
                self.accept()
    
    def get_selected_month(self):
        return self.selected_month
    
    def get_selected_year(self):
        return self.selected_year


class TransferDistrictDialog(QDialog):
    """Dialog for selecting district to transfer"""
    
    def __init__(self, parent=None, employee_name=""):
        super().__init__(parent)
        self.setWindowTitle("Add to Transfer List")
        self.setMinimumWidth(400)
        self.setFont(QFont("Calibri", 10))
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        
        # Employee info label
        info_label = QLabel(f"Employee: {employee_name}")
        info_label.setFont(QFont("Calibri", 11, QFont.Bold))
        layout.addWidget(info_label)
        
        # District selection
        form_layout = QFormLayout()
        
        self.district_combo = QComboBox()
        self.district_combo.setEditable(True)  # Allow typing to search
        self.district_combo.setInsertPolicy(QComboBox.NoInsert)
        
        districts = ["Thiruvananthapuram", "Kollam", "Pathanamthitta", "Alappuzha",
                    "Kottayam", "Idukki", "Ernakulam", "Thrissur", "Palakkad",
                    "Malappuram", "Kozhikode", "Wayanad", "Kannur", "Kasaragod"]
        self.district_combo.addItems(districts)
        self.district_combo.setCurrentIndex(-1)  # No default selection
        self.district_combo.lineEdit().setPlaceholderText("Type to search district...")
        
        form_layout.addRow("District to be Transferred:", self.district_combo)
        layout.addLayout(form_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.ok_btn = QPushButton("Add to List")
        self.ok_btn.clicked.connect(self.accept)
        self.ok_btn.setDefault(True)
        
        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(self.ok_btn)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(button_layout)
    
    def get_district(self):
        return self.district_combo.currentText()


class ApplicationDetailsDialog(QDialog):
    """Dialog for entering application details when marking employees as applied"""
    
    def __init__(self, parent=None, employee_count=1, employee_names=None, districts=None, current_district=None):
        super().__init__(parent)
        self.setWindowTitle("Transfer Application Details")
        self.setWindowState(Qt.WindowMaximized)  # Full screen
        self.setFont(QFont("Calibri", 10))
        
        self.districts = districts or []
        self.current_district = current_district
        # Filter out current district from available options
        if current_district:
            self.all_districts = [d for d in self.districts if d != current_district]
        else:
            self.all_districts = self.districts
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Title
        if employee_count == 1 and employee_names:
            title = QLabel(f"Application Details for: {employee_names[0]}")
        else:
            title = QLabel(f"Application Details for {employee_count} Employee(s)")
        title.setFont(QFont("Calibri", 16, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        # Main content in horizontal layout
        main_content = QHBoxLayout()
        main_content.setSpacing(20)
        
        # Left side - Receipt, Date, Weightage
        left_frame = QFrame()
        left_frame.setFrameStyle(QFrame.StyledPanel)
        left_layout = QVBoxLayout(left_frame)
        
        # Receipt Numbers Section
        receipt_group = QGroupBox("Receipt Numbers")
        receipt_layout = QVBoxLayout(receipt_group)
        
        # Receipt list widget
        self.receipt_list = QTableWidget()
        self.receipt_list.setColumnCount(2)
        self.receipt_list.setHorizontalHeaderLabels(["Receipt Number", "Remove"])
        self.receipt_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.receipt_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.receipt_list.setMinimumHeight(100) # Ensure it's visible
        receipt_layout.addWidget(self.receipt_list)
        
        # Add receipt row
        add_receipt_layout = QHBoxLayout()
        self.receipt_input = QLineEdit()
        self.receipt_input.setPlaceholderText("Enter receipt number...")
        self.receipt_input.returnPressed.connect(self.add_receipt)
        add_receipt_layout.addWidget(self.receipt_input)
        
        self.add_receipt_btn = QPushButton("➕ Add Receipt")
        self.add_receipt_btn.clicked.connect(self.add_receipt)
        add_receipt_layout.addWidget(self.add_receipt_btn)
        receipt_layout.addLayout(add_receipt_layout)
        
        left_layout.addWidget(receipt_group)
        
        # Application Date
        date_group = QGroupBox("Application Date")
        date_layout = QVBoxLayout(date_group)
        self.app_date_edit = NoScrollDateEdit()
        self.app_date_edit.setCalendarPopup(True)
        self.app_date_edit.setDate(QDate.currentDate())
        self.app_date_edit.setDisplayFormat("dd-MM-yyyy")
        self.app_date_edit.setMinimumHeight(35)
        date_layout.addWidget(self.app_date_edit)
        left_layout.addWidget(date_group)
        
        # Weightage Section
        weightage_group = QGroupBox("Weightage (Optional)")
        weightage_layout = QVBoxLayout(weightage_group)
        
        self.weightage_check = QCheckBox("Has Weightage")
        self.weightage_check.stateChanged.connect(self.toggle_weightage_details)
        weightage_layout.addWidget(self.weightage_check)
        
        # Weightage Priority
        priority_layout = QHBoxLayout()
        self.weightage_priority_label = QLabel("Weightage Priority:")
        self.weightage_priority_label.setEnabled(False)
        priority_layout.addWidget(self.weightage_priority_label)
        
        self.weightage_priority_combo = NoScrollComboBox()
        self.weightage_priority_combo.addItems(["1 - Highest", "2 - High", "3 - Medium", "4 - Low", "5 - Lowest"])
        self.weightage_priority_combo.setCurrentIndex(0)  # Default to highest priority
        self.weightage_priority_combo.setEnabled(False)
        self.weightage_priority_combo.setMinimumWidth(150)
        priority_layout.addWidget(self.weightage_priority_combo)
        priority_layout.addStretch()
        weightage_layout.addLayout(priority_layout)
        
        self.weightage_details_label = QLabel("Weightage Details:")
        self.weightage_details_label.setEnabled(False)
        weightage_layout.addWidget(self.weightage_details_label)
        
        self.weightage_details = QTextEdit()
        self.weightage_details.setPlaceholderText("Enter weightage details (e.g., Spouse working in different district, Medical reasons, etc.)")
        self.weightage_details.setMaximumHeight(80)
        self.weightage_details.setEnabled(False)
        weightage_layout.addWidget(self.weightage_details)
        
        left_layout.addWidget(weightage_group)
        
        # Special Priority Section (HIGHEST PRIORITY - processed before weightage)
        special_group = QGroupBox("Special Priority")
        special_group.setStyleSheet("QGroupBox { font-weight: bold; color: #c0392b; }")
        special_layout = QVBoxLayout(special_group)
        
        self.special_priority_check = QCheckBox("Grant Special Priority")
        self.special_priority_check.setStyleSheet("font-weight: bold; color: #c0392b;")
        special_layout.addWidget(self.special_priority_check)
        
        special_info = QLabel("⚠️ Special Priority employees are processed FIRST,\nbefore weightage employees, for their preferred district.")
        special_info.setStyleSheet("color: #666; font-size: 9pt; font-style: italic;")
        special_layout.addWidget(special_info)
        
        left_layout.addWidget(special_group)
        left_layout.addStretch()
        
        main_content.addWidget(left_frame, 1)
        
        # Right side - District Preferences
        right_frame = QFrame()
        right_frame.setFrameStyle(QFrame.StyledPanel)
        right_layout = QVBoxLayout(right_frame)
        
        pref_group = QGroupBox("District Preferences (Select up to 8 preferences in order of priority)")
        pref_group.setFont(QFont("Calibri", 11, QFont.Bold))
        pref_layout = QVBoxLayout(pref_group)
        
        if current_district:
            district_note = QLabel(f"Note: Current District '{current_district}' is excluded from preferences")
            district_note.setStyleSheet("color: #dc3545; font-style: italic;")
            pref_layout.addWidget(district_note)
        
        # Create preference combos in a grid
        pref_form = QFormLayout()
        pref_form.setSpacing(12)
        
        self.pref_combos = []
        self._updating_combos = False
        
        for i in range(8):
            combo = SearchableComboBox()
            combo.setMinimumHeight(30)
            combo.setMinimumWidth(250)
            combo.addItem("")
            combo.addItems(self.all_districts)
            combo.setCurrentIndex(0)
            combo.currentIndexChanged.connect(self.update_available_districts)
            self.pref_combos.append(combo)
            pref_form.addRow(f"Preference {i + 1}:", combo)
        
        pref_layout.addLayout(pref_form)
        
        # Quick clear button
        clear_pref_btn = QPushButton("Clear All Preferences")
        clear_pref_btn.clicked.connect(self.clear_all_preferences)
        pref_layout.addWidget(clear_pref_btn)
        
        right_layout.addWidget(pref_group)
        right_layout.addStretch()
        
        main_content.addWidget(right_frame, 1)
        
        layout.addLayout(main_content)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        self.ok_btn = QPushButton("✅ Mark as Applied")
        self.ok_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold; padding: 10px 30px; font-size: 12pt;")
        self.ok_btn.clicked.connect(self.validate_and_accept)
        self.ok_btn.setDefault(True)
        button_layout.addWidget(self.ok_btn)
        
        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.setStyleSheet("padding: 10px 30px; font-size: 12pt;")
        self.cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(button_layout)
        
        # Initialize with one empty receipt row
        self.receipts = []
    
    def update_available_districts(self):
        """Update each combo box to show only unselected districts"""
        if self._updating_combos:
            return
        
        self._updating_combos = True
        
        # Get currently selected districts (excluding empty)
        selected_districts = set()
        for combo in self.pref_combos:
            text = combo.currentText().strip()
            if text:
                selected_districts.add(text)
        
        # Update each combo box
        for combo in self.pref_combos:
            current_text = combo.currentText().strip()
            
            # Block signals while updating
            combo.blockSignals(True)
            combo.clear()
            combo.addItem("")  # Empty option
            
            # Add districts that are either not selected or are the current selection
            for district in self.all_districts:
                if district not in selected_districts or district == current_text:
                    combo.addItem(district)
            
            # Restore selection
            if current_text:
                idx = combo.findText(current_text)
                if idx >= 0:
                    combo.setCurrentIndex(idx)
            
            # Refresh completer after items update
            combo.refresh_completer()
            
            combo.blockSignals(False)
        
        self._updating_combos = False
    
    def clear_all_preferences(self):
        """Clear all preference selections"""
        self._updating_combos = True
        for combo in self.pref_combos:
            combo.setCurrentIndex(0)
        self._updating_combos = False
        self.update_available_districts()
    
    def validate_and_accept(self):
        """Validate that no district is selected more than once before accepting"""
        selected = []
        for i, combo in enumerate(self.pref_combos):
            text = combo.currentText().strip()
            if text:
                if text in selected:
                    QMessageBox.warning(self, "Duplicate Selection", 
                                       f"'{text}' is already selected as a preference. Each district can only be selected once.")
                    return
                selected.append(text)
        self.accept()
    
    def add_receipt(self):
        """Add a receipt number to the list"""
        receipt_num = self.receipt_input.text().strip()
        if receipt_num:
            if receipt_num in self.receipts:
                QMessageBox.warning(self, "Duplicate", "This receipt number is already added!")
                return
            
            self.receipts.append(receipt_num)
            row = self.receipt_list.rowCount()
            self.receipt_list.insertRow(row)
            
            # Receipt number
            item = QTableWidgetItem(receipt_num)
            self.receipt_list.setItem(row, 0, item)
            
            # Remove button
            remove_btn = QPushButton("❌")
            remove_btn.setMaximumWidth(40)
            remove_btn.clicked.connect(lambda checked, r=receipt_num: self.remove_receipt(r))
            self.receipt_list.setCellWidget(row, 1, remove_btn)
            
            self.receipt_input.clear()
            self.receipt_input.setFocus()
    
    def remove_receipt(self, receipt_num):
        """Remove a receipt from the list"""
        if receipt_num in self.receipts:
            self.receipts.remove(receipt_num)
            # Rebuild the table
            self.receipt_list.setRowCount(0)
            for r in self.receipts:
                row = self.receipt_list.rowCount()
                self.receipt_list.insertRow(row)
                item = QTableWidgetItem(r)
                self.receipt_list.setItem(row, 0, item)
                remove_btn = QPushButton("❌")
                remove_btn.setMaximumWidth(40)
                remove_btn.clicked.connect(lambda checked, rn=r: self.remove_receipt(rn))
                self.receipt_list.setCellWidget(row, 1, remove_btn)
    
    def toggle_weightage_details(self, state):
        """Enable/disable weightage details based on checkbox"""
        enabled = state == Qt.CheckState.Checked.value
        self.weightage_priority_label.setEnabled(enabled)
        self.weightage_priority_combo.setEnabled(enabled)
        self.weightage_details_label.setEnabled(enabled)
        self.weightage_details.setEnabled(enabled)
        if not enabled:
            self.weightage_details.clear()
            self.weightage_priority_combo.setCurrentIndex(0)
    
    def get_data(self):
        """Return the entered data"""
        # Get preferences
        preferences = []
        for combo in self.pref_combos:
            text = combo.currentText().strip()
            preferences.append(text if text else None)
        
        # Get weightage priority (1-5, extract number from combo text)
        weightage_priority = self.weightage_priority_combo.currentIndex() + 1 if self.weightage_check.isChecked() else 5
        
        return {
            'receipt_numbers': ", ".join(self.receipts) if self.receipts else "",
            'application_date': self.app_date_edit.date().toString("dd-MM-yyyy"),
            'has_weightage': self.weightage_check.isChecked(),
            'weightage_priority': weightage_priority,
            'weightage_details': self.weightage_details.toPlainText().strip() if self.weightage_check.isChecked() else "",
            'special_priority': self.special_priority_check.isChecked(),
            'special_priority_reason': "",
            'preferences': preferences
        }


class TransferListPreviewDialog(QDialog):
    """Dialog for previewing and printing the transfer list in official format"""
    
    def __init__(self, parent=None, transfer_data=None, districts=None, transfer_type="general", order_number=""):
        super().__init__(parent)
        self.setWindowTitle("Transfer List Preview - Government of Kerala")
        self.setMinimumSize(900, 700)
        self.setFont(QFont("Calibri", 10))
        
        self.transfer_data = transfer_data or []
        self.districts = districts or []
        self.transfer_type = transfer_type
        self.order_number = order_number
        
        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        
        # Toolbar
        toolbar_layout = QHBoxLayout()
        
        self.print_btn = QPushButton("🖨️ Print")
        self.print_btn.setStyleSheet("background-color: #007bff; color: white; font-weight: bold;")
        self.print_btn.clicked.connect(self.print_document)
        toolbar_layout.addWidget(self.print_btn)
        
        self.export_pdf_btn = QPushButton("📄 Export to PDF")
        self.export_pdf_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold;")
        self.export_pdf_btn.clicked.connect(self.export_to_pdf)
        toolbar_layout.addWidget(self.export_pdf_btn)
        
        self.export_html_btn = QPushButton("🌐 Export to HTML")
        self.export_html_btn.clicked.connect(self.export_to_html)
        toolbar_layout.addWidget(self.export_html_btn)
        
        self.export_excel_btn = QPushButton("📊 Export to Excel")
        self.export_excel_btn.setStyleSheet("background-color: #217346; color: white; font-weight: bold;")
        self.export_excel_btn.clicked.connect(self.export_to_excel)
        toolbar_layout.addWidget(self.export_excel_btn)
        
        self.export_word_btn = QPushButton("📝 Export to Word")
        self.export_word_btn.setStyleSheet("background-color: #2b579a; color: white; font-weight: bold;")
        self.export_word_btn.clicked.connect(self.export_to_word)
        toolbar_layout.addWidget(self.export_word_btn)
        
        toolbar_layout.addStretch()
        
        self.close_btn = QPushButton("Close")
        self.close_btn.clicked.connect(self.reject)
        toolbar_layout.addWidget(self.close_btn)
        
        self.confirm_btn = QPushButton("✅ Confirm & Save")
        self.confirm_btn.setStyleSheet("background-color: #dc3545; color: white; font-weight: bold;")
        self.confirm_btn.clicked.connect(self.accept)
        toolbar_layout.addWidget(self.confirm_btn)
        
        layout.addLayout(toolbar_layout)
        
        # Preview area
        self.preview_browser = QTextBrowser()
        self.preview_browser.setOpenExternalLinks(False)
        self.preview_browser.setFont(QFont("Times New Roman", 11))
        layout.addWidget(self.preview_browser)
        
        # Generate the HTML content
        self.html_content = self.generate_html()
        self.preview_browser.setHtml(self.html_content)
    
    def generate_html(self):
        """Generate HTML content for the transfer list"""
        if self.transfer_type == "regular":
            return self._generate_regular_transfer_html()
        else:
            return self._generate_general_transfer_html()
    
    def _generate_regular_transfer_html(self):
        """Generate HTML for Regular Transfer in Malayalam format - matching Word template exactly"""
        total_records = len(self.transfer_data)
        
        html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4 portrait;
                    margin: 1in;
                }}
                @font-face {{
                    font-family: 'MANDARAM';
                    src: url('mandaram.ttf');
                }}
                body {{
                    font-family: 'MANDARAM', 'Noto Sans Malayalam', 'Manjari', Arial, sans-serif;
                    font-size: 14pt;
                    margin: 30px;
                    line-height: 1.15;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 12pt;
                }}
                .header-text {{
                    font-size: 17pt;
                    font-weight: bold;
                    text-decoration: underline;
                    line-height: 1.15;
                }}
                .order-info {{
                    text-align: center;
                    margin-top: 12pt;
                    margin-bottom: 12pt;
                    font-size: 16pt;
                    font-weight: bold;
                    text-decoration: underline;
                }}
                .subject-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 12pt;
                    margin-bottom: 0;
                }}
                .subject-table td {{
                    padding: 6pt 8pt;
                    vertical-align: top;
                    font-size: 14pt;
                    border: none;
                }}
                .subject-table td:first-child {{
                    width: 85px;
                    white-space: nowrap;
                }}
                .subject-table td:last-child {{
                    text-align: justify;
                }}
                .body-text {{
                    text-align: justify;
                    margin-top: 12pt;
                    margin-bottom: 12pt;
                    text-indent: 36pt;
                    font-size: 14pt;
                    line-height: 1.15;
                }}
                table.data-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 12pt auto;
                }}
                table.data-table th, table.data-table td {{
                    border: 1px solid black;
                    padding: 6pt 8pt;
                    font-size: 14pt;
                }}
                table.data-table th {{
                    text-align: center;
                    font-weight: normal;
                }}
                table.data-table td {{
                    text-align: justify;
                }}
                .signature {{
                    margin-top: 24pt;
                    margin-left: 55%;
                }}
                .signature p {{
                    margin: 0;
                    font-weight: bold;
                    font-size: 14pt;
                }}
                .recipients {{
                    margin-top: 12pt;
                }}
                .recipients p {{
                    margin-top: 12pt;
                    margin-bottom: 6pt;
                    font-size: 14pt;
                }}
                .indent {{
                    margin-left: 36pt;
                }}
                .bold {{
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <span class="header-text">തിരുവനന്തപുരം, ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ കാര്യാലയത്തിലെ<br>അഡീഷണൽ ഡയറക്ടർ (എ & റ്റി) – യുടെ നടപടിക്രമം</span>
            </div>
            
            <table class="subject-table">
                <tr>
                    <td style="text-align: right;">വിഷയം :</td>
                    <td>ആ.വ. – ആ.വ.ഡ - ജൂനിയർ പബ്ലിക് ഹെൽത്ത് നഴ്‌സ് ഗ്രേഡ് 1 തസ്തികയിലെ ജീവനക്കാർക്ക് സ്ഥലംമാറ്റം അനുവദിച്ച് ഉത്തരവ് പുറപ്പെടുവിക്കുന്നു.</td>
                </tr>
                <tr>
                    <td style="text-align: right;">പരാമർശം :</td>
                    <td>വിവിധ ജില്ലാ മെഡിക്കൽ ഓഫീസ് (ആരോഗ്യം) - ൽ നിന്നും ലഭ്യമായ ജൂനിയർ പബ്ലിക്ക് ഹെൽത്ത് നഴ്സ് ഗ്രേഡ് I ജീവനക്കാരുടെ അപേക്ഷകൾ.</td>
                </tr>
            </table>
            
            <div class="order-info">
                ഉത്തരവ് നം. {self.order_number}, തീയതി: #ApprovedDate#
            </div>
            
            <p class="body-text">
                വിവിധ ജില്ലാ മെഡിക്കൽ ഓഫീസ് (ആരോഗ്യം) - ൽ നിന്നും മേൽ പരാമർശ പ്രകാരം ലഭ്യമായ ജൂനിയർ പബ്ലിക്ക് ഹെൽത്ത് നഴ്സ് ഗ്രേഡ് I – ജീവനക്കാരുടെ സ്ഥലംമാറ്റത്തിനുള്ള അപേക്ഷകൾ പരിഗണിച്ച്, ചുവടെ ചേർക്കുന്ന ജീവനക്കാർക്ക് അവരുടെ പേരിന് നേരെ രേഖപ്പെടുത്തിയിട്ടുള്ള ജില്ലകളിലേക്ക് സ്ഥലം മാറ്റം നൽകി ഉത്തരവാകുന്നു.
            </p>
            
            <table class="data-table">
                <tr>
                    <th style="width: 6%;">Sl. No.</th>
                    <th style="width: 10%;">PEN</th>
                    <th style="width: 22%;">Name</th>
                    <th style="width: 24%;">Present Institution</th>
                    <th style="width: 19%;">Present District</th>
                    <th style="width: 19%;">Allotted District</th>
                </tr>
        '''
        
        for idx, row in enumerate(self.transfer_data):
            html += f'''
                <tr>
                    <td style="text-align: center;">{idx + 1}</td>
                    <td>{row['pen']}</td>
                    <td>{row['name']}</td>
                    <td>{row['institution']}</td>
                    <td>{row['from_district'].upper()}</td>
                    <td>{row['to_district'].upper()}</td>
                </tr>
            '''
        
        html += '''
            </table>
            
            <p class="body-text">
                ബന്ധപ്പെട്ട ജില്ലാ മെഡിക്കൽ ഓഫീസർമാർ ജീവനക്കാർക്ക് നിയമന ഉത്തരവ് നൽകേണ്ടതും നിയമന ഉത്തരവ് ലഭിച്ച ശേഷം സ്ഥാപന മേധാവികൾ ജീവനക്കാരെ വിടുതൽ ചെയ്യേണ്ടതും, ടി വിവരം യഥാസമയേ ഈ കാര്യാലയത്തിൽ റിപ്പോർട്ട് ചെയ്യേണ്ടതുമാണ്.
            </p>
            
            <p class="body-text">
                ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ ഔദ്യോഗിക വെബ്സൈറ്റിൽ പ്രസിദ്ധീകരിച്ചിട്ടുള്ള ഈ ഉത്തരവിന്റെ പകർപ്പ് ഔദ്യോഗിക രേഖയായി ഉപയോഗിക്കാവുന്നതാണ്.
            </p>
            
            <div class="signature">
                <p>#ApprovedByName#</p>
                <p>#ApprovedByDesignation#</p>
            </div>
            
            <div class="recipients">
                <p>സ്വീകർത്താവ്</p>
                <p class="indent bold">ബന്ധപ്പെട്ട ജില്ലാ മെഡിക്കൽ ഓഫീസർ (ആരോഗ്യം) മാർ</p>
                
                <p>പകർപ്പ്:</p>
                <p class="indent">1. ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ ഔദ്യോഗിക വെബ്സൈറ്റ്</p>
                <p class="indent">2. ഫയൽ/കരുതൽ ഫയൽ</p>
            </div>
        </body>
        </html>
        '''
        
        return html
    
    def _generate_general_transfer_html(self):
        """Generate HTML for General Transfer in English format - matching Word template"""
        # Group data by From District
        district_groups = {}
        for row in self.transfer_data:
            from_district = row['from_district']
            if from_district not in district_groups:
                district_groups[from_district] = []
            district_groups[from_district].append(row)
        
        # Calculate total pages (roughly)
        total_records = len(self.transfer_data)
        
        html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                @page {{
                    size: A4 landscape;
                    margin: 0.5in;
                }}
                body {{
                    font-family: "Times New Roman", Times, serif;
                    font-size: 12pt;
                    margin: 20px;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 20px;
                }}
                .header h2 {{
                    margin: 5px 0;
                    font-size: 14pt;
                    font-weight: bold;
                }}
                .header h3 {{
                    margin: 5px 0;
                    font-size: 12pt;
                    font-weight: bold;
                }}
                .header h4 {{
                    margin: 5px 0;
                    font-size: 12pt;
                    font-weight: bold;
                    text-decoration: underline;
                }}
                .post-cadre {{
                    margin: 15px 0;
                    font-weight: bold;
                    font-size: 10pt;
                }}
                .district-header {{
                    margin: 15px 0 10px 0;
                    font-weight: bold;
                    font-size: 11pt;
                }}
                table.data-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 15px;
                }}
                table.data-table th, table.data-table td {{
                    border: 1px solid black;
                    padding: 5px 6px;
                    font-size: 10pt;
                }}
                table.data-table th {{
                    text-align: center;
                    font-weight: bold;
                    background-color: #f0f0f0;
                }}
                table.data-table td {{
                    text-align: left;
                }}
                .page-info {{
                    text-align: right;
                    font-size: 10pt;
                    margin-bottom: 10px;
                }}
            </style>
        </head>
        <body>
            <div class="page-info">Page 1 of {max(1, (total_records // 15) + 1)}</div>
            
            <div class="header">
                <h2>Government of Kerala</h2>
                <h3>Department: Health Services</h3>
                <h4>Norms Based General Transfer for Junior Public Health Nurse Gr. I</h4>
            </div>
            
            <div class="post-cadre">Post/Cadre Name: Junior Public Health Nurse</div>
        '''
        
        sl_no = 1
        for district in self.districts:
            if district not in district_groups:
                continue
            
            records = district_groups[district]
            
            html += f'''
            <div class="district-header">District: {district.upper()}</div>
            <table class="data-table">
                <tr>
                    <th style="width: 5%;">Sl. No.</th>
                    <th style="width: 8%;">PEN</th>
                    <th style="width: 18%;">Name</th>
                    <th style="width: 12%;">Designation</th>
                    <th style="width: 22%;">Office Transferred from</th>
                    <th style="width: 11%;">From District</th>
                    <th style="width: 11%;">To District</th>
                    <th style="width: 13%;">Protection If Any</th>
                </tr>
            '''
            
            for row in records:
                protection = ""
                if row.get('weightage') == 'Yes' and row.get('weightage_details'):
                    protection = row['weightage_details']
                
                html += f'''
                <tr>
                    <td style="text-align: center;">{sl_no}</td>
                    <td>{row['pen']}</td>
                    <td>{row['name']}</td>
                    <td>{row['designation']}</td>
                    <td>{row['institution']}</td>
                    <td>{row['from_district']}</td>
                    <td>{row['to_district']}</td>
                    <td>{protection}</td>
                </tr>
                '''
                sl_no += 1
            
            html += '</table>'
        
        html += '''
            <br><br>
            <div style="margin-top: 30px;">
                <div style="text-align: right; margin-right: 50px;">
                    <p>_________________________</p>
                    <p><strong>Authorized Signatory</strong></p>
                </div>
            </div>
        </body>
        </html>
        '''
        
        return html
    
    def print_document(self):
        """Print the transfer list"""
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageSize(QPageSize(QPageSize.A4))
        if self.transfer_type == "regular":
            printer.setPageOrientation(QPageLayout.Orientation.Portrait)
        else:
            printer.setPageOrientation(QPageLayout.Orientation.Landscape)
        
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.Accepted:
            document = QTextDocument()
            document.setHtml(self.html_content)
            document.print_(printer)
    
    def export_to_pdf(self):
        """Export to PDF file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export to PDF", 
            f"transfer_list_{datetime.now().strftime('%Y%m%d')}.pdf",
            "PDF Files (*.pdf)"
        )
        
        if file_path:
            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(file_path)
            printer.setPageSize(QPageSize(QPageSize.A4))
            if self.transfer_type == "regular":
                printer.setPageOrientation(QPageLayout.Orientation.Portrait)
            else:
                printer.setPageOrientation(QPageLayout.Orientation.Landscape)
            
            document = QTextDocument()
            document.setHtml(self.html_content)
            document.setHtml(self.html_content)
            document.print_(printer)
            
            QMessageBox.information(self, "Success", f"PDF exported to:\n{file_path}")
    
    def export_to_html(self):
        """Export to HTML file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export to HTML", 
            f"transfer_list_{datetime.now().strftime('%Y%m%d')}.html",
            "HTML Files (*.html)"
        )
        
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.html_content)
            QMessageBox.information(self, "Success", f"HTML exported to:\n{file_path}")
    
    def export_to_excel(self):
        """Export to Excel file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export to Excel", 
            f"transfer_list_{datetime.now().strftime('%Y%m%d')}.xlsx",
            "Excel Files (*.xlsx)"
        )
        
        if file_path:
            try:
                import openpyxl
                from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
                
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Transfer List"
                
                # Set column widths
                ws.column_dimensions['A'].width = 8   # Sl. No.
                ws.column_dimensions['B'].width = 12  # PEN
                ws.column_dimensions['C'].width = 25  # Name
                ws.column_dimensions['D'].width = 20  # Designation
                ws.column_dimensions['E'].width = 35  # Office
                ws.column_dimensions['F'].width = 18  # From District
                ws.column_dimensions['G'].width = 18  # To District
                ws.column_dimensions['H'].width = 20  # Protection
                
                # Styles
                header_font = Font(bold=True, size=14)
                subheader_font = Font(bold=True, size=12)
                district_font = Font(bold=True, size=11)
                table_header_font = Font(bold=True, size=10)
                table_header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
                thin_border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                center_align = Alignment(horizontal='center', vertical='center')
                left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
                
                # Header
                ws.merge_cells('A1:H1')
                ws['A1'] = 'Government of Kerala'
                ws['A1'].font = header_font
                ws['A1'].alignment = center_align
                
                ws.merge_cells('A2:H2')
                ws['A2'] = 'Department: Health Services'
                ws['A2'].font = subheader_font
                ws['A2'].alignment = center_align
                
                ws.merge_cells('A3:H3')
                ws['A3'] = 'Norms Based General Transfer for Junior Public Health Nurse Gr. I'
                ws['A3'].font = subheader_font
                ws['A3'].alignment = center_align
                
                ws.merge_cells('A4:H4')
                ws['A4'] = 'Post/Cadre Name: Junior Public Health Nurse'
                ws['A4'].font = Font(bold=True, size=10)
                ws['A4'].alignment = Alignment(horizontal='left')
                
                row_num = 6
                sl_no = 1
                
                # Group data by district
                district_groups = {}
                for row in self.transfer_data:
                    from_district = row['from_district']
                    if from_district not in district_groups:
                        district_groups[from_district] = []
                    district_groups[from_district].append(row)
                
                headers = ['Sl. No.', 'PEN', 'Name', 'Designation', 'Office Transferred from', 
                          'From District', 'To District', 'Protection If Any']
                
                for district in self.districts:
                    if district not in district_groups:
                        continue
                    
                    records = district_groups[district]
                    
                    # District header
                    ws.merge_cells(f'A{row_num}:H{row_num}')
                    ws[f'A{row_num}'] = f'District: {district.upper()}'
                    ws[f'A{row_num}'].font = district_font
                    row_num += 1
                    
                    # Table headers
                    for col_num, header in enumerate(headers, 1):
                        cell = ws.cell(row=row_num, column=col_num, value=header)
                        cell.font = table_header_font
                        cell.fill = table_header_fill
                        cell.border = thin_border
                        cell.alignment = center_align
                    row_num += 1
                    
                    # Data rows
                    for record in records:
                        protection = ""
                        if record.get('weightage') == 'Yes' and record.get('weightage_details'):
                            protection = record['weightage_details']
                        
                        data = [sl_no, record['pen'], record['name'], record['designation'],
                               record['institution'], record['from_district'], record['to_district'], protection]
                        
                        for col_num, value in enumerate(data, 1):
                            cell = ws.cell(row=row_num, column=col_num, value=value)
                            cell.border = thin_border
                            if col_num == 1:
                                cell.alignment = center_align
                            else:
                                cell.alignment = left_align
                        
                        sl_no += 1
                        row_num += 1
                    
                    row_num += 1  # Empty row between districts
                
                wb.save(file_path)
                QMessageBox.information(self, "Success", f"Excel file exported to:\n{file_path}")
                
            except ImportError:
                QMessageBox.warning(self, "Missing Library", 
                    "openpyxl library is required for Excel export.\n\n"
                    "Install it using:\npip install openpyxl")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export Excel:\n{str(e)}")
    
    def export_to_word(self):
        """Export to Word document"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export to Word", 
            f"transfer_list_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        
        if file_path:
            try:
                from docx import Document
                from docx.shared import Inches, Pt, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                
                doc = Document()
                
                # Set page margins
                section = doc.sections[0]
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                
                if self.transfer_type == "regular":
                    # Regular Transfer - Malayalam format
                    self._export_regular_transfer_word(doc)
                else:
                    # General Transfer - English format with landscape
                    section.page_width, section.page_height = section.page_height, section.page_width
                    section.left_margin = Cm(1.5)
                    section.right_margin = Cm(1.5)
                    section.top_margin = Cm(1.5)
                    section.bottom_margin = Cm(1.5)
                    self._export_general_transfer_word(doc)
                
                doc.save(file_path)
                QMessageBox.information(self, "Success", f"Word document exported to:\n{file_path}")
                
            except ImportError:
                QMessageBox.warning(self, "Missing Library", 
                    "python-docx library is required for Word export.\n\n"
                    "Install it using:\npip install python-docx")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export Word document:\n{str(e)}")
    
    def _export_regular_transfer_word(self, doc):
        """Export Regular Transfer in Malayalam format - matching template exactly"""
        from docx.shared import Pt, Cm, Inches, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Set font name for Malayalam (MANDARAM as in template)
        font_name = 'MANDARAM'
        
        # Helper function to set paragraph spacing
        def set_paragraph_spacing(para, space_before_pt=12, space_after_pt=12, line_spacing=1.15):
            para.paragraph_format.space_before = Pt(space_before_pt)
            para.paragraph_format.space_after = Pt(space_after_pt)
            para.paragraph_format.line_spacing = line_spacing
        
        # Helper function to add tab stop
        def add_tab_stop(para, position_cm):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = para._p.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'left')
            tab.set(qn('w:pos'), str(int(position_cm * 567)))  # Convert cm to twips (1cm = 567 twips)
            tabs.append(tab)
        
        # Title - Malayalam (17pt Bold Underline Center, both lines in same paragraph with line break)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run1 = title.add_run('തിരുവനന്തപുരം, ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ കാര്യാലയത്തിലെ')
        title_run1.bold = True
        title_run1.underline = True
        title_run1.font.size = Pt(17)
        title_run1.font.name = font_name
        title.add_run('\n')
        title_run2 = title.add_run('അഡീഷണൽ ഡയറക്ടർ (എ & റ്റി) – യുടെ നടപടിക്രമം')
        title_run2.bold = True
        title_run2.underline = True
        title_run2.font.size = Pt(17)
        title_run2.font.name = font_name
        set_paragraph_spacing(title, space_before_pt=0, space_after_pt=12)
        
        # Subject/Reference table (No borders, 2 columns)
        subject_table = doc.add_table(rows=2, cols=2)
        subject_table.autofit = True
        
        # Remove borders by setting style to 'Normal Table'
        subject_table.style = 'Normal Table'
        
        # Set column widths - first column narrow, second column wide
        from docx.shared import Cm
        for row in subject_table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(14)
        
        # Set table indent (move entire table 1pt to the right)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tbl = subject_table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '20')  # 20 twips = ~1pt
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
        
        # Subject row
        cell00 = subject_table.cell(0, 0)
        cell00.text = 'വിഷയം :'
        cell00.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(cell00.paragraphs[0], space_before_pt=12, space_after_pt=0)
        for run in cell00.paragraphs[0].runs:
            run.font.name = font_name
            run.font.size = Pt(14)
        
        cell01 = subject_table.cell(0, 1)
        cell01.text = 'ആ.വ. – ആ.വ.ഡ - ജൂനിയർ പബ്ലിക് ഹെൽത്ത് നഴ്‌സ് ഗ്രേഡ് 1 തസ്തികയിലെ ജീവനക്കാർക്ക് സ്ഥലംമാറ്റം അനുവദിച്ച് ഉത്തരവ് പുറപ്പെടുവിക്കുന്നു.'
        cell01.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(cell01.paragraphs[0], space_before_pt=12, space_after_pt=0)
        for run in cell01.paragraphs[0].runs:
            run.font.name = font_name
            run.font.size = Pt(14)
        
        # Reference row
        cell10 = subject_table.cell(1, 0)
        cell10.text = 'പരാമർശം :'
        cell10.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(cell10.paragraphs[0], space_before_pt=12, space_after_pt=0)
        for run in cell10.paragraphs[0].runs:
            run.font.name = font_name
            run.font.size = Pt(14)
        
        cell11 = subject_table.cell(1, 1)
        cell11.text = 'വിവിധ ജില്ലാ മെഡിക്കൽ ഓഫീസ് (ആരോഗ്യം) - ൽ നിന്നും ലഭ്യമായ ജൂനിയർ പബ്ലിക്ക് ഹെൽത്ത് നഴ്സ് ഗ്രേഡ് I ജീവനക്കാരുടെ അപേക്ഷകൾ.'
        cell11.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(cell11.paragraphs[0], space_before_pt=12, space_after_pt=0)
        for run in cell11.paragraphs[0].runs:
            run.font.name = font_name
            run.font.size = Pt(14)
        
        # Order number and date (16pt Bold Underline Center)
        order_para = doc.add_paragraph()
        order_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(order_para, space_before_pt=12, space_after_pt=12)
        order_run = order_para.add_run(f'ഉത്തരവ് നം. {self.order_number}, തീയതി: #ApprovedDate#')
        order_run.bold = True
        order_run.underline = True
        order_run.font.size = Pt(16)
        order_run.font.name = font_name
        
        # Body text (14pt, JUSTIFY, with indent)
        body1 = doc.add_paragraph()
        body1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body1.paragraph_format.first_line_indent = Pt(36)
        set_paragraph_spacing(body1, space_before_pt=12, space_after_pt=12)
        body1_run = body1.add_run('വിവിധ ജില്ലാ മെഡിക്കൽ ഓഫീസ് (ആരോഗ്യം) - ൽ നിന്നും മേൽ പരാമർശ പ്രകാരം ലഭ്യമായ ജൂനിയർ പബ്ലിക്ക് ഹെൽത്ത് നഴ്സ് ഗ്രേഡ് I – ജീവനക്കാരുടെ സ്ഥലംമാറ്റത്തിനുള്ള അപേക്ഷകൾ പരിഗണിച്ച്, ചുവടെ ചേർക്കുന്ന ജീവനക്കാർക്ക് അവരുടെ പേരിന് നേരെ രേഖപ്പെടുത്തിയിട്ടുള്ള ജില്ലകളിലേക്ക് സ്ഥലം മാറ്റം നൽകി ഉത്തരവാകുന്നു.')
        body1_run.font.size = Pt(14)
        body1_run.font.name = font_name
        
        # Data table - 6 columns as per template (14pt, headers CENTER, data JUSTIFY)
        headers = ['Sl. No.', 'PEN', 'Name', 'Present Institution', 'Present District', 'Allotted District']
        
        total_records = len(self.transfer_data)
        
        table = doc.add_table(rows=1 + total_records, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row (14pt, CENTER aligned)
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = font_name
                run.font.size = Pt(14)
        
        # Data rows (14pt, Sl. No. CENTER, others JUSTIFY aligned)
        for row_idx, record in enumerate(self.transfer_data):
            data = [
                str(row_idx + 1),
                record['pen'],
                record['name'],
                record['institution'],
                record['from_district'].upper(),
                record['to_district'].upper()
            ]
            
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(data):
                cell = row.cells[col_idx]
                cell.text = str(value)
                # Center align the Sl. No. column (first column)
                if col_idx == 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in cell.paragraphs[0].runs:
                    run.font.name = font_name
                    run.font.size = Pt(14)
        
        # Instructions paragraph 1 (14pt, JUSTIFY)
        instr1 = doc.add_paragraph()
        instr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        instr1.paragraph_format.first_line_indent = Pt(36)
        set_paragraph_spacing(instr1, space_before_pt=12, space_after_pt=12)
        instr1_run = instr1.add_run('ബന്ധപ്പെട്ട ജില്ലാ മെഡിക്കൽ ഓഫീസർമാർ ജീവനക്കാർക്ക് നിയമന ഉത്തരവ് നൽകേണ്ടതും നിയമന ഉത്തരവ് ലഭിച്ച ശേഷം സ്ഥാപന മേധാവികൾ ജീവനക്കാരെ വിടുതൽ ചെയ്യേണ്ടതും, ടി വിവരം യഥാസമയേ ഈ കാര്യാലയത്തിൽ റിപ്പോർട്ട് ചെയ്യേണ്ടതുമാണ്.')
        instr1_run.font.size = Pt(14)
        instr1_run.font.name = font_name
        
        # Instructions paragraph 2 (14pt, JUSTIFY)
        instr2 = doc.add_paragraph()
        instr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        instr2.paragraph_format.first_line_indent = Pt(36)
        set_paragraph_spacing(instr2, space_before_pt=12, space_after_pt=12)
        instr2_run = instr2.add_run('ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ ഔദ്യോഗിക വെബ്സൈറ്റിൽ പ്രസിദ്ധീകരിച്ചിട്ടുള്ള ഈ ഉത്തരവിന്റെ പകർപ്പ് ഔദ്യോഗിക രേഖയായി ഉപയോഗിക്കാവുന്നതാണ്.')
        instr2_run.font.size = Pt(14)
        instr2_run.font.name = font_name
        
        # Empty paragraphs before signature
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature area (14pt Bold, with indent - JUSTIFY aligned as per template)
        sig_name = doc.add_paragraph()
        sig_name.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sig_name.paragraph_format.first_line_indent = Pt(252)  # Large indent to push right
        set_paragraph_spacing(sig_name, space_before_pt=12, space_after_pt=0)
        sig_name_run = sig_name.add_run('#ApprovedByName#')
        sig_name_run.bold = True
        sig_name_run.font.size = Pt(14)
        sig_name_run.font.name = font_name
        
        sig_desig = doc.add_paragraph()
        sig_desig.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sig_desig.paragraph_format.first_line_indent = Pt(252)
        set_paragraph_spacing(sig_desig, space_before_pt=0, space_after_pt=12)
        sig_desig_run = sig_desig.add_run('#ApprovedByDesignation#')
        sig_desig_run.bold = True
        sig_desig_run.font.size = Pt(14)
        sig_desig_run.font.name = font_name
        
        # Recipients (14pt, NOT bold as per template)
        recv = doc.add_paragraph()
        recv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(recv, space_before_pt=12, space_after_pt=6)
        recv_run = recv.add_run('സ്വീകർത്താവ്')
        recv_run.font.size = Pt(14)
        recv_run.font.name = font_name
        
        recv_list = doc.add_paragraph()
        recv_list.paragraph_format.first_line_indent = Pt(36)
        set_paragraph_spacing(recv_list, space_before_pt=12, space_after_pt=12)
        recv_list_run = recv_list.add_run('ബന്ധപ്പെട്ട ജില്ലാ മെഡിക്കൽ ഓഫീസർ (ആരോഗ്യം) മാർ')
        recv_list_run.bold = True
        recv_list_run.font.size = Pt(14)
        recv_list_run.font.name = font_name
        
        # Copy to (14pt)
        copy = doc.add_paragraph()
        set_paragraph_spacing(copy, space_before_pt=12, space_after_pt=6)
        copy_run = copy.add_run('പകർപ്പ്:')
        copy_run.font.size = Pt(14)
        copy_run.font.name = font_name
        
        # Numbered list items
        copy1 = doc.add_paragraph()
        copy1.paragraph_format.left_indent = Pt(36)
        set_paragraph_spacing(copy1, space_before_pt=6, space_after_pt=6)
        copy1_run = copy1.add_run('1. ആരോഗ്യവകുപ്പ് ഡയറക്ടറുടെ ഔദ്യോഗിക വെബ്സൈറ്റ്')
        copy1_run.font.size = Pt(14)
        copy1_run.font.name = font_name
        
        copy2 = doc.add_paragraph()
        copy2.paragraph_format.left_indent = Pt(36)
        set_paragraph_spacing(copy2, space_before_pt=6, space_after_pt=12)
        copy2_run = copy2.add_run('2. ഫയൽ/കരുതൽ ഫയൽ')
        copy2_run.font.size = Pt(14)
        copy2_run.font.name = font_name

    def _export_general_transfer_word(self, doc):
        """Export General Transfer in English format"""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Header
        header1 = doc.add_paragraph('Government of Kerala')
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header1.runs[0].bold = True
        header1.runs[0].font.size = Pt(14)
        
        header2 = doc.add_paragraph('Department: Health Services')
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header2.runs[0].bold = True
        header2.runs[0].font.size = Pt(12)
        
        header3 = doc.add_paragraph('Norms Based General Transfer for Junior Public Health Nurse Gr. I')
        header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header3.runs[0].bold = True
        header3.runs[0].font.size = Pt(12)
        header3.runs[0].underline = True
        
        cadre = doc.add_paragraph('Post/Cadre Name: Junior Public Health Nurse')
        cadre.runs[0].bold = True
        cadre.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Group data by district
        district_groups = {}
        for row in self.transfer_data:
            from_district = row['from_district']
            if from_district not in district_groups:
                district_groups[from_district] = []
            district_groups[from_district].append(row)
        
        sl_no = 1
        headers = ['Sl. No.', 'PEN', 'Name', 'Designation', 'Office Transferred from', 
                  'From District', 'To District', 'Protection If Any']
        
        # Column widths in inches
        col_widths = [0.5, 0.8, 1.8, 1.2, 2.5, 1.2, 1.2, 1.3]
        
        for district in self.districts:
            if district not in district_groups:
                continue
            
            records = district_groups[district]
            
            # District header
            district_para = doc.add_paragraph(f'District: {district.upper()}')
            district_para.runs[0].bold = True
            district_para.runs[0].font.size = Pt(11)
            
            # Create table
            table = doc.add_table(rows=1 + len(records), cols=8)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
            
            # Header row
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for row_idx, record in enumerate(records):
                protection = ""
                if record.get('weightage') == 'Yes' and record.get('weightage_details'):
                    protection = record['weightage_details']
                
                data = [str(sl_no), record['pen'], record['name'], record['designation'],
                       record['institution'], record['from_district'], record['to_district'], protection]
                
                row = table.rows[row_idx + 1]
                for col_idx, value in enumerate(data):
                    cell = row.cells[col_idx]
                    cell.text = str(value)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if col_idx == 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                sl_no += 1
            
            doc.add_paragraph()  # Space between districts
        
        # Signature area
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_para.add_run('_________________________\n')
        sig_run.font.size = Pt(10)
        sig_run2 = sig_para.add_run('Authorized Signatory')
        sig_run2.bold = True
        sig_run2.font.size = Pt(10)


class PreferenceDialog(QDialog):
    """Dialog for setting 8 district preferences for an employee"""
    
    def __init__(self, parent=None, employee_name="", pen="", current_prefs=None, districts=None, current_district=None):
        super().__init__(parent)
        self.setWindowTitle("Set District Preferences")
        self.setMinimumWidth(500)
        self.setFont(QFont("Calibri", 10))
        
        if current_prefs is None:
            current_prefs = [""] * 8
        if districts is None:
            districts = []
        
        # Filter out current district from available options
        self.current_district = current_district
        if current_district:
            self.all_districts = [d for d in districts if d != current_district]
        else:
            self.all_districts = districts
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        
        # Employee info
        info_frame = QFrame()
        info_frame.setFrameStyle(QFrame.StyledPanel)
        info_layout = QVBoxLayout(info_frame)
        
        name_label = QLabel(f"Employee: {employee_name}")
        name_label.setFont(QFont("mandaram.ttf", 11, QFont.Bold))
        info_layout.addWidget(name_label)
        
        pen_label = QLabel(f"PEN: {pen}")
        pen_label.setFont(QFont("Calibri", 10))
        info_layout.addWidget(pen_label)
        
        if current_district:
            district_label = QLabel(f"Current District: {current_district} (excluded from preferences)")
            district_label.setFont(QFont("Calibri", 10))
            district_label.setStyleSheet("color: #dc3545;")
            info_layout.addWidget(district_label)
        
        layout.addWidget(info_frame)
        
        # Instructions
        inst_label = QLabel("Select up to 8 district preferences in order of priority (each district can only be selected once):")
        inst_label.setFont(QFont("Calibri", 10))
        inst_label.setStyleSheet("color: #666;")
        inst_label.setWordWrap(True)
        layout.addWidget(inst_label)
        
        # Preferences form
        form_layout = QFormLayout()
        form_layout.setSpacing(10)
        
        self.pref_combos = []
        self._updating_combos = False  # Flag to prevent recursive updates
        
        for i in range(8):
            combo = SearchableComboBox()
            combo.setMinimumHeight(30)
            combo.addItem("")  # Empty option
            combo.addItems(districts)
            
            # Set current value if exists
            if current_prefs[i]:
                idx = combo.findText(current_prefs[i])
                if idx >= 0:
                    combo.setCurrentIndex(idx)
                else:
                    combo.setCurrentIndex(0)
            else:
                combo.setCurrentIndex(0)  # Empty
            
            # Connect to update available districts when selection changes
            combo.currentIndexChanged.connect(self.update_available_districts)
            
            self.pref_combos.append(combo)
            form_layout.addRow(f"Preference {i + 1}:", combo)
        
        layout.addLayout(form_layout)
        
        # Update available districts after all combos are created
        self.update_available_districts()
        
        # Quick clear button
        clear_btn = QPushButton("Clear All Preferences")
        clear_btn.clicked.connect(self.clear_all)
        layout.addWidget(clear_btn)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.ok_btn = QPushButton("Save Preferences")
        self.ok_btn.clicked.connect(self.validate_and_accept)
        self.ok_btn.setDefault(True)
        self.ok_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold;")
        
        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(self.ok_btn)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(button_layout)
    
    def update_available_districts(self):
        """Update each combo box to show only unselected districts"""
        if self._updating_combos:
            return
        
        self._updating_combos = True
        
        # Get currently selected districts (excluding empty)
        selected_districts = set()
        for combo in self.pref_combos:
            text = combo.currentText().strip()
            if text:
                selected_districts.add(text)
        
        # Update each combo box
        for combo in self.pref_combos:
            current_text = combo.currentText().strip()
            
            # Block signals while updating
            combo.blockSignals(True)
            combo.clear()
            combo.addItem("")  # Empty option
            
            # Add districts that are either not selected or are the current selection
            for district in self.all_districts:
                if district not in selected_districts or district == current_text:
                    combo.addItem(district)
            
            # Restore selection
            if current_text:
                idx = combo.findText(current_text)
                if idx >= 0:
                    combo.setCurrentIndex(idx)
            
            # Refresh completer after items update
            combo.refresh_completer()
            
            combo.blockSignals(False)
        
        self._updating_combos = False
    
    def validate_and_accept(self):
        """Validate that no district is selected more than once before accepting"""
        selected = []
        for i, combo in enumerate(self.pref_combos):
            text = combo.currentText().strip()
            if text:
                if text in selected:
                    from PySide6.QtWidgets import QMessageBox
                    QMessageBox.warning(self, "Duplicate Selection", 
                                       f"'{text}' is already selected as a preference. Each district can only be selected once.")
                    return
                selected.append(text)
        self.accept()
    
    def clear_all(self):
        """Clear all preference selections"""
        self._updating_combos = True
        for combo in self.pref_combos:
            combo.setCurrentIndex(0)
        self._updating_combos = False
        self.update_available_districts()
    
    def get_preferences(self):
        """Return list of 8 preferences"""
        prefs = []
        for combo in self.pref_combos:
            text = combo.currentText().strip()
            prefs.append(text if text else None)
        return prefs


class JPHNDialog(QDialog):
    """Dialog for adding/editing JPHN records"""
    
    def __init__(self, parent=None, record=None):
        super().__init__(parent)
        self.record = record
        self.setWindowTitle("Add JPHN Record" if record is None else "Edit JPHN Record")
        self.setMinimumWidth(500)
        self.setup_ui()
        
        if record:
            self.load_record(record)
    
    def setup_ui(self):
        layout = QFormLayout()
        
        # Name
        self.name_edit = QLineEdit()
        self.name_edit.setFont(QFont("mandaram.ttf", 10))
        layout.addRow("Name:", self.name_edit)
        
        # PEN
        self.pen_edit = QLineEdit()
        layout.addRow("PEN:", self.pen_edit)
        
        # Designation
        self.designation_edit = QLineEdit()
        self.designation_edit.setText("JPHN Gr I")
        layout.addRow("Designation:", self.designation_edit)
        
        # Present Institution
        self.institution_edit = QLineEdit()
        self.institution_edit.setFont(QFont("mandaram.ttf", 10))
        layout.addRow("Present Institution:", self.institution_edit)
        
        # District
        self.district_combo = NoScrollComboBox()
        districts = ["Thiruvananthapuram", "Kollam", "Pathanamthitta", "Alappuzha", 
                    "Kottayam", "Idukki", "Ernakulam", "Thrissur", "Palakkad", 
                    "Malappuram", "Kozhikode", "Wayanad", "Kannur", "Kasaragod"]
        self.district_combo.addItems(districts)
        layout.addRow("District:", self.district_combo)
        
        # Date of Entry
        self.entry_date = NoScrollDateEdit()
        self.entry_date.setCalendarPopup(True)
        self.entry_date.setDisplayFormat("dd-MM-yyyy")
        self.entry_date.setDate(QDate.currentDate())
        layout.addRow("Date of Entry:", self.entry_date)
        
        # Date of Retirement
        self.retirement_date = NoScrollDateEdit()
        self.retirement_date.setCalendarPopup(True)
        self.retirement_date.setDisplayFormat("dd-MM-yyyy")
        self.retirement_date.setDate(QDate.currentDate().addYears(30))
        layout.addRow("Date of Retirement:", self.retirement_date)
        
        # Date of Joining in Present District
        self.district_join_date = NoScrollDateEdit()
        self.district_join_date.setCalendarPopup(True)
        self.district_join_date.setDisplayFormat("dd-MM-yyyy")
        self.district_join_date.setDate(QDate.currentDate())
        layout.addRow("Date of Joining in Present District:", self.district_join_date)
        
        # Date of Joining in Present Institution
        self.institution_join_date = NoScrollDateEdit()
        self.institution_join_date.setCalendarPopup(True)
        self.institution_join_date.setDisplayFormat("dd-MM-yyyy")
        self.institution_join_date.setDate(QDate.currentDate())
        layout.addRow("Date of Joining in Present Institution:", self.institution_join_date)
        
        # Contact Number
        self.contact_edit = QLineEdit()
        self.contact_edit.setPlaceholderText("10-digit mobile number")
        layout.addRow("Contact:", self.contact_edit)
        
        # Weightage
        self.weightage_check = QCheckBox("Yes")
        layout.addRow("Weightage:", self.weightage_check)
        
        # Weightage Details
        self.weightage_details = QTextEdit()
        self.weightage_details.setMaximumHeight(80)
        self.weightage_details.setFont(QFont("mandaram.ttf", 10))
        self.weightage_details.setEnabled(False)
        layout.addRow("Weightage Details:", self.weightage_details)
        
        # Connect weightage checkbox
        self.weightage_check.stateChanged.connect(self.on_weightage_changed)
        
        # Buttons
        button_layout = QHBoxLayout()
        self.save_btn = QPushButton("Save")
        self.cancel_btn = QPushButton("Cancel")
        
        self.save_btn.clicked.connect(self.accept)
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(self.save_btn)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addRow(button_layout)
        self.setLayout(layout)
    
    def on_weightage_changed(self, state):
        """Enable/disable weightage details based on checkbox state"""
        is_checked = state == Qt.CheckState.Checked.value
        self.weightage_details.setEnabled(is_checked)
        if not is_checked:
            self.weightage_details.clear()
    
    def load_record(self, record):
        """Load existing record into form"""
        # New order: pen, name, designation, institution, district, entry_date, 
        # retirement_date, district_join_date, duration_days, institution_join_date, weightage, weightage_details, contact
        self.pen_edit.setText(str(record[0]) if record[0] else "")
        self.name_edit.setText(str(record[1]) if record[1] else "")
        self.designation_edit.setText(str(record[2]) if record[2] else "")
        self.institution_edit.setText(str(record[3]) if record[3] else "")
        if record[4]:
            self.district_combo.setCurrentText(str(record[4]))
        
        # Handle dates with proper null checking
        if record[5]:
            date = QDate.fromString(str(record[5]), "dd-MM-yyyy")
            if date.isValid():
                self.entry_date.setDate(date)
        if record[6]:
            date = QDate.fromString(str(record[6]), "dd-MM-yyyy")
            if date.isValid():
                self.retirement_date.setDate(date)
        if record[7]:
            date = QDate.fromString(str(record[7]), "dd-MM-yyyy")
            if date.isValid():
                self.district_join_date.setDate(date)
        if record[9]:
            date = QDate.fromString(str(record[9]), "dd-MM-yyyy")
            if date.isValid():
                self.institution_join_date.setDate(date)
        
        # Contact (index 12 when using explicit column order in SELECT)
        if len(record) > 12 and record[12]:
            contact_val = str(record[12])
            # Skip if it looks like weightage_priority (single digit 1-5) due to old schema
            if not (len(contact_val) == 1 and contact_val.isdigit()):
                self.contact_edit.setText(contact_val)
        
        weightage = record[10] == "Yes" if record[10] else False
        self.weightage_check.setChecked(weightage)
        # Manually enable and set details since signal may not trigger correctly during load
        if weightage:
            self.weightage_details.setEnabled(True)
            if record[11]:
                self.weightage_details.setText(str(record[11]))
    
    def get_data(self):
        """Get form data as dictionary"""
        return {
            'name': self.name_edit.text(),
            'pen': self.pen_edit.text(),
            'designation': self.designation_edit.text(),
            'institution': self.institution_edit.text(),
            'district': self.district_combo.currentText(),
            'entry_date': self.entry_date.date().toString("dd-MM-yyyy"),
            'retirement_date': self.retirement_date.date().toString("dd-MM-yyyy"),
            'district_join_date': self.district_join_date.date().toString("dd-MM-yyyy"),
            'institution_join_date': self.institution_join_date.date().toString("dd-MM-yyyy"),
            'weightage': "Yes" if self.weightage_check.isChecked() else "No",
            'weightage_details': self.weightage_details.toPlainText() if self.weightage_check.isChecked() else "",
            'contact': self.contact_edit.text().strip()
        }


class JPHNManagementSystem(QMainWindow):
    def __init__(self, transfer_type="general", year=None, month=None):
        super().__init__()
        
        # Store transfer type, year, and month, set database name accordingly
        self.transfer_type = transfer_type
        self.year = year
        self.month = month
        
        if transfer_type == "general":
            self.db_name = f'jphn_general_{year}.db'
            self.title_suffix = f"General Transfer {year}"
        else:
            self.db_name = f'jphn_regular_{month}_{year}.db'
            self.title_suffix = f"Regular Transfer - {month} {year}"
        
        self.setWindowTitle(f"JPHN Gr I Management System - {self.title_suffix} - Kerala Health Services")
        self.setMinimumSize(1400, 700)
        
        # Set window icon
        icon_path = resource_path("favicon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # District list
        self.districts = ["Thiruvananthapuram", "Kollam", "Pathanamthitta", "Alappuzha", 
                         "Kottayam", "Idukki", "Ernakulam", "Thrissur", "Palakkad", 
                         "Malappuram", "Kozhikode", "Wayanad", "Kannur", "Kasaragod"]
        
        # Nearby districts mapping (geographically adjacent districts)
        self.nearby_districts = {
            "Thiruvananthapuram": ["Kollam", "Pathanamthitta", "Kottayam"],
            "Kollam": ["Thiruvananthapuram", "Pathanamthitta", "Alappuzha"],
            "Pathanamthitta": ["Kollam", "Alappuzha", "Kottayam", "Idukki"],
            "Alappuzha": ["Kollam", "Pathanamthitta", "Kottayam", "Ernakulam"],
            "Kottayam": ["Pathanamthitta", "Alappuzha", "Idukki", "Ernakulam"],
            "Idukki": ["Pathanamthitta", "Kottayam", "Ernakulam", "Thrissur"],
            "Ernakulam": ["Alappuzha", "Kottayam", "Idukki", "Thrissur"],
            "Thrissur": ["Ernakulam", "Idukki", "Palakkad", "Malappuram"],
            "Palakkad": ["Thrissur", "Malappuram", "Coimbatore"],
            "Malappuram": ["Thrissur", "Palakkad", "Kozhikode", "Wayanad"],
            "Kozhikode": ["Malappuram", "Wayanad", "Kannur"],
            "Wayanad": ["Malappuram", "Kozhikode", "Kannur"],
            "Kannur": ["Kozhikode", "Wayanad", "Kasaragod"],
            "Kasaragod": ["Kannur"]
        }
        
        # Initialize database
        self.init_database()
        
        # Setup UI
        self.setup_ui()
        self.apply_styles()
        
        # Load data - Vacancy data first (mandatory)
        self.load_vacancy_data()
        self.load_data()
        self.load_transfer_list()
    
    def init_database(self):
        """Initialize SQLite database"""
        self.conn = sqlite3.connect(self.db_name)
        self.cursor = self.conn.cursor()
        
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS jphn (
                pen TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                designation TEXT,
                institution TEXT,
                district TEXT,
                entry_date TEXT,
                retirement_date TEXT,
                district_join_date TEXT,
                duration_days INTEGER,
                institution_join_date TEXT,
                weightage TEXT,
                weightage_details TEXT,
                contact TEXT
            )
        ''')
        
        # Create transfer draft table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS transfer_draft (
                pen TEXT PRIMARY KEY,
                transfer_to_district TEXT NOT NULL,
                added_date TEXT,
                against_info TEXT
            )
        ''')
        
        # Add against_info column if it doesn't exist (for existing databases)
        try:
            self.cursor.execute('ALTER TABLE transfer_draft ADD COLUMN against_info TEXT')
        except:
            pass  # Column already exists
        
        # Add contact column if it doesn't exist (for existing databases)
        try:
            self.cursor.execute('ALTER TABLE jphn ADD COLUMN contact TEXT')
        except:
            pass  # Column already exists
        
        # Create transfer final (confirmed) table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS transfer_final (
                pen TEXT PRIMARY KEY,
                transfer_to_district TEXT NOT NULL,
                confirmed_date TEXT
            )
        ''')
        
        # Create vacancy table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS vacancy (
                district TEXT PRIMARY KEY,
                total_strength INTEGER DEFAULT 0,
                vacancy_reported INTEGER DEFAULT 0
            )
        ''')
        
        # Create transfer_applied table for voluntary transfer applications with preferences
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS transfer_applied (
                pen TEXT PRIMARY KEY,
                applied_to_district TEXT,
                applied_date TEXT,
                receipt_numbers TEXT,
                pref1 TEXT,
                pref2 TEXT,
                pref3 TEXT,
                pref4 TEXT,
                pref5 TEXT,
                pref6 TEXT,
                pref7 TEXT,
                pref8 TEXT
            )
        ''')
        
        # Add preference columns if they don't exist (for existing databases)
        try:
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref1 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref2 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref3 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref4 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref5 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref6 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref7 TEXT')
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN pref8 TEXT')
        except:
            pass  # Columns already exist
        
        # Add receipt_numbers column if it doesn't exist
        try:
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN receipt_numbers TEXT')
        except:
            pass  # Column already exists
        
        # Add special_priority column if it doesn't exist
        try:
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN special_priority TEXT DEFAULT "No"')
        except:
            pass  # Column already exists
        
        # Add special_priority_reason column if it doesn't exist
        try:
            self.cursor.execute('ALTER TABLE transfer_applied ADD COLUMN special_priority_reason TEXT')
        except:
            pass  # Column already exists
        
        # Add weightage_priority column if it doesn't exist
        try:
            self.cursor.execute('ALTER TABLE jphn ADD COLUMN weightage_priority INTEGER DEFAULT 5')
        except:
            pass  # Column already exists
        
        # Add contact column if it doesn't exist
        try:
            self.cursor.execute('ALTER TABLE jphn ADD COLUMN contact TEXT')
        except:
            pass  # Column already exists
        
        # Create settings table to store order number and other settings
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        ''')
        
        # Initialize vacancy table with all districts if empty
        self.cursor.execute('SELECT COUNT(*) FROM vacancy')
        if self.cursor.fetchone()[0] == 0:
            for district in self.districts:
                self.cursor.execute('INSERT INTO vacancy (district, total_strength, vacancy_reported) VALUES (?, 0, 0)', (district,))
        
        self.conn.commit()
    
    def get_saved_order_number(self):
        """Get the saved order number from settings"""
        try:
            self.cursor.execute("SELECT value FROM settings WHERE key = 'order_number'")
            result = self.cursor.fetchone()
            if result:
                return result[0]
        except:
            pass
        return "DMOH-TVM/___/2025-A6"  # Default value
    
    def save_order_number(self, order_number):
        """Save the order number to settings"""
        try:
            self.cursor.execute('''
                INSERT OR REPLACE INTO settings (key, value) VALUES ('order_number', ?)
            ''', (order_number,))
            self.conn.commit()
        except Exception as e:
            print(f"Error saving order number: {e}")
    
    def setup_ui(self):
        """Setup the user interface"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # Header with Go Back button
        header_layout = QHBoxLayout()
        
        # Go Back button on left
        self.go_back_btn = QPushButton("← Go Back")
        self.go_back_btn.setMinimumSize(120, 35)
        self.go_back_btn.setFont(QFont("Calibri", 10, QFont.Bold))
        self.go_back_btn.setCursor(Qt.PointingHandCursor)
        self.go_back_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1a5276, stop:1 #27ae60);
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px 15px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #154360, stop:1 #1e8449);
            }
        """)
        self.go_back_btn.clicked.connect(self.go_back_to_selection)
        header_layout.addWidget(self.go_back_btn)
        
        header_layout.addStretch()
        
        # Header title in center
        header_label = QLabel("JPHN Gr I Transfer & Posting Management")
        header_label.setFont(QFont("Calibri", 16, QFont.Bold))
        header_label.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(header_label)
        
        header_layout.addStretch()
        
        # Spacer on right to balance the layout
        spacer_widget = QWidget()
        spacer_widget.setFixedWidth(120)
        header_layout.addWidget(spacer_widget)
        
        main_layout.addLayout(header_layout)
        
        # Tab Widget
        self.tab_widget = QTabWidget()
        self.tab_widget.setFont(QFont("Calibri", 10))
        main_layout.addWidget(self.tab_widget)
        
        # Create Transfer Applied Tab FIRST (for voluntary applications)
        self.setup_transfer_applied_tab()
        
        # Create Vacancy Position Tab (second)
        self.setup_vacancy_tab()
        
        # Create Transfer Draft List Tab
        self.setup_draft_list_tab()
        
        # Create Transfer List Tab
        self.setup_transfer_tab()
        
        # Create Cadre Tab (last)
        self.setup_cadre_tab()
        
        # Connect tab change signal
        self.tab_widget.currentChanged.connect(self.on_tab_changed)
        
        # Status bar
        self.statusBar().showMessage("Ready")
    
    def setup_cadre_tab(self):
        """Setup the Cadre tab with existing functionality"""
        cadre_widget = QWidget()
        cadre_layout = QVBoxLayout(cadre_widget)
        cadre_layout.setSpacing(10)
        cadre_layout.setContentsMargins(10, 10, 10, 10)
        
        # Search and Filter Section
        search_layout = QHBoxLayout()
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search by Name, PEN, Institution...")
        self.search_input.setFont(QFont("Calibri", 10))
        self.search_input.textChanged.connect(self.filter_data)
        search_layout.addWidget(QLabel("Search:"))
        search_layout.addWidget(self.search_input)
        
        self.district_filter = NoScrollComboBox()
        self.district_filter.addItem("All Districts")
        self.district_filter.addItems(self.districts)
        self.district_filter.currentTextChanged.connect(self.filter_data)
        search_layout.addWidget(QLabel("District:"))
        search_layout.addWidget(self.district_filter)
        
        cadre_layout.addLayout(search_layout)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(10)
        headers = ["PEN", "Name", "Designation", "Present Institution", 
                  "District", "Date of Entry", "Date of Retirement", 
                  "Date of Joining in District", "Duration", "Contact"]
        self.table.setHorizontalHeaderLabels(headers)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSortingEnabled(True)
        self.table.setAlternatingRowColors(True)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.ExtendedSelection)  # Enable multi-row selection
        self.table.setFont(QFont("Calibri", 9))
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.table.horizontalHeader().setFont(header_font)
        
        # Enable context menu (right-click)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_cadre_context_menu)
        
        # Enable double-click to edit
        self.table.doubleClicked.connect(self.on_table_double_click)
        
        cadre_layout.addWidget(self.table)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.select_all_btn = QPushButton("☑️ Select All")
        self.select_all_btn.clicked.connect(self.select_all_records)
        
        self.add_btn = QPushButton("➕ Add New JPHN")
        self.add_btn.clicked.connect(self.add_record)
        
        self.edit_btn = QPushButton("✏️ Edit Selected")
        self.edit_btn.clicked.connect(self.edit_record)
        
        self.delete_btn = QPushButton("🗑️ Delete Selected")
        self.delete_btn.clicked.connect(self.delete_record)
        
        self.refresh_btn = QPushButton("🔄 Refresh")
        self.refresh_btn.clicked.connect(self.load_data)
        
        self.export_btn = QPushButton("📊 Export to CSV")
        self.export_btn.clicked.connect(self.export_to_csv)
        
        self.import_btn = QPushButton("📥 Import from File")
        self.import_btn.clicked.connect(self.import_from_file)
        
        button_layout.addWidget(self.select_all_btn)
        button_layout.addWidget(self.add_btn)
        button_layout.addWidget(self.edit_btn)
        button_layout.addWidget(self.delete_btn)
        button_layout.addWidget(self.refresh_btn)
        button_layout.addWidget(self.export_btn)
        button_layout.addWidget(self.import_btn)
        button_layout.addStretch()
        
        cadre_layout.addLayout(button_layout)
        
        self.tab_widget.addTab(cadre_widget, "📋 Cadre List")
    
    def setup_transfer_applied_tab(self):
        """Setup the Transfer Applied tab with sub-tabs for Application and Preference Selection"""
        applied_widget = QWidget()
        applied_layout = QVBoxLayout(applied_widget)
        applied_layout.setSpacing(5)
        applied_layout.setContentsMargins(5, 5, 5, 5)
        
        # Create sub-tab widget
        self.applied_sub_tabs = QTabWidget()
        self.applied_sub_tabs.setFont(QFont("Calibri", 10))
        
        # Setup Application Selection tab
        self.setup_application_selection_tab()
        
        # Setup Preference Selection tab
        self.setup_preference_selection_tab()
        
        applied_layout.addWidget(self.applied_sub_tabs)
        
        self.tab_widget.addTab(applied_widget, "📝 Transfer Application")
    
    def setup_application_selection_tab(self):
        """Setup the Application Selection sub-tab"""
        app_widget = QWidget()
        app_layout = QVBoxLayout(app_widget)
        app_layout.setSpacing(10)
        app_layout.setContentsMargins(10, 10, 10, 10)
        
        # Title
        title_label = QLabel("Application Entry - Mark Employees Who Applied for Transfer")
        title_label.setFont(QFont("Calibri", 12, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        app_layout.addWidget(title_label)
        
        # Filter Section
        filter_frame = QFrame()
        filter_frame.setFrameStyle(QFrame.StyledPanel)
        filter_layout = QHBoxLayout(filter_frame)
        
        # Current District filter
        filter_layout.addWidget(QLabel("Current District:"))
        self.applied_district_filter = NoScrollComboBox()
        self.applied_district_filter.addItem("Select District")
        self.applied_district_filter.addItems(self.districts)
        self.applied_district_filter.currentTextChanged.connect(self.load_employees_for_application)
        filter_layout.addWidget(self.applied_district_filter)
        
        # Search input
        filter_layout.addWidget(QLabel("Search:"))
        self.applied_search_input = QLineEdit()
        self.applied_search_input.setPlaceholderText("Search by Name, PEN, Institution...")
        self.applied_search_input.setFont(QFont("Calibri", 10))
        self.applied_search_input.textChanged.connect(self.filter_applied_table)
        filter_layout.addWidget(self.applied_search_input)
        
        filter_layout.addStretch()
        
        # Refresh button
        self.refresh_applied_btn = QPushButton("🔄 Refresh")
        self.refresh_applied_btn.clicked.connect(self.load_employees_for_application)
        filter_layout.addWidget(self.refresh_applied_btn)
        
        app_layout.addWidget(filter_frame)
        
        # Employee List Table
        self.applied_table = QTableWidget()
        self.applied_table.setColumnCount(5)
        applied_headers = ["Select", "Name", "PEN", "Institution", "Duration"]
        self.applied_table.setHorizontalHeaderLabels(applied_headers)
        self.applied_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.applied_table.horizontalHeader().setStretchLastSection(True)
        self.applied_table.setAlternatingRowColors(True)
        self.applied_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.applied_table.setFont(QFont("Calibri", 10))
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.applied_table.horizontalHeader().setFont(header_font)
        
        app_layout.addWidget(self.applied_table)
        
        # Buttons Section
        button_layout = QHBoxLayout()
        
        self.add_new_employee_btn = QPushButton("➕ Add New Employee")
        self.add_new_employee_btn.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold;")
        self.add_new_employee_btn.clicked.connect(self.add_new_employee_for_application)
        button_layout.addWidget(self.add_new_employee_btn)
        
        self.select_all_applied_btn = QPushButton("☑️ Select All")
        self.select_all_applied_btn.clicked.connect(self.select_all_for_application)
        button_layout.addWidget(self.select_all_applied_btn)
        
        self.deselect_all_applied_btn = QPushButton("☐ Deselect All")
        self.deselect_all_applied_btn.clicked.connect(self.deselect_all_for_application)
        button_layout.addWidget(self.deselect_all_applied_btn)
        
        button_layout.addStretch()
        
        self.mark_applied_btn = QPushButton("✅ Mark as Applied")
        self.mark_applied_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold;")
        self.mark_applied_btn.clicked.connect(self.mark_employees_as_applied)
        button_layout.addWidget(self.mark_applied_btn)
        
        app_layout.addLayout(button_layout)
        
        # Summary
        summary_layout = QHBoxLayout()
        self.applied_summary_label = QLabel("Select a district to view employees")
        self.applied_summary_label.setFont(QFont("Calibri", 10, QFont.Bold))
        summary_layout.addWidget(self.applied_summary_label)
        
        summary_layout.addStretch()
        
        app_layout.addLayout(summary_layout)
        
        # Instructions
        instructions = QLabel("💡 Select employees who have voluntarily applied for transfer and click 'Mark as Applied'")
        instructions.setFont(QFont("Calibri", 9))
        instructions.setStyleSheet("color: #666;")
        app_layout.addWidget(instructions)
        
        self.applied_sub_tabs.addTab(app_widget, "📋 Application Entry")
    
    def setup_preference_selection_tab(self):
        """Setup the Applied Employees List sub-tab"""
        pref_widget = QWidget()
        pref_layout = QVBoxLayout(pref_widget)
        pref_layout.setSpacing(10)
        pref_layout.setContentsMargins(10, 10, 10, 10)
        
        # Title
        title_label = QLabel("Applied Employees List - View All Employees Who Applied for Transfer")
        title_label.setFont(QFont("Calibri", 12, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        pref_layout.addWidget(title_label)
        
        # Filter Section
        filter_frame = QFrame()
        filter_frame.setFrameStyle(QFrame.StyledPanel)
        filter_layout = QHBoxLayout(filter_frame)
        
        # Current District filter
        filter_layout.addWidget(QLabel("Current District:"))
        self.pref_district_filter = NoScrollComboBox()
        self.pref_district_filter.addItem("All Districts")
        self.pref_district_filter.addItems(self.districts)
        self.pref_district_filter.currentTextChanged.connect(self.load_preference_list)
        filter_layout.addWidget(self.pref_district_filter)
        
        # To District filter
        filter_layout.addWidget(QLabel("Most Preferred:"))
        self.pref_to_district_filter = NoScrollComboBox()
        self.pref_to_district_filter.addItem("All Districts")
        self.pref_to_district_filter.addItems(self.districts)
        self.pref_to_district_filter.currentTextChanged.connect(self.load_preference_list)
        filter_layout.addWidget(self.pref_to_district_filter)
        
        # Search input
        filter_layout.addWidget(QLabel("Search:"))
        self.pref_search_input = QLineEdit()
        self.pref_search_input.setPlaceholderText("Search by Name, PEN, Institution...")
        self.pref_search_input.setFont(QFont("Calibri", 10))
        self.pref_search_input.textChanged.connect(self.filter_preference_table)
        filter_layout.addWidget(self.pref_search_input)
        
        filter_layout.addStretch()
        
        # Refresh button
        self.refresh_pref_btn = QPushButton("🔄 Refresh")
        self.refresh_pref_btn.clicked.connect(self.load_preference_list)
        filter_layout.addWidget(self.refresh_pref_btn)
        
        pref_layout.addWidget(filter_frame)
        
        # Applied Employees Table - shows applied employees
        self.pref_table = QTableWidget()
        self.pref_table.setColumnCount(10)
        pref_headers = ["Name", "PEN", "Institution", "Current District", "DOJ in District", "Duration", "Preferred District", "Weightage", "Receipt No.", "App. Date"]
        self.pref_table.setHorizontalHeaderLabels(pref_headers)
        self.pref_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.pref_table.horizontalHeader().setStretchLastSection(True)
        self.pref_table.setAlternatingRowColors(True)
        self.pref_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.pref_table.setSelectionMode(QTableWidget.ExtendedSelection)
        self.pref_table.setFont(QFont("Calibri", 10))
        self.pref_table.setMouseTracking(True)  # Enable mouse tracking for tooltips
        
        # Enable double-click to edit
        self.pref_table.doubleClicked.connect(self.edit_applied_employee)
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.pref_table.horizontalHeader().setFont(header_font)
        
        pref_layout.addWidget(self.pref_table)
        
        # Buttons Section
        button_layout = QHBoxLayout()
        
        self.edit_applied_btn = QPushButton("✏️ Edit")
        self.edit_applied_btn.setStyleSheet("background-color: #ffc107; color: black; font-weight: bold;")
        self.edit_applied_btn.clicked.connect(self.edit_selected_applied_employee)
        button_layout.addWidget(self.edit_applied_btn)
        
        self.remove_from_applied_btn = QPushButton("❌ Remove from Applied")
        self.remove_from_applied_btn.setStyleSheet("background-color: #dc3545; color: white; font-weight: bold;")
        self.remove_from_applied_btn.clicked.connect(self.remove_from_applied_status)
        button_layout.addWidget(self.remove_from_applied_btn)
        
        button_layout.addStretch()
        
        # Auto Fill button
        self.auto_fill_btn = QPushButton("🚀 Auto Fill Vacancies")
        self.auto_fill_btn.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold; padding: 8px 15px;")
        self.auto_fill_btn.clicked.connect(self.auto_fill_vacancies)
        button_layout.addWidget(self.auto_fill_btn)
        
        # Export to Excel button
        self.export_applied_excel_btn = QPushButton("📊 Export to Excel")
        self.export_applied_excel_btn.setStyleSheet("background-color: #217346; color: white; font-weight: bold;")
        self.export_applied_excel_btn.clicked.connect(self.export_applied_employees_to_excel)
        button_layout.addWidget(self.export_applied_excel_btn)
        
        button_layout.addStretch()
        
        # Summary
        self.pref_summary_label = QLabel("Applied Employees: 0")
        self.pref_summary_label.setFont(QFont("Calibri", 10, QFont.Bold))
        button_layout.addWidget(self.pref_summary_label)
        
        pref_layout.addLayout(button_layout)
        
        # Instructions
        instructions = QLabel("💡 Double-click a row to edit details. Hover over a row to see preferences. Use 'Remove from Applied' to remove selected employees.")
        instructions.setFont(QFont("Calibri", 9))
        instructions.setStyleSheet("color: #666;")
        pref_layout.addWidget(instructions)
        
        self.applied_sub_tabs.addTab(pref_widget, "📋 Applied Employees List")
        
        # Connect sub-tab change to load preference list
        self.applied_sub_tabs.currentChanged.connect(self.on_applied_sub_tab_changed)
    
    def on_applied_sub_tab_changed(self, index):
        """Handle sub-tab change in Transfer Applied"""
        if index == 1:  # Applied Employees List tab
            self.load_preference_list()
    
    def load_preference_list(self):
        """Load applied employees list"""
        self.pref_table.setUpdatesEnabled(False)
        self.pref_table.setSortingEnabled(False)
        self.pref_table.setRowCount(0)
        
        current_district = self.pref_district_filter.currentText()
        to_district = self.pref_to_district_filter.currentText()
        
        # Build query with both current and to district filters
        where_clauses = []
        params = []
        
        if current_district != "All Districts":
            where_clauses.append("j.district = ?")
            params.append(current_district)
        
        if to_district != "All Districts":
            # Filter by first preference only (Most Preferred District)
            where_clauses.append("t.pref1 = ?")
            params.append(to_district)
        
        where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
        
        query = f'''
            SELECT j.pen, j.name, j.institution, j.district,
                   t.receipt_numbers, t.applied_date, j.duration_days,
                   t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8,
                   t.special_priority, j.weightage, j.weightage_details, j.district_join_date
            FROM jphn j
            INNER JOIN transfer_applied t ON j.pen = t.pen
            WHERE {where_sql}
            ORDER BY j.district, j.duration_days DESC
        '''
        self.cursor.execute(query, params)
        
        records = self.cursor.fetchall()
        self.pref_table.setRowCount(len(records))
        
        for row_idx, record in enumerate(records):
            pen = record[0]
            special_priority = record[15] if len(record) > 15 and record[15] else None
            is_special = special_priority and str(special_priority).lower() == 'yes'
            
            # Build preferences tooltip
            prefs = [record[i] for i in range(7, 15) if record[i]]
            if prefs:
                tooltip = "District Preferences:\n" + "\n".join([f"{i+1}. {p}" for i, p in enumerate(prefs)])
            else:
                tooltip = "No preferences set"
            if is_special:
                tooltip = "🔴 SPECIAL PRIORITY\n\n" + tooltip
            
            # Name
            name_item = QTableWidgetItem(str(record[1]) if record[1] else "")
            name_item.setFont(QFont("mandaram.ttf", 10))
            name_item.setToolTip(tooltip)
            name_item.setData(Qt.UserRole, pen)  # Store PEN for editing
            self.pref_table.setItem(row_idx, 0, name_item)
            
            # PEN
            pen_item = QTableWidgetItem(str(pen) if pen else "")
            pen_item.setToolTip(tooltip)
            self.pref_table.setItem(row_idx, 1, pen_item)
            
            # Institution
            inst_item = QTableWidgetItem(str(record[2]) if record[2] else "")
            inst_item.setFont(QFont("mandaram.ttf", 10))
            inst_item.setToolTip(tooltip)
            self.pref_table.setItem(row_idx, 2, inst_item)
            
            # Current District
            district_item = QTableWidgetItem(str(record[3]) if record[3] else "")
            district_item.setToolTip(tooltip)
            self.pref_table.setItem(row_idx, 3, district_item)
            
            # DOJ in District (Date of Joining in Present District)
            doj_district = record[18] if len(record) > 18 and record[18] else ""  # district_join_date is at index 18
            doj_item = QTableWidgetItem(str(doj_district))
            doj_item.setTextAlignment(Qt.AlignCenter)
            doj_item.setToolTip(tooltip)
            doj_item.setBackground(QColor("#e2e3e5"))  # Light gray
            self.pref_table.setItem(row_idx, 4, doj_item)
            
            # Duration
            duration = record[6] if record[6] else 0
            duration_item = QTableWidgetItem(self.format_duration(duration))
            duration_item.setTextAlignment(Qt.AlignCenter)
            duration_item.setToolTip(tooltip)
            # Highlight long durations
            if duration > 1825:  # More than 5 years
                duration_item.setBackground(QColor("#ffcccc"))
            elif duration > 1095:  # More than 3 years
                duration_item.setBackground(QColor("#fff3cd"))
            self.pref_table.setItem(row_idx, 5, duration_item)
            
            # Preferred District (First Preference - pref1)
            pref_district = record[7] if record[7] else ""  # pref1 is at index 7
            pref_district_item = QTableWidgetItem(str(pref_district))
            pref_district_item.setToolTip(tooltip)
            if is_special:
                # Red bold for special priority
                pref_district_item.setForeground(QColor("#c0392b"))  # Red color
                pref_district_item.setFont(QFont("Calibri", 10, QFont.Bold))
                pref_district_item.setBackground(QColor("#fadbd8"))  # Light red background
            else:
                pref_district_item.setBackground(QColor("#d4edda"))  # Light green
            self.pref_table.setItem(row_idx, 6, pref_district_item)
            
            # Weightage
            weightage_val = record[16] if len(record) > 16 else None
            weightage_details = record[17] if len(record) > 17 and record[17] else ""
            weightage_display = "Yes" if weightage_val and str(weightage_val).lower() in ["yes", "y", "1", "true"] else "No"
            weightage_item = QTableWidgetItem(weightage_display)
            weightage_item.setTextAlignment(Qt.AlignCenter)
            # Show weightage details in tooltip
            if weightage_display == "Yes" and weightage_details:
                weightage_item.setToolTip(f"Weightage Details:\n{weightage_details}")
                weightage_item.setBackground(QColor("#d4edda"))  # Light green
                weightage_item.setFont(QFont("Calibri", 10, QFont.Bold))
            elif weightage_display == "Yes":
                weightage_item.setToolTip("Weightage: Yes (No details provided)")
                weightage_item.setBackground(QColor("#d4edda"))  # Light green
                weightage_item.setFont(QFont("Calibri", 10, QFont.Bold))
            else:
                weightage_item.setToolTip("No weightage")
            self.pref_table.setItem(row_idx, 7, weightage_item)
            
            # Receipt Numbers
            receipt_item = QTableWidgetItem(str(record[4]) if record[4] else "")
            receipt_item.setBackground(QColor("#fff3cd"))  # Light yellow
            receipt_item.setToolTip(tooltip)
            self.pref_table.setItem(row_idx, 8, receipt_item)
            
            # Application Date
            date_item = QTableWidgetItem(str(record[5]) if record[5] else "")
            date_item.setTextAlignment(Qt.AlignCenter)
            date_item.setBackground(QColor("#cce5ff"))  # Light blue
            date_item.setToolTip(tooltip)
            self.pref_table.setItem(row_idx, 9, date_item)
        
        self.pref_table.setUpdatesEnabled(True)
        self.pref_table.setSortingEnabled(True)
        
        self.pref_summary_label.setText(f"Applied Employees: {len(records)}")
    
    def export_applied_employees_to_excel(self):
        """Export Applied Employees List to Excel"""
        if self.pref_table.rowCount() == 0:
            QMessageBox.warning(self, "No Data", "No applied employees to export!")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Applied Employees to Excel",
            f"Applied_Employees_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            "Excel Files (*.xlsx)"
        )
        
        if not file_path:
            return
        
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Applied Employees"
            
            # Set column widths
            ws.column_dimensions['A'].width = 8    # Sl. No.
            ws.column_dimensions['B'].width = 25   # Name
            ws.column_dimensions['C'].width = 12   # PEN
            ws.column_dimensions['D'].width = 30   # Institution
            ws.column_dimensions['E'].width = 18   # Current District
            ws.column_dimensions['F'].width = 15   # DOJ in District
            ws.column_dimensions['G'].width = 12   # Duration
            ws.column_dimensions['H'].width = 18   # Preferred District
            ws.column_dimensions['I'].width = 10   # Weightage
            ws.column_dimensions['J'].width = 15   # Receipt No.
            ws.column_dimensions['K'].width = 12   # App. Date
            
            # Styles
            header_font = Font(bold=True, size=14)
            subheader_font = Font(bold=True, size=11)
            table_header_font = Font(bold=True, size=10)
            table_header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            table_header_font_white = Font(bold=True, size=10, color="FFFFFF")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_align = Alignment(horizontal='center', vertical='center')
            left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
            
            # Title
            ws.merge_cells('A1:K1')
            ws['A1'] = 'Applied Employees List - JPHN Gr I Transfer'
            ws['A1'].font = header_font
            ws['A1'].alignment = center_align
            
            # Filters info
            current_filter = self.pref_district_filter.currentText()
            to_filter = self.pref_to_district_filter.currentText()
            ws.merge_cells('A2:K2')
            ws['A2'] = f'Current District: {current_filter} | Preferred District: {to_filter} | Export Date: {datetime.now().strftime("%d-%m-%Y %H:%M")}'
            ws['A2'].font = Font(size=10, italic=True)
            ws['A2'].alignment = center_align
            
            # Table headers
            headers = ["Sl. No.", "Name", "PEN", "Institution", "Current District", "DOJ in District", "Duration", "Preferred District", "Weightage", "Receipt No.", "App. Date"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = table_header_font_white
                cell.fill = table_header_fill
                cell.border = thin_border
                cell.alignment = center_align
            
            # Data rows
            for row_idx in range(self.pref_table.rowCount()):
                excel_row = row_idx + 5
                
                # Sl. No.
                cell = ws.cell(row=excel_row, column=1, value=row_idx + 1)
                cell.border = thin_border
                cell.alignment = center_align
                
                # Data columns
                for col_idx in range(self.pref_table.columnCount()):
                    item = self.pref_table.item(row_idx, col_idx)
                    value = item.text() if item else ""
                    cell = ws.cell(row=excel_row, column=col_idx + 2, value=value)
                    cell.border = thin_border
                    if col_idx in [4, 5, 7, 8, 9]:  # DOJ, Duration, Weightage, Receipt, Date - center align
                        cell.alignment = center_align
                    else:
                        cell.alignment = left_align
            
            # Summary row
            summary_row = self.pref_table.rowCount() + 6
            ws.merge_cells(f'A{summary_row}:K{summary_row}')
            ws[f'A{summary_row}'] = f'Total Applied Employees: {self.pref_table.rowCount()}'
            ws[f'A{summary_row}'].font = Font(bold=True, size=11)
            ws[f'A{summary_row}'].alignment = center_align
            
            wb.save(file_path)
            QMessageBox.information(self, "Success", f"Applied Employees exported to:\n{file_path}")
            
        except ImportError:
            QMessageBox.critical(self, "Error", "openpyxl module not installed.\nPlease install it using: pip install openpyxl")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export to Excel:\n{str(e)}")
    
    def remove_from_applied_status(self):
        """Remove selected employees from applied status (moves back to Application Selection)"""
        selected_rows = self.pref_table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select employee(s) to remove from applied status!")
            return
        
        # Collect selected employees
        selected_employees = []
        for index in selected_rows:
            row = index.row()
            pen = self.pref_table.item(row, 1).text()
            name = self.pref_table.item(row, 0).text()
            selected_employees.append((pen, name))
        
        if len(selected_employees) == 1:
            msg = f"Remove '{selected_employees[0][1]}' from applied status?\n\nThis will delete all preferences and move the employee back to Application Selection tab."
        else:
            msg = f"Remove {len(selected_employees)} employees from applied status?\n\nThis will delete all their preferences and move them back to Application Selection tab."
        
        reply = QMessageBox.question(self, "Confirm", msg, QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            try:
                for pen, name in selected_employees:
                    self.cursor.execute('DELETE FROM transfer_applied WHERE pen = ?', (pen,))
                
                self.conn.commit()
                self.load_preference_list()
                
                if len(selected_employees) == 1:
                    QMessageBox.information(self, "Success", f"'{selected_employees[0][1]}' removed from applied status.\nEmployee will now appear in Application Selection tab.")
                else:
                    QMessageBox.information(self, "Success", f"{len(selected_employees)} employees removed from applied status.\nThey will now appear in Application Selection tab.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to remove from applied: {str(e)}")
    
    def edit_selected_applied_employee(self):
        """Edit the selected applied employee (Edit button handler)"""
        selected_rows = self.pref_table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select an employee to edit!")
            return
        
        if len(selected_rows) > 1:
            QMessageBox.warning(self, "Warning", "Please select only one employee to edit!")
            return
        
        # Call the existing edit function with the selected index
        self.edit_applied_employee(selected_rows[0])
    
    def edit_applied_employee(self, index):
        """Edit an applied employee's details including preferences (double-click handler)"""
        row = index.row()
        
        # Get PEN from the first column's UserRole data
        name_item = self.pref_table.item(row, 0)
        if name_item:
            pen = name_item.data(Qt.UserRole)
        else:
            pen = self.pref_table.item(row, 1).text()
        
        if not pen:
            return
        
        # Fetch employee details from jphn table
        self.cursor.execute('''
            SELECT name, pen, institution, district
            FROM jphn WHERE pen = ?
        ''', (pen,))
        employee = self.cursor.fetchone()
        
        if not employee:
            QMessageBox.warning(self, "Error", "Employee not found in database!")
            return
        
        # Fetch current preferences from transfer_applied
        self.cursor.execute('''
            SELECT pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8,
                   receipt_numbers, applied_date, special_priority
            FROM transfer_applied WHERE pen = ?
        ''', (pen,))
        prefs = self.cursor.fetchone()
        
        # Fetch weightage and weightage_details from jphn table
        self.cursor.execute('SELECT weightage, weightage_details FROM jphn WHERE pen = ?', (pen,))
        weightage_result = self.cursor.fetchone()
        weightage_value = weightage_result[0] if weightage_result else 'No'
        weightage_details_value = weightage_result[1] if weightage_result and len(weightage_result) > 1 else ''
        
        if not prefs:
            QMessageBox.warning(self, "Error", "Application record not found!")
            return
        
        # Create edit dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Edit Application - {employee[0]}")
        dialog.setMinimumWidth(600)
        dialog.setFont(QFont("Calibri", 10))
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(10)
        
        # Employee Info Section (read-only)
        info_frame = QFrame()
        info_frame.setFrameStyle(QFrame.StyledPanel)
        info_frame.setStyleSheet("QFrame { background-color: #e8f4fd; border-radius: 5px; }")
        info_layout = QFormLayout(info_frame)
        
        info_layout.addRow("Name:", QLabel(str(employee[0]) if employee[0] else ""))
        info_layout.addRow("PEN:", QLabel(str(employee[1]) if employee[1] else ""))
        info_layout.addRow("Institution:", QLabel(str(employee[2]) if employee[2] else ""))
        info_layout.addRow("Current District:", QLabel(str(employee[3]) if employee[3] else ""))
        
        layout.addWidget(info_frame)
        
        # Application Details Section
        details_frame = QFrame()
        details_frame.setFrameStyle(QFrame.StyledPanel)
        details_layout = QFormLayout(details_frame)
        
        # Receipt Numbers
        receipt_edit = QLineEdit()
        receipt_edit.setText(str(prefs[8]) if prefs[8] else "")
        receipt_edit.setPlaceholderText("Enter receipt number(s)")
        details_layout.addRow("Receipt No:", receipt_edit)
        
        # Applied Date
        date_edit = NoScrollDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDisplayFormat("dd-MM-yyyy")
        if prefs[9]:
            try:
                date_obj = QDate.fromString(str(prefs[9]), "dd-MM-yyyy")
                if date_obj.isValid():
                    date_edit.setDate(date_obj)
                else:
                    date_edit.setDate(QDate.currentDate())
            except:
                date_edit.setDate(QDate.currentDate())
        else:
            date_edit.setDate(QDate.currentDate())
        details_layout.addRow("Applied Date:", date_edit)
        
        # Special Priority checkbox
        special_priority_check = QCheckBox("Grant Special Priority (Highest Priority)")
        special_priority_check.setStyleSheet("font-weight: bold; color: #c0392b;")
        special_priority_check.setChecked(prefs[10] == 'Yes' if len(prefs) > 10 and prefs[10] else False)
        details_layout.addRow("", special_priority_check)
        
        # Weightage checkbox
        weightage_check = QCheckBox("Has Weightage")
        weightage_check.setChecked(weightage_value == 'Yes' if weightage_value else False)
        details_layout.addRow("", weightage_check)
        
        # Weightage Details
        weightage_details_edit = QTextEdit()
        weightage_details_edit.setMaximumHeight(60)
        weightage_details_edit.setText(str(weightage_details_value) if weightage_details_value else "")
        weightage_details_edit.setPlaceholderText("Enter weightage details (e.g., Spouse working in different district, Medical reasons, etc.)")
        weightage_details_edit.setEnabled(weightage_check.isChecked())
        weightage_check.stateChanged.connect(lambda s: weightage_details_edit.setEnabled(s == Qt.CheckState.Checked.value))
        details_layout.addRow("Weightage Details:", weightage_details_edit)
        
        layout.addWidget(details_frame)
        
        # Preferences Section
        pref_frame = QFrame()
        pref_frame.setFrameStyle(QFrame.StyledPanel)
        pref_frame.setStyleSheet("QFrame { background-color: #f8f9fa; border-radius: 5px; }")
        pref_layout = QVBoxLayout(pref_frame)
        
        pref_title = QLabel("District Preferences (Order of Priority)")
        pref_title.setFont(QFont("Calibri", 10, QFont.Bold))
        pref_layout.addWidget(pref_title)
        
        # Get list of districts for combo boxes
        self.cursor.execute('SELECT DISTINCT district FROM jphn WHERE district IS NOT NULL ORDER BY district')
        districts = [row[0] for row in self.cursor.fetchall()]
        
        # Create 8 preference combo boxes
        pref_combos = []
        pref_grid = QGridLayout()
        
        for i in range(8):
            label = QLabel(f"Preference {i+1}:")
            combo = SearchableComboBox()
            combo.setMinimumWidth(200)
            combo.addItem("Select District")
            combo.addItems(districts)
            
            # Set current value if exists
            if prefs[i]:
                idx = combo.findText(str(prefs[i]))
                if idx >= 0:
                    combo.setCurrentIndex(idx)
            
            pref_combos.append(combo)
            row_num = i // 2
            col_num = (i % 2) * 2
            pref_grid.addWidget(label, row_num, col_num)
            pref_grid.addWidget(combo, row_num, col_num + 1)
        
        pref_layout.addLayout(pref_grid)
        layout.addWidget(pref_frame)
        
        # Buttons
        button_layout = QHBoxLayout()
        save_btn = QPushButton("Save Changes")
        save_btn.setStyleSheet("background-color: #28a745; color: white; padding: 8px 20px;")
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("background-color: #6c757d; color: white; padding: 8px 20px;")
        
        button_layout.addStretch()
        button_layout.addWidget(save_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        def save_changes():
            # Get preference values
            preferences = []
            for combo in pref_combos:
                val = combo.currentText()
                preferences.append(val if val != "Select District" else None)
            
            # Validate - at least one preference required
            if not any(preferences):
                QMessageBox.warning(dialog, "Warning", "Please select at least one district preference!")
                return
            
            try:
                # Update transfer_applied table
                self.cursor.execute('''
                    UPDATE transfer_applied 
                    SET pref1=?, pref2=?, pref3=?, pref4=?, pref5=?, pref6=?, pref7=?, pref8=?,
                        receipt_numbers=?, applied_date=?, special_priority=?
                    WHERE pen=?
                ''', (
                    preferences[0], preferences[1], preferences[2], preferences[3],
                    preferences[4], preferences[5], preferences[6], preferences[7],
                    receipt_edit.text() if receipt_edit.text() else None,
                    date_edit.date().toString("dd-MM-yyyy"),
                    'Yes' if special_priority_check.isChecked() else 'No',
                    pen
                ))
                
                # Update weightage and weightage_details in jphn table
                self.cursor.execute('''
                    UPDATE jphn SET weightage=?, weightage_details=? WHERE pen=?
                ''', (
                    'Yes' if weightage_check.isChecked() else 'No',
                    weightage_details_edit.toPlainText().strip() if weightage_check.isChecked() else '',
                    pen
                ))
                
                self.conn.commit()
                dialog.accept()
                self.load_preference_list()
                QMessageBox.information(self, "Success", "Application updated successfully!")
            except Exception as e:
                QMessageBox.critical(dialog, "Error", f"Failed to update: {str(e)}")
        
        save_btn.clicked.connect(save_changes)
        cancel_btn.clicked.connect(dialog.reject)
        
        dialog.exec()
    
    def auto_fill_vacancies(self):
        """Automatically fill vacancies based on employee preferences, weightage, and seniority"""
        # Check if there are any applied employees with preferences
        self.cursor.execute('''
            SELECT COUNT(*) FROM transfer_applied 
            WHERE pref1 IS NOT NULL OR pref2 IS NOT NULL OR pref3 IS NOT NULL OR pref4 IS NOT NULL
               OR pref5 IS NOT NULL OR pref6 IS NOT NULL OR pref7 IS NOT NULL OR pref8 IS NOT NULL
        ''')
        has_preferences = self.cursor.fetchone()[0] > 0
        
        if not has_preferences:
            QMessageBox.warning(self, "No Preferences", 
                              "No employees have set their district preferences yet.\n\n"
                              "Please ensure employees have filled their preferences before auto-filling.")
            return
        
        # Create custom dialog with options
        dialog = QDialog(self)
        dialog.setWindowTitle("Auto Fill Vacancies")
        dialog.setMinimumWidth(500)
        dialog.setFont(QFont("Calibri", 10))
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        
        # Title
        title = QLabel("Auto Fill Vacancy Options")
        title.setFont(QFont("Calibri", 12, QFont.Bold))
        layout.addWidget(title)
        
        # Info section
        info_frame = QFrame()
        info_frame.setFrameStyle(QFrame.StyledPanel)
        info_layout = QVBoxLayout(info_frame)
        
        info_text = QLabel(
            "<b>This will automatically fill vacancies based on:</b><br><br>"
            "<b style='color: #c0392b;'>PRIORITY 0:</b> Employees with SPECIAL PRIORITY<br>"
            "   → Processed FIRST, get preferences in order<br><br>"
            "<b>PRIORITY 1:</b> Employees with WEIGHTAGE (Yes)<br>"
            "   → Get preferences in order if vacancy available<br><br>"
            "<b>PRIORITY 2:</b> Employees WITHOUT weightage<br>"
            "   → Junior employees first (least duration)<br>"
            "   → Preferences checked in order (Pref 1 to 8)<br><br>"
            "<i>Employees already in Draft List will be skipped.</i>"
        )
        info_text.setWordWrap(True)
        info_layout.addWidget(info_text)
        layout.addWidget(info_frame)
        
        # Options section
        options_group = QGroupBox("Additional Options")
        options_layout = QVBoxLayout(options_group)
        
        # Against transfer checkbox
        self.enable_against_check = QCheckBox("Enable 'Against' Transfers (Create Additional Vacancies)")
        self.enable_against_check.setChecked(False)  # Default OFF
        options_layout.addWidget(self.enable_against_check)
        
        against_info = QLabel(
            "<i>If enabled: When an applicant cannot get ANY of their 8 preferences,<br>"
            "the senior-most employee in their Pref 1 district will be displaced<br>"
            "to a nearby district to create a vacancy.</i>"
        )
        against_info.setStyleSheet("color: #666; margin-left: 20px;")
        options_layout.addWidget(against_info)
        
        layout.addWidget(options_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        proceed_btn = QPushButton("🚀 Proceed with Auto Fill")
        proceed_btn.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold; padding: 8px 15px;")
        proceed_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(proceed_btn)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
        
        if dialog.exec() != QDialog.Accepted:
            return
        
        # Get the option value
        enable_against_transfer = self.enable_against_check.isChecked()
        
        try:
            # Get current vacancy status for all districts
            vacancy_status = {}
            for district in self.districts:
                self.cursor.execute('SELECT vacancy_reported FROM vacancy WHERE district = ?', (district,))
                result = self.cursor.fetchone()
                vacancy_reported = result[0] if result else 0
                
                self.cursor.execute('SELECT COUNT(*) FROM transfer_draft WHERE transfer_to_district = ?', (district,))
                filled_count = self.cursor.fetchone()[0]
                
                vacancy_status[district] = {
                    'reported': vacancy_reported,
                    'filled': filled_count,
                    'remaining': vacancy_reported - filled_count if vacancy_reported > 0 else 999  # 999 = unlimited
                }
            
            # Track allocations
            allocated_count = 0
            special_allocated = 0
            weightage_allocated = 0
            normal_allocated = 0
            against_transfers = 0
            not_allocated = []
            allocation_details = []
            
            # STEP 0: Process employees WITH SPECIAL PRIORITY first (HIGHEST PRIORITY)
            self.cursor.execute('''
                SELECT t.pen, j.name, j.district, j.duration_days, t.special_priority,
                       t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
                FROM transfer_applied t
                INNER JOIN jphn j ON t.pen = j.pen
                WHERE t.special_priority = 'Yes'
                  AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                       OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
                ORDER BY j.duration_days DESC
            ''')
            
            special_priority_employees = self.cursor.fetchall()
            
            for emp in special_priority_employees:
                pen, name, current_district, duration, special_priority, pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8 = emp
                preferences = [pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8]
                
                # Check if already in draft list
                self.cursor.execute('SELECT 1 FROM transfer_draft WHERE pen = ?', (pen,))
                if self.cursor.fetchone():
                    continue  # Skip if already in draft
                
                # Try each preference in order for special priority employees
                allocated = False
                for pref_idx, pref_district in enumerate(preferences):
                    if not pref_district:
                        continue
                    
                    # Check if vacancy available
                    if pref_district in vacancy_status:
                        if vacancy_status[pref_district]['remaining'] > 0 or vacancy_status[pref_district]['reported'] == 0:
                            # Allocate this employee to this district
                            self.cursor.execute('''
                                INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
                                VALUES (?, ?, ?, NULL)
                            ''', (pen, pref_district, datetime.now().strftime("%d-%m-%Y")))
                            
                            # Update vacancy status
                            if vacancy_status[pref_district]['reported'] > 0:
                                vacancy_status[pref_district]['filled'] += 1
                                vacancy_status[pref_district]['remaining'] -= 1
                            
                            allocated = True
                            allocated_count += 1
                            special_allocated += 1
                            allocation_details.append(f"🔴 {name} → {pref_district} (Pref {pref_idx + 1}, SPECIAL PRIORITY)")
                            break
                
                # For special priority, try against transfer if enabled and no preference worked
                if not allocated and pref1 and enable_against_transfer:
                    against_result = self.try_against_transfer(pen, name, pref1, vacancy_status)
                    if against_result:
                        allocated = True
                        allocated_count += 1
                        special_allocated += 1
                        against_transfers += 1
                        allocation_details.append(f"🔴 {name} → {pref1} (Against {against_result}, SPECIAL PRIORITY)")
                
                if not allocated:
                    not_allocated.append(f"{name} (SPECIAL PRIORITY)")
            
            # STEP 1: Process employees WITH WEIGHTAGE (they get priority after special)
            # Order by weightage_priority (1=highest, 5=lowest), then by duration (junior first)
            self.cursor.execute('''
                SELECT t.pen, j.name, j.district, j.duration_days, j.weightage, j.weightage_priority,
                       t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
                FROM transfer_applied t
                INNER JOIN jphn j ON t.pen = j.pen
                WHERE j.weightage = 'Yes'
                  AND (t.special_priority IS NULL OR t.special_priority != 'Yes')
                  AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                       OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
                ORDER BY COALESCE(j.weightage_priority, 5) ASC, j.duration_days ASC
            ''')
            
            weightage_employees = self.cursor.fetchall()
            
            for emp in weightage_employees:
                pen, name, current_district, duration, weightage, weightage_priority, pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8 = emp
                preferences = [pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8]
                priority_text = f"P{weightage_priority}" if weightage_priority else "P5"
                
                # Check if already in draft list
                self.cursor.execute('SELECT 1 FROM transfer_draft WHERE pen = ?', (pen,))
                if self.cursor.fetchone():
                    continue  # Skip if already in draft
                
                # Try each preference in order for weightage employees (normal allocation only)
                allocated = False
                for pref_idx, pref_district in enumerate(preferences):
                    if not pref_district:
                        continue
                    
                    # Check if vacancy available
                    if pref_district in vacancy_status:
                        if vacancy_status[pref_district]['remaining'] > 0 or vacancy_status[pref_district]['reported'] == 0:
                            # Allocate this employee to this district
                            self.cursor.execute('''
                                INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
                                VALUES (?, ?, ?, NULL)
                            ''', (pen, pref_district, datetime.now().strftime("%d-%m-%Y")))
                            
                            # Update vacancy status
                            if vacancy_status[pref_district]['reported'] > 0:
                                vacancy_status[pref_district]['filled'] += 1
                                vacancy_status[pref_district]['remaining'] -= 1
                            
                            allocated = True
                            allocated_count += 1
                            weightage_allocated += 1
                            allocation_details.append(f"⭐ {name} → {pref_district} (Pref {pref_idx + 1}, Weightage {priority_text})")
                            break
                
                # ONLY if NO preference could be allocated AND against transfer is enabled, try on Pref 1
                if not allocated and pref1 and enable_against_transfer:
                    against_result = self.try_against_transfer(pen, name, pref1, vacancy_status)
                    if against_result:
                        allocated = True
                        allocated_count += 1
                        weightage_allocated += 1
                        against_transfers += 1
                        allocation_details.append(f"⭐ {name} → {pref1} (Against {against_result}, Weightage {priority_text})")
                
                if not allocated:
                    not_allocated.append(f"{name} (Weightage {priority_text})")
            
            # STEP 2: Process employees WITHOUT WEIGHTAGE and WITHOUT SPECIAL PRIORITY (ordered by junior first)
            self.cursor.execute('''
                SELECT t.pen, j.name, j.district, j.duration_days, j.weightage,
                       t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
                FROM transfer_applied t
                INNER JOIN jphn j ON t.pen = j.pen
                WHERE (j.weightage IS NULL OR j.weightage != 'Yes')
                  AND (t.special_priority IS NULL OR t.special_priority != 'Yes')
                  AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                       OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
                ORDER BY j.duration_days ASC
            ''')
            
            normal_employees = self.cursor.fetchall()
            
            for emp in normal_employees:
                pen, name, current_district, duration, weightage, pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8 = emp
                preferences = [pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8]
                
                # Check if already in draft list
                self.cursor.execute('SELECT 1 FROM transfer_draft WHERE pen = ?', (pen,))
                if self.cursor.fetchone():
                    continue  # Skip if already in draft
                
                # Try each preference in order (normal allocation only)
                allocated = False
                for pref_idx, pref_district in enumerate(preferences):
                    if not pref_district:
                        continue
                    
                    # Check if vacancy available
                    if pref_district in vacancy_status:
                        if vacancy_status[pref_district]['remaining'] > 0 or vacancy_status[pref_district]['reported'] == 0:
                            # Allocate this employee to this district
                            self.cursor.execute('''
                                INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
                                VALUES (?, ?, ?, NULL)
                            ''', (pen, pref_district, datetime.now().strftime("%d-%m-%Y")))
                            
                            # Update vacancy status
                            if vacancy_status[pref_district]['reported'] > 0:
                                vacancy_status[pref_district]['filled'] += 1
                                vacancy_status[pref_district]['remaining'] -= 1
                            
                            allocated = True
                            allocated_count += 1
                            normal_allocated += 1
                            allocation_details.append(f"{name} → {pref_district} (Pref {pref_idx + 1})")
                            break
                
                # ONLY if NO preference could be allocated AND against transfer is enabled, try on Pref 1
                if not allocated and pref1 and enable_against_transfer:
                    against_result = self.try_against_transfer(pen, name, pref1, vacancy_status)
                    if against_result:
                        allocated = True
                        allocated_count += 1
                        normal_allocated += 1
                        against_transfers += 1
                        allocation_details.append(f"{name} → {pref1} (Against {against_result})")
                
                if not allocated:
                    not_allocated.append(name)
            
            self.conn.commit()
            
            # Prepare result message
            msg = f"Auto-Fill Complete!\n\n"
            msg += f"✅ Total Allocated: {allocated_count} employee(s)\n"
            msg += f"   🔴 Special Priority: {special_allocated}\n"
            msg += f"   ⭐ With Weightage: {weightage_allocated}\n"
            msg += f"   📋 Without Weightage: {normal_allocated}\n"
            if enable_against_transfer:
                msg += f"   🔄 Against Transfers: {against_transfers}\n"
            else:
                msg += f"   ℹ️ Against Transfers: Disabled\n"
            
            if not_allocated:
                msg += f"\n❌ Not Allocated: {len(not_allocated)} employee(s)\n"
                msg += f"   (No vacancy in any of their preferred districts)\n"
            
            if allocation_details:
                msg += f"\n📋 Allocation Details:\n"
                for detail in allocation_details[:20]:  # Show first 20
                    msg += f"   • {detail}\n"
                if len(allocation_details) > 20:
                    msg += f"   ... and {len(allocation_details) - 20} more\n"
            
            msg += f"\nPlease check the Draft List tab for details."
            
            QMessageBox.information(self, "Auto Fill Results", msg)
            
            # Refresh the preference list and draft list
            self.load_preference_list()
            self.load_draft_list()
            self.load_vacancy_data()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to auto-fill vacancies: {str(e)}")
    
    def try_against_transfer(self, applicant_pen, applicant_name, pref_district, vacancy_status):
        """Try to create an 'Against' transfer by displacing the most senior employee in the preferred district"""
        # Find the most senior employee (longest duration) currently in the preferred district
        # who is NOT already in the draft list and NOT the applicant
        self.cursor.execute('''
            SELECT j.pen, j.name, j.duration_days
            FROM jphn j
            WHERE j.district = ?
              AND j.pen != ?
              AND j.pen NOT IN (SELECT pen FROM transfer_draft)
              AND j.pen NOT IN (SELECT pen FROM transfer_applied)
            ORDER BY j.duration_days DESC
            LIMIT 1
        ''', (pref_district, applicant_pen))
        
        senior_employee = self.cursor.fetchone()
        
        if not senior_employee:
            return None
        
        senior_pen, senior_name, senior_duration = senior_employee
        
        # Find a nearby district with vacancy for the senior employee
        nearby_districts = self.nearby_districts.get(pref_district, [])
        target_district = None
        
        for nearby in nearby_districts:
            if nearby in vacancy_status:
                if vacancy_status[nearby]['remaining'] > 0 or vacancy_status[nearby]['reported'] == 0:
                    target_district = nearby
                    break
        
        if not target_district:
            # Try any district with vacancy
            for district in self.districts:
                if district != pref_district and district in vacancy_status:
                    if vacancy_status[district]['remaining'] > 0:
                        target_district = district
                        break
        
        if not target_district:
            return None
        
        # Add the senior employee to draft list (being transferred out)
        against_info_for_senior = f"Displaced for {applicant_name} ({applicant_pen})"
        self.cursor.execute('''
            INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
            VALUES (?, ?, ?, ?)
        ''', (senior_pen, target_district, datetime.now().strftime("%d-%m-%Y"), against_info_for_senior))
        
        # Update vacancy status for the nearby district
        if vacancy_status[target_district]['reported'] > 0:
            vacancy_status[target_district]['filled'] += 1
            vacancy_status[target_district]['remaining'] -= 1
        
        # Add the applicant to draft list with against info
        against_info = f"Against {senior_name} ({senior_pen})"
        self.cursor.execute('''
            INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
            VALUES (?, ?, ?, ?)
        ''', (applicant_pen, pref_district, datetime.now().strftime("%d-%m-%Y"), against_info))
        
        # The preferred district doesn't need vacancy update since one person left and one came
        
        return f"{senior_name} ({senior_pen})"
    
    def load_employees_for_application(self):
        """Load employees from selected district for transfer application (excluding already applied)"""
        current_district = self.applied_district_filter.currentText()
        
        if current_district == "Select District":
            self.applied_table.setRowCount(0)
            self.applied_summary_label.setText("Select a district to view employees")
            return
        
        # Disable updates while loading for performance
        self.applied_table.setUpdatesEnabled(False)
        self.applied_table.setSortingEnabled(False)
        self.applied_table.setRowCount(0)
        
        # Reset headers
        self.applied_table.setColumnCount(5)
        self.applied_table.setHorizontalHeaderLabels(["Select", "Name", "PEN", "Institution", "Duration"])
        
        # Get employees from the selected district who have NOT applied yet
        self.cursor.execute('''
            SELECT j.pen, j.name, j.institution, j.duration_days
            FROM jphn j
            WHERE j.district = ?
            AND j.pen NOT IN (SELECT pen FROM transfer_applied)
            ORDER BY j.duration_days DESC
        ''', (current_district,))
        
        records = self.cursor.fetchall()
        
        # Also get count of applied employees for summary
        self.cursor.execute('''
            SELECT COUNT(*) FROM transfer_applied t
            INNER JOIN jphn j ON t.pen = j.pen
            WHERE j.district = ?
        ''', (current_district,))
        applied_count = self.cursor.fetchone()[0]
        
        # Batch insert rows
        self.applied_table.setRowCount(len(records))
        
        for row_idx, record in enumerate(records):
            # Checkbox for selection
            checkbox = QCheckBox()
            checkbox_widget = QWidget()
            checkbox_layout = QHBoxLayout(checkbox_widget)
            checkbox_layout.addWidget(checkbox)
            checkbox_layout.setAlignment(Qt.AlignCenter)
            checkbox_layout.setContentsMargins(0, 0, 0, 0)
            self.applied_table.setCellWidget(row_idx, 0, checkbox_widget)
            
            # Name
            name_item = QTableWidgetItem(str(record[1]) if record[1] else "")
            name_item.setFont(QFont("mandaram.ttf", 10))
            self.applied_table.setItem(row_idx, 1, name_item)
            
            # PEN
            pen_item = QTableWidgetItem(str(record[0]) if record[0] else "")
            self.applied_table.setItem(row_idx, 2, pen_item)
            
            # Institution
            inst_item = QTableWidgetItem(str(record[2]) if record[2] else "")
            inst_item.setFont(QFont("mandaram.ttf", 10))
            self.applied_table.setItem(row_idx, 3, inst_item)
            
            # Duration
            duration = record[3] if record[3] else 0
            duration_item = QTableWidgetItem(self.format_duration(duration))
            duration_item.setTextAlignment(Qt.AlignCenter)
            # Highlight long durations
            if duration > 1825:  # More than 5 years
                duration_item.setBackground(QColor("#ffcccc"))
            elif duration > 1095:  # More than 3 years
                duration_item.setBackground(QColor("#fff3cd"))
            self.applied_table.setItem(row_idx, 4, duration_item)
        
        # Re-enable updates
        self.applied_table.setUpdatesEnabled(True)
        self.applied_table.setSortingEnabled(True)
        
        # Update summary
        summary = f"Not Applied: {len(records)} | Already Applied: {applied_count} | {current_district}"
        self.applied_summary_label.setText(summary)
    
    def load_transfer_applied_list(self):
        """Load all transfer applications"""
        # Disable updates while loading for performance
        self.applied_table.setUpdatesEnabled(False)
        self.applied_table.setSortingEnabled(False)
        self.applied_table.setRowCount(0)
        
        self.applied_district_filter.setCurrentText("Select District")
        
        # Update headers for this view
        self.applied_table.setColumnCount(5)
        self.applied_table.setHorizontalHeaderLabels(["Current District", "Name", "PEN", "Institution", "Duration"])
        
        # Get all applications
        self.cursor.execute('''
            SELECT j.pen, j.name, j.institution, j.district, j.duration_days
            FROM jphn j
            INNER JOIN transfer_applied t ON j.pen = t.pen
            ORDER BY j.district, j.duration_days DESC
        ''')
        
        records = self.cursor.fetchall()
        
        # Batch insert rows
        self.applied_table.setRowCount(len(records))
        
        for row_idx, record in enumerate(records):
            # Current District
            district_item = QTableWidgetItem(str(record[3]) if record[3] else "")
            district_item.setBackground(QColor("#d4edda"))
            self.applied_table.setItem(row_idx, 0, district_item)
            
            # Name
            name_item = QTableWidgetItem(str(record[1]) if record[1] else "")
            name_item.setFont(QFont("mandaram.ttf", 10))
            name_item.setBackground(QColor("#d4edda"))
            self.applied_table.setItem(row_idx, 1, name_item)
            
            # PEN
            pen_item = QTableWidgetItem(str(record[0]) if record[0] else "")
            pen_item.setBackground(QColor("#d4edda"))
            self.applied_table.setItem(row_idx, 2, pen_item)
            
            # Institution
            inst_item = QTableWidgetItem(str(record[2]) if record[2] else "")
            inst_item.setFont(QFont("mandaram.ttf", 10))
            inst_item.setBackground(QColor("#d4edda"))
            self.applied_table.setItem(row_idx, 3, inst_item)
            
            # Duration
            duration = record[4] if record[4] else 0
            duration_item = QTableWidgetItem(self.format_duration(duration))
            duration_item.setTextAlignment(Qt.AlignCenter)
            duration_item.setBackground(QColor("#d4edda"))
            # Highlight long durations
            if duration > 1825:  # More than 5 years
                duration_item.setBackground(QColor("#ffcccc"))
            elif duration > 1095:  # More than 3 years
                duration_item.setBackground(QColor("#fff3cd"))
            self.applied_table.setItem(row_idx, 4, duration_item)
        
        # Re-enable updates
        self.applied_table.setUpdatesEnabled(True)
        self.applied_table.setSortingEnabled(True)
        
        self.applied_summary_label.setText(f"Total Applications: {len(records)}")
    
    def add_new_employee_for_application(self):
        """Add a new employee who is not in the cadre list"""
        # Ask for confirmation first
        reply = QMessageBox.question(self, "Add New Employee",
            "This will add a new employee to the Cadre List.\n"
            "The employee will then be available for marking as applied.\n\n"
            "Do you want to continue?",
            QMessageBox.Yes | QMessageBox.No)
        
        if reply != QMessageBox.Yes:
            return
        
        # Open the same dialog used for adding JPHN records
        dialog = JPHNDialog(self)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.get_data()
            
            try:
                duration = self.calculate_duration(data['district_join_date'])
                
                self.cursor.execute('''
                    INSERT INTO jphn (pen, name, designation, institution, district,
                                     entry_date, retirement_date, district_join_date,
                                     duration_days, institution_join_date, weightage,
                                     weightage_details)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (data['pen'], data['name'], data['designation'], data['institution'],
                      data['district'], data['entry_date'], data['retirement_date'],
                      data['district_join_date'], duration, data['institution_join_date'],
                      data['weightage'], data['weightage_details']))
                
                self.conn.commit()
                
                # Refresh the cadre list
                self.load_data()
                
                # Set the district filter to the new employee's district and refresh application list
                idx = self.applied_district_filter.findText(data['district'])
                if idx >= 0:
                    self.applied_district_filter.setCurrentIndex(idx)
                else:
                    self.load_employees_for_application()
                
                QMessageBox.information(self, "Success", 
                    f"Employee '{data['name']}' (PEN: {data['pen']}) has been added to the Cadre List.\n\n"
                    f"The employee is now available in district '{data['district']}' for marking as applied.")
                
            except sqlite3.IntegrityError:
                QMessageBox.warning(self, "Error", "PEN already exists in database!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to add employee: {str(e)}")
    
    def select_all_for_application(self):
        """Select all checkboxes in the application table"""
        for row in range(self.applied_table.rowCount()):
            widget = self.applied_table.cellWidget(row, 0)
            if widget:
                checkbox = widget.findChild(QCheckBox)
                if checkbox:
                    checkbox.setChecked(True)
    
    def deselect_all_for_application(self):
        """Deselect all checkboxes in the application table"""
        for row in range(self.applied_table.rowCount()):
            widget = self.applied_table.cellWidget(row, 0)
            if widget:
                checkbox = widget.findChild(QCheckBox)
                if checkbox:
                    checkbox.setChecked(False)
    
    def mark_employees_as_applied(self):
        """Mark selected employees as having applied for transfer"""
        # Get selected employees
        selected_employees = []
        for row in range(self.applied_table.rowCount()):
            widget = self.applied_table.cellWidget(row, 0)
            if widget:
                checkbox = widget.findChild(QCheckBox)
                if checkbox and checkbox.isChecked():
                    pen = self.applied_table.item(row, 2).text()
                    name = self.applied_table.item(row, 1).text()
                    selected_employees.append((pen, name))
        
        if not selected_employees:
            QMessageBox.warning(self, "Warning", "Please select at least one employee!")
            return
        
        # Check if any already applied
        already_applied = []
        new_employees = []
        for pen, name in selected_employees:
            self.cursor.execute('SELECT pen FROM transfer_applied WHERE pen = ?', (pen,))
            if self.cursor.fetchone():
                already_applied.append(name)
            else:
                new_employees.append((pen, name))
        
        if not new_employees:
            QMessageBox.information(self, "Info", 
                f"All {len(already_applied)} selected employee(s) have already applied for transfer.")
            return
        
        # Get the current district from filter
        current_district = self.applied_district_filter.currentText()
        if current_district == "Select District":
            current_district = None
        
        # Open dialog to enter application details (fullscreen with preferences)
        employee_names = [name for pen, name in new_employees]
        dialog = ApplicationDetailsDialog(self, len(new_employees), employee_names, self.districts, current_district)
        
        if dialog.exec() == QDialog.Accepted:
            data = dialog.get_data()
            preferences = data.get('preferences', [None] * 8)
            weightage_priority = data.get('weightage_priority', 5)
            special_priority = 'Yes' if data.get('special_priority', False) else 'No'
            special_priority_reason = data.get('special_priority_reason', '')
            
            try:
                added_count = 0
                for pen, name in new_employees:
                    # Insert new application record with preferences and special priority
                    self.cursor.execute('''
                        INSERT INTO transfer_applied (pen, applied_to_district, applied_date, receipt_numbers,
                                                      pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8,
                                                      special_priority, special_priority_reason)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (pen, "", data['application_date'], data['receipt_numbers'],
                          preferences[0], preferences[1], preferences[2], preferences[3],
                          preferences[4], preferences[5], preferences[6], preferences[7],
                          special_priority, special_priority_reason))
                    
                    # Update weightage in jphn table if provided
                    if data['has_weightage']:
                        self.cursor.execute('''
                            UPDATE jphn SET weightage = 'Yes', weightage_details = ?, weightage_priority = ? WHERE pen = ?
                        ''', (data['weightage_details'], weightage_priority, pen))
                    
                    added_count += 1
                
                self.conn.commit()
                self.load_employees_for_application()
                self.load_data()  # Refresh cadre list to update weightage
                
                msg = f"Transfer application recorded!\n\n"
                msg += f"✅ Newly marked: {added_count}\n"
                if data['receipt_numbers']:
                    msg += f"📄 Receipt Numbers: {data['receipt_numbers']}\n"
                msg += f"📅 Application Date: {data['application_date']}"
                if special_priority == 'Yes':
                    msg += f"\n🔴 Special Priority: Yes"
                if data['has_weightage']:
                    msg += f"\n⭐ Weightage: Yes (Priority {weightage_priority})"
                # Show preferences summary
                pref_count = sum(1 for p in preferences if p)
                if pref_count > 0:
                    msg += f"\n🎯 Preferences Set: {pref_count}"
                if already_applied:
                    msg += f"\n\nℹ️ Already applied (skipped): {len(already_applied)}"
                
                QMessageBox.information(self, "Success", msg)
            except Exception as e:
                self.conn.rollback()
                QMessageBox.critical(self, "Error", f"Failed to record application: {str(e)}")
    
    def setup_draft_list_tab(self):
        """Setup the Transfer - Draft List tab"""
        draft_widget = QWidget()
        draft_layout = QVBoxLayout(draft_widget)
        draft_layout.setSpacing(10)
        draft_layout.setContentsMargins(10, 10, 10, 10)
        
        # Filter Section
        filter_layout = QHBoxLayout()
        
        # Current District filter for draft list
        filter_layout.addWidget(QLabel("Current District:"))
        self.draft_current_district_filter = NoScrollComboBox()
        self.draft_current_district_filter.addItem("All Districts")
        self.draft_current_district_filter.addItems(self.districts)
        self.draft_current_district_filter.currentTextChanged.connect(self.load_draft_list)
        filter_layout.addWidget(self.draft_current_district_filter)
        
        # District filter for draft list (Transfer To District)
        filter_layout.addWidget(QLabel("To District:"))
        self.draft_district_filter = NoScrollComboBox()
        self.draft_district_filter.addItem("All Districts")
        self.draft_district_filter.addItems(self.districts)
        self.draft_district_filter.currentTextChanged.connect(self.load_draft_list)
        filter_layout.addWidget(self.draft_district_filter)
        
        # Search input
        filter_layout.addWidget(QLabel("Search:"))
        self.draft_search_input = QLineEdit()
        self.draft_search_input.setPlaceholderText("Search by Name, PEN, Institution...")
        self.draft_search_input.setFont(QFont("Calibri", 10))
        self.draft_search_input.textChanged.connect(self.filter_draft_table)
        filter_layout.addWidget(self.draft_search_input)
        
        filter_layout.addStretch()
        
        # Refresh button
        self.draft_refresh_btn = QPushButton("🔄 Refresh List")
        self.draft_refresh_btn.clicked.connect(self.load_draft_list)
        filter_layout.addWidget(self.draft_refresh_btn)
        
        draft_layout.addLayout(filter_layout)
        
        # Draft List Table - now with 11 columns including Application Status
        self.draft_table = QTableWidget()
        self.draft_table.setColumnCount(11)
        draft_headers = ["Sl.No", "Name", "PEN", "Present Institution", 
                        "Current District", "Transfer To District",
                        "Application Status", "Date of Joining in District", 
                        "Duration", "Weightage", "Weightage Details"]
        self.draft_table.setHorizontalHeaderLabels(draft_headers)
        self.draft_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.draft_table.horizontalHeader().setStretchLastSection(True)
        self.draft_table.setSortingEnabled(True)
        self.draft_table.setAlternatingRowColors(True)
        self.draft_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.draft_table.setSelectionMode(QTableWidget.ExtendedSelection)
        self.draft_table.setFont(QFont("Calibri", 9))
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.draft_table.horizontalHeader().setFont(header_font)
        
        draft_layout.addWidget(self.draft_table)
        
        # Buttons and Summary
        summary_layout = QHBoxLayout()
        
        self.draft_summary_label = QLabel("Total: 0 records")
        self.draft_summary_label.setFont(QFont("Calibri", 10, QFont.Bold))
        summary_layout.addWidget(self.draft_summary_label)
        
        summary_layout.addStretch()
        
        self.select_all_draft_btn = QPushButton("☑️ Select All")
        self.select_all_draft_btn.clicked.connect(self.select_all_draft_records)
        summary_layout.addWidget(self.select_all_draft_btn)
        
        self.reevaluate_draft_btn = QPushButton("🔄 Re-evaluate Priorities")
        self.reevaluate_draft_btn.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold;")
        self.reevaluate_draft_btn.clicked.connect(self.reevaluate_draft_priorities)
        summary_layout.addWidget(self.reevaluate_draft_btn)
        
        self.remove_from_draft_btn = QPushButton("❌ Remove from List")
        self.remove_from_draft_btn.clicked.connect(self.remove_from_draft_list)
        summary_layout.addWidget(self.remove_from_draft_btn)
        
        self.export_draft_btn = QPushButton("📊 Export Draft List")
        self.export_draft_btn.clicked.connect(self.export_draft_list)
        summary_layout.addWidget(self.export_draft_btn)
        
        self.confirm_transfer_btn = QPushButton("✅ Confirm Transfer")
        self.confirm_transfer_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold;")
        self.confirm_transfer_btn.clicked.connect(self.confirm_transfer_list)
        summary_layout.addWidget(self.confirm_transfer_btn)
        
        draft_layout.addLayout(summary_layout)
        
        self.tab_widget.addTab(draft_widget, "📝 Transfer - Draft List")
    
    def setup_transfer_tab(self):
        """Setup the Transfer List tab - shows confirmed transfers"""
        transfer_widget = QWidget()
        transfer_layout = QVBoxLayout(transfer_widget)
        transfer_layout.setSpacing(10)
        transfer_layout.setContentsMargins(10, 10, 10, 10)
        
        # Filter Section
        filter_layout = QHBoxLayout()
        
        # Current District filter for transfer list
        filter_layout.addWidget(QLabel("Current District:"))
        self.transfer_current_district_filter = NoScrollComboBox()
        self.transfer_current_district_filter.addItem("All Districts")
        self.transfer_current_district_filter.addItems(self.districts)
        self.transfer_current_district_filter.currentTextChanged.connect(self.load_transfer_list)
        filter_layout.addWidget(self.transfer_current_district_filter)
        
        # District filter for transfer list (Transfer To District)
        filter_layout.addWidget(QLabel("To District:"))
        self.transfer_district_filter = NoScrollComboBox()
        self.transfer_district_filter.addItem("All Districts")
        self.transfer_district_filter.addItems(self.districts)
        self.transfer_district_filter.currentTextChanged.connect(self.load_transfer_list)
        filter_layout.addWidget(self.transfer_district_filter)
        
        # Search input
        filter_layout.addWidget(QLabel("Search:"))
        self.transfer_search_input = QLineEdit()
        self.transfer_search_input.setPlaceholderText("Search by Name, PEN, Institution...")
        self.transfer_search_input.setFont(QFont("Calibri", 10))
        self.transfer_search_input.textChanged.connect(self.filter_transfer_table)
        filter_layout.addWidget(self.transfer_search_input)
        
        filter_layout.addStretch()
        
        # Refresh button
        self.transfer_refresh_btn = QPushButton("🔄 Refresh List")
        self.transfer_refresh_btn.clicked.connect(self.load_transfer_list)
        filter_layout.addWidget(self.transfer_refresh_btn)
        
        transfer_layout.addLayout(filter_layout)
        
        # Transfer List Table - shows confirmed transfers
        self.transfer_table = QTableWidget()
        self.transfer_table.setColumnCount(10)
        transfer_headers = ["Sl.No", "Name", "PEN", "Present Institution", 
                           "Current District", "Transfer To District",
                           "Date of Joining in District", 
                           "Duration", "Weightage", "Weightage Details"]
        self.transfer_table.setHorizontalHeaderLabels(transfer_headers)
        self.transfer_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.transfer_table.horizontalHeader().setStretchLastSection(True)
        self.transfer_table.setSortingEnabled(True)
        self.transfer_table.setAlternatingRowColors(True)
        self.transfer_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.transfer_table.setSelectionMode(QTableWidget.ExtendedSelection)  # Enable multi-row selection
        self.transfer_table.setFont(QFont("Calibri", 9))
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.transfer_table.horizontalHeader().setFont(header_font)
        
        transfer_layout.addWidget(self.transfer_table)
        
        # Button Section for Transfer List
        transfer_button_layout = QHBoxLayout()
        
        self.transfer_select_all_btn = QPushButton("☑️ Select All")
        self.transfer_select_all_btn.clicked.connect(self.select_all_transfer_records)
        transfer_button_layout.addWidget(self.transfer_select_all_btn)
        
        self.transfer_delete_btn = QPushButton("🗑️ Delete Selected")
        self.transfer_delete_btn.clicked.connect(self.delete_transfer_records)
        transfer_button_layout.addWidget(self.transfer_delete_btn)
        
        transfer_button_layout.addStretch()
        
        transfer_layout.addLayout(transfer_button_layout)
        
        # Export/Preview Section
        export_layout = QHBoxLayout()
        
        self.transfer_summary_label = QLabel("Total: 0 records")
        self.transfer_summary_label.setFont(QFont("Calibri", 10, QFont.Bold))
        export_layout.addWidget(self.transfer_summary_label)
        
        export_layout.addStretch()
        
        # Preview button
        self.transfer_preview_btn = QPushButton("👁️ Preview")
        self.transfer_preview_btn.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold;")
        self.transfer_preview_btn.clicked.connect(self.preview_transfer_list)
        export_layout.addWidget(self.transfer_preview_btn)
        
        # Print button
        self.transfer_print_btn = QPushButton("🖨️ Print")
        self.transfer_print_btn.clicked.connect(self.print_transfer_list_direct)
        export_layout.addWidget(self.transfer_print_btn)
        
        # Export to Word button
        self.transfer_export_word_btn = QPushButton("📝 Word")
        self.transfer_export_word_btn.clicked.connect(self.export_transfer_list_word)
        export_layout.addWidget(self.transfer_export_word_btn)
        
        # Export to Excel button  
        self.transfer_export_excel_btn = QPushButton("📊 Excel")
        self.transfer_export_excel_btn.clicked.connect(self.export_transfer_list_excel)
        export_layout.addWidget(self.transfer_export_excel_btn)
        
        # Export to PDF button
        self.transfer_export_pdf_btn = QPushButton("📄 PDF")
        self.transfer_export_pdf_btn.clicked.connect(self.export_transfer_list_pdf)
        export_layout.addWidget(self.transfer_export_pdf_btn)
        
        # Export to HTML button
        self.transfer_export_html_btn = QPushButton("🌐 HTML")
        self.transfer_export_html_btn.clicked.connect(self.export_transfer_list_html)
        export_layout.addWidget(self.transfer_export_html_btn)
        
        transfer_layout.addLayout(export_layout)
        
        self.tab_widget.addTab(transfer_widget, "📝 Transfer - Final List")
    
    def setup_vacancy_tab(self):
        """Setup the Vacancy Position tab"""
        vacancy_widget = QWidget()
        vacancy_layout = QVBoxLayout(vacancy_widget)
        vacancy_layout.setSpacing(10)
        vacancy_layout.setContentsMargins(10, 10, 10, 10)
        
        # Title
        title_label = QLabel("Vacancy Position - JPHN Gr I")
        title_label.setFont(QFont("Calibri", 14, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        vacancy_layout.addWidget(title_label)
        
        # Search Section
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search:"))
        self.vacancy_search_input = QLineEdit()
        self.vacancy_search_input.setPlaceholderText("Search by District...")
        self.vacancy_search_input.setFont(QFont("Calibri", 10))
        self.vacancy_search_input.textChanged.connect(self.filter_vacancy_data)
        search_layout.addWidget(self.vacancy_search_input)
        search_layout.addStretch()
        vacancy_layout.addLayout(search_layout)
        
        # Vacancy Table
        self.vacancy_table = QTableWidget()
        self.vacancy_table.setColumnCount(6)
        vacancy_headers = ["District", "Total Strength", "Vacancy Reported", "Additional Vacancy Created", "Filled Now", "Remaining Vacancy"]
        self.vacancy_table.setHorizontalHeaderLabels(vacancy_headers)
        self.vacancy_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.vacancy_table.setAlternatingRowColors(True)
        self.vacancy_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.vacancy_table.setFont(QFont("Calibri", 10))
        self.vacancy_table.verticalHeader().setVisible(False)
        
        # Set header font
        header_font = QFont("Calibri", 10, QFont.Bold)
        self.vacancy_table.horizontalHeader().setFont(header_font)
        
        vacancy_layout.addWidget(self.vacancy_table)
        
        # Buttons Section
        button_layout = QHBoxLayout()
        
        self.save_vacancy_btn = QPushButton("💾 Save Changes")
        self.save_vacancy_btn.setStyleSheet("background-color: #28a745; color: white; font-weight: bold;")
        self.save_vacancy_btn.clicked.connect(self.save_vacancy_data)
        button_layout.addWidget(self.save_vacancy_btn)
        
        self.refresh_vacancy_btn = QPushButton("🔄 Refresh")
        self.refresh_vacancy_btn.clicked.connect(self.load_vacancy_data)
        button_layout.addWidget(self.refresh_vacancy_btn)
        
        button_layout.addStretch()
        
        # Summary
        self.vacancy_summary_label = QLabel("Total Vacancy: 0 | Total Filled: 0 | Remaining: 0")
        self.vacancy_summary_label.setFont(QFont("Calibri", 10, QFont.Bold))
        button_layout.addWidget(self.vacancy_summary_label)
        
        vacancy_layout.addLayout(button_layout)
        
        # Instructions
        instructions = QLabel("⚠️ Enter Total Strength and Vacancy Reported for districts and SAVE before proceeding to Draft List or Transfer List.")
        instructions.setFont(QFont("Calibri", 9, QFont.Bold))
        instructions.setStyleSheet("color: #dc3545;")
        vacancy_layout.addWidget(instructions)
        
        self.tab_widget.addTab(vacancy_widget, "📊 Vacancy Position")
    
    def load_vacancy_data(self):
        """Load vacancy data into the table"""
        self.vacancy_table.setUpdatesEnabled(False)
        self.vacancy_table.setRowCount(0)
        
        total_vacancy = 0
        total_filled = 0
        
        for row_idx, district in enumerate(self.districts):
            self.vacancy_table.insertRow(row_idx)
            
            # Get vacancy data from database
            self.cursor.execute('SELECT total_strength, vacancy_reported FROM vacancy WHERE district = ?', (district,))
            result = self.cursor.fetchone()
            total_strength = result[0] if result else 0
            vacancy_reported = result[1] if result else 0
            
            # Get filled count (employees added to draft list for this district)
            self.cursor.execute('SELECT COUNT(*) FROM transfer_draft WHERE transfer_to_district = ?', (district,))
            filled_count = self.cursor.fetchone()[0]
            
            # Get additional vacancies created (employees displaced FROM this district via against transfer)
            # These are employees whose against_info starts with "Displaced for" and they were moved FROM this district
            self.cursor.execute('''
                SELECT COUNT(*) FROM transfer_draft d
                INNER JOIN jphn j ON d.pen = j.pen
                WHERE j.district = ? AND d.against_info LIKE 'Displaced for%'
            ''', (district,))
            additional_vacancy = self.cursor.fetchone()[0]
            
            # Calculate remaining vacancy (including additional vacancies created)
            total_available = vacancy_reported + additional_vacancy
            remaining = total_available - filled_count
            if remaining < 0:
                remaining = 0
            
            total_vacancy += vacancy_reported
            total_filled += filled_count
            
            # District (read-only)
            district_item = QTableWidgetItem(district)
            district_item.setFlags(district_item.flags() & ~Qt.ItemIsEditable)
            district_item.setBackground(QColor("#f0f0f0"))
            self.vacancy_table.setItem(row_idx, 0, district_item)
            
            # Total Strength (editable)
            strength_item = QTableWidgetItem(str(total_strength))
            strength_item.setTextAlignment(Qt.AlignCenter)
            self.vacancy_table.setItem(row_idx, 1, strength_item)
            
            # Vacancy Reported (editable)
            vacancy_item = QTableWidgetItem(str(vacancy_reported))
            vacancy_item.setTextAlignment(Qt.AlignCenter)
            self.vacancy_table.setItem(row_idx, 2, vacancy_item)
            
            # Additional Vacancy Created (read-only) - from forced transfers
            additional_item = QTableWidgetItem(str(additional_vacancy))
            additional_item.setTextAlignment(Qt.AlignCenter)
            additional_item.setFlags(additional_item.flags() & ~Qt.ItemIsEditable)
            if additional_vacancy > 0:
                additional_item.setBackground(QColor("#ffcccb"))  # Light red - forced vacancy
            self.vacancy_table.setItem(row_idx, 3, additional_item)
            
            # Filled Now (read-only)
            filled_item = QTableWidgetItem(str(filled_count))
            filled_item.setTextAlignment(Qt.AlignCenter)
            filled_item.setFlags(filled_item.flags() & ~Qt.ItemIsEditable)
            filled_item.setBackground(QColor("#cce5ff"))
            self.vacancy_table.setItem(row_idx, 4, filled_item)
            
            # Remaining Vacancy (read-only)
            remaining_item = QTableWidgetItem(str(remaining))
            remaining_item.setTextAlignment(Qt.AlignCenter)
            remaining_item.setFlags(remaining_item.flags() & ~Qt.ItemIsEditable)
            if remaining == 0 and total_available > 0:
                remaining_item.setBackground(QColor("#d4edda"))  # Green - fully filled
            elif remaining > 0:
                remaining_item.setBackground(QColor("#fff3cd"))  # Yellow - has vacancy
            self.vacancy_table.setItem(row_idx, 5, remaining_item)
        
        # Re-enable updates and force refresh
        self.vacancy_table.setUpdatesEnabled(True)
        self.vacancy_table.viewport().update()
        
        # Update summary
        total_remaining = total_vacancy - total_filled
        if total_remaining < 0:
            total_remaining = 0
        self.vacancy_summary_label.setText(f"Total Vacancy: {total_vacancy} | Total Filled: {total_filled} | Remaining: {total_remaining}")
    
    def save_vacancy_data(self):
        """Save vacancy data from table to database"""
        try:
            for row in range(self.vacancy_table.rowCount()):
                district = self.vacancy_table.item(row, 0).text()
                
                # Get values from editable cells
                strength_text = self.vacancy_table.item(row, 1).text()
                vacancy_text = self.vacancy_table.item(row, 2).text()
                
                # Validate and convert to integers
                try:
                    total_strength = int(strength_text) if strength_text else 0
                    vacancy_reported = int(vacancy_text) if vacancy_text else 0
                except ValueError:
                    QMessageBox.warning(self, "Invalid Input", f"Please enter valid numbers for {district}")
                    return
                
                # Update database
                self.cursor.execute('''
                    UPDATE vacancy SET total_strength = ?, vacancy_reported = ? WHERE district = ?
                ''', (total_strength, vacancy_reported, district))
            
            self.conn.commit()
            self.load_vacancy_data()  # Refresh to show updated calculations
            QMessageBox.information(self, "Success", "Vacancy data saved successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save vacancy data: {str(e)}")
    
    def check_vacancy_available(self, district):
        """Check if vacancy is available for a district. Returns (is_available, remaining_count)"""
        # Get vacancy reported for the district
        self.cursor.execute('SELECT vacancy_reported FROM vacancy WHERE district = ?', (district,))
        result = self.cursor.fetchone()
        vacancy_reported = result[0] if result else 0
        
        # If no vacancy is set (0), allow unlimited additions
        if vacancy_reported == 0:
            return True, -1  # -1 means unlimited
        
        # Get current filled count
        self.cursor.execute('SELECT COUNT(*) FROM transfer_draft WHERE transfer_to_district = ?', (district,))
        filled_count = self.cursor.fetchone()[0]
        
        remaining = vacancy_reported - filled_count
        return remaining > 0, remaining
    
    def is_vacancy_data_filled(self):
        """Check if vacancy data has been filled for all districts"""
        self.cursor.execute('SELECT COUNT(*) FROM vacancy WHERE vacancy_reported > 0')
        count = self.cursor.fetchone()[0]
        return count > 0  # At least one district should have vacancy data
    
    def on_tab_changed(self, index):
        """Handle tab change events"""
        # Tab order: 0=Transfer Applied, 1=Vacancy, 2=Draft List, 3=Transfer List, 4=Cadre
        
        # Vacancy data is only required for Draft List (index 2) and Transfer List (index 3)
        # Transfer Application (index 0) and Cadre (index 4) can work without vacancy data
        if index in [2, 3] and not self.is_vacancy_data_filled():
            QMessageBox.warning(self, "Vacancy Data Required", 
                              "Please fill the Vacancy Position data first before accessing Draft List or Transfer List.\n\n"
                              "Enter Total Strength and Vacancy Reported for at least one district and save.")
            # Switch back to Vacancy tab (index 1)
            self.tab_widget.blockSignals(True)
            self.tab_widget.setCurrentIndex(1)
            self.tab_widget.blockSignals(False)
            return
        
        # Tab order: 0=Transfer Applied, 1=Vacancy, 2=Draft List, 3=Transfer List, 4=Cadre
        if index == 0:  # Transfer Applied tab
            self.load_preference_list()  # Load Applied Employees List
        elif index == 1:  # Vacancy Position tab
            self.load_vacancy_data()
        elif index == 2:  # Draft List tab
            self.load_draft_list()
        elif index == 3:  # Transfer List tab
            self.load_transfer_list()
    
    def go_back_to_selection(self):
        """Go back to transfer type selection screen"""
        reply = QMessageBox.question(self, "Go Back", 
            "Are you sure you want to go back to transfer selection?\n\nAny unsaved changes will be lost.",
            QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            # Close database connection
            if hasattr(self, 'conn') and self.conn:
                self.conn.close()
            # Set a flag to indicate restart is needed
            self.restart_requested = True
            self.close()
    
    def load_transfer_list(self):
        """Load confirmed transfer list from transfer_final table"""
        self.transfer_table.setSortingEnabled(False)
        self.transfer_table.setRowCount(0)
        
        current_district = self.transfer_current_district_filter.currentText()
        to_district = self.transfer_district_filter.currentText()
        
        # Build query - join jphn with transfer_final, filter by both Current and To District
        where_clauses = []
        params = []
        
        if current_district != "All Districts":
            where_clauses.append("j.district = ?")
            params.append(current_district)
        
        if to_district != "All Districts":
            where_clauses.append("t.transfer_to_district = ?")
            params.append(to_district)
        
        where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
        
        query = f'''SELECT j.pen, j.name, j.institution, j.district, 
                          t.transfer_to_district, j.district_join_date, 
                          j.duration_days, j.weightage, j.weightage_details 
                   FROM jphn j
                   INNER JOIN transfer_final t ON j.pen = t.pen
                   WHERE {where_sql}
                   ORDER BY t.transfer_to_district, j.duration_days DESC'''
        self.cursor.execute(query, params)
        
        records = self.cursor.fetchall()
        
        for row_idx, record in enumerate(records):
            self.transfer_table.insertRow(row_idx)
            
            # Serial Number
            sl_item = QTableWidgetItem(str(row_idx + 1))
            sl_item.setTextAlignment(Qt.AlignCenter)
            self.transfer_table.setItem(row_idx, 0, sl_item)
            
            # Name
            name_item = QTableWidgetItem(str(record[1]) if record[1] else "")
            name_item.setFont(QFont("mandaram.ttf", 9))
            self.transfer_table.setItem(row_idx, 1, name_item)
            
            # PEN
            self.transfer_table.setItem(row_idx, 2, QTableWidgetItem(str(record[0]) if record[0] else ""))
            
            # Institution
            inst_item = QTableWidgetItem(str(record[2]) if record[2] else "")
            inst_item.setFont(QFont("mandaram.ttf", 9))
            self.transfer_table.setItem(row_idx, 3, inst_item)
            
            # Current District
            self.transfer_table.setItem(row_idx, 4, QTableWidgetItem(str(record[3]) if record[3] else ""))
            
            # Transfer To District - highlight in green
            transfer_to_item = QTableWidgetItem(str(record[4]) if record[4] else "")
            transfer_to_item.setBackground(QColor("#d4edda"))
            transfer_to_item.setFont(QFont("Calibri", 9, QFont.Bold))
            self.transfer_table.setItem(row_idx, 5, transfer_to_item)
            
            # Date of Joining in District
            self.transfer_table.setItem(row_idx, 6, QTableWidgetItem(str(record[5]) if record[5] else ""))
            
            # Duration - format as Years, Months, Days
            duration = record[6] if record[6] else 0
            duration_item = QTableWidgetItem(self.format_duration(duration))
            duration_item.setTextAlignment(Qt.AlignCenter)
            # Highlight long durations
            if duration > 1825:  # More than 5 years
                duration_item.setBackground(QColor("#ffcccc"))
            elif duration > 1095:  # More than 3 years
                duration_item.setBackground(QColor("#fff3cd"))
            self.transfer_table.setItem(row_idx, 7, duration_item)
            
            # Weightage
            weightage_val = record[7]
            weightage_display = "Yes" if weightage_val and str(weightage_val).lower() in ["yes", "y", "1", "true"] else "No"
            weightage_item = QTableWidgetItem(weightage_display)
            weightage_item.setTextAlignment(Qt.AlignCenter)
            if weightage_display == "Yes":
                weightage_item.setBackground(QColor("#d4edda"))
            self.transfer_table.setItem(row_idx, 8, weightage_item)
            
            # Weightage Details
            details_item = QTableWidgetItem(str(record[8]) if record[8] else "")
            details_item.setFont(QFont("mandaram.ttf", 9))
            self.transfer_table.setItem(row_idx, 9, details_item)
        
        self.transfer_table.setSortingEnabled(True)
        
        # Update summary with both filter values
        filter_info = []
        if current_district != "All Districts":
            filter_info.append(f"From: {current_district}")
        if to_district != "All Districts":
            filter_info.append(f"To: {to_district}")
        filter_str = " | ".join(filter_info) if filter_info else "All Districts"
        self.transfer_summary_label.setText(f"Total: {len(records)} confirmed transfers | {filter_str}")
        self.statusBar().showMessage(f"Transfer list loaded: {len(records)} records")
    
    def export_transfer_list(self):
        """Export transfer list to CSV"""
        district = self.transfer_district_filter.currentText()
        default_name = f"transfer_list_{district.lower().replace(' ', '_')}.csv"
        
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Transfer List", 
                                                    default_name, "CSV Files (*.csv)")
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8-sig') as file:
                    writer = csv.writer(file)
                    
                    # Write headers
                    headers = []
                    for col in range(self.transfer_table.columnCount()):
                        headers.append(self.transfer_table.horizontalHeaderItem(col).text())
                    writer.writerow(headers)
                    
                    # Write data
                    for row in range(self.transfer_table.rowCount()):
                        row_data = []
                        for col in range(self.transfer_table.columnCount()):
                            item = self.transfer_table.item(row, col)
                            row_data.append(item.text() if item else "")
                        writer.writerow(row_data)
                
                QMessageBox.information(self, "Success", f"Transfer list exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export: {str(e)}")
    
    def _get_transfer_data_for_export(self):
        """Get transfer data from confirmed transfers for export/preview"""
        self.cursor.execute('''
            SELECT f.pen, j.name, j.designation, j.institution, j.district,
                   f.transfer_to_district, j.weightage, j.weightage_details
            FROM transfer_final f
            INNER JOIN jphn j ON f.pen = j.pen
            ORDER BY j.district, j.name
        ''')
        
        records = self.cursor.fetchall()
        
        transfer_data = []
        for row in records:
            transfer_data.append({
                'pen': row[0],
                'name': row[1],
                'designation': row[2] or 'Junior PH Nurse Gr I',
                'institution': row[3],
                'from_district': row[4],
                'to_district': row[5],
                'weightage': row[6],
                'weightage_details': row[7]
            })
        
        return transfer_data
    
    def preview_transfer_list(self):
        """Preview the confirmed transfer list"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to preview!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Show preview dialog
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.exec()
    
    def print_transfer_list_direct(self):
        """Print the transfer list directly"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to print!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Create preview dialog and call print directly
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.print_document()
    
    def export_transfer_list_word(self):
        """Export the confirmed transfer list to Word"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to export!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Create preview dialog and call export
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.export_to_word()
    
    def export_transfer_list_excel(self):
        """Export the confirmed transfer list to Excel"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to export!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Create preview dialog and call export
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.export_to_excel()
    
    def export_transfer_list_pdf(self):
        """Export the confirmed transfer list to PDF"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to export!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Create preview dialog and call export
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.export_to_pdf()
    
    def export_transfer_list_html(self):
        """Export the confirmed transfer list to HTML"""
        transfer_data = self._get_transfer_data_for_export()
        
        if not transfer_data:
            QMessageBox.warning(self, "Warning", "No confirmed transfers to export!")
            return
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Create preview dialog and call export
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        preview_dialog.export_to_html()

    def load_draft_list(self):
        """Load draft transfer list - only employees added to transfer list"""
        self.draft_table.setSortingEnabled(False)
        self.draft_table.setRowCount(0)
        
        current_district = self.draft_current_district_filter.currentText()
        to_district = self.draft_district_filter.currentText()
        
        # Build query - join jphn with transfer_draft, check if employee EXISTS in transfer_applied
        # Apply both current district and to district filters
        where_clauses = []
        params = []
        
        if current_district != "All Districts":
            where_clauses.append("j.district = ?")
            params.append(current_district)
        
        if to_district != "All Districts":
            where_clauses.append("t.transfer_to_district = ?")
            params.append(to_district)
        
        where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
        
        query = f'''SELECT j.pen, j.name, j.institution, j.district, 
                          t.transfer_to_district, j.district_join_date, 
                          j.duration_days, j.weightage, j.weightage_details,
                          (SELECT 1 FROM transfer_applied WHERE pen = j.pen) as has_applied,
                          t.against_info,
                          (SELECT special_priority FROM transfer_applied WHERE pen = j.pen) as special_priority
                   FROM jphn j
                   INNER JOIN transfer_draft t ON j.pen = t.pen
                   WHERE {where_sql}
                   ORDER BY t.transfer_to_district, j.duration_days DESC'''
        self.cursor.execute(query, params)
        
        records = self.cursor.fetchall()
        
        for row_idx, record in enumerate(records):
            self.draft_table.insertRow(row_idx)
            
            # Serial Number
            sl_item = QTableWidgetItem(str(row_idx + 1))
            sl_item.setTextAlignment(Qt.AlignCenter)
            self.draft_table.setItem(row_idx, 0, sl_item)
            
            # Name
            name_item = QTableWidgetItem(str(record[1]) if record[1] else "")
            name_item.setFont(QFont("mandaram.ttf", 9))
            self.draft_table.setItem(row_idx, 1, name_item)
            
            # PEN
            self.draft_table.setItem(row_idx, 2, QTableWidgetItem(str(record[0]) if record[0] else ""))
            
            # Institution
            inst_item = QTableWidgetItem(str(record[2]) if record[2] else "")
            inst_item.setFont(QFont("mandaram.ttf", 9))
            self.draft_table.setItem(row_idx, 3, inst_item)
            
            # Current District
            self.draft_table.setItem(row_idx, 4, QTableWidgetItem(str(record[3]) if record[3] else ""))
            
            # Transfer To District - highlight in blue
            transfer_to_item = QTableWidgetItem(str(record[4]) if record[4] else "")
            transfer_to_item.setBackground(QColor("#cce5ff"))
            transfer_to_item.setFont(QFont("Calibri", 9, QFont.Bold))
            self.draft_table.setItem(row_idx, 5, transfer_to_item)
            
            # Application Status - check if applied and show against_info if present
            has_applied = record[9] if len(record) > 9 and record[9] else None
            against_info = record[10] if len(record) > 10 and record[10] else None
            special_priority = record[11] if len(record) > 11 and record[11] else None
            is_special = special_priority and str(special_priority).lower() == 'yes'
            
            if against_info:
                # Show Against info with special formatting
                status_text = against_info
                status_item = QTableWidgetItem(status_text)
                status_item.setBackground(QColor("#fff3cd"))  # Yellow for against transfer
                status_item.setFont(QFont("Calibri", 9, QFont.Bold))
            elif has_applied:
                if is_special:
                    status_text = "Special Priority"
                    status_item = QTableWidgetItem(status_text)
                    status_item.setForeground(QColor("#c0392b"))  # Red text
                    status_item.setBackground(QColor("#fadbd8"))  # Light red background
                    status_item.setFont(QFont("Calibri", 9, QFont.Bold))
                else:
                    status_text = "Applied"
                    status_item = QTableWidgetItem(status_text)
                    status_item.setBackground(QColor("#d4edda"))  # Green for applied
                    status_item.setFont(QFont("Calibri", 9, QFont.Bold))
            else:
                status_text = "Not Applied"
                status_item = QTableWidgetItem(status_text)
                status_item.setBackground(QColor("#f8d7da"))  # Light red for not applied
            status_item.setTextAlignment(Qt.AlignCenter)
            self.draft_table.setItem(row_idx, 6, status_item)
            
            # Date of Joining in District
            self.draft_table.setItem(row_idx, 7, QTableWidgetItem(str(record[5]) if record[5] else ""))
            
            # Duration - format as Years, Months, Days
            duration = record[6] if record[6] else 0
            duration_item = QTableWidgetItem(self.format_duration(duration))
            duration_item.setTextAlignment(Qt.AlignCenter)
            # Highlight long durations
            if duration > 1825:  # More than 5 years
                duration_item.setBackground(QColor("#ffcccc"))
            elif duration > 1095:  # More than 3 years
                duration_item.setBackground(QColor("#fff3cd"))
            self.draft_table.setItem(row_idx, 8, duration_item)
            
            # Weightage
            weightage_val = record[7]
            weightage_display = "Yes" if weightage_val and str(weightage_val).lower() in ["yes", "y", "1", "true"] else "No"
            weightage_item = QTableWidgetItem(weightage_display)
            weightage_item.setTextAlignment(Qt.AlignCenter)
            if weightage_display == "Yes":
                weightage_item.setBackground(QColor("#d4edda"))
            self.draft_table.setItem(row_idx, 9, weightage_item)
            
            # Weightage Details
            details_item = QTableWidgetItem(str(record[8]) if record[8] else "")
            details_item.setFont(QFont("mandaram.ttf", 9))
            self.draft_table.setItem(row_idx, 10, details_item)
        
        self.draft_table.setSortingEnabled(True)
        
        # Update summary with both filters
        filter_info = []
        if current_district != "All Districts":
            filter_info.append(f"From: {current_district}")
        if to_district != "All Districts":
            filter_info.append(f"To: {to_district}")
        filter_str = " | ".join(filter_info) if filter_info else "All Districts"
        self.draft_summary_label.setText(f"Total: {len(records)} records | {filter_str}")
        self.statusBar().showMessage(f"Draft list loaded: {len(records)} records")
    
    def select_all_draft_records(self):
        """Select all records in the Draft List table"""
        self.draft_table.selectAll()
        selected_count = len(self.draft_table.selectionModel().selectedRows())
        self.statusBar().showMessage(f"Selected {selected_count} records")
    
    def remove_from_draft_list(self):
        """Remove selected records from Draft List (removes from transfer_draft table only)"""
        selected_rows = self.draft_table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select record(s) to remove!")
            return
        
        count = len(selected_rows)
        reply = QMessageBox.question(self, "Confirm Remove", 
                                     f"Are you sure you want to remove {count} selected record(s) from transfer list?",
                                     QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            try:
                removed_count = 0
                removed_districts = set()
                pens_to_delete = []
                
                # First collect all PENs to delete
                for index in selected_rows:
                    row = index.row()
                    pen_item = self.draft_table.item(row, 2)  # PEN is at column 2
                    transfer_to_item = self.draft_table.item(row, 5)  # Transfer To District at column 5
                    
                    if pen_item:
                        pen = pen_item.text()
                        pens_to_delete.append(pen)
                        if transfer_to_item:
                            removed_districts.add(transfer_to_item.text())
                
                # Now delete each PEN
                for pen in pens_to_delete:
                    self.cursor.execute('DELETE FROM transfer_draft WHERE pen = ?', (pen,))
                    removed_count += 1
                
                self.conn.commit()
                
                # Force refresh of both draft list and vacancy data
                self.load_draft_list()
                self.load_vacancy_data()
                
                district_info = ", ".join(removed_districts) if removed_districts else ""
                QMessageBox.information(self, "Success", 
                    f"{removed_count} record(s) removed from transfer list!\n\n"
                    f"PENs deleted: {', '.join(pens_to_delete)}\n"
                    f"Districts affected: {district_info}\n"
                    f"Vacancy position has been updated.")
            except Exception as e:
                self.conn.rollback()
                QMessageBox.critical(self, "Error", f"Failed to remove record(s): {str(e)}")
    
    def reevaluate_draft_priorities(self):
        """Re-evaluate draft list to ensure weightage employees get priority for their preferences"""
        try:
            # Get current vacancy status
            vacancy_status = {}
            for district in self.districts:
                self.cursor.execute('SELECT vacancy_reported FROM vacancy WHERE district = ?', (district,))
                result = self.cursor.fetchone()
                vacancy_reported = result[0] if result else 0
                
                self.cursor.execute('SELECT COUNT(*) FROM transfer_draft WHERE transfer_to_district = ?', (district,))
                filled_count = self.cursor.fetchone()[0]
                
                vacancy_status[district] = {
                    'reported': vacancy_reported,
                    'filled': filled_count,
                    'remaining': vacancy_reported - filled_count if vacancy_reported > 0 else 999
                }
            
            # Get all employees with weightage who are in draft but NOT in their first preference
            self.cursor.execute('''
                SELECT d.pen, j.name, d.transfer_to_district, t.pref1, t.pref2, t.pref3, t.pref4, 
                       t.pref5, t.pref6, t.pref7, t.pref8, j.weightage
                FROM transfer_draft d
                INNER JOIN jphn j ON d.pen = j.pen
                INNER JOIN transfer_applied t ON d.pen = t.pen
                WHERE j.weightage = 'Yes'
                  AND t.pref1 IS NOT NULL
                  AND d.transfer_to_district != t.pref1
            ''')
            
            weightage_not_in_pref1 = self.cursor.fetchall()
            
            if not weightage_not_in_pref1:
                QMessageBox.information(self, "Re-evaluate", 
                    "All employees with weightage already have their best possible preference.\n"
                    "No changes needed.")
                return
            
            swaps_made = 0
            against_made = 0
            swap_details = []
            
            for emp in weightage_not_in_pref1:
                pen, name, current_alloc, pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8, weightage = emp
                preferences = [pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8]
                
                # Find the index of current allocation in preferences
                current_pref_idx = -1
                for idx, pref in enumerate(preferences):
                    if pref == current_alloc:
                        current_pref_idx = idx
                        break
                
                # Try to get a better preference (lower index = higher priority)
                for pref_idx, pref_district in enumerate(preferences):
                    if not pref_district:
                        continue
                    if pref_idx >= current_pref_idx and current_pref_idx != -1:
                        break  # No better preference available
                    
                    # First check if someone WITHOUT weightage is already in draft for this district
                    self.cursor.execute('''
                        SELECT d.pen, j.name, j.weightage
                        FROM transfer_draft d
                        INNER JOIN jphn j ON d.pen = j.pen
                        WHERE d.transfer_to_district = ?
                          AND (j.weightage IS NULL OR j.weightage != 'Yes')
                        ORDER BY j.duration_days ASC
                        LIMIT 1
                    ''', (pref_district,))
                    
                    swap_candidate = self.cursor.fetchone()
                    
                    if swap_candidate:
                        swap_pen, swap_name, swap_weightage = swap_candidate
                        
                        # Check if the person being swapped out can go to another district
                        self.cursor.execute('''
                            SELECT t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
                            FROM transfer_applied t WHERE t.pen = ?
                        ''', (swap_pen,))
                        swap_prefs_row = self.cursor.fetchone()
                        
                        if swap_prefs_row:
                            swap_prefs = list(swap_prefs_row)
                            new_spot_found = False
                            new_district = None
                            
                            # First check if they can go to the weightage employee's old spot
                            if current_alloc in swap_prefs:
                                new_district = current_alloc
                                new_spot_found = True
                            else:
                                # Find any other available district from their preferences
                                for sp in swap_prefs:
                                    if not sp or sp == pref_district:
                                        continue
                                    if sp in vacancy_status:
                                        if vacancy_status[sp]['remaining'] > 0:
                                            new_district = sp
                                            new_spot_found = True
                                            break
                            
                            if new_spot_found and new_district:
                                # Perform the swap with against info
                                self.cursor.execute('''
                                    UPDATE transfer_draft 
                                    SET transfer_to_district = ?, against_info = ? 
                                    WHERE pen = ?
                                ''', (pref_district, f"Against {swap_name} ({swap_pen})", pen))
                                
                                self.cursor.execute('''
                                    UPDATE transfer_draft 
                                    SET transfer_to_district = ?, against_info = ? 
                                    WHERE pen = ?
                                ''', (new_district, f"Displaced for {name} ({pen})", swap_pen))
                                
                                swaps_made += 1
                                swap_details.append(f"⭐ {name} → {pref_district} (Against {swap_name})")
                                swap_details.append(f"   ↔️ {swap_name} → {new_district}")
                                break
                    else:
                        # No one in draft to swap - try "Against" transfer from the district itself
                        # Remove current allocation first
                        self.cursor.execute('DELETE FROM transfer_draft WHERE pen = ?', (pen,))
                        
                        # Try against transfer
                        against_result = self.try_against_transfer(pen, name, pref_district, vacancy_status)
                        if against_result:
                            against_made += 1
                            swap_details.append(f"⭐ {name} → {pref_district} (Against {against_result})")
                            break
                        else:
                            # Restore original allocation if against transfer failed
                            self.cursor.execute('''
                                INSERT INTO transfer_draft (pen, transfer_to_district, added_date, against_info)
                                VALUES (?, ?, ?, NULL)
                            ''', (pen, current_alloc, datetime.now().strftime("%d-%m-%Y")))
            
            self.conn.commit()
            self.load_draft_list()
            
            total_changes = swaps_made + against_made
            if total_changes > 0:
                details_text = "\n".join(swap_details[:20])
                if len(swap_details) > 20:
                    details_text += f"\n... and {len(swap_details) - 20} more"
                QMessageBox.information(self, "Re-evaluate Complete", 
                    f"Re-evaluation complete!\n\n"
                    f"Changes made: {total_changes}\n"
                    f"  - Swaps: {swaps_made}\n"
                    f"  - Against transfers: {against_made}\n\n"
                    f"Details:\n{details_text}")
            else:
                QMessageBox.information(self, "Re-evaluate", 
                    "No changes could be made.\n\n"
                    "This may be because:\n"
                    "- All preferred districts have only weightage employees\n"
                    "- No suitable swap/against partners found")
            
            # Update vacancy data after re-evaluation
            self.load_vacancy_data()
                    
        except Exception as e:
            self.conn.rollback()
            QMessageBox.critical(self, "Error", f"Re-evaluation failed: {str(e)}")
    
    def confirm_transfer_list(self):
        """Confirm all draft transfers and show preview in official format"""
        # Check if there are any records in draft list
        self.cursor.execute('SELECT COUNT(*) FROM transfer_draft')
        count = self.cursor.fetchone()[0]
        
        if count == 0:
            QMessageBox.warning(self, "Warning", "No records in Draft List to confirm!")
            return
        
        # Fetch all draft data with employee details
        self.cursor.execute('''
            SELECT d.pen, j.name, j.designation, j.institution, j.district,
                   d.transfer_to_district, j.weightage, j.weightage_details, d.against_info
            FROM transfer_draft d
            INNER JOIN jphn j ON d.pen = j.pen
            ORDER BY j.district, j.name
        ''')
        
        records = self.cursor.fetchall()
        
        # Prepare data for preview
        transfer_data = []
        for row in records:
            transfer_data.append({
                'pen': row[0],
                'name': row[1],
                'designation': row[2] or 'Junior PH Nurse Gr I',
                'institution': row[3],
                'from_district': row[4],
                'to_district': row[5],
                'weightage': row[6],
                'weightage_details': row[7],
                'against_info': row[8]
            })
        
        # For Regular Transfer, ask for order number
        order_number = ""
        if self.transfer_type == "regular":
            saved_order = self.get_saved_order_number()
            order_number, ok = QInputDialog.getText(
                self, "Order Number", 
                "Enter Order Number (ഉത്തരവ് നം.):",
                text=saved_order
            )
            if not ok:
                return
            self.save_order_number(order_number)
        
        # Show preview dialog
        preview_dialog = TransferListPreviewDialog(self, transfer_data, self.districts, self.transfer_type, order_number)
        
        if preview_dialog.exec() == QDialog.Accepted:
            try:
                # Clear previous final transfers
                self.cursor.execute('DELETE FROM transfer_final')
                
                # Copy all records from draft to final (draft list remains intact)
                self.cursor.execute('''
                    INSERT INTO transfer_final (pen, transfer_to_district, confirmed_date)
                    SELECT pen, transfer_to_district, ? FROM transfer_draft
                ''', (datetime.now().strftime("%d-%m-%Y"),))
                
                self.conn.commit()
                
                # Switch to Transfer List tab
                self.tab_widget.setCurrentIndex(2)
                self.load_transfer_list()
                
                QMessageBox.information(self, "Success", 
                                       f"{count} transfer(s) confirmed successfully!\n\n"
                                       "View the confirmed list in Transfer List tab.\n"
                                       "Draft List has been maintained.")
            except Exception as e:
                self.conn.rollback()
                QMessageBox.critical(self, "Error", f"Failed to confirm transfers: {str(e)}")
    
    def show_cadre_context_menu(self, position):
        """Show context menu on right-click in cadre table"""
        # Get the item at the clicked position
        item = self.table.itemAt(position)
        if item is None:
            return
        
        # Get the row at clicked position
        clicked_row = item.row()
        
        # Check if clicked row is already selected (part of multi-selection)
        selected_rows = self.table.selectionModel().selectedRows()
        clicked_is_selected = any(index.row() == clicked_row for index in selected_rows)
        
        # Only select the single row if it's not part of current selection
        if not clicked_is_selected:
            self.table.selectRow(clicked_row)
            
        menu = QMenu(self)
        
        # Show count if multiple selected
        selected_count = len(self.table.selectionModel().selectedRows())
        if selected_count > 1:
            add_to_transfer_action = menu.addAction(f"📋 Add {selected_count} to Transfer List")
        else:
            add_to_transfer_action = menu.addAction("📋 Add to Transfer List")
        
        # Show menu and get selected action
        global_pos = self.table.viewport().mapToGlobal(position)
        action = menu.exec(global_pos)
        
        if action == add_to_transfer_action:
            self.add_to_transfer_list()
    
    def add_to_transfer_list(self):
        """Add selected employees to transfer list"""
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select record(s) to add!")
            return
        
        # Get selected employees
        employees = []
        for index in selected_rows:
            row = index.row()
            pen = self.table.item(row, 0).text()  # PEN is at column 0
            name = self.table.item(row, 1).text()  # Name is at column 1
            employees.append((pen, name))
        
        # Check for duplicates first
        duplicate_pens = []
        for pen, name in employees:
            self.cursor.execute('SELECT pen FROM transfer_draft WHERE pen = ?', (pen,))
            if self.cursor.fetchone():
                duplicate_pens.append(name)
        
        if duplicate_pens:
            if len(duplicate_pens) == len(employees):
                QMessageBox.warning(self, "Duplicate Entry", 
                                   f"The following employee(s) are already in the transfer list:\n\n" + 
                                   "\n".join(duplicate_pens))
                return
            else:
                # Some are duplicates, ask if user wants to continue with others
                reply = QMessageBox.question(self, "Duplicate Entry", 
                                            f"The following employee(s) are already in the transfer list:\n\n" + 
                                            "\n".join(duplicate_pens) + 
                                            "\n\nDo you want to add the remaining employees?",
                                            QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.No:
                    return
                # Filter out duplicates
                employees = [(pen, name) for pen, name in employees if name not in duplicate_pens]
        
        if not employees:
            return
        
        # If single employee, show dialog directly
        # If multiple, show dialog for each or ask for common district
        if len(employees) == 1:
            pen, name = employees[0]
            dialog = TransferDistrictDialog(self, name)
            if dialog.exec() == QDialog.Accepted:
                transfer_district = dialog.get_district()
                if not transfer_district:
                    QMessageBox.warning(self, "Warning", "Please select a district!")
                    return
                
                # Check vacancy before adding
                is_available, remaining = self.check_vacancy_available(transfer_district)
                if not is_available:
                    QMessageBox.critical(self, "Vacancy Full", 
                                        f"Cannot add to {transfer_district}!\n\n"
                                        "The vacancy for this district is completely filled.\n"
                                        "Please check the Vacancy Position tab.")
                    return
                
                try:
                    self.cursor.execute('''
                        INSERT INTO transfer_draft (pen, transfer_to_district, added_date)
                        VALUES (?, ?, ?)
                    ''', (pen, transfer_district, datetime.now().strftime("%d-%m-%Y")))
                    self.conn.commit()
                    QMessageBox.information(self, "Success", f"{name} added to transfer list!")
                    if hasattr(self, 'draft_table'):
                        self.load_draft_list()
                    self.load_vacancy_data()  # Update vacancy position
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to add to transfer list: {str(e)}")
        else:
            # Multiple employees - ask for common district
            reply = QMessageBox.question(self, "Multiple Selection", 
                                        f"You have selected {len(employees)} employees.\n\n" +
                                        "Do you want to assign the same transfer district for all?",
                                        QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                dialog = TransferDistrictDialog(self, f"{len(employees)} employees")
                if dialog.exec() == QDialog.Accepted:
                    transfer_district = dialog.get_district()
                    if not transfer_district:
                        QMessageBox.warning(self, "Warning", "Please select a district!")
                        return
                    
                    # Check vacancy before adding
                    is_available, remaining = self.check_vacancy_available(transfer_district)
                    if not is_available:
                        QMessageBox.critical(self, "Vacancy Full", 
                                            f"Cannot add to {transfer_district}!\n\n"
                                            "The vacancy for this district is completely filled.\n"
                                            "Please check the Vacancy Position tab.")
                        return
                    
                    # Check if enough vacancy for all employees
                    if remaining != -1 and remaining < len(employees):
                        QMessageBox.critical(self, "Insufficient Vacancy", 
                                            f"Cannot add {len(employees)} employees to {transfer_district}!\n\n"
                                            f"Only {remaining} vacancy position(s) remaining.\n"
                                            "Please reduce the number of selections or choose a different district.")
                        return
                    
                    try:
                        added_count = 0
                        for pen, name in employees:
                            self.cursor.execute('''
                                INSERT INTO transfer_draft (pen, transfer_to_district, added_date)
                                VALUES (?, ?, ?)
                            ''', (pen, transfer_district, datetime.now().strftime("%d-%m-%Y")))
                            added_count += 1
                        
                        self.conn.commit()
                        QMessageBox.information(self, "Success", f"{added_count} employee(s) added to transfer list!")
                        if hasattr(self, 'draft_table'):
                            self.load_draft_list()
                        self.load_vacancy_data()  # Update vacancy position
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Failed to add to transfer list: {str(e)}")
            else:
                # Add each one separately
                added_count = 0
                skipped_districts = []
                for pen, name in employees:
                    dialog = TransferDistrictDialog(self, name)
                    if dialog.exec() == QDialog.Accepted:
                        transfer_district = dialog.get_district()
                        if transfer_district:
                            # Check vacancy before adding
                            is_available, remaining = self.check_vacancy_available(transfer_district)
                            if not is_available:
                                QMessageBox.warning(self, "Vacancy Full", 
                                                   f"Cannot add {name} to {transfer_district}!\n\n"
                                                   "The vacancy for this district is completely filled.")
                                skipped_districts.append(f"{name} -> {transfer_district}")
                                continue
                            
                            try:
                                self.cursor.execute('''
                                    INSERT INTO transfer_draft (pen, transfer_to_district, added_date)
                                    VALUES (?, ?, ?)
                                ''', (pen, transfer_district, datetime.now().strftime("%d-%m-%Y")))
                                added_count += 1
                            except Exception as e:
                                QMessageBox.critical(self, "Error", f"Failed to add {name}: {str(e)}")
                
                if added_count > 0:
                    self.conn.commit()
                    msg = f"{added_count} employee(s) added to transfer list!"
                    if skipped_districts:
                        msg += f"\n\nSkipped due to full vacancy:\n" + "\n".join(skipped_districts)
                    QMessageBox.information(self, "Success", msg)
                    if hasattr(self, 'draft_table'):
                        self.load_draft_list()
                    self.load_vacancy_data()  # Update vacancy position
                elif skipped_districts:
                    QMessageBox.warning(self, "No Additions", 
                                       "No employees were added.\n\nSkipped due to full vacancy:\n" + 
                                       "\n".join(skipped_districts))
    
    def export_draft_list(self):
        """Export draft list to CSV"""
        district = self.draft_district_filter.currentText()
        default_name = f"draft_list_{district.lower().replace(' ', '_')}.csv"
        
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Draft List", 
                                                    default_name, "CSV Files (*.csv)")
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8-sig') as file:
                    writer = csv.writer(file)
                    
                    # Write headers
                    headers = []
                    for col in range(self.draft_table.columnCount()):
                        headers.append(self.draft_table.horizontalHeaderItem(col).text())
                    writer.writerow(headers)
                    
                    # Write data
                    for row in range(self.draft_table.rowCount()):
                        row_data = []
                        for col in range(self.draft_table.columnCount()):
                            item = self.draft_table.item(row, col)
                            row_data.append(item.text() if item else "")
                        writer.writerow(row_data)
                
                QMessageBox.information(self, "Success", f"Draft list exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export: {str(e)}")
    
    def apply_styles(self):
        """Apply modern stylesheet"""
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QLabel {
                color: #2c3e50;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
            QLineEdit, QComboBox, QDateEdit, QSpinBox {
                padding: 6px;
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QLineEdit:focus, QComboBox:focus, QDateEdit:focus, QSpinBox:focus {
                border: 2px solid #3498db;
            }
            QTableWidget {
                background-color: white;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                gridline-color: #ecf0f1;
            }
            QTableWidget::item {
                padding: 5px;
            }
            QTableWidget::item:selected {
                background-color: #3498db;
                color: white;
            }
            QHeaderView::section {
                background-color: #34495e;
                color: white;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QTextEdit {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QTabWidget::pane {
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #ecf0f1;
                color: #2c3e50;
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                font-weight: bold;
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
            }
            QTabBar::tab:hover:!selected {
                background-color: #bdc3c7;
            }
        """)
    
    def calculate_duration(self, join_date_str):
        """Calculate duration in days from join date to today"""
        try:
            join_date = datetime.strptime(join_date_str, "%d-%m-%Y")
            today = datetime.now()
            duration = (today - join_date).days
            return duration
        except:
            return 0
    
    def format_duration(self, days):
        """Format duration in days as Years, Months, Days"""
        if days <= 0:
            return "0D"
        years = days // 365
        remaining = days % 365
        months = remaining // 30
        days_left = remaining % 30
        
        parts = []
        if years > 0:
            parts.append(f"{years}Y")
        if months > 0:
            parts.append(f"{months}M")
        if days_left > 0 or not parts:
            parts.append(f"{days_left}D")
        return " ".join(parts)
    
    def load_data(self):
        """Load data from database into table"""
        self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        
        # Get list of PENs who have applied for transfer
        self.cursor.execute('SELECT pen FROM transfer_applied')
        applied_pens = set(row[0] for row in self.cursor.fetchall())
        
        # Select specific columns in order we need for display
        # Table columns: PEN, Name, Designation, Institution, District, Entry Date, 
        #                Retirement Date, District Join Date, Duration, Contact
        self.cursor.execute('''
            SELECT pen, name, designation, institution, district, entry_date, 
                   retirement_date, district_join_date, duration_days, 
                   COALESCE(contact, '') as contact
            FROM jphn ORDER BY pen
        ''')
        records = self.cursor.fetchall()
        
        # Light red color for applied employees
        light_red = QColor(255, 200, 200)  # Light red/pink background
        
        for row_idx, record in enumerate(records):
            # Update duration (district_join_date is at index 7)
            duration = self.calculate_duration(record[7])
            self.cursor.execute('UPDATE jphn SET duration_days = ? WHERE pen = ?', 
                              (duration, record[0]))
            
            self.table.insertRow(row_idx)
            # Set row number in vertical header (starting from 1)
            self.table.setVerticalHeaderItem(row_idx, QTableWidgetItem(str(row_idx + 1)))
            
            # Check if this employee has applied for transfer
            is_applied = str(record[0]) in applied_pens
            
            # Direct column mapping - SELECT returns exactly what we need for display
            # record: 0-pen, 1-name, 2-designation, 3-institution, 4-district, 5-entry_date,
            #         6-retirement_date, 7-district_join_date, 8-duration_days, 9-contact
            for col_idx, value in enumerate(record):
                if col_idx == 8:  # Duration column - format as Y M D
                    value = self.format_duration(duration)
                item = QTableWidgetItem(str(value) if value else "")
                if col_idx in [1, 3]:  # Name, Institution
                    item.setFont(QFont("mandaram.ttf", 9))
                # Highlight applied employees with light red background
                if is_applied:
                    item.setBackground(light_red)
                self.table.setItem(row_idx, col_idx, item)
        
        self.conn.commit()
        self.table.setSortingEnabled(True)
        self.statusBar().showMessage(f"Loaded {len(records)} records | Applied for transfer: {len(applied_pens)} (shown in light red)")
        
        # Refresh transfer list if it exists
        if hasattr(self, 'transfer_table'):
            self.load_transfer_list()
    
    def filter_data(self):
        """Filter table based on search and district filter"""
        search_text = self.search_input.text().lower()
        district = self.district_filter.currentText()
        
        visible_count = 0
        for row in range(self.table.rowCount()):
            show_row = True
            
            # Check search text
            if search_text:
                row_text = ""
                for col in range(self.table.columnCount()):
                    item = self.table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "
                
                if search_text not in row_text:
                    show_row = False
            
            # Check district filter (District is at column 4)
            if district != "All Districts":
                district_item = self.table.item(row, 4)
                if district_item and district_item.text() != district:
                    show_row = False
            
            self.table.setRowHidden(row, not show_row)
            
            # Update vertical header (row number) for visible rows
            if show_row:
                visible_count += 1
                self.table.setVerticalHeaderItem(row, QTableWidgetItem(str(visible_count)))
    
    def filter_vacancy_data(self):
        """Filter vacancy table based on search text"""
        search_text = self.vacancy_search_input.text().lower()
        
        for row in range(self.vacancy_table.rowCount()):
            show_row = True
            
            if search_text:
                district_item = self.vacancy_table.item(row, 0)
                if district_item and search_text not in district_item.text().lower():
                    show_row = False
            
            self.vacancy_table.setRowHidden(row, not show_row)
    
    def filter_applied_table(self):
        """Filter application selection table based on search text"""
        search_text = self.applied_search_input.text().lower()
        
        # Check if the searched text (if it looks like a PEN) is already applied
        if search_text and len(search_text) >= 3:
            # Check if an employee with this PEN/Name has already applied
            self.cursor.execute('''
                SELECT j.pen, j.name, j.district, t.applied_date, t.receipt_numbers
                FROM transfer_applied t
                INNER JOIN jphn j ON t.pen = j.pen
                WHERE LOWER(j.pen) LIKE ? OR LOWER(j.name) LIKE ?
            ''', (f'%{search_text}%', f'%{search_text}%'))
            
            already_applied = self.cursor.fetchall()
            if already_applied:
                # Update the summary label to show already applied info
                names = [f"{row[1]} (PEN: {row[0]}, District: {row[2]})" for row in already_applied[:3]]
                if len(already_applied) > 3:
                    names.append(f"...and {len(already_applied) - 3} more")
                self.applied_summary_label.setText(f"⚠️ Already Applied: {', '.join(names)}")
                self.applied_summary_label.setStyleSheet("color: #dc3545; font-weight: bold;")
            else:
                # Reset the label style
                district = self.applied_district_filter.currentText()
                if district != "Select District":
                    self.cursor.execute('SELECT COUNT(*) FROM transfer_applied t INNER JOIN jphn j ON t.pen = j.pen WHERE j.district = ?', (district,))
                    applied_count = self.cursor.fetchone()[0]
                    visible_count = sum(1 for row in range(self.applied_table.rowCount()) if not self.applied_table.isRowHidden(row))
                    self.applied_summary_label.setText(f"Not Applied: {visible_count} | Already Applied: {applied_count} | {district}")
                self.applied_summary_label.setStyleSheet("")
        else:
            # Reset the label style when search is cleared
            district = self.applied_district_filter.currentText()
            if district != "Select District":
                self.cursor.execute('SELECT COUNT(*) FROM transfer_applied t INNER JOIN jphn j ON t.pen = j.pen WHERE j.district = ?', (district,))
                applied_count = self.cursor.fetchone()[0]
                self.applied_summary_label.setText(f"Not Applied: {self.applied_table.rowCount()} | Already Applied: {applied_count} | {district}")
            self.applied_summary_label.setStyleSheet("")
        
        for row in range(self.applied_table.rowCount()):
            show_row = True
            
            if search_text:
                row_text = ""
                for col in range(self.applied_table.columnCount()):
                    item = self.applied_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "
                
                if search_text not in row_text:
                    show_row = False
            
            self.applied_table.setRowHidden(row, not show_row)
    
    def filter_preference_table(self):
        """Filter preference selection table based on search text"""
        search_text = self.pref_search_input.text().lower()
        
        for row in range(self.pref_table.rowCount()):
            show_row = True
            
            if search_text:
                row_text = ""
                # Search in Name, PEN, Institution, Current District columns (0-3)
                for col in range(4):
                    item = self.pref_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "
                
                if search_text not in row_text:
                    show_row = False
            
            self.pref_table.setRowHidden(row, not show_row)
    
    def filter_draft_table(self):
        """Filter draft list table based on search text"""
        search_text = self.draft_search_input.text().lower()
        
        for row in range(self.draft_table.rowCount()):
            show_row = True
            
            if search_text:
                row_text = ""
                for col in range(self.draft_table.columnCount()):
                    item = self.draft_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "
                
                if search_text not in row_text:
                    show_row = False
            
            self.draft_table.setRowHidden(row, not show_row)
    
    def filter_transfer_table(self):
        """Filter transfer list table based on search text"""
        search_text = self.transfer_search_input.text().lower()
        
        for row in range(self.transfer_table.rowCount()):
            show_row = True
            
            if search_text:
                row_text = ""
                for col in range(self.transfer_table.columnCount()):
                    item = self.transfer_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "
                
                if search_text not in row_text:
                    show_row = False
            
            self.transfer_table.setRowHidden(row, not show_row)
    
    def add_record(self):
        """Add new JPHN record"""
        dialog = JPHNDialog(self)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.get_data()
            
            try:
                duration = self.calculate_duration(data['district_join_date'])
                
                self.cursor.execute('''
                    INSERT INTO jphn (pen, name, designation, institution, district,
                                     entry_date, retirement_date, district_join_date,
                                     duration_days, institution_join_date, weightage,
                                     weightage_details, contact)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (data['pen'], data['name'], data['designation'], data['institution'],
                      data['district'], data['entry_date'], data['retirement_date'],
                      data['district_join_date'], duration, data['institution_join_date'],
                      data['weightage'], data['weightage_details'], data.get('contact', '')))
                
                self.conn.commit()
                self.load_data()
                QMessageBox.information(self, "Success", "Record added successfully!")
            except sqlite3.IntegrityError:
                QMessageBox.warning(self, "Error", "PEN already exists in database!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to add record: {str(e)}")
    
    def on_table_double_click(self, index):
        """Handle double-click on cadre table to open edit dialog"""
        if index.isValid():
            self.table.selectRow(index.row())
            self.edit_record()
    
    def edit_record(self):
        """Edit selected JPHN record"""
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select a record to edit!")
            return
        
        row = selected_rows[0].row()
        pen_item = self.table.item(row, 0)  # PEN is at column 0
        if not pen_item:
            QMessageBox.warning(self, "Error", "Could not read PEN from selected row!")
            return
        
        pen = pen_item.text()
        
        # Fetch full record with explicit column order
        self.cursor.execute('''
            SELECT pen, name, designation, institution, district, entry_date,
                   retirement_date, district_join_date, duration_days, institution_join_date,
                   weightage, weightage_details, contact
            FROM jphn WHERE pen = ?
        ''', (pen,))
        record = self.cursor.fetchone()
        
        if not record:
            QMessageBox.warning(self, "Error", f"Record with PEN '{pen}' not found in database!\n\nPlease refresh the data and try again.")
            self.load_data()
            return
        
        dialog = JPHNDialog(self, record)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.get_data()
            
            try:
                duration = self.calculate_duration(data['district_join_date'])
                old_pen = pen
                
                # If PEN changed, need to delete old and insert new
                if data['pen'] != old_pen:
                    self.cursor.execute('DELETE FROM jphn WHERE pen = ?', (old_pen,))
                    self.cursor.execute('''
                        INSERT INTO jphn (pen, name, designation, institution, district,
                                         entry_date, retirement_date, district_join_date,
                                         duration_days, institution_join_date, weightage,
                                         weightage_details, contact)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (data['pen'], data['name'], data['designation'], data['institution'],
                          data['district'], data['entry_date'], data['retirement_date'],
                          data['district_join_date'], duration, data['institution_join_date'],
                          data['weightage'], data['weightage_details'], data.get('contact', '')))
                else:
                    self.cursor.execute('''
                        UPDATE jphn SET name=?, designation=?, institution=?,
                                       district=?, entry_date=?, retirement_date=?,
                                       district_join_date=?, duration_days=?,
                                       institution_join_date=?, weightage=?,
                                       weightage_details=?, contact=?
                        WHERE pen=?
                    ''', (data['name'], data['designation'], data['institution'],
                          data['district'], data['entry_date'], data['retirement_date'],
                          data['district_join_date'], duration, data['institution_join_date'],
                          data['weightage'], data['weightage_details'], data.get('contact', ''), pen))
                
                self.conn.commit()
                self.load_data()
                QMessageBox.information(self, "Success", "Record updated successfully!")
            except sqlite3.IntegrityError:
                QMessageBox.warning(self, "Error", "New PEN already exists in database!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to update record: {str(e)}")
    
    def select_all_records(self):
        """Select all visible records in the Cadre table"""
        self.table.selectAll()
        selected_count = len(self.table.selectionModel().selectedRows())
        self.statusBar().showMessage(f"Selected {selected_count} records")
    
    def select_all_transfer_records(self):
        """Select all records in the Transfer List table"""
        self.transfer_table.selectAll()
        selected_count = len(self.transfer_table.selectionModel().selectedRows())
        self.statusBar().showMessage(f"Selected {selected_count} records")
    
    def delete_record(self):
        """Delete selected JPHN record(s)"""
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select record(s) to delete!")
            return
        
        count = len(selected_rows)
        reply = QMessageBox.question(self, "Confirm Delete", 
                                     f"Are you sure you want to delete {count} selected record(s)?",
                                     QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            try:
                deleted_count = 0
                for index in selected_rows:
                    row = index.row()
                    pen = self.table.item(row, 0).text()  # PEN is at column 0
                    self.cursor.execute('DELETE FROM jphn WHERE pen = ?', (pen,))
                    deleted_count += 1
                
                self.conn.commit()
                self.load_data()
                QMessageBox.information(self, "Success", f"{deleted_count} record(s) deleted successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to delete record(s): {str(e)}")
    
    def delete_transfer_records(self):
        """Delete selected records from Transfer List (deletes from database)"""
        selected_rows = self.transfer_table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select record(s) to delete!")
            return
        
        count = len(selected_rows)
        reply = QMessageBox.question(self, "Confirm Delete", 
                                     f"Are you sure you want to delete {count} selected record(s)?\n\n"
                                     "This will permanently remove them from the database.",
                                     QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            try:
                deleted_count = 0
                # Get PENs from selected rows (column 2 is PEN)
                for index in selected_rows:
                    row = index.row()
                    pen = self.transfer_table.item(row, 2).text()
                    self.cursor.execute('DELETE FROM jphn WHERE pen = ?', (pen,))
                    deleted_count += 1
                
                self.conn.commit()
                self.load_data()  # This also refreshes transfer list
                QMessageBox.information(self, "Success", f"{deleted_count} record(s) deleted successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to delete record(s): {str(e)}")
    
    def export_to_csv(self):
        """Export table data to CSV"""
        file_path, _ = QFileDialog.getSaveFileName(self, "Export to CSV", 
                                                    "jphn_data.csv", "CSV Files (*.csv)")
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8-sig') as file:
                    writer = csv.writer(file)
                    
                    # Write headers
                    headers = []
                    for col in range(self.table.columnCount()):
                        headers.append(self.table.horizontalHeaderItem(col).text())
                    writer.writerow(headers)
                    
                    # Write data
                    for row in range(self.table.rowCount()):
                        if not self.table.isRowHidden(row):
                            row_data = []
                            for col in range(self.table.columnCount()):
                                item = self.table.item(row, col)
                                row_data.append(item.text() if item else "")
                            writer.writerow(row_data)
                
                QMessageBox.information(self, "Success", f"Data exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export: {str(e)}")
    
    def import_from_file(self):
        """Import data from Excel, CSV, or PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, 
            "Import Data", 
            "", 
            "All Supported Files (*.csv *.xlsx *.xls *.pdf);;CSV Files (*.csv);;Excel Files (*.xlsx *.xls);;PDF Files (*.pdf)"
        )
        
        if not file_path:
            return
        
        try:
            data = []
            sheet_info = []  # To track which sheets were read
            
            if file_path.lower().endswith('.csv'):
                # Read CSV file
                with open(file_path, 'r', encoding='utf-8-sig') as file:
                    reader = csv.reader(file)
                    headers = next(reader)  # Skip header row
                    for row in reader:
                        if row:  # Skip empty rows
                            data.append(row)
                sheet_info.append(("CSV File", len(data)))
            
            elif file_path.lower().endswith('.pdf'):
                # Read PDF file
                data, sheet_info = self.import_from_pdf(file_path)
                if data is None:
                    return  # Error already shown
            
            elif file_path.lower().endswith(('.xlsx', '.xls')):
                # Read Excel file using openpyxl - ALL SHEETS
                try:
                    import openpyxl
                except ImportError:
                    QMessageBox.critical(
                        self, 
                        "Error", 
                        "openpyxl library is required to read Excel files.\n\n"
                        "Please install it using:\npip install openpyxl"
                    )
                    return
                
                wb = openpyxl.load_workbook(file_path)
                
                # Iterate through ALL sheets in the workbook
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheet_row_count = 0
                    
                    rows = list(ws.iter_rows(values_only=True))
                    if rows:
                        # Skip header row (first row of each sheet)
                        for row in rows[1:]:
                            if any(cell is not None for cell in row):  # Skip empty rows
                                data.append([str(cell) if cell is not None else "" for cell in row])
                                sheet_row_count += 1
                    
                    if sheet_row_count > 0:
                        sheet_info.append((sheet_name, sheet_row_count))
                
                wb.close()
            
            if not data:
                QMessageBox.warning(self, "Warning", "No data found in the file!")
                return
            
            # Build sheet summary for confirmation dialog
            sheet_summary = "\n".join([f"  • {name}: {count} records" for name, count in sheet_info])
            
            # Show confirmation dialog
            reply = QMessageBox.question(
                self, 
                "Confirm Import", 
                f"Found {len(data)} total records to import.\n\n"
                f"Sheets found:\n{sheet_summary}\n\n"
                "Expected columns (in order):\n"
                "PEN, Name, Designation, Institution, District,\n"
                "Date of Entry, Date of Retirement,\n"
                "Date of Joining in District, Contact\n\n"
                "Note: Duration is calculated automatically.\n"
                "Note: Existing records with same PEN will be UPDATED.\n\n"
                "Do you want to proceed?",
                QMessageBox.Yes | QMessageBox.No
            )
            
            if reply != QMessageBox.Yes:
                return
            
            # Import data
            success_count = 0
            error_count = 0
            errors = []
            
            progress = QProgressDialog("Importing records...", "Cancel", 0, len(data), self)
            progress.setWindowModality(Qt.WindowModal)
            progress.setWindowTitle("Import Progress")
            
            for idx, row in enumerate(data):
                progress.setValue(idx)
                if progress.wasCanceled():
                    break
                
                try:
                    # Parse row data
                    # Expected: PEN, Name, Designation, Institution, District, 
                    #           Entry Date, Retirement Date, District Join Date, [Duration-skip], Contact
                    if len(row) < 8:
                        row.extend([''] * (10 - len(row)))  # Pad with empty strings
                    
                    # Column order matches display: PEN, Name, Designation, Institution, District, etc.
                    pen = str(row[0]).strip()
                    name = str(row[1]).strip()
                    designation = str(row[2]).strip() if row[2] else "JPHN Gr I"
                    institution = str(row[3]).strip()
                    district = str(row[4]).strip()
                    entry_date = self.parse_date(row[5])
                    retirement_date = self.parse_date(row[6])
                    district_join_date = self.parse_date(row[7])
                    
                    # Column 8 could be Duration (skip it) or Contact
                    # Check if column 8 looks like a duration (contains Y, M, D) or is numeric with Y/M/D
                    col8_value = str(row[8]).strip() if len(row) > 8 and row[8] else ""
                    col9_value = str(row[9]).strip() if len(row) > 9 and row[9] else ""
                    
                    # If col8 looks like duration (e.g., "5Y 3M 2D" or just contains Y/M/D pattern)
                    # then contact is in col9, otherwise contact is in col8
                    if col8_value and ('Y' in col8_value.upper() or 'M' in col8_value.upper() or 'D' in col8_value.upper()):
                        # col8 is Duration (calculated, so skip), col9 is Contact
                        contact_raw = col9_value
                    else:
                        # col8 is Contact (no Duration column in import)
                        contact_raw = col8_value
                    
                    # Clean up contact
                    contact = contact_raw if contact_raw and contact_raw.lower() != 'none' else ""
                    
                    # Validate required fields
                    if not name or not pen:
                        errors.append(f"Row {idx + 2}: Name and PEN are required")
                        error_count += 1
                        continue
                    
                    # Calculate duration from district_join_date (not imported)
                    duration = self.calculate_duration(district_join_date)
                    
                    # Check if PEN already exists
                    self.cursor.execute('SELECT pen FROM jphn WHERE pen = ?', (pen,))
                    existing = self.cursor.fetchone()
                    
                    if existing:
                        # Update existing record (preserve existing weightage data)
                        self.cursor.execute('''
                            UPDATE jphn SET name=?, designation=?, institution=?, district=?,
                                           entry_date=?, retirement_date=?, district_join_date=?,
                                           duration_days=?, contact=?
                            WHERE pen=?
                        ''', (name, designation, institution, district,
                              entry_date, retirement_date, district_join_date,
                              duration, contact, pen))
                    else:
                        # Insert new record
                        self.cursor.execute('''
                            INSERT INTO jphn (pen, name, designation, institution, district,
                                             entry_date, retirement_date, district_join_date,
                                             duration_days, contact)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (pen, name, designation, institution, district,
                              entry_date, retirement_date, district_join_date,
                              duration, contact))
                    
                    success_count += 1
                    
                except Exception as e:
                    errors.append(f"Row {idx + 2}: {str(e)}")
                    error_count += 1
            
            progress.setValue(len(data))
            self.conn.commit()
            self.load_data()
            
            # Show result
            message = f"Import completed!\n\nSuccessfully imported: {success_count}\nFailed: {error_count}"
            if errors and len(errors) <= 10:
                message += "\n\nErrors:\n" + "\n".join(errors)
            elif errors:
                message += f"\n\nFirst 10 errors:\n" + "\n".join(errors[:10])
                message += f"\n... and {len(errors) - 10} more errors"
            
            if error_count > 0:
                QMessageBox.warning(self, "Import Result", message)
            else:
                QMessageBox.information(self, "Import Result", message)
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to import file: {str(e)}")
    
    def parse_date(self, date_value):
        """Parse date from various formats and return DD-MM-YYYY format only (no time)"""
        if not date_value:
            return datetime.now().strftime("%d-%m-%Y")
        
        date_str = str(date_value).strip()
        
        # Handle datetime objects from Excel (which may include time)
        if isinstance(date_value, datetime):
            return date_value.strftime("%d-%m-%Y")
        
        # Handle date objects
        if isinstance(date_value, date):
            return date_value.strftime("%d-%m-%Y")
        
        # Remove any time component if present (e.g., "01-01-2020 00:00:00")
        if ' ' in date_str:
            date_str = date_str.split(' ')[0]
        
        # Try various date formats
        formats = [
            "%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%Y/%m/%d",
            "%d-%m-%y", "%d/%m/%y", "%m-%d-%Y", "%m/%d/%Y",
            "%d.%m.%Y", "%d.%m.%y", "%Y%m%d"
        ]
        
        for fmt in formats:
            try:
                parsed_date = datetime.strptime(date_str, fmt)
                return parsed_date.strftime("%d-%m-%Y")
            except ValueError:
                continue
        
        # If all formats fail, return the original string or current date
        return date_str if date_str else datetime.now().strftime("%d-%m-%Y")
    
    def import_from_pdf(self, file_path):
        """Import data from PDF file containing cadre register - one employee per page"""
        try:
            import pdfplumber
            import re
        except ImportError:
            QMessageBox.critical(
                self, 
                "Error", 
                "pdfplumber library is required to read PDF files.\n\n"
                "Please install it using:\npip install pdfplumber"
            )
            return None, []
        
        data = []
        sheet_info = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                progress = QProgressDialog(f"Reading PDF ({total_pages} pages)...", "Cancel", 0, total_pages, self)
                progress.setWindowModality(Qt.WindowModal)
                progress.setWindowTitle("Reading PDF")
                
                for page_num, page in enumerate(pdf.pages):
                    progress.setValue(page_num)
                    if progress.wasCanceled():
                        return None, []
                    
                    # Extract tables from page - each page has one employee
                    tables = page.extract_tables()
                    
                    if not tables:
                        continue
                    
                    # Process the first table on the page (one employee per page)
                    table = tables[0]
                    if not table or len(table) < 10:
                        continue
                    
                    # Initialize variables for extracted data
                    pen = ""
                    name = ""
                    designation = ""
                    institution = ""
                    district = ""
                    entry_date = ""
                    retirement_date = ""
                    district_join_date = ""
                    contact = ""
                    
                    # PDF structure: Label in column 1, Data in column 4
                    # Row 4: PEN & Name
                    # Row 5: Designation
                    # Row 6: Name of Present Institution
                    # Row 7: District
                    # Row 8: Date of Entry in Service
                    # Row 9: Date of retirement
                    # Row 11: Date of joining in the present district
                    # Last row with 'Contact Number': Contact
                    
                    for row in table:
                        if not row or len(row) < 5:
                            continue
                        
                        # Get label from column 1 and data from column 4
                        label = str(row[1]).strip().lower() if row[1] else ""
                        data_value = str(row[4]).strip() if row[4] else ""
                        
                        # Also check column 0 for Contact Number (it spans differently)
                        col0 = str(row[0]).strip().lower() if row[0] else ""
                        
                        if not data_value:
                            continue
                        
                        # PEN & Name - numbers are PEN (6-7 digits), rest is Name
                        if 'pen' in label and 'name' in label:
                            # Extract PEN (6 or 7 digit number at the start)
                            pen_match = re.match(r'^(\d{6,7})\s*', data_value)
                            if pen_match:
                                pen = pen_match.group(1)
                                # Rest is the name (convert to uppercase)
                                name = data_value[len(pen):].strip().upper()
                            else:
                                # Try to find PEN anywhere in the string
                                pen_match = re.search(r'(\d{6,7})', data_value)
                                if pen_match:
                                    pen = pen_match.group(1)
                                    name = data_value.replace(pen, "").strip().upper()
                                else:
                                    name = data_value.upper()
                        
                        elif label == 'designation':
                            designation = data_value
                        
                        elif 'name of present institution' in label:
                            institution = data_value
                        
                        elif label == 'district':
                            district = data_value
                        
                        elif 'entry' in label and 'service' in label:
                            entry_date = data_value
                        
                        elif 'retirement' in label:
                            retirement_date = data_value
                        
                        elif 'joining' in label and 'present district' in label:
                            district_join_date = data_value
                        
                        elif 'contact' in col0 or 'contact' in label:
                            # Extract only digits and check if it's a 10 digit number
                            contact_clean = re.sub(r'[^0-9]', '', data_value)
                            if len(contact_clean) == 10 and contact_clean.isdigit():
                                contact = contact_clean
                    
                    # Skip if we don't have essential data
                    if not pen and not name:
                        continue
                    
                    # Skip header rows
                    if name.lower() in ['name', 'pen', 'sl no', 'serial', 'pen & name', 'pen and name']:
                        continue
                    
                    # Create processed row
                    processed_row = [
                        name,
                        pen,
                        designation if designation else "JPHN Gr I",
                        institution,
                        district,
                        entry_date,
                        retirement_date,
                        district_join_date,
                        "",  # institution_join_date
                        "No",  # weightage
                        "",  # weightage_details
                        contact  # contact number
                    ]
                    data.append(processed_row)
                
                progress.setValue(total_pages)
                
            if data:
                sheet_info.append((f"PDF ({total_pages} pages)", len(data)))
            
            return data, sheet_info
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to read PDF file: {str(e)}")
            return None, []
    
    def closeEvent(self, event):
        """Clean up database connection on close"""
        self.conn.close()
        event.accept()


def main():
    app = QApplication(sys.argv)
    
    # Set application icon
    icon_path = resource_path("favicon.ico")
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))
    
    # Set default font
    try:
        app.setFont(QFont("Calibri", 10))
    except:
        pass  # Font will fall back to system default
    
    # Show login screen first
    login = LoginDialog()
    if login.exec() != QDialog.Accepted or not login.is_authenticated():
        sys.exit(0)  # User closed or failed authentication
    
    while True:
        # Show welcome dialog to choose transfer type (fullscreen)
        welcome = WelcomeDialog()
        if welcome.exec() != QDialog.Accepted:
            sys.exit(0)  # User closed the dialog without selecting
        
        transfer_type = welcome.get_selected_type()
        if not transfer_type:
            sys.exit(0)
        
        year = None
        month = None
        
        # If General Transfer, show year selection
        if transfer_type == "general":
            year_dialog = YearSelectionDialog()
            if year_dialog.exec() != QDialog.Accepted:
                # User pressed back, go back to welcome screen
                continue
            year = year_dialog.get_selected_year()
            if not year:
                continue
        else:
            # Regular Transfer - show month/year selection
            regular_dialog = RegularTransferSelectionDialog()
            if regular_dialog.exec() != QDialog.Accepted:
                # User pressed back, go back to welcome screen
                continue
            month = regular_dialog.get_selected_month()
            year = regular_dialog.get_selected_year()
            if not month or not year:
                continue
        
        # Create and show main window
        window = JPHNManagementSystem(transfer_type, year, month)
        window.restart_requested = False
        window.showMaximized()
        
        # Run the event loop for this window
        app.exec()
        
        # Check if restart was requested
        if hasattr(window, 'restart_requested') and window.restart_requested:
            continue  # Go back to welcome screen
        else:
            break  # Exit the application
    
    sys.exit(0)


if __name__ == '__main__':
    main()