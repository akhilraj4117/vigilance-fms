"""
JPHN Transfer Management System - Flask Web Application
Department of Health Services, Kerala
"""
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date
from functools import wraps
import os
import io
import csv
import time

from config import config

# Initialize Flask app
app = Flask(__name__)
app.config.from_object(config[os.environ.get('FLASK_CONFIG', 'default')])

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

# Districts and other constants from config
DISTRICTS = app.config['DISTRICTS']
NEARBY_DISTRICTS = app.config['NEARBY_DISTRICTS']
MONTHS = app.config['MONTHS']
VALID_USERS = app.config['VALID_USERS']


# ==================== HEALTH CHECK & ERROR HANDLING ====================
@app.route('/health')
def health_check():
    """Health check endpoint for Render"""
    try:
        # Test database connection
        db.session.execute(db.text('SELECT 1'))
        db.session.commit()
        return jsonify({'status': 'healthy', 'database': 'connected'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 500


@app.errorhandler(502)
def bad_gateway_error(error):
    """Handle 502 errors gracefully"""
    return render_template('error.html', 
                          error='Service temporarily unavailable. Please refresh the page.'), 502


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors and attempt database reconnection"""
    try:
        db.session.rollback()
    except:
        pass
    return render_template('error.html', 
                          error='Internal server error. Please try again.'), 500


# ==================== USER MODEL ======================================
class User(UserMixin):
    """Simple user model for authentication"""
    def __init__(self, user_id):
        self.id = user_id


@login_manager.user_loader
def load_user(user_id):
    if user_id in VALID_USERS:
        return User(user_id)
    return None


# ==================== DATABASE HELPER FUNCTIONS ====================
def get_table_prefix():
    """Get the table prefix based on session transfer type and year/month"""
    transfer_type = session.get('transfer_type', 'general')
    year = session.get('year', datetime.now().year)
    month = session.get('month', '')
    
    if transfer_type == 'general':
        return f'general_{year}_'
    else:
        return f'regular_{month.lower()}_{year}_'


def get_nearby_districts(district):
    """Get list of nearby districts for a given district (for against transfers)"""
    return NEARBY_DISTRICTS.get(district, [])


def check_autofill_ran():
    """Check if autofill has been run for current session"""
    if not session.get('transfer_type'):
        return False
    
    # Check session flag first (set when autofill is run)
    if session.get('autofill_ran'):
        return True
    
    # Also check database for draft entries (for existing sessions)
    try:
        prefix = get_table_prefix()
        result = db.session.execute(db.text(f"""
            SELECT COUNT(*) FROM {prefix}transfer_draft 
            WHERE remarks LIKE 'Vacancy by Transfer%%' 
               OR remarks LIKE 'Pref %%' 
               OR against_info LIKE 'Displaced for%%'
        """))
        return result.fetchone()[0] > 0
    except:
        return False


def check_final_exists():
    """Check if final list has been created (draft confirmed)"""
    if not session.get('transfer_type'):
        return False
    
    try:
        prefix = get_table_prefix()
        result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final"))
        return result.fetchone()[0] > 0
    except:
        return False


@app.context_processor
def inject_autofill_status():
    """Inject autofill and final status into all templates"""
    return {
        'autofill_ran': check_autofill_ran(),
        'final_exists': check_final_exists()
    }


def execute_query(query, params=None, fetch_all=True):
    """Execute a raw SQL query with table prefix substitution"""
    prefix = get_table_prefix()
    
    # Replace table placeholders
    query = query.replace('{jphn}', f'{prefix}jphn')
    query = query.replace('{transfer_draft}', f'{prefix}transfer_draft')
    query = query.replace('{transfer_final}', f'{prefix}transfer_final')
    query = query.replace('{transfer_applied}', f'{prefix}transfer_applied')
    query = query.replace('{vacancy}', f'{prefix}vacancy')
    query = query.replace('{settings}', f'{prefix}settings')
    
    result = db.session.execute(db.text(query), params or {})
    
    if query.strip().upper().startswith('SELECT'):
        if fetch_all:
            return result.fetchall()
        else:
            return result.fetchone()
    else:
        db.session.commit()
        return result


def ensure_tables():
    """Create tables if they don't exist"""
    prefix = get_table_prefix()
    
    tables_sql = f"""
    -- JPHN main table
    CREATE TABLE IF NOT EXISTS {prefix}jphn (
        pen TEXT PRIMARY KEY,
        name TEXT NOT NULL,
        designation TEXT,
        institution TEXT,
        district TEXT,
        entry_date TEXT,
        retirement_date TEXT,
        district_join_date TEXT,
        duration_days INTEGER,
        institution_join_date TEXT,
        weightage TEXT,
        weightage_details TEXT,
        contact TEXT,
        weightage_priority INTEGER DEFAULT 5,
        last_modified TEXT
    );
    
    -- Transfer draft table
    CREATE TABLE IF NOT EXISTS {prefix}transfer_draft (
        pen TEXT PRIMARY KEY,
        transfer_to_district TEXT NOT NULL,
        added_date TEXT,
        against_info TEXT,
        remarks TEXT,
        last_modified TEXT
    );
    
    -- Transfer final table
    CREATE TABLE IF NOT EXISTS {prefix}transfer_final (
        pen TEXT PRIMARY KEY,
        transfer_to_district TEXT NOT NULL,
        confirmed_date TEXT,
        last_modified TEXT
    );
    
    -- Transfer applied table
    CREATE TABLE IF NOT EXISTS {prefix}transfer_applied (
        pen TEXT PRIMARY KEY,
        applied_to_district TEXT,
        applied_date TEXT,
        receipt_numbers TEXT,
        pref1 TEXT,
        pref2 TEXT,
        pref3 TEXT,
        pref4 TEXT,
        pref5 TEXT,
        pref6 TEXT,
        pref7 TEXT,
        pref8 TEXT,
        special_priority TEXT DEFAULT 'No',
        special_priority_reason TEXT,
        locked TEXT DEFAULT 'No',
        last_modified TEXT
    );
    
    -- Vacancy table
    CREATE TABLE IF NOT EXISTS {prefix}vacancy (
        district TEXT PRIMARY KEY,
        total_strength INTEGER DEFAULT 0,
        vacancy_reported INTEGER DEFAULT 0,
        last_modified TEXT
    );
    
    -- Settings table
    CREATE TABLE IF NOT EXISTS {prefix}settings (
        key TEXT PRIMARY KEY,
        value TEXT,
        last_modified TEXT
    );
    """
    
    # Execute each statement separately
    for statement in tables_sql.split(';'):
        statement = statement.strip()
        if statement:
            try:
                db.session.execute(db.text(statement))
            except Exception as e:
                print(f"Table creation warning: {e}")
    
    db.session.commit()
    
    # Initialize vacancy table with districts
    for district in DISTRICTS:
        try:
            db.session.execute(db.text(f"""
                INSERT INTO {prefix}vacancy (district, total_strength, vacancy_reported, last_modified)
                VALUES (:district, 0, 0, :now)
                ON CONFLICT (district) DO NOTHING
            """), {'district': district, 'now': datetime.now().isoformat()})
        except:
            pass
    
    db.session.commit()
    
    # Ensure locked column exists in transfer_applied (for existing databases)
    try:
        db.session.execute(db.text(f"""
            ALTER TABLE {prefix}transfer_applied ADD COLUMN IF NOT EXISTS locked TEXT DEFAULT 'No'
        """))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Note: locked column may already exist or couldn't be added: {e}")


def calculate_duration(join_date_str):
    """Calculate duration in days from join date to today"""
    try:
        join_date = datetime.strptime(join_date_str, "%d-%m-%Y")
        today = datetime.now()
        return (today - join_date).days
    except:
        return 0


def format_duration(days):
    """Format duration in days as Years, Months, Days"""
    if days is None or days <= 0:
        return "0D"
    years = days // 365
    remaining = days % 365
    months = remaining // 30
    days_left = remaining % 30
    
    parts = []
    if years > 0:
        parts.append(f"{years}Y")
    if months > 0:
        parts.append(f"{months}M")
    if days_left > 0 or not parts:
        parts.append(f"{days_left}D")
    return " ".join(parts)


def requires_transfer_session(f):
    """Decorator to ensure transfer session is set"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'transfer_type' not in session:
            flash('Please select a transfer type first.', 'warning')
            return redirect(url_for('select_transfer'))
        return f(*args, **kwargs)
    return decorated_function


# ==================== ROUTES ====================

@app.route('/')
def index():
    """Home page - redirect to login or main page"""
    if current_user.is_authenticated:
        return redirect(url_for('select_transfer'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if current_user.is_authenticated:
        return redirect(url_for('select_transfer'))
    
    if request.method == 'POST':
        user_id = request.form.get('user_id', '').strip()
        password = request.form.get('password', '').strip()
        
        if user_id in VALID_USERS and VALID_USERS[user_id] == password:
            user = User(user_id)
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('select_transfer'))
        else:
            flash('Invalid User ID or Password.', 'error')
    
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    """Logout"""
    logout_user()
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


@app.route('/select-transfer')
@login_required
def select_transfer():
    """Welcome page to select transfer type"""
    return render_template('select_transfer.html')


@app.route('/manage-databases')
@login_required
def manage_databases():
    """Manage existing transfer databases"""
    databases = []
    
    try:
        # Get General Transfer databases
        result = db.session.execute(db.text("""
            SELECT table_name FROM information_schema.tables 
            WHERE table_schema = 'public' AND table_name LIKE 'general_%_jphn'
            ORDER BY table_name DESC
        """))
        for row in result:
            try:
                year = int(row[0].replace('general_', '').replace('_jphn', ''))
                if 2000 <= year <= 2100:
                    prefix = f'general_{year}_'
                    # Get counts
                    jphn_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}jphn")).fetchone()[0] or 0
                    applied_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_applied")).fetchone()[0] or 0
                    draft_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_draft")).fetchone()[0] or 0
                    final_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final")).fetchone()[0] or 0
                    
                    databases.append({
                        'type': 'General',
                        'year': year,
                        'month': '',
                        'prefix': prefix,
                        'display_name': f'General Transfer {year}',
                        'jphn_count': jphn_count,
                        'applied_count': applied_count,
                        'draft_count': draft_count,
                        'final_count': final_count
                    })
            except:
                pass
        
        # Get Regular Transfer databases
        result = db.session.execute(db.text("""
            SELECT table_name FROM information_schema.tables 
            WHERE table_schema = 'public' AND table_name LIKE 'regular_%_jphn'
            ORDER BY table_name DESC
        """))
        for row in result:
            try:
                parts = row[0].replace('_jphn', '').replace('regular_', '').rsplit('_', 1)
                if len(parts) == 2:
                    month = parts[0].capitalize()
                    year = int(parts[1])
                    if 2000 <= year <= 2100:
                        prefix = f'regular_{month.lower()}_{year}_'
                        # Get counts
                        jphn_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}jphn")).fetchone()[0] or 0
                        applied_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_applied")).fetchone()[0] or 0
                        draft_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_draft")).fetchone()[0] or 0
                        final_count = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final")).fetchone()[0] or 0
                        
                        databases.append({
                            'type': 'Regular',
                            'year': year,
                            'month': month,
                            'prefix': prefix,
                            'display_name': f'Regular Transfer - {month} {year}',
                            'jphn_count': jphn_count,
                            'applied_count': applied_count,
                            'draft_count': draft_count,
                            'final_count': final_count
                        })
            except:
                pass
    except Exception as e:
        flash(f'Error loading databases: {str(e)}', 'error')
    
    return render_template('manage_databases.html', databases=databases)


@app.route('/delete-database', methods=['POST'])
@login_required
def delete_database():
    """Delete a transfer database"""
    prefix = request.form.get('prefix', '')
    display_name = request.form.get('display_name', 'Unknown')
    
    if not prefix:
        flash('Invalid database prefix!', 'error')
        return redirect(url_for('manage_databases'))
    
    # Security check - only allow valid prefixes
    if not (prefix.startswith('general_') or prefix.startswith('regular_')):
        flash('Invalid database prefix format!', 'error')
        return redirect(url_for('manage_databases'))
    
    try:
        # Drop all tables with this prefix
        tables = ['jphn', 'transfer_applied', 'transfer_draft', 'transfer_final', 'vacancy']
        for table in tables:
            full_table = f'{prefix}{table}'
            db.session.execute(db.text(f'DROP TABLE IF EXISTS {full_table} CASCADE'))
        
        db.session.commit()
        flash(f'Successfully deleted: {display_name}', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting database: {str(e)}', 'error')
    
    return redirect(url_for('manage_databases'))


@app.route('/select-year/<transfer_type>')
@login_required
def select_year(transfer_type):
    """Select year for General Transfer or month/year for Regular Transfer"""
    if transfer_type == 'general':
        # Get available years from existing tables
        available_years = []
        try:
            result = db.session.execute(db.text("""
                SELECT table_name FROM information_schema.tables 
                WHERE table_schema = 'public' AND table_name LIKE 'general_%_jphn'
            """))
            for row in result:
                try:
                    year = int(row[0].replace('general_', '').replace('_jphn', ''))
                    if 2000 <= year <= 2100:
                        available_years.append(year)
                except:
                    pass
        except:
            pass
        
        return render_template('select_year.html', 
                             transfer_type=transfer_type,
                             available_years=sorted(available_years, reverse=True),
                             current_year=datetime.now().year)
    else:
        # Regular Transfer - select month/year
        available_lists = []
        try:
            result = db.session.execute(db.text("""
                SELECT table_name FROM information_schema.tables 
                WHERE table_schema = 'public' AND table_name LIKE 'regular_%_jphn'
            """))
            for row in result:
                try:
                    parts = row[0].replace('_jphn', '').replace('regular_', '').rsplit('_', 1)
                    if len(parts) == 2:
                        month = parts[0].capitalize()
                        year = int(parts[1])
                        if month in MONTHS and 2000 <= year <= 2100:
                            available_lists.append((month, year))
                except:
                    pass
        except:
            pass
        
        return render_template('select_month.html',
                             transfer_type=transfer_type,
                             available_lists=sorted(available_lists, key=lambda x: (x[1], MONTHS.index(x[0])), reverse=True),
                             months=MONTHS,
                             current_month=MONTHS[datetime.now().month - 1],
                             current_year=datetime.now().year)


@app.route('/set-session', methods=['POST'])
@login_required
def set_session():
    """Set transfer session (type, year, month)"""
    transfer_type = request.form.get('transfer_type', 'general')
    year = request.form.get('year', datetime.now().year)
    month = request.form.get('month', '')
    
    session['transfer_type'] = transfer_type
    session['year'] = int(year)
    session['month'] = month
    
    # Clear autofill flag when changing session
    session.pop('autofill_ran', None)
    
    # Ensure tables exist
    ensure_tables()
    
    flash(f'Session set: {transfer_type.title()} Transfer - {month + " " if month else ""}{year}', 'success')
    return redirect(url_for('dashboard'))


@app.route('/dashboard')
@login_required
@requires_transfer_session
def dashboard():
    """Main dashboard"""
    prefix = get_table_prefix()
    
    # Get statistics
    stats = {
        'total_employees': 0,
        'applied_employees': 0,
        'draft_transfers': 0,
        'confirmed_transfers': 0,
        'total_vacancy': 0,
        'filled_vacancy': 0
    }
    
    try:
        result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}jphn"))
        stats['total_employees'] = result.fetchone()[0] or 0
        
        result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_applied"))
        stats['applied_employees'] = result.fetchone()[0] or 0
        
        result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_draft"))
        stats['draft_transfers'] = result.fetchone()[0] or 0
        
        result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final"))
        stats['confirmed_transfers'] = result.fetchone()[0] or 0
        
        result = db.session.execute(db.text(f"SELECT SUM(vacancy_reported) FROM {prefix}vacancy"))
        stats['total_vacancy'] = result.fetchone()[0] or 0
        
        # Check if autofill has been run (look for autofill-specific remarks)
        # Autofill sets remarks like "Vacancy by Transfer", "Pref N:", or against_info "Displaced for X"
        result = db.session.execute(db.text(f"""
            SELECT COUNT(*) FROM {prefix}transfer_draft 
            WHERE remarks LIKE 'Vacancy by Transfer%%' 
               OR remarks LIKE 'Pref %%' 
               OR against_info LIKE 'Displaced for%%'
        """))
        autofill_ran = result.fetchone()[0] > 0
        
        # Only count filled vacancy if autofill has been run
        if autofill_ran:
            stats['filled_vacancy'] = stats['draft_transfers']
        else:
            stats['filled_vacancy'] = 0
    except Exception as e:
        print(f"Dashboard stats error: {e}")
    
    return render_template('dashboard.html',
                         stats=stats,
                         transfer_type=session.get('transfer_type'),
                         year=session.get('year'),
                         month=session.get('month'),
                         districts=DISTRICTS)


# ==================== CADRE LIST ROUTES ====================

@app.route('/cadre')
@login_required
@requires_transfer_session
def cadre_list():
    """View cadre list"""
    prefix = get_table_prefix()
    district_filter = request.args.get('district', '')
    search = request.args.get('search', '').lower()
    
    query = f"""
        SELECT pen, name, designation, institution, district, entry_date,
               retirement_date, district_join_date, duration_days, contact,
               weightage, weightage_details
        FROM {prefix}jphn
        WHERE 1=1
    """
    params = {}
    
    if district_filter:
        query += " AND district = :district"
        params['district'] = district_filter
    
    if search:
        query += """ AND (LOWER(name) LIKE :search OR LOWER(pen) LIKE :search 
                    OR LOWER(institution) LIKE :search)"""
        params['search'] = f'%{search}%'
    
    query += " ORDER BY district, name"
    
    result = db.session.execute(db.text(query), params)
    employees = result.fetchall()
    
    # Get applied PENs
    applied_result = db.session.execute(db.text(f"SELECT pen FROM {prefix}transfer_applied"))
    applied_pens = set(row[0] for row in applied_result.fetchall())
    
    return render_template('cadre_list.html',
                         employees=employees,
                         applied_pens=applied_pens,
                         districts=DISTRICTS,
                         district_filter=district_filter,
                         search=search,
                         format_duration=format_duration)


@app.route('/cadre/upload', methods=['POST'])
@login_required
@requires_transfer_session
def upload_cadre_data():
    """Upload cadre data from PDF, CSV, or Excel file"""
    import re
    
    prefix = get_table_prefix()
    
    if 'file' not in request.files:
        flash('No file selected!', 'error')
        return redirect(url_for('cadre_list'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected!', 'error')
        return redirect(url_for('cadre_list'))
    
    replace_existing = request.form.get('replace_existing') == 'on'
    filename = file.filename.lower()
    
    try:
        data = []
        
        if filename.endswith('.pdf'):
            # Import PDF using pdfplumber
            try:
                import pdfplumber
            except ImportError:
                flash('pdfplumber library not installed. Please install with: pip install pdfplumber', 'error')
                return redirect(url_for('cadre_list'))
            
            # Save file temporarily
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                with pdfplumber.open(tmp_path) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        if not tables:
                            continue
                        
                        table = tables[0]
                        if not table or len(table) < 10:
                            continue
                        
                        # Initialize variables
                        pen = ""
                        name = ""
                        designation = ""
                        institution = ""
                        district = ""
                        entry_date = ""
                        retirement_date = ""
                        district_join_date = ""
                        contact = ""
                        
                        for row in table:
                            if not row or len(row) < 5:
                                continue
                            
                            label = str(row[1]).strip().lower() if row[1] else ""
                            data_value = str(row[4]).strip() if row[4] else ""
                            col0 = str(row[0]).strip().lower() if row[0] else ""
                            
                            if not data_value:
                                continue
                            
                            if 'pen' in label and 'name' in label:
                                pen_match = re.match(r'^(\d{6,7})\s*', data_value)
                                if pen_match:
                                    pen = pen_match.group(1)
                                    name = data_value[len(pen):].strip().upper()
                                else:
                                    pen_match = re.search(r'(\d{6,7})', data_value)
                                    if pen_match:
                                        pen = pen_match.group(1)
                                        name = data_value.replace(pen, "").strip().upper()
                                    else:
                                        name = data_value.upper()
                            elif label == 'designation':
                                designation = data_value
                            elif 'name of present institution' in label:
                                institution = data_value
                            elif label == 'district':
                                district = data_value
                            elif 'entry' in label and 'service' in label:
                                entry_date = data_value
                            elif 'retirement' in label:
                                retirement_date = data_value
                            elif 'joining' in label and 'present district' in label:
                                district_join_date = data_value
                            elif 'contact' in col0 or 'contact' in label:
                                contact_clean = re.sub(r'[^0-9]', '', data_value)
                                if len(contact_clean) == 10:
                                    contact = contact_clean
                        
                        if pen and name:
                            data.append({
                                'pen': pen, 'name': name, 'designation': designation or 'JPHN Gr I',
                                'institution': institution, 'district': district,
                                'entry_date': entry_date, 'retirement_date': retirement_date,
                                'district_join_date': district_join_date, 'contact': contact
                            })
            finally:
                import os
                os.unlink(tmp_path)
        
        elif filename.endswith('.csv'):
            # Import CSV
            import csv as csv_module
            content = file.read().decode('utf-8')
            reader = csv_module.reader(content.splitlines())
            header = next(reader, None)  # Skip header
            
            for row in reader:
                if len(row) >= 5:
                    data.append({
                        'name': row[0] if len(row) > 0 else '',
                        'pen': row[1] if len(row) > 1 else '',
                        'designation': row[2] if len(row) > 2 else 'JPHN Gr I',
                        'institution': row[3] if len(row) > 3 else '',
                        'district': row[4] if len(row) > 4 else '',
                        'entry_date': row[5] if len(row) > 5 else '',
                        'retirement_date': row[6] if len(row) > 6 else '',
                        'district_join_date': row[7] if len(row) > 7 else '',
                        'contact': row[11] if len(row) > 11 else ''
                    })
        
        elif filename.endswith(('.xlsx', '.xls')):
            # Import Excel
            try:
                import openpyxl
            except ImportError:
                flash('openpyxl library not installed. Please install with: pip install openpyxl', 'error')
                return redirect(url_for('cadre_list'))
            
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                wb = openpyxl.load_workbook(tmp_path)
                ws = wb.active
                rows = list(ws.iter_rows(min_row=2, values_only=True))  # Skip header
                
                for row in rows:
                    if row and len(row) >= 5 and row[1]:  # Check PEN exists
                        data.append({
                            'name': str(row[0] or ''),
                            'pen': str(row[1] or ''),
                            'designation': str(row[2] or 'JPHN Gr I'),
                            'institution': str(row[3] or ''),
                            'district': str(row[4] or ''),
                            'entry_date': str(row[5] or '') if len(row) > 5 else '',
                            'retirement_date': str(row[6] or '') if len(row) > 6 else '',
                            'district_join_date': str(row[7] or '') if len(row) > 7 else '',
                            'contact': str(row[11] or '') if len(row) > 11 else ''
                        })
            finally:
                import os
                os.unlink(tmp_path)
        
        else:
            flash('Unsupported file format. Please use PDF, CSV, or Excel files.', 'error')
            return redirect(url_for('cadre_list'))
        
        if not data:
            flash('No valid data found in the file!', 'warning')
            return redirect(url_for('cadre_list'))
        
        # Clear existing data if replace option is selected
        if replace_existing:
            db.session.execute(db.text(f"DELETE FROM {prefix}jphn"))
        
        # Insert data
        imported = 0
        skipped = 0
        
        for emp in data:
            if not emp['pen'] or not emp['name']:
                skipped += 1
                continue
            
            duration = calculate_duration(emp['district_join_date']) if emp.get('district_join_date') else 0
            
            try:
                db.session.execute(db.text(f"""
                    INSERT INTO {prefix}jphn 
                    (pen, name, designation, institution, district, entry_date, retirement_date,
                     district_join_date, duration_days, weightage, contact, last_modified)
                    VALUES (:pen, :name, :designation, :institution, :district, :entry_date,
                            :retirement_date, :district_join_date, :duration, 'No', :contact, :now)
                    ON CONFLICT (pen) DO UPDATE SET
                    name = :name, designation = :designation, institution = :institution,
                    district = :district, entry_date = :entry_date, retirement_date = :retirement_date,
                    district_join_date = :district_join_date, duration_days = :duration,
                    contact = :contact, last_modified = :now
                """), {
                    'pen': emp['pen'], 'name': emp['name'], 
                    'designation': emp.get('designation', 'JPHN Gr I'),
                    'institution': emp.get('institution', ''), 'district': emp.get('district', ''),
                    'entry_date': emp.get('entry_date', ''), 
                    'retirement_date': emp.get('retirement_date', ''),
                    'district_join_date': emp.get('district_join_date', ''),
                    'duration': duration, 'contact': emp.get('contact', ''),
                    'now': datetime.now().isoformat()
                })
                imported += 1
            except Exception as e:
                skipped += 1
                print(f"Error importing {emp['pen']}: {e}")
        
        db.session.commit()
        
        # Check if AJAX request
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or not request.accept_mimetypes.accept_html:
            return jsonify({'success': True, 'message': f'Import complete! {imported} employees imported, {skipped} skipped.'})
        
        flash(f'Import complete! {imported} employees imported, {skipped} skipped.', 'success')
        
    except Exception as e:
        db.session.rollback()
        
        # Check if AJAX request
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or not request.accept_mimetypes.accept_html:
            return jsonify({'success': False, 'message': f'Error importing data: {str(e)}'}), 500
        
        flash(f'Error importing data: {str(e)}', 'error')
    
    return redirect(url_for('cadre_list'))


@app.route('/cadre/add', methods=['GET', 'POST'])
@login_required
@requires_transfer_session
def add_employee():
    """Add new employee"""
    if request.method == 'POST':
        prefix = get_table_prefix()
        
        pen = request.form.get('pen', '').strip()
        name = request.form.get('name', '').strip()
        designation = request.form.get('designation', 'JPHN Gr I').strip()
        institution = request.form.get('institution', '').strip()
        district = request.form.get('district', '')
        entry_date = request.form.get('entry_date', '')
        retirement_date = request.form.get('retirement_date', '')
        district_join_date = request.form.get('district_join_date', '')
        institution_join_date = request.form.get('institution_join_date', '')
        contact = request.form.get('contact', '').strip()
        weightage = 'Yes' if request.form.get('weightage') else 'No'
        weightage_details = request.form.get('weightage_details', '').strip()
        
        if not pen or not name:
            flash('PEN and Name are required!', 'error')
            return redirect(url_for('add_employee'))
        
        duration = calculate_duration(district_join_date) if district_join_date else 0
        
        try:
            db.session.execute(db.text(f"""
                INSERT INTO {prefix}jphn 
                (pen, name, designation, institution, district, entry_date, retirement_date,
                 district_join_date, duration_days, institution_join_date, weightage, 
                 weightage_details, contact, last_modified)
                VALUES (:pen, :name, :designation, :institution, :district, :entry_date,
                        :retirement_date, :district_join_date, :duration, :institution_join_date,
                        :weightage, :weightage_details, :contact, :now)
            """), {
                'pen': pen, 'name': name, 'designation': designation, 'institution': institution,
                'district': district, 'entry_date': entry_date, 'retirement_date': retirement_date,
                'district_join_date': district_join_date, 'duration': duration,
                'institution_join_date': institution_join_date, 'weightage': weightage,
                'weightage_details': weightage_details, 'contact': contact,
                'now': datetime.now().isoformat()
            })
            db.session.commit()
            flash('Employee added successfully!', 'success')
            return redirect(url_for('cadre_list'))
        except Exception as e:
            db.session.rollback()
            if 'duplicate' in str(e).lower() or 'unique' in str(e).lower():
                flash('PEN already exists!', 'error')
            else:
                flash(f'Error: {str(e)}', 'error')
    
    return render_template('employee_form.html', 
                         employee=None, 
                         districts=DISTRICTS,
                         action='Add')


@app.route('/cadre/edit/<pen>', methods=['GET', 'POST'])
@login_required
@requires_transfer_session
def edit_employee(pen):
    """Edit employee"""
    prefix = get_table_prefix()
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        designation = request.form.get('designation', 'JPHN Gr I').strip()
        institution = request.form.get('institution', '').strip()
        district = request.form.get('district', '')
        entry_date = request.form.get('entry_date', '')
        retirement_date = request.form.get('retirement_date', '')
        district_join_date = request.form.get('district_join_date', '')
        institution_join_date = request.form.get('institution_join_date', '')
        contact = request.form.get('contact', '').strip()
        weightage = 'Yes' if request.form.get('weightage') else 'No'
        weightage_details = request.form.get('weightage_details', '').strip()
        
        duration = calculate_duration(district_join_date) if district_join_date else 0
        
        try:
            db.session.execute(db.text(f"""
                UPDATE {prefix}jphn SET
                name = :name, designation = :designation, institution = :institution,
                district = :district, entry_date = :entry_date, retirement_date = :retirement_date,
                district_join_date = :district_join_date, duration_days = :duration,
                institution_join_date = :institution_join_date, weightage = :weightage,
                weightage_details = :weightage_details, contact = :contact,
                last_modified = :now
                WHERE pen = :pen
            """), {
                'name': name, 'designation': designation, 'institution': institution,
                'district': district, 'entry_date': entry_date, 'retirement_date': retirement_date,
                'district_join_date': district_join_date, 'duration': duration,
                'institution_join_date': institution_join_date, 'weightage': weightage,
                'weightage_details': weightage_details, 'contact': contact,
                'now': datetime.now().isoformat(), 'pen': pen
            })
            db.session.commit()
            flash('Employee updated successfully!', 'success')
            return redirect(url_for('cadre_list'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error: {str(e)}', 'error')
    
    # Get employee data
    result = db.session.execute(db.text(f"""
        SELECT pen, name, designation, institution, district, entry_date,
               retirement_date, district_join_date, duration_days, institution_join_date,
               weightage, weightage_details, contact
        FROM {prefix}jphn WHERE pen = :pen
    """), {'pen': pen})
    employee = result.fetchone()
    
    if not employee:
        flash('Employee not found!', 'error')
        return redirect(url_for('cadre_list'))
    
    return render_template('employee_form.html',
                         employee=employee,
                         districts=DISTRICTS,
                         action='Edit')


@app.route('/cadre/delete/<pen>', methods=['POST'])
@login_required
@requires_transfer_session
def delete_employee(pen):
    """Delete employee"""
    prefix = get_table_prefix()
    
    try:
        db.session.execute(db.text(f"DELETE FROM {prefix}jphn WHERE pen = :pen"), {'pen': pen})
        db.session.commit()
        flash('Employee deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('cadre_list'))


# ==================== VACANCY ROUTES ====================

@app.route('/vacancy')
@login_required
@requires_transfer_session
def vacancy_list():
    """View and manage vacancy positions"""
    prefix = get_table_prefix()
    autofill_ran = check_autofill_ran()
    
    vacancy_data = {}
    filled_counts = {}
    cascade_counts = {}
    displaced_counts = {}
    applied_from_counts = {}
    applied_to_counts = {}
    
    # Get all vacancy data in one query
    try:
        vacancy_result = db.session.execute(db.text(f"""
            SELECT district, total_strength, vacancy_reported FROM {prefix}vacancy
        """))
        vacancy_data = {row[0]: {'total_strength': row[1] or 0, 'vacancy_reported': row[2] or 0} 
                       for row in vacancy_result.fetchall()}
    except Exception as e:
        print(f"Error fetching vacancy data: {e}")
    
    # Get applied from counts (employees who applied FROM each district)
    try:
        applied_from_result = db.session.execute(db.text(f"""
            SELECT j.district, COUNT(*) FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
            GROUP BY j.district
        """))
        applied_from_counts = {row[0]: row[1] for row in applied_from_result.fetchall()}
    except Exception as e:
        print(f"Error fetching applied from counts: {e}")
    
    # Get applied to counts (employees who applied TO each district - all preferences)
    applied_to_pref1_counts = {}  # First preference counts
    applied_to_other_counts = {}  # Non-first preference counts
    try:
        # First preference counts
        applied_to_pref1_result = db.session.execute(db.text(f"""
            SELECT t.pref1, COUNT(*) FROM {prefix}transfer_applied t
            WHERE t.pref1 IS NOT NULL AND t.pref1 != ''
            GROUP BY t.pref1
        """))
        applied_to_pref1_counts = {row[0]: row[1] for row in applied_to_pref1_result.fetchall()}
        
        # Non-first preference counts (pref2 to pref8)
        for pref_col in ['pref2', 'pref3', 'pref4', 'pref5', 'pref6', 'pref7', 'pref8']:
            pref_result = db.session.execute(db.text(f"""
                SELECT t.{pref_col}, COUNT(*) FROM {prefix}transfer_applied t
                WHERE t.{pref_col} IS NOT NULL AND t.{pref_col} != ''
                GROUP BY t.{pref_col}
            """))
            for row in pref_result.fetchall():
                district = row[0]
                count = row[1]
                if district not in applied_to_other_counts:
                    applied_to_other_counts[district] = 0
                applied_to_other_counts[district] += count
    except Exception as e:
        print(f"Error fetching applied to counts: {e}")
    
    # Only fetch filled/cascade/displaced if autofill has been run
    if autofill_ran:
        try:
            # Get filled counts (transfers TO each district) in one query
            filled_result = db.session.execute(db.text(f"""
                SELECT transfer_to_district, COUNT(*) FROM {prefix}transfer_draft 
                GROUP BY transfer_to_district
            """))
            filled_counts = {row[0]: row[1] for row in filled_result.fetchall()}
            
            # Get cascade counts (transfers OUT of each district) in one query
            cascade_result = db.session.execute(db.text(f"""
                SELECT j.district, COUNT(*) FROM {prefix}transfer_draft d
                INNER JOIN {prefix}jphn j ON d.pen = j.pen
                WHERE j.district != d.transfer_to_district
                GROUP BY j.district
            """))
            cascade_counts = {row[0]: row[1] for row in cascade_result.fetchall()}
            
            # Get displaced counts (against transfers) in one query
            displaced_result = db.session.execute(db.text(f"""
                SELECT j.district, COUNT(*) FROM {prefix}transfer_draft d
                INNER JOIN {prefix}jphn j ON d.pen = j.pen
                WHERE d.against_info LIKE 'Displaced for%%'
                GROUP BY j.district
            """))
            displaced_counts = {row[0]: row[1] for row in displaced_result.fetchall()}
        except Exception as e:
            print(f"Error fetching transfer draft data: {e}")
    
    vacancies = []
    for district in DISTRICTS:
        vd = vacancy_data.get(district, {'total_strength': 0, 'vacancy_reported': 0})
        total_strength = vd['total_strength']
        vacancy_reported = vd['vacancy_reported']
        
        filled = filled_counts.get(district, 0)
        cascade = cascade_counts.get(district, 0)
        displaced = displaced_counts.get(district, 0)
        applied_from = applied_from_counts.get(district, 0)
        applied_to_pref1 = applied_to_pref1_counts.get(district, 0)
        applied_to_other = applied_to_other_counts.get(district, 0)
        
        # Total available = reported + cascade vacancies
        total_available = vacancy_reported + cascade
        
        # Remaining = total available - filled
        remaining = total_available - filled
        if remaining < 0:
            remaining = 0
        
        vacancies.append({
            'district': district,
            'total_strength': total_strength,
            'vacancy_reported': vacancy_reported,
            'displaced': displaced,
            'applied_from': applied_from,
            'cascade': cascade,
            'total_available': total_available,
            'applied_to_pref1': applied_to_pref1,
            'applied_to_other': applied_to_other,
            'filled': filled,
            'remaining': remaining
        })
    
    return render_template('vacancy.html', vacancies=vacancies, autofill_ran=autofill_ran)


@app.route('/vacancy/save', methods=['POST'])
@login_required
@requires_transfer_session
def save_vacancy():
    """Save vacancy data"""
    prefix = get_table_prefix()
    
    try:
        for district in DISTRICTS:
            strength = request.form.get(f'strength_{district}', 0)
            vacancy = request.form.get(f'vacancy_{district}', 0)
            
            db.session.execute(db.text(f"""
                UPDATE {prefix}vacancy SET 
                total_strength = :strength, vacancy_reported = :vacancy, last_modified = :now
                WHERE district = :district
            """), {
                'strength': int(strength) if strength else 0,
                'vacancy': int(vacancy) if vacancy else 0,
                'now': datetime.now().isoformat(),
                'district': district
            })
        
        db.session.commit()
        flash('Vacancy data saved successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error saving vacancy data: {str(e)}', 'error')
    
    return redirect(url_for('vacancy_list'))


# ==================== TRANSFER APPLICATION ROUTES ====================

@app.route('/application')
@login_required
@requires_transfer_session
def application_list():
    """View transfer applications - employees who haven't applied yet"""
    prefix = get_table_prefix()
    district_filter = request.args.get('district', 'All Districts')  # Default to All Districts
    
    if not district_filter:
        district_filter = 'All Districts'
    
    try:
        # Get last applied employee (by last_modified timestamp)
        last_applied_result = db.session.execute(db.text(f"""
            SELECT j.name, j.pen, j.district
            FROM {prefix}transfer_applied t
            INNER JOIN {prefix}jphn j ON t.pen = j.pen
            ORDER BY t.last_modified DESC NULLS LAST, t.pen DESC LIMIT 1
        """))
        last_applied = last_applied_result.fetchone()
        
        # Check if "All Districts" is selected
        if district_filter == 'All Districts':
            # Get employees NOT in transfer_applied from ALL districts
            result = db.session.execute(db.text(f"""
                SELECT j.pen, j.name, j.institution, j.duration_days, j.district
                FROM {prefix}jphn j
                WHERE j.pen NOT IN (SELECT pen FROM {prefix}transfer_applied)
                ORDER BY j.district, j.duration_days DESC
            """))
            
            employees = result.fetchall()
            
            # Get count of already applied from all districts
            applied_result = db.session.execute(db.text(f"""
                SELECT COUNT(*) FROM {prefix}transfer_applied
            """))
            applied_count = applied_result.fetchone()[0] or 0
        else:
            # Get employees NOT in transfer_applied for this district
            result = db.session.execute(db.text(f"""
                SELECT j.pen, j.name, j.institution, j.duration_days
                FROM {prefix}jphn j
                WHERE j.district = :district
                AND j.pen NOT IN (SELECT pen FROM {prefix}transfer_applied)
                ORDER BY j.duration_days DESC
            """), {'district': district_filter})
            
            employees = result.fetchall()
            
            # Get count of already applied
            applied_result = db.session.execute(db.text(f"""
                SELECT COUNT(*) FROM {prefix}transfer_applied t
                INNER JOIN {prefix}jphn j ON t.pen = j.pen
                WHERE j.district = :district
            """), {'district': district_filter})
            applied_count = applied_result.fetchone()[0] or 0
        
        # Get all applied PENs for comparison feature
        applied_pens_result = db.session.execute(db.text(f"""
            SELECT pen FROM {prefix}transfer_applied
        """))
        applied_pens = [row[0] for row in applied_pens_result.fetchall()]
        
    except Exception as e:
        db.session.rollback()
        employees = []
        applied_count = 0
        last_applied = None
        applied_pens = []
        flash(f'Error loading employees: {str(e)}', 'error')
    
    return render_template('application.html',
                         employees=employees,
                         districts=DISTRICTS,
                         district_filter=district_filter,
                         applied_count=applied_count,
                         last_applied=last_applied,
                         applied_pens=applied_pens,
                         format_duration=format_duration)


@app.route('/application/check/<pen>')
@login_required
@requires_transfer_session
def check_applied(pen):
    """Check if an employee has already applied"""
    prefix = get_table_prefix()
    try:
        result = db.session.execute(db.text(f"""
            SELECT t.pen, j.name, j.district 
            FROM {prefix}transfer_applied t
            INNER JOIN {prefix}jphn j ON t.pen = j.pen
            WHERE t.pen = :pen
        """), {'pen': pen})
        emp = result.fetchone()
        if emp:
            return jsonify({'applied': True, 'name': emp[1], 'district': emp[2]})
        return jsonify({'applied': False})
    except:
        return jsonify({'applied': False})


@app.route('/application/mark', methods=['POST'])
@login_required
@requires_transfer_session
def mark_applied():
    """Mark employees as having applied for transfer"""
    prefix = get_table_prefix()
    
    pens = request.form.getlist('pens')
    receipt_numbers = request.form.get('receipt_numbers', '')
    applied_date = request.form.get('applied_date', datetime.now().strftime('%d-%m-%Y'))
    special_priority = 'Yes' if request.form.get('special_priority') else 'No'
    has_weightage = request.form.get('has_weightage')
    weightage_details = request.form.get('weightage_details', '')
    weightage_priority = int(request.form.get('weightage_priority', 5))
    
    # Get preferences
    prefs = []
    for i in range(1, 9):
        prefs.append(request.form.get(f'pref{i}', ''))
    
    if not pens:
        flash('Please select at least one employee!', 'warning')
        return redirect(url_for('application_list'))
    
    try:
        for pen in pens:
            # Insert into transfer_applied
            db.session.execute(db.text(f"""
                INSERT INTO {prefix}transfer_applied
                (pen, applied_date, receipt_numbers, special_priority,
                 pref1, pref2, pref3, pref4, pref5, pref6, pref7, pref8, last_modified)
                VALUES (:pen, :applied_date, :receipt_numbers, :special_priority,
                        :pref1, :pref2, :pref3, :pref4, :pref5, :pref6, :pref7, :pref8, :now)
                ON CONFLICT (pen) DO UPDATE SET
                applied_date = :applied_date, receipt_numbers = :receipt_numbers,
                special_priority = :special_priority,
                pref1 = :pref1, pref2 = :pref2, pref3 = :pref3, pref4 = :pref4,
                pref5 = :pref5, pref6 = :pref6, pref7 = :pref7, pref8 = :pref8,
                last_modified = :now
            """), {
                'pen': pen, 'applied_date': applied_date, 'receipt_numbers': receipt_numbers,
                'special_priority': special_priority,
                'pref1': prefs[0] or None, 'pref2': prefs[1] or None, 'pref3': prefs[2] or None,
                'pref4': prefs[3] or None, 'pref5': prefs[4] or None, 'pref6': prefs[5] or None,
                'pref7': prefs[6] or None, 'pref8': prefs[7] or None,
                'now': datetime.now().isoformat()
            })
            
            # Update weightage in jphn if needed
            if has_weightage:
                db.session.execute(db.text(f"""
                    UPDATE {prefix}jphn SET weightage = 'Yes', 
                    weightage_details = :details, weightage_priority = :priority
                    WHERE pen = :pen
                """), {
                    'details': weightage_details, 'priority': weightage_priority, 'pen': pen
                })
        
        db.session.commit()
        flash(f'{len(pens)} employee(s) marked as applied!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    # Redirect to All Districts with focus_search parameter to focus search field
    return redirect(url_for('application_list', focus_search='1'))


@app.route('/applied')
@login_required
@requires_transfer_session
def applied_employees():
    """View employees who have applied for transfer"""
    prefix = get_table_prefix()
    district_filter = request.args.get('district', '')
    pref_district = request.args.get('pref_district', '')
    sort_by = request.args.get('sort', '')  # 'to_district' for To District Sort
    
    query = f"""
        SELECT j.pen, j.name, j.institution, j.district, j.duration_days,
               j.district_join_date, j.weightage, j.weightage_details,
               t.receipt_numbers, t.applied_date, t.special_priority,
               t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8,
               j.weightage_priority, COALESCE(t.locked, 'No') as locked
        FROM {prefix}jphn j
        INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
        WHERE 1=1
    """
    params = {}
    
    if district_filter:
        query += " AND j.district = :district"
        params['district'] = district_filter
    
    if pref_district:
        # Find which preference column matches this district for ordering
        query += """ AND (t.pref1 = :pref_district OR t.pref2 = :pref_district OR t.pref3 = :pref_district 
                    OR t.pref4 = :pref_district OR t.pref5 = :pref_district OR t.pref6 = :pref_district 
                    OR t.pref7 = :pref_district OR t.pref8 = :pref_district)"""
        params['pref_district'] = pref_district
    
    # Order: Different sorting options
    if pref_district:
        # When filtering by pref_district, order by preference number (1 first, then 2, etc.), then by seniority
        query += f""" ORDER BY 
            CASE 
                WHEN t.pref1 = :pref_district THEN 1
                WHEN t.pref2 = :pref_district THEN 2
                WHEN t.pref3 = :pref_district THEN 3
                WHEN t.pref4 = :pref_district THEN 4
                WHEN t.pref5 = :pref_district THEN 5
                WHEN t.pref6 = :pref_district THEN 6
                WHEN t.pref7 = :pref_district THEN 7
                WHEN t.pref8 = :pref_district THEN 8
                ELSE 9
            END,
            j.duration_days DESC"""
    elif sort_by == 'to_district':
        # To District Sort: by Pref 1 district (South to North), then by seniority
        query += """ ORDER BY CASE t.pref1
            WHEN 'Thiruvananthapuram' THEN 1
            WHEN 'Kollam' THEN 2
            WHEN 'Pathanamthitta' THEN 3
            WHEN 'Alappuzha' THEN 4
            WHEN 'Kottayam' THEN 5
            WHEN 'Idukki' THEN 6
            WHEN 'Ernakulam' THEN 7
            WHEN 'Thrissur' THEN 8
            WHEN 'Palakkad' THEN 9
            WHEN 'Malappuram' THEN 10
            WHEN 'Kozhikode' THEN 11
            WHEN 'Wayanad' THEN 12
            WHEN 'Kannur' THEN 13
            WHEN 'Kasaragod' THEN 14
            ELSE 15 END,
            j.duration_days DESC"""
    else:
        # Default sort: by current district from South to North, then Special Priority, Weightage Priority, Duration
        query += """ ORDER BY CASE j.district
        WHEN 'Thiruvananthapuram' THEN 1
        WHEN 'Kollam' THEN 2
        WHEN 'Pathanamthitta' THEN 3
        WHEN 'Alappuzha' THEN 4
        WHEN 'Kottayam' THEN 5
        WHEN 'Idukki' THEN 6
        WHEN 'Ernakulam' THEN 7
        WHEN 'Thrissur' THEN 8
        WHEN 'Palakkad' THEN 9
        WHEN 'Malappuram' THEN 10
        WHEN 'Kozhikode' THEN 11
        WHEN 'Wayanad' THEN 12
        WHEN 'Kannur' THEN 13
        WHEN 'Kasaragod' THEN 14
        ELSE 15 END,
        CASE WHEN t.special_priority = 'Yes' THEN 0 ELSE 1 END,
        CASE WHEN j.weightage = 'Yes' THEN 0 ELSE 1 END,
        COALESCE(j.weightage_priority, 5),
        j.duration_days DESC"""
    
    result = db.session.execute(db.text(query), params)
    employees = result.fetchall()
    
    # If filtering by pref_district, add the matched preference number for each employee
    employees_with_pref = []
    pref_counts = {}  # Count employees per preference number
    if pref_district:
        for emp in employees:
            # emp[11-18] are pref1-pref8
            matched_pref = 0
            for i, pref_idx in enumerate(range(11, 19)):
                if emp[pref_idx] == pref_district:
                    matched_pref = i + 1
                    break
            employees_with_pref.append((*emp, matched_pref))
            # Count by preference
            pref_counts[matched_pref] = pref_counts.get(matched_pref, 0) + 1
        employees = employees_with_pref
    
    return render_template('applied_employees.html',
                         employees=employees,
                         districts=DISTRICTS,
                         district_filter=district_filter,
                         pref_district=pref_district,
                         sort_by=sort_by,
                         pref_counts=pref_counts,
                         format_duration=format_duration,
                         now=datetime.now)


@app.route('/applied/remove/<pen>', methods=['POST'])
@login_required
@requires_transfer_session
def remove_from_applied(pen):
    """Remove employee from applied status"""
    prefix = get_table_prefix()
    
    try:
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_applied WHERE pen = :pen"), {'pen': pen})
        db.session.commit()
        flash('Employee removed from applied list!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('applied_employees'))


@app.route('/applied/clear', methods=['POST'])
@login_required
@requires_transfer_session
def clear_applied_list():
    """Clear all applied employees"""
    prefix = get_table_prefix()
    
    try:
        # Count before clearing
        count_result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_applied"))
        count = count_result.fetchone()[0] or 0
        
        # Clear all applied employees
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_applied"))
        db.session.commit()
        
        flash(f'{count} applied employee(s) cleared successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('applied_employees'))


@app.route('/applied/edit/<pen>', methods=['GET', 'POST'])
@login_required
@requires_transfer_session
def edit_applied(pen):
    """Edit applied employee details"""
    prefix = get_table_prefix()
    
    if request.method == 'POST':
        try:
            receipt_numbers = request.form.get('receipt_numbers', '')
            applied_date = request.form.get('applied_date', '')
            special_priority = 'Yes' if request.form.get('special_priority') else 'No'
            has_weightage = request.form.get('has_weightage')
            weightage_details = request.form.get('weightage_details', '')
            weightage_priority = int(request.form.get('weightage_priority', 5))
            
            # Get preferences
            prefs = []
            for i in range(1, 9):
                prefs.append(request.form.get(f'pref{i}', ''))
            
            # Update transfer_applied
            db.session.execute(db.text(f"""
                UPDATE {prefix}transfer_applied
                SET receipt_numbers = :receipt, applied_date = :applied_date,
                    special_priority = :special_priority,
                    pref1 = :p1, pref2 = :p2, pref3 = :p3, pref4 = :p4,
                    pref5 = :p5, pref6 = :p6, pref7 = :p7, pref8 = :p8
                WHERE pen = :pen
            """), {
                'receipt': receipt_numbers, 'applied_date': applied_date,
                'special_priority': special_priority,
                'p1': prefs[0], 'p2': prefs[1], 'p3': prefs[2], 'p4': prefs[3],
                'p5': prefs[4], 'p6': prefs[5], 'p7': prefs[6], 'p8': prefs[7],
                'pen': pen
            })
            
            # Update weightage in jphn table
            if has_weightage:
                db.session.execute(db.text(f"""
                    UPDATE {prefix}jphn
                    SET weightage = 'Yes', weightage_details = :details, weightage_priority = :priority
                    WHERE pen = :pen
                """), {'details': weightage_details, 'priority': weightage_priority, 'pen': pen})
            else:
                db.session.execute(db.text(f"""
                    UPDATE {prefix}jphn
                    SET weightage = 'No', weightage_details = '', weightage_priority = 5
                    WHERE pen = :pen
                """), {'pen': pen})
            
            db.session.commit()
            flash('Application updated successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error: {str(e)}', 'error')
        
        return redirect(url_for('applied_employees'))
    
    # GET request - return employee data as JSON for modal
    try:
        result = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, j.institution, j.district, j.weightage, j.weightage_details, j.weightage_priority,
                   t.receipt_numbers, t.applied_date, t.special_priority,
                   t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
            WHERE j.pen = :pen
        """), {'pen': pen})
        emp = result.fetchone()
        
        if emp:
            return jsonify({
                'success': True,
                'data': {
                    'pen': emp[0], 'name': emp[1], 'institution': emp[2], 'district': emp[3],
                    'weightage': emp[4], 'weightage_details': emp[5] or '', 'weightage_priority': emp[6] or 5,
                    'receipt_numbers': emp[7] or '', 'applied_date': emp[8] or '', 'special_priority': emp[9],
                    'pref1': emp[10] or '', 'pref2': emp[11] or '', 'pref3': emp[12] or '', 'pref4': emp[13] or '',
                    'pref5': emp[14] or '', 'pref6': emp[15] or '', 'pref7': emp[16] or '', 'pref8': emp[17] or ''
                }
            })
        return jsonify({'success': False, 'error': 'Employee not found'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/applied/toggle-lock/<pen>', methods=['POST'])
@login_required
@requires_transfer_session
def toggle_lock_applied(pen):
    """Toggle lock status of an applied employee"""
    prefix = get_table_prefix()
    
    try:
        # Get current lock status
        result = db.session.execute(db.text(f"""
            SELECT COALESCE(locked, 'No') FROM {prefix}transfer_applied WHERE pen = :pen
        """), {'pen': pen})
        row = result.fetchone()
        
        if not row:
            return jsonify({'success': False, 'error': 'Employee not found'})
        
        current_status = row[0]
        new_status = 'No' if current_status == 'Yes' else 'Yes'
        
        db.session.execute(db.text(f"""
            UPDATE {prefix}transfer_applied 
            SET locked = :status, last_modified = :now 
            WHERE pen = :pen
        """), {'status': new_status, 'now': datetime.now().isoformat(), 'pen': pen})
        db.session.commit()
        
        return jsonify({'success': True, 'locked': new_status == 'Yes'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})


@app.route('/applied/unlock-all', methods=['POST'])
@login_required
@requires_transfer_session
def unlock_all_applied():
    """Unlock all applied employees - make them available for Auto-fill"""
    try:
        prefix = get_table_prefix()
        
        # Count how many were locked before unlocking
        result = db.session.execute(db.text(f"""
            SELECT COUNT(*) FROM {prefix}transfer_applied WHERE locked = 'Yes'
        """))
        count = result.fetchone()[0]
        
        # Unlock all
        db.session.execute(db.text(f"""
            UPDATE {prefix}transfer_applied 
            SET locked = 'No', last_modified = :now 
            WHERE locked = 'Yes'
        """), {'now': datetime.now().isoformat()})
        db.session.commit()
        
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})


@app.route('/applied/export-excel')
@login_required
@requires_transfer_session
def export_applied_excel():
    """Export applied employees list to Excel with proper formatting (respects current filters and sort)"""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        flash('openpyxl library not installed', 'error')
        return redirect(url_for('applied_employees'))
    
    prefix = get_table_prefix()
    
    # Get filter and sort parameters from request (same as applied_employees view)
    district_filter = request.args.get('district', '')
    pref_district = request.args.get('pref_district', '')
    sort_by = request.args.get('sort', '')  # 'to_district' for To District Sort
    
    query = f"""
        SELECT j.pen, j.name, j.institution, j.district, j.duration_days,
               j.weightage, j.weightage_details, t.receipt_numbers,
               t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8,
               t.special_priority, j.weightage_priority
        FROM {prefix}jphn j
        INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
        WHERE 1=1
    """
    params = {}
    
    # Apply the same filters as the Applied Employees view
    if district_filter:
        query += " AND j.district = :district"
        params['district'] = district_filter
    
    if pref_district:
        # Match across ALL preference columns (same as the view)
        query += """ AND (t.pref1 = :pref_district OR t.pref2 = :pref_district OR t.pref3 = :pref_district 
                    OR t.pref4 = :pref_district OR t.pref5 = :pref_district OR t.pref6 = :pref_district 
                    OR t.pref7 = :pref_district OR t.pref8 = :pref_district)"""
        params['pref_district'] = pref_district
    
    # Order: Apply same sorting as the view
    if pref_district:
        # When filtering by pref_district, order by preference number (1 first, then 2, etc.), then by seniority
        query += f""" ORDER BY 
            CASE 
                WHEN t.pref1 = :pref_district THEN 1
                WHEN t.pref2 = :pref_district THEN 2
                WHEN t.pref3 = :pref_district THEN 3
                WHEN t.pref4 = :pref_district THEN 4
                WHEN t.pref5 = :pref_district THEN 5
                WHEN t.pref6 = :pref_district THEN 6
                WHEN t.pref7 = :pref_district THEN 7
                WHEN t.pref8 = :pref_district THEN 8
                ELSE 9
            END,
            j.duration_days DESC"""
    elif sort_by == 'to_district':
        # To District Sort: by Pref 1 district (South to North), then by seniority
        query += """ ORDER BY CASE t.pref1
            WHEN 'Thiruvananthapuram' THEN 1
            WHEN 'Kollam' THEN 2
            WHEN 'Pathanamthitta' THEN 3
            WHEN 'Alappuzha' THEN 4
            WHEN 'Kottayam' THEN 5
            WHEN 'Idukki' THEN 6
            WHEN 'Ernakulam' THEN 7
            WHEN 'Thrissur' THEN 8
            WHEN 'Palakkad' THEN 9
            WHEN 'Malappuram' THEN 10
            WHEN 'Kozhikode' THEN 11
            WHEN 'Wayanad' THEN 12
            WHEN 'Kannur' THEN 13
            WHEN 'Kasaragod' THEN 14
            ELSE 15 END,
            j.duration_days DESC"""
    else:
        # Default (From District Sort): by current district (South to North), then Special Priority, Weightage, Duration
        query += """ ORDER BY CASE j.district
            WHEN 'Thiruvananthapuram' THEN 1
            WHEN 'Kollam' THEN 2
            WHEN 'Pathanamthitta' THEN 3
            WHEN 'Alappuzha' THEN 4
            WHEN 'Kottayam' THEN 5
            WHEN 'Idukki' THEN 6
            WHEN 'Ernakulam' THEN 7
            WHEN 'Thrissur' THEN 8
            WHEN 'Palakkad' THEN 9
            WHEN 'Malappuram' THEN 10
            WHEN 'Kozhikode' THEN 11
            WHEN 'Wayanad' THEN 12
            WHEN 'Kannur' THEN 13
            WHEN 'Kasaragod' THEN 14
            ELSE 15 END,
            CASE WHEN t.special_priority = 'Yes' THEN 0 ELSE 1 END,
            CASE WHEN j.weightage = 'Yes' THEN 0 ELSE 1 END,
            COALESCE(j.weightage_priority, 5),
            j.duration_days DESC
    """
    
    result = db.session.execute(db.text(query), params)
    employees = result.fetchall()
    
    # If filtering by pref_district, calculate matched preference number for each employee
    if pref_district:
        employees_with_pref = []
        for emp in employees:
            # emp[8-15] are pref1-pref8
            matched_pref = 0
            for i, pref_idx in enumerate(range(8, 16)):
                if emp[pref_idx] == pref_district:
                    matched_pref = i + 1
                    break
            employees_with_pref.append((*emp, matched_pref))
        employees = employees_with_pref
    
    # Create workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Applied Employees'
    
    # Styles
    header_font = Font(bold=True, size=14)
    subheader_font = Font(bold=True, size=11)
    table_header_font = Font(bold=True, size=10)
    table_header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # Set column widths
    ws.column_dimensions['A'].width = 8    # Sl. No.
    ws.column_dimensions['B'].width = 12   # PEN
    ws.column_dimensions['C'].width = 25   # Name
    ws.column_dimensions['D'].width = 30   # Institution
    ws.column_dimensions['E'].width = 18   # District
    ws.column_dimensions['F'].width = 12   # Duration
    ws.column_dimensions['G'].width = 10   # Weightage
    ws.column_dimensions['H'].width = 20   # Reason
    ws.column_dimensions['I'].width = 12   # Receipt
    ws.column_dimensions['J'].width = 15   # Pref 1
    ws.column_dimensions['K'].width = 15   # Pref 2
    ws.column_dimensions['L'].width = 15   # Pref 3
    ws.column_dimensions['M'].width = 15   # Pref 4
    ws.column_dimensions['N'].width = 15   # Pref 5
    ws.column_dimensions['O'].width = 15   # Pref 6
    ws.column_dimensions['P'].width = 15   # Pref 7
    ws.column_dimensions['Q'].width = 15   # Pref 8
    if pref_district:
        ws.column_dimensions['R'].width = 10   # Pref #
    
    # Determine last column for merge
    last_col = 'R' if pref_district else 'Q'
    
    # Title rows
    ws.merge_cells(f'A1:{last_col}1')
    ws['A1'] = 'JPHN Transfer Management System'
    ws['A1'].font = header_font
    ws['A1'].alignment = center_align
    
    # Build subtitle based on filters and sort
    subtitle = 'Applied Employees List'
    filter_parts = []
    if district_filter:
        filter_parts.append(f'From District: {district_filter}')
    if pref_district:
        filter_parts.append(f'To District: {pref_district}')
    if sort_by == 'to_district':
        filter_parts.append('Sorted by: To District')
    elif not pref_district:
        filter_parts.append('Sorted by: From District')
    if filter_parts:
        subtitle += ' (' + ', '.join(filter_parts) + ')'
    
    ws.merge_cells(f'A2:{last_col}2')
    ws['A2'] = subtitle
    ws['A2'].font = subheader_font
    ws['A2'].alignment = center_align
    
    ws.merge_cells(f'A3:{last_col}3')
    ws['A3'] = f'Generated: {datetime.now().strftime("%d-%m-%Y %I:%M %p")}'
    ws['A3'].alignment = center_align
    
    # Column headers (row 5)
    headers = ['Sl. No.', 'PEN', 'Name', 'Institution', 'Current District', 'Duration',
               'Weightage', 'Reason', 'Receipt No.', 'Pref 1', 'Pref 2', 'Pref 3', 'Pref 4',
               'Pref 5', 'Pref 6', 'Pref 7', 'Pref 8']
    if pref_district:
        headers.append('Pref #')
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col_num, value=header)
        cell.font = table_header_font
        cell.fill = table_header_fill
        cell.border = thin_border
        cell.alignment = center_align
    
    # Data rows
    for idx, emp in enumerate(employees, 1):
        row_num = idx + 5
        duration_days = emp[4] or 0
        years = duration_days // 365
        months = (duration_days % 365) // 30
        days = duration_days % 30
        duration_str = ''
        if years > 0:
            duration_str += f'{years}Y '
        if months > 0:
            duration_str += f'{months}M '
        if days > 0:
            duration_str += f'{days}D'
        duration_str = duration_str.strip() or '-'
        
        data = [
            idx,                       # Sl. No.
            emp[0],                    # PEN
            emp[1],                    # Name
            emp[2],                    # Institution
            emp[3],                    # District
            duration_str,              # Duration
            emp[5] or 'No',            # Weightage
            emp[6] or 'NA',            # Reason
            emp[7] or '',              # Receipt
            emp[8] or '',              # Pref 1
            emp[9] or '',              # Pref 2
            emp[10] or '',             # Pref 3
            emp[11] or '',             # Pref 4
            emp[12] or '',             # Pref 5
            emp[13] or '',             # Pref 6
            emp[14] or '',             # Pref 7
            emp[15] or ''              # Pref 8
        ]
        if pref_district:
            data.append(emp[18])       # Pref # (matched preference number)
        
        for col_num, value in enumerate(data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.border = thin_border
            if col_num == 1 or (pref_district and col_num == len(data)):  # Sl. No. and Pref # - center aligned
                cell.alignment = center_align
            else:
                cell.alignment = left_align
    
    # Save to BytesIO
    from io import BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Build filename based on filters
    filename_parts = ['applied_employees']
    if district_filter:
        filename_parts.append(district_filter.replace(' ', '_'))
    if pref_district:
        filename_parts.append(f'pref1_{pref_district.replace(" ", "_")}')
    filename_parts.append(datetime.now().strftime("%Y%m%d"))
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'{"_".join(filename_parts)}.xlsx'
    )


# ==================== DRAFT LIST ROUTES ====================

@app.route('/draft')
@login_required
@requires_transfer_session
def draft_list():
    """View draft transfer list"""
    prefix = get_table_prefix()
    
    # Check if autofill has been run
    if not check_autofill_ran():
        flash('Please run Auto-fill Vacancies first from the Applied Employees page before viewing Draft List.', 'warning')
        return redirect(url_for('applied_employees'))
    
    from_district = request.args.get('from_district', '')
    to_district = request.args.get('to_district', '')
    sort_by = request.args.get('sort', '')  # 'to_district' for To District Sort
    
    query = f"""
        SELECT j.pen, j.name, j.institution, j.district, d.transfer_to_district,
               j.district_join_date, j.duration_days, j.weightage, j.weightage_details,
               d.against_info, d.remarks,
               t.special_priority, t.pref1, t.pref2, t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
        FROM {prefix}jphn j
        INNER JOIN {prefix}transfer_draft d ON j.pen = d.pen
        LEFT JOIN {prefix}transfer_applied t ON j.pen = t.pen
        WHERE 1=1
    """
    params = {}
    
    if from_district:
        query += " AND j.district = :from_district"
        params['from_district'] = from_district
    
    if to_district:
        query += " AND d.transfer_to_district = :to_district"
        params['to_district'] = to_district
    
    # Order based on sort option
    if sort_by == 'to_district':
        # To District Sort: by transfer_to_district (South to North), then by seniority
        query += """ ORDER BY CASE d.transfer_to_district
            WHEN 'Thiruvananthapuram' THEN 1
            WHEN 'Kollam' THEN 2
            WHEN 'Pathanamthitta' THEN 3
            WHEN 'Alappuzha' THEN 4
            WHEN 'Kottayam' THEN 5
            WHEN 'Idukki' THEN 6
            WHEN 'Ernakulam' THEN 7
            WHEN 'Thrissur' THEN 8
            WHEN 'Palakkad' THEN 9
            WHEN 'Malappuram' THEN 10
            WHEN 'Kozhikode' THEN 11
            WHEN 'Wayanad' THEN 12
            WHEN 'Kannur' THEN 13
            WHEN 'Kasaragod' THEN 14
            ELSE 15 END, j.duration_days DESC"""
    else:
        # Default: by current district (South to North), then by seniority
        query += """ ORDER BY CASE j.district
            WHEN 'Thiruvananthapuram' THEN 1
            WHEN 'Kollam' THEN 2
            WHEN 'Pathanamthitta' THEN 3
            WHEN 'Alappuzha' THEN 4
            WHEN 'Kottayam' THEN 5
            WHEN 'Idukki' THEN 6
            WHEN 'Ernakulam' THEN 7
            WHEN 'Thrissur' THEN 8
            WHEN 'Palakkad' THEN 9
            WHEN 'Malappuram' THEN 10
            WHEN 'Kozhikode' THEN 11
            WHEN 'Wayanad' THEN 12
            WHEN 'Kannur' THEN 13
            WHEN 'Kasaragod' THEN 14
            ELSE 15 END, j.duration_days DESC"""
    
    result = db.session.execute(db.text(query), params)
    transfers = result.fetchall()
    
    return render_template('draft_list.html',
                         transfers=transfers,
                         districts=DISTRICTS,
                         from_district=from_district,
                         to_district=to_district,
                         sort_by=sort_by,
                         format_duration=format_duration,
                         now=datetime.now)


@app.route('/draft/add', methods=['POST'])
@login_required
@requires_transfer_session
def add_to_draft():
    """Add employee(s) to draft transfer list"""
    prefix = get_table_prefix()
    
    pens = request.form.getlist('pens')
    transfer_district = request.form.get('transfer_district', '')
    
    if not pens or not transfer_district:
        flash('Please select employee(s) and transfer district!', 'warning')
        return redirect(url_for('cadre_list'))
    
    # Check vacancy
    vacancy_result = db.session.execute(db.text(f"""
        SELECT vacancy_reported FROM {prefix}vacancy WHERE district = :district
    """), {'district': transfer_district})
    vacancy_row = vacancy_result.fetchone()
    vacancy_reported = vacancy_row[0] if vacancy_row else 0
    
    filled_result = db.session.execute(db.text(f"""
        SELECT COUNT(*) FROM {prefix}transfer_draft WHERE transfer_to_district = :district
    """), {'district': transfer_district})
    filled = filled_result.fetchone()[0] or 0
    
    remaining = vacancy_reported - filled
    
    if vacancy_reported > 0 and remaining < len(pens):
        flash(f'Insufficient vacancy! Only {remaining} position(s) available in {transfer_district}.', 'error')
        return redirect(url_for('cadre_list'))
    
    try:
        added = 0
        for pen in pens:
            # Check if already in draft
            check = db.session.execute(db.text(f"""
                SELECT pen FROM {prefix}transfer_draft WHERE pen = :pen
            """), {'pen': pen})
            
            if check.fetchone():
                continue
            
            db.session.execute(db.text(f"""
                INSERT INTO {prefix}transfer_draft (pen, transfer_to_district, added_date, last_modified)
                VALUES (:pen, :district, :added_date, :now)
            """), {
                'pen': pen, 'district': transfer_district,
                'added_date': datetime.now().strftime('%d-%m-%Y'),
                'now': datetime.now().isoformat()
            })
            added += 1
        
        db.session.commit()
        flash(f'{added} employee(s) added to draft list!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('draft_list'))


@app.route('/draft/remove/<pen>', methods=['POST'])
@login_required
@requires_transfer_session
def remove_from_draft(pen):
    """Remove employee from draft list"""
    prefix = get_table_prefix()
    
    try:
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_draft WHERE pen = :pen"), {'pen': pen})
        db.session.commit()
        flash('Employee removed from draft list!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('draft_list'))


@app.route('/draft/auto-fill', methods=['POST'])
@login_required
@requires_transfer_session
def auto_fill_vacancies():
    """Auto-fill vacancies based on preferences and priorities"""
    prefix = get_table_prefix()
    enable_against = request.form.get('enable_against') == 'on'
    
    try:
        # Check if there are any applied employees with preferences
        pref_check = db.session.execute(db.text(f"""
            SELECT COUNT(*) FROM {prefix}transfer_applied 
            WHERE pref1 IS NOT NULL OR pref2 IS NOT NULL OR pref3 IS NOT NULL OR pref4 IS NOT NULL
               OR pref5 IS NOT NULL OR pref6 IS NOT NULL OR pref7 IS NOT NULL OR pref8 IS NOT NULL
        """))
        has_preferences = pref_check.fetchone()[0] > 0
        
        if not has_preferences:
            flash('No employees have set their district preferences yet. Please ensure employees have filled their preferences before auto-filling.', 'warning')
            return redirect(url_for('applied_employees'))
        
        # Get vacancy status
        vacancy_status = {}
        for district in DISTRICTS:
            v_result = db.session.execute(db.text(f"""
                SELECT vacancy_reported FROM {prefix}vacancy WHERE district = :district
            """), {'district': district})
            v_row = v_result.fetchone()
            vacancy_reported = v_row[0] if v_row else 0
            
            f_result = db.session.execute(db.text(f"""
                SELECT COUNT(*) FROM {prefix}transfer_draft WHERE transfer_to_district = :district
            """), {'district': district})
            filled = f_result.fetchone()[0] or 0
            
            vacancy_status[district] = {
                'reported': vacancy_reported,
                'filled': filled,
                'remaining': vacancy_reported - filled if vacancy_reported > 0 else 0,
                'cascade': 0
            }
        
        allocated_count = 0
        special_count = 0
        weightage_count = 0
        normal_count = 0
        cascade_count = 0
        against_count = 0
        not_allocated_count = 0
        
        def allocate(pen, name, current_district, to_district, remarks=None):
            nonlocal allocated_count, cascade_count
            
            # Check if already allocated (ON CONFLICT DO NOTHING handles this but we need accurate count)
            check_result = db.session.execute(db.text(f"""
                SELECT 1 FROM {prefix}transfer_draft WHERE pen = :pen
            """), {'pen': pen})
            if check_result.fetchone():
                return False
            
            # Determine if it's a cascade allocation
            is_cascade = vacancy_status[to_district]['remaining'] <= 0 and vacancy_status[to_district]['cascade'] > 0
            if is_cascade and remarks is None:
                remarks = "Vacancy by Transfer"
            
            db.session.execute(db.text(f"""
                INSERT INTO {prefix}transfer_draft (pen, transfer_to_district, added_date, remarks, last_modified)
                VALUES (:pen, :to_district, :date, :remarks, :now)
            """), {
                'pen': pen, 'to_district': to_district,
                'date': datetime.now().strftime('%d-%m-%Y'),
                'remarks': remarks, 'now': datetime.now().isoformat()
            })
            
            # Update vacancy status for destination
            if vacancy_status[to_district]['remaining'] > 0:
                vacancy_status[to_district]['remaining'] -= 1
            elif vacancy_status[to_district]['cascade'] > 0:
                vacancy_status[to_district]['cascade'] -= 1
                cascade_count += 1
            
            vacancy_status[to_district]['filled'] += 1
            
            # Create cascading vacancy in source district
            if current_district and current_district != to_district:
                if current_district in vacancy_status:
                    vacancy_status[current_district]['cascade'] += 1
            
            allocated_count += 1
            return True
        
        def try_allocate(pen, name, current_district, preferences):
            for pref_idx, pref in enumerate(preferences):
                if not pref:
                    continue
                if pref in vacancy_status:
                    available = vacancy_status[pref]['remaining'] + vacancy_status[pref]['cascade']
                    if available > 0 or vacancy_status[pref]['reported'] == 0:
                        is_cascade = vacancy_status[pref]['remaining'] <= 0 and vacancy_status[pref]['cascade'] > 0
                        remarks = "Vacancy by Transfer" if is_cascade else None
                        if allocate(pen, name, current_district, pref, remarks):
                            return True, pref_idx
            return False, -1
        
        def try_against_transfer(pen, name, pref_district):
            """Try to create an 'Against' transfer by displacing senior employee"""
            nonlocal against_count
            
            # Find the most senior employee in the preferred district who is NOT in draft or applied
            senior_result = db.session.execute(db.text(f"""
                SELECT j.pen, j.name
                FROM {prefix}jphn j
                WHERE j.district = :district
                  AND j.pen != :applicant_pen
                  AND j.pen NOT IN (SELECT pen FROM {prefix}transfer_draft)
                  AND j.pen NOT IN (SELECT pen FROM {prefix}transfer_applied)
                ORDER BY j.duration_days DESC
                LIMIT 1
            """), {'district': pref_district, 'applicant_pen': pen})
            
            senior = senior_result.fetchone()
            if not senior:
                return False, None
            
            senior_pen, senior_name = senior[0], senior[1]
            
            # Find a nearby district with vacancy for the senior employee
            nearby_districts = get_nearby_districts(pref_district)
            target_district = None
            
            for nearby in nearby_districts:
                if nearby in vacancy_status:
                    avail = vacancy_status[nearby]['remaining'] + vacancy_status[nearby]['cascade']
                    if avail > 0:
                        target_district = nearby
                        break
            
            if not target_district:
                return False, None
            
            # Transfer senior employee to nearby district
            allocate(senior_pen, senior_name, pref_district, target_district, f"Displaced for {name}")
            
            # Now allocate the applicant to their preferred district (cascade vacancy created)
            allocate(pen, name, None, pref_district, f"Against {senior_name}")
            against_count += 1
            
            return True, senior_name
        
        # STEP 0: Process employees WITH SPECIAL PRIORITY first (HIGHEST PRIORITY)
        special_result = db.session.execute(db.text(f"""
            SELECT t.pen, j.name, j.district, t.pref1, t.pref2, t.pref3, t.pref4,
                   t.pref5, t.pref6, t.pref7, t.pref8
            FROM {prefix}transfer_applied t
            INNER JOIN {prefix}jphn j ON t.pen = j.pen
            WHERE t.special_priority = 'Yes'
            AND (t.locked IS NULL OR t.locked != 'Yes')
            AND t.pen NOT IN (SELECT pen FROM {prefix}transfer_draft)
            AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                 OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
            ORDER BY j.duration_days DESC
        """))
        
        for row in special_result.fetchall():
            pen, name, district = row[0], row[1], row[2]
            prefs = [row[i] for i in range(3, 11)]
            allocated, _ = try_allocate(pen, name, district, prefs)
            if allocated:
                special_count += 1
            elif prefs[0] and enable_against:
                success, _ = try_against_transfer(pen, name, prefs[0])
                if success:
                    special_count += 1
                else:
                    not_allocated_count += 1
            else:
                not_allocated_count += 1
        
        # STEP 1: Process WEIGHTAGE employees
        weightage_result = db.session.execute(db.text(f"""
            SELECT t.pen, j.name, j.district, j.weightage_priority, t.pref1, t.pref2, 
                   t.pref3, t.pref4, t.pref5, t.pref6, t.pref7, t.pref8
            FROM {prefix}transfer_applied t
            INNER JOIN {prefix}jphn j ON t.pen = j.pen
            WHERE j.weightage = 'Yes'
            AND (t.special_priority IS NULL OR t.special_priority != 'Yes')
            AND (t.locked IS NULL OR t.locked != 'Yes')
            AND t.pen NOT IN (SELECT pen FROM {prefix}transfer_draft)
            AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                 OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
            ORDER BY COALESCE(j.weightage_priority, 5) ASC, j.duration_days ASC
        """))
        
        for row in weightage_result.fetchall():
            pen, name, district = row[0], row[1], row[2]
            prefs = [row[i] for i in range(4, 12)]
            allocated, _ = try_allocate(pen, name, district, prefs)
            if allocated:
                weightage_count += 1
            elif prefs[0] and enable_against:
                success, _ = try_against_transfer(pen, name, prefs[0])
                if success:
                    weightage_count += 1
                else:
                    not_allocated_count += 1
            else:
                not_allocated_count += 1
        
        # STEP 2: Process NORMAL employees (junior first)
        normal_result = db.session.execute(db.text(f"""
            SELECT t.pen, j.name, j.district, t.pref1, t.pref2, t.pref3, t.pref4,
                   t.pref5, t.pref6, t.pref7, t.pref8
            FROM {prefix}transfer_applied t
            INNER JOIN {prefix}jphn j ON t.pen = j.pen
            WHERE (j.weightage IS NULL OR j.weightage != 'Yes')
            AND (t.special_priority IS NULL OR t.special_priority != 'Yes')
            AND (t.locked IS NULL OR t.locked != 'Yes')
            AND t.pen NOT IN (SELECT pen FROM {prefix}transfer_draft)
            AND (t.pref1 IS NOT NULL OR t.pref2 IS NOT NULL OR t.pref3 IS NOT NULL OR t.pref4 IS NOT NULL
                 OR t.pref5 IS NOT NULL OR t.pref6 IS NOT NULL OR t.pref7 IS NOT NULL OR t.pref8 IS NOT NULL)
            ORDER BY j.duration_days ASC
        """))
        
        for row in normal_result.fetchall():
            pen, name, district = row[0], row[1], row[2]
            prefs = [row[i] for i in range(3, 11)]
            allocated, _ = try_allocate(pen, name, district, prefs)
            if allocated:
                normal_count += 1
            elif prefs[0] and enable_against:
                success, _ = try_against_transfer(pen, name, prefs[0])
                if success:
                    normal_count += 1
                else:
                    not_allocated_count += 1
            else:
                not_allocated_count += 1
        
        # Commit all changes
        try:
            db.session.commit()
        except Exception as commit_error:
            db.session.rollback()
            flash(f'Error committing changes: {str(commit_error)}', 'error')
            return redirect(url_for('applied_employees'))
        
        # Mark autofill as run in session
        session['autofill_ran'] = True
        
        # Build result message
        msg = f'Auto-fill complete! Total Allocated: {allocated_count}'
        msg += f' (Special: {special_count}, Weightage: {weightage_count}, Normal: {normal_count}'
        if cascade_count > 0:
            msg += f', Cascade: {cascade_count}'
        if against_count > 0:
            msg += f', Against: {against_count}'
        msg += ')'
        if not_allocated_count > 0:
            msg += f'. Not allocated: {not_allocated_count} (no vacancy in preferred districts)'
        
        flash(msg, 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error during auto-fill: {str(e)}', 'error')
    
    return redirect(url_for('draft_list'))


@app.route('/draft/clear', methods=['POST'])
@login_required
@requires_transfer_session
def clear_draft_list():
    """Clear all draft transfers"""
    prefix = get_table_prefix()
    
    try:
        # Count before clearing
        count_result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_draft"))
        count = count_result.fetchone()[0] or 0
        
        # Clear all draft transfers
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_draft"))
        db.session.commit()
        
        # Clear autofill flag
        session.pop('autofill_ran', None)
        
        flash(f'{count} draft transfer(s) cleared successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('draft_list'))


@app.route('/final/clear', methods=['POST'])
@login_required
@requires_transfer_session
def clear_final_list():
    """Clear all final transfers"""
    prefix = get_table_prefix()
    
    try:
        # Count before clearing
        count_result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final"))
        count = count_result.fetchone()[0] or 0
        
        # Clear all final transfers
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_final"))
        db.session.commit()
        
        flash(f'{count} final transfer(s) cleared successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('final_list'))


@app.route('/final/revert', methods=['POST'])
@login_required
@requires_transfer_session
def revert_to_draft():
    """Revert from final list back to draft list"""
    prefix = get_table_prefix()
    
    try:
        # Count final transfers
        count_result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final"))
        count = count_result.fetchone()[0] or 0
        
        # Clear the final list (draft list is still intact)
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_final"))
        db.session.commit()
        
        flash(f'Reverted to draft list. {count} final transfer(s) removed. You can now make changes in the Draft List.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('draft_list'))


@app.route('/draft/confirm', methods=['POST'])
@login_required
@requires_transfer_session
def confirm_transfers():
    """Confirm draft transfers and move to final list"""
    prefix = get_table_prefix()
    
    try:
        # Clear previous final
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_final"))
        
        # Copy from draft to final
        db.session.execute(db.text(f"""
            INSERT INTO {prefix}transfer_final (pen, transfer_to_district, confirmed_date, last_modified)
            SELECT pen, transfer_to_district, :date, :now FROM {prefix}transfer_draft
        """), {
            'date': datetime.now().strftime('%d-%m-%Y'),
            'now': datetime.now().isoformat()
        })
        
        db.session.commit()
        
        # Count transfers
        count_result = db.session.execute(db.text(f"SELECT COUNT(*) FROM {prefix}transfer_final"))
        count = count_result.fetchone()[0] or 0
        
        flash(f'{count} transfer(s) confirmed successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('final_list'))


# ==================== FINAL LIST ROUTES ====================

@app.route('/final')
@login_required
@requires_transfer_session
def final_list():
    """View confirmed transfer list"""
    prefix = get_table_prefix()
    
    # Check if final list has been created
    if not check_final_exists():
        flash('Please confirm the Draft List first to create the Final List.', 'warning')
        return redirect(url_for('draft_list') if check_autofill_ran() else url_for('applied_employees'))
    
    from_district = request.args.get('from_district', '')
    to_district = request.args.get('to_district', '')
    
    query = f"""
        SELECT j.pen, j.name, j.institution, j.district, f.transfer_to_district,
               j.district_join_date, j.duration_days, j.weightage, j.weightage_details
        FROM {prefix}jphn j
        INNER JOIN {prefix}transfer_final f ON j.pen = f.pen
        WHERE 1=1
    """
    params = {}
    
    if from_district:
        query += " AND j.district = :from_district"
        params['from_district'] = from_district
    
    if to_district:
        query += " AND f.transfer_to_district = :to_district"
        params['to_district'] = to_district
    
    # Order by transfer_to_district from South to North (Kerala geography)
    query += """ ORDER BY CASE f.transfer_to_district
        WHEN 'Thiruvananthapuram' THEN 1
        WHEN 'Kollam' THEN 2
        WHEN 'Pathanamthitta' THEN 3
        WHEN 'Alappuzha' THEN 4
        WHEN 'Kottayam' THEN 5
        WHEN 'Idukki' THEN 6
        WHEN 'Ernakulam' THEN 7
        WHEN 'Thrissur' THEN 8
        WHEN 'Palakkad' THEN 9
        WHEN 'Malappuram' THEN 10
        WHEN 'Kozhikode' THEN 11
        WHEN 'Wayanad' THEN 12
        WHEN 'Kannur' THEN 13
        WHEN 'Kasaragod' THEN 14
        ELSE 15 END, j.duration_days DESC"""
    
    result = db.session.execute(db.text(query), params)
    transfers = result.fetchall()
    
    return render_template('final_list.html',
                         transfers=transfers,
                         districts=DISTRICTS,
                         from_district=from_district,
                         to_district=to_district,
                         format_duration=format_duration)


@app.route('/final/delete/<pen>', methods=['POST'])
@login_required
@requires_transfer_session
def delete_from_final(pen):
    """Delete employee from final list"""
    prefix = get_table_prefix()
    
    try:
        db.session.execute(db.text(f"DELETE FROM {prefix}transfer_final WHERE pen = :pen"), {'pen': pen})
        db.session.commit()
        flash('Transfer deleted from final list!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    
    return redirect(url_for('final_list'))


# ==================== EXCLUDED EMPLOYEES ROUTES ====================

@app.route('/draft/excluded')
@login_required
@requires_transfer_session
def get_excluded_from_draft():
    """Get employees who applied but were NOT included in draft list"""
    prefix = get_table_prefix()
    
    try:
        result = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, j.institution, j.district, j.district_join_date, 
                   j.duration_days, j.weightage, t.pref1, t.pref2, t.pref3
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
            WHERE j.pen NOT IN (SELECT pen FROM {prefix}transfer_draft)
            ORDER BY j.district, j.duration_days DESC
        """))
        employees = result.fetchall()
        
        data = [{
            'pen': emp[0],
            'name': emp[1],
            'institution': emp[2],
            'district': emp[3],
            'district_join_date': emp[4] or '',
            'duration': format_duration(emp[5]) if emp[5] else '',
            'weightage': emp[6] or 'No',
            'pref1': emp[7] or '',
            'pref2': emp[8] or '',
            'pref3': emp[9] or ''
        } for emp in employees]
        
        return jsonify({'success': True, 'employees': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/final/excluded')
@login_required
@requires_transfer_session
def get_excluded_from_final():
    """Get employees who applied but were NOT included in final list"""
    prefix = get_table_prefix()
    
    try:
        result = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, j.institution, j.district, j.district_join_date, 
                   j.duration_days, j.weightage, t.pref1, t.pref2, t.pref3
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
            WHERE j.pen NOT IN (SELECT pen FROM {prefix}transfer_final)
            ORDER BY j.district, j.duration_days DESC
        """))
        employees = result.fetchall()
        
        data = [{
            'pen': emp[0],
            'name': emp[1],
            'institution': emp[2],
            'district': emp[3],
            'district_join_date': emp[4] or '',
            'duration': format_duration(emp[5]) if emp[5] else '',
            'weightage': emp[6] or 'No',
            'pref1': emp[7] or '',
            'pref2': emp[8] or '',
            'pref3': emp[9] or ''
        } for emp in employees]
        
        return jsonify({'success': True, 'employees': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


# ==================== EXPORT ROUTES ====================

@app.route('/export/csv/<list_type>')
@login_required
@requires_transfer_session
def export_csv(list_type):
    """Export list to CSV"""
    prefix = get_table_prefix()
    
    if list_type == 'cadre':
        result = db.session.execute(db.text(f"""
            SELECT pen, name, designation, institution, district, entry_date,
                   retirement_date, district_join_date, duration_days, contact
            FROM {prefix}jphn ORDER BY district, name
        """))
        headers = ['PEN', 'Name', 'Designation', 'Institution', 'District', 
                  'Entry Date', 'Retirement Date', 'District Join Date', 'Duration', 'Contact']
    elif list_type == 'draft':
        result = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, j.institution, j.district, d.transfer_to_district,
                   j.duration_days, j.weightage
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_draft d ON j.pen = d.pen
            ORDER BY d.transfer_to_district, j.name
        """))
        headers = ['PEN', 'Name', 'Institution', 'From District', 'To District', 'Duration', 'Weightage']
    elif list_type == 'final':
        result = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, j.institution, j.district, f.transfer_to_district,
                   j.duration_days, j.weightage
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_final f ON j.pen = f.pen
            ORDER BY f.transfer_to_district, j.name
        """))
        headers = ['PEN', 'Name', 'Institution', 'From District', 'To District', 'Duration', 'Weightage']
    else:
        flash('Invalid export type!', 'error')
        return redirect(url_for('dashboard'))
    
    rows = result.fetchall()
    
    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([str(cell) if cell else '' for cell in row])
    
    output.seek(0)
    
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8-sig')),
        mimetype='text/csv',
        as_attachment=True,
        download_name=f'{list_type}_list_{datetime.now().strftime("%Y%m%d")}.csv'
    )


@app.route('/export/excel/<list_type>')
@login_required
@requires_transfer_session
def export_excel(list_type):
    """Export list to Excel - matching desktop app format"""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        flash('openpyxl library not installed. Please install with: pip install openpyxl', 'error')
        return redirect(url_for('dashboard'))
    
    prefix = get_table_prefix()
    transfer_type = session.get('transfer_type', 'general')
    year = session.get('year', '')
    month = session.get('month', '')
    
    if list_type == 'draft':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, d.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_draft d ON j.pen = d.pen
            ORDER BY j.district, j.name
        """
    elif list_type == 'final':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, f.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_final f ON j.pen = f.pen
            ORDER BY j.district, j.name
        """
    else:
        flash('Invalid export type!', 'error')
        return redirect(url_for('dashboard'))
    
    result = db.session.execute(db.text(query))
    rows = result.fetchall()
    
    # Group by district
    district_groups = {}
    for row in rows:
        from_district = row[4]
        if from_district not in district_groups:
            district_groups[from_district] = []
        district_groups[from_district].append(row)
    
    # Create Excel workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f'{list_type.title()} Transfer List'
    
    # Set column widths
    ws.column_dimensions['A'].width = 8   # Sl. No.
    ws.column_dimensions['B'].width = 12  # PEN
    ws.column_dimensions['C'].width = 25  # Name
    ws.column_dimensions['D'].width = 18  # Designation
    ws.column_dimensions['E'].width = 35  # Office
    ws.column_dimensions['F'].width = 15  # From District
    ws.column_dimensions['G'].width = 15  # To District
    ws.column_dimensions['H'].width = 20  # Protection
    
    # Styles
    header_font = Font(bold=True, size=14)
    subheader_font = Font(bold=True, size=12)
    district_font = Font(bold=True, size=11)
    table_header_font = Font(bold=True, size=10)
    table_header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # Header
    ws.merge_cells('A1:H1')
    ws['A1'] = 'Government of Kerala'
    ws['A1'].font = header_font
    ws['A1'].alignment = center_align
    
    ws.merge_cells('A2:H2')
    ws['A2'] = 'Department: Health Services'
    ws['A2'].font = subheader_font
    ws['A2'].alignment = center_align
    
    ws.merge_cells('A3:H3')
    if transfer_type == 'regular':
        ws['A3'] = f'Regular Transfer for Junior Public Health Nurse Gr. I - {month} {year}'
    else:
        ws['A3'] = f'Norms Based General Transfer for Junior Public Health Nurse Gr. I - {year}'
    ws['A3'].font = subheader_font
    ws['A3'].alignment = center_align
    
    ws.merge_cells('A4:H4')
    ws['A4'] = f'{list_type.title()} Transfer List'
    ws['A4'].font = Font(bold=True, size=10)
    ws['A4'].alignment = Alignment(horizontal='left')
    
    row_num = 6
    sl_no = 1
    
    headers = ['Sl. No.', 'PEN', 'Name', 'Designation', 'Office Transferred from', 
              'From District', 'To District', 'Protection If Any']
    
    for district in DISTRICTS:
        if district not in district_groups:
            continue
        
        records = district_groups[district]
        
        # District header
        ws.merge_cells(f'A{row_num}:H{row_num}')
        ws[f'A{row_num}'] = f'District: {district.upper()}'
        ws[f'A{row_num}'].font = district_font
        row_num += 1
        
        # Table headers
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=row_num, column=col_num, value=header)
            cell.font = table_header_font
            cell.fill = table_header_fill
            cell.border = thin_border
            cell.alignment = center_align
        row_num += 1
        
        # Data rows
        for record in records:
            protection = ""
            if record[6] == 'Yes' and record[7]:
                protection = record[7]
            
            data = [sl_no, record[0], record[1], record[2] or 'JPHN Gr I',
                   record[3], record[4], record[5], protection]
            
            for col_num, value in enumerate(data, 1):
                cell = ws.cell(row=row_num, column=col_num, value=value)
                cell.border = thin_border
                if col_num == 1:
                    cell.alignment = center_align
                else:
                    cell.alignment = left_align
            
            sl_no += 1
            row_num += 1
        
        row_num += 1  # Empty row between districts
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'{list_type}_transfer_list_{datetime.now().strftime("%Y%m%d")}.xlsx'
    )


@app.route('/export/word/<list_type>')
@login_required
@requires_transfer_session
def export_word(list_type):
    """Export list to Word document - matching desktop app format"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.table import Wd_Table_Alignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        flash('python-docx library not installed. Please install with: pip install python-docx', 'error')
        return redirect(url_for('dashboard'))
    
    prefix = get_table_prefix()
    transfer_type = session.get('transfer_type', 'general')
    year = session.get('year', '')
    month = session.get('month', '')
    
    if list_type == 'draft':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, d.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_draft d ON j.pen = d.pen
            ORDER BY j.district, j.name
        """
    elif list_type == 'final':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, f.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_final f ON j.pen = f.pen
            ORDER BY j.district, j.name
        """
    else:
        flash('Invalid export type!', 'error')
        return redirect(url_for('dashboard'))
    
    result = db.session.execute(db.text(query))
    rows = result.fetchall()
    
    # Group by district
    district_groups = {}
    for row in rows:
        from_district = row[4]
        if from_district not in district_groups:
            district_groups[from_district] = []
        district_groups[from_district].append(row)
    
    # Create Word document
    doc = Document()
    
    # Set page margins and landscape for general transfer
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    if transfer_type == 'general':
        # Landscape orientation
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Header
    header1 = doc.add_paragraph('Government of Kerala')
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.runs[0].bold = True
    header1.runs[0].font.size = Pt(14)
    
    header2 = doc.add_paragraph('Department: Health Services')
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.runs[0].bold = True
    header2.runs[0].font.size = Pt(12)
    
    if transfer_type == 'regular':
        title_text = f'Regular Transfer for Junior Public Health Nurse Gr. I - {month} {year}'
    else:
        title_text = f'Norms Based General Transfer for Junior Public Health Nurse Gr. I - {year}'
    
    header3 = doc.add_paragraph(title_text)
    header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header3.runs[0].bold = True
    header3.runs[0].font.size = Pt(12)
    header3.runs[0].underline = True
    
    cadre = doc.add_paragraph(f'{list_type.title()} Transfer List')
    cadre.runs[0].bold = True
    cadre.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    sl_no = 1
    headers = ['Sl. No.', 'PEN', 'Name', 'Designation', 'Office Transferred from', 
              'From District', 'To District', 'Protection If Any']
    
    # Column widths in inches
    col_widths = [0.5, 0.8, 1.8, 1.2, 2.5, 1.2, 1.2, 1.3]
    
    for district in DISTRICTS:
        if district not in district_groups:
            continue
        
        records = district_groups[district]
        
        # District header
        district_para = doc.add_paragraph(f'District: {district.upper()}')
        district_para.runs[0].bold = True
        district_para.runs[0].font.size = Pt(11)
        
        # Create table
        table = doc.add_table(rows=1 + len(records), cols=8)
        table.style = 'Table Grid'
        table.alignment = Wd_Table_Alignment.CENTER
        
        # Set column widths
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        
        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, record in enumerate(records):
            protection = ""
            if record[6] == 'Yes' and record[7]:
                protection = record[7]
            
            data = [str(sl_no), record[0], record[1], record[2] or 'JPHN Gr I',
                   record[3], record[4], record[5], protection]
            
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(data):
                cell = row.cells[col_idx]
                cell.text = str(value) if value else ''
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if col_idx == 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            sl_no += 1
        
        doc.add_paragraph()  # Space between districts
    
    # Signature area
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_para.add_run('_________________________\n')
    sig_run.font.size = Pt(10)
    sig_run2 = sig_para.add_run('Authorized Signatory')
    sig_run2.bold = True
    sig_run2.font.size = Pt(10)
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{list_type}_transfer_list_{datetime.now().strftime("%Y%m%d")}.docx'
    )


@app.route('/export/pdf/<list_type>')
@login_required
@requires_transfer_session
def export_pdf(list_type):
    """Export list to PDF - matching desktop app format"""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
    except ImportError:
        flash('reportlab library not installed. Please install with: pip install reportlab', 'error')
        return redirect(url_for('dashboard'))
    
    prefix = get_table_prefix()
    transfer_type = session.get('transfer_type', 'general')
    year = session.get('year', '')
    month = session.get('month', '')
    
    if list_type == 'draft':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, d.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_draft d ON j.pen = d.pen
            ORDER BY j.district, j.name
        """
    elif list_type == 'final':
        query = f"""
            SELECT j.pen, j.name, j.designation, j.institution, j.district, f.transfer_to_district,
                   j.weightage, j.weightage_details
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_final f ON j.pen = f.pen
            ORDER BY j.district, j.name
        """
    else:
        flash('Invalid export type!', 'error')
        return redirect(url_for('dashboard'))
    
    result = db.session.execute(db.text(query))
    rows = result.fetchall()
    
    # Group by district
    district_groups = {}
    for row in rows:
        from_district = row[4]
        if from_district not in district_groups:
            district_groups[from_district] = []
        district_groups[from_district].append(row)
    
    # Create PDF
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4), 
                           rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, fontSize=14, spaceAfter=6)
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=1, fontSize=12, spaceAfter=6)
    district_style = ParagraphStyle('District', parent=styles['Heading3'], fontSize=11, spaceBefore=12, spaceAfter=6)
    
    # Header
    elements.append(Paragraph('Government of Kerala', title_style))
    elements.append(Paragraph('Department: Health Services', subtitle_style))
    
    if transfer_type == 'regular':
        title_text = f'Regular Transfer for Junior Public Health Nurse Gr. I - {month} {year}'
    else:
        title_text = f'Norms Based General Transfer for Junior Public Health Nurse Gr. I - {year}'
    
    elements.append(Paragraph(title_text, subtitle_style))
    elements.append(Paragraph(f'{list_type.title()} Transfer List', 
                             ParagraphStyle('ListTitle', parent=styles['Normal'], fontSize=10, spaceBefore=6, spaceAfter=12)))
    
    sl_no = 1
    headers = ['Sl.No', 'PEN', 'Name', 'Designation', 'Office', 'From', 'To', 'Protection']
    col_widths = [0.4*inch, 0.7*inch, 1.6*inch, 1.0*inch, 2.2*inch, 1.0*inch, 1.0*inch, 1.2*inch]
    
    for district in DISTRICTS:
        if district not in district_groups:
            continue
        
        records = district_groups[district]
        
        # District header
        elements.append(Paragraph(f'District: {district.upper()}', district_style))
        
        # Table data
        table_data = [headers]
        for record in records:
            protection = ""
            if record[6] == 'Yes' and record[7]:
                protection = record[7][:30] + '...' if len(str(record[7])) > 30 else record[7]
            
            table_data.append([
                str(sl_no), 
                record[0] or '', 
                record[1] or '', 
                record[2] or 'JPHN Gr I',
                record[3] or '', 
                record[4] or '', 
                record[5] or '', 
                protection
            ])
            sl_no += 1
        
        # Create table
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9D9D9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (2, 1), (4, -1), 'LEFT'),  # Name, Designation, Office left align
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')])
        ]))
        
        elements.append(table)
        elements.append(Spacer(1, 12))
    
    # Signature
    elements.append(Spacer(1, 30))
    sig_style = ParagraphStyle('Signature', parent=styles['Normal'], alignment=2, fontSize=10)
    elements.append(Paragraph('_________________________<br/>Authorized Signatory', sig_style))
    
    doc.build(elements)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'{list_type}_transfer_list_{datetime.now().strftime("%Y%m%d")}.pdf'
    )


# ==================== API ROUTES ====================

@app.route('/api/employees/<district>')
@login_required
@requires_transfer_session
def api_employees(district):
    """API endpoint to get employees by district"""
    prefix = get_table_prefix()
    
    result = db.session.execute(db.text(f"""
        SELECT pen, name, institution, duration_days FROM {prefix}jphn 
        WHERE district = :district ORDER BY name
    """), {'district': district})
    
    employees = [{'pen': r[0], 'name': r[1], 'institution': r[2], 
                 'duration': format_duration(r[3])} for r in result.fetchall()]
    
    return jsonify(employees)


@app.route('/api/vacancy-stats')
@login_required
@requires_transfer_session
def api_vacancy_stats():
    """API endpoint to get vacancy statistics"""
    prefix = get_table_prefix()
    
    stats = []
    for district in DISTRICTS:
        try:
            v_result = db.session.execute(db.text(f"""
                SELECT vacancy_reported FROM {prefix}vacancy WHERE district = :district
            """), {'district': district})
            v_row = v_result.fetchone()
            vacancy = v_row[0] if v_row and v_row[0] is not None else 0
        except:
            vacancy = 0
        
        try:
            f_result = db.session.execute(db.text(f"""
                SELECT COUNT(*) FROM {prefix}transfer_draft WHERE transfer_to_district = :district
            """), {'district': district})
            filled = f_result.fetchone()[0] or 0
        except:
            filled = 0
        
        stats.append({
            'district': district,
            'vacancy': vacancy,
            'filled': filled,
            'remaining': max(0, vacancy - filled)
        })
    
    return jsonify(stats)


@app.route('/api/debug/vacancy')
@login_required
@requires_transfer_session
def api_debug_vacancy():
    """Debug: Show raw vacancy table data"""
    prefix = get_table_prefix()
    
    try:
        result = db.session.execute(db.text(f"""
            SELECT district, total_strength, vacancy_reported FROM {prefix}vacancy ORDER BY district
        """))
        rows = result.fetchall()
        data = [{'district': r[0], 'total_strength': r[1], 'vacancy_reported': r[2]} for r in rows]
        return jsonify({
            'table': f'{prefix}vacancy',
            'count': len(data),
            'data': data
        })
    except Exception as e:
        return jsonify({'error': str(e), 'table': f'{prefix}vacancy'})


@app.route('/api/debug/district/<district>')
@login_required
@requires_transfer_session
def api_debug_district(district):
    """Debug: Show detailed breakdown for a specific district"""
    prefix = get_table_prefix()
    
    try:
        # Applied from this district (in transfer_applied)
        applied_from = db.session.execute(db.text(f"""
            SELECT j.pen, j.name, t.pref1, t.pref2, t.pref3
            FROM {prefix}jphn j
            INNER JOIN {prefix}transfer_applied t ON j.pen = t.pen
            WHERE j.district = :district
        """), {'district': district}).fetchall()
        
        # In draft FROM this district (leaving)
        draft_leaving = db.session.execute(db.text(f"""
            SELECT d.pen, j.name, j.district as from_district, d.transfer_to_district, d.remarks
            FROM {prefix}transfer_draft d
            INNER JOIN {prefix}jphn j ON d.pen = j.pen
            WHERE j.district = :district AND j.district != d.transfer_to_district
        """), {'district': district}).fetchall()
        
        # In draft TO this district (filling)
        draft_filling = db.session.execute(db.text(f"""
            SELECT d.pen, j.name, j.district as from_district, d.transfer_to_district
            FROM {prefix}transfer_draft d
            INNER JOIN {prefix}jphn j ON d.pen = j.pen
            WHERE d.transfer_to_district = :district
        """), {'district': district}).fetchall()
        
        return jsonify({
            'district': district,
            'applied_from': {
                'count': len(applied_from),
                'employees': [{'pen': r[0], 'name': r[1], 'pref1': r[2], 'pref2': r[3], 'pref3': r[4]} for r in applied_from]
            },
            'draft_leaving_cascade': {
                'count': len(draft_leaving),
                'employees': [{'pen': r[0], 'name': r[1], 'from': r[2], 'to': r[3], 'remarks': r[4]} for r in draft_leaving]
            },
            'draft_filling': {
                'count': len(draft_filling),
                'employees': [{'pen': r[0], 'name': r[1], 'from': r[2], 'to': r[3]} for r in draft_filling]
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)})


# ==================== ERROR HANDLERS ====================

@app.errorhandler(404)
def not_found_error(error):
    return render_template('error.html', error='Page not found'), 404


@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('error.html', error='Internal server error'), 500


# ==================== RUN ====================

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    with app.app_context():
        # Create tables for default session
        session_data = {'transfer_type': 'general', 'year': datetime.now().year, 'month': ''}
    app.run(debug=True, host='0.0.0.0', port=port)
